"""Workspace file movement and indexing helpers."""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass, replace
from datetime import date, datetime, timedelta
from functools import lru_cache
import hashlib
import hmac
from io import BytesIO
import json
import math
import mimetypes
import os
from pathlib import Path
import re
import secrets
import shutil
import socket
import subprocess
import tempfile
import time as time_module
from typing import Any, Callable, Dict, FrozenSet, Iterable, List, Mapping, Optional, Tuple
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, parse_qsl, quote, urlencode, urlparse, urlunparse
from urllib.request import urlopen
import zipfile

from docx import Document

from uplands_site_command_centre import config
from uplands_site_command_centre.permits.ingestion_engine import IngestionEngine
from uplands_site_command_centre.permits.models import (
    BaseDocument,
    BroadcastDispatchDocument,
    COMMON_CONSTRUCTION_EWC_CODES,
    CarrierComplianceDocument,
    CarrierComplianceDocumentType,
    COSHHDocument,
    DailyAttendanceEntryDocument,
    DocumentStatus,
    FileGroup,
    InductionDocument,
    LadderPermit,
    LadderStabilisationMethod,
    PlantAssetDocument,
    PlantInspectionType,
    PLANT_PENDING_INSPECTION_TEXT,
    RAMSDocument,
    SafetyAsset,
    SITE_CHECK_WEEKDAY_KEYS,
    SiteDiaryDocument,
    SiteAttendanceRegister,
    SiteAttendanceRecord,
    SiteCheckRegister,
    SiteWorker,
    ToolboxTalkDocument,
    ToolboxTalkCompletionDocument,
    TemplateRegistry,
    ValidationError,
    WeeklySiteCheck,
    WeeklySiteCheckRowDefinition,
    WasteRegister,
    WasteTransferNoteDocument,
    WasteTransferNoteSourceSnapshot,
    format_plant_inspection_reference,
    get_weekly_site_check_frequency_for_row,
    infer_plant_inspection_type,
    is_pending_plant_inspection_reference,
)
from uplands_site_command_centre.permits.repository import (
    DocumentNotFoundError,
    DocumentRepository,
    IndexedFileRecord,
)
from uplands_site_command_centre.permits.template_manager import (
    TemplateManager,
    TemplateValidationError,
)


ABUCS_PDF_PATTERN = re.compile(r"^\d+(?:-\d+)?\.pdf$", re.IGNORECASE)
TEMPLATE_TAG_PATTERN = re.compile(r"{{\s*([a-zA-Z0-9_\.]+)\s*}}")
DATE_VALUE_PATTERN = (
    r"(?:\d{1,2}[./-]+\d{1,2}[./-]+\d{2,4}|"
    r"\d{1,2}(?:st|nd|rd|th)?(?:\s+|-)"
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|"
    r"Dec(?:ember)?)(?:\s+|-)\d{4})"
)
NUMERIC_DATE_PATTERN = re.compile(r"\b(\d{1,2})[./-]+(\d{1,2})[./-]+(\d{2,4})\b")
TEXTUAL_DATE_PATTERN = re.compile(
    r"\b(\d{1,2})(?:st|nd|rd|th)?(?:\s+|-)"
    r"(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|"
    r"Dec(?:ember)?)(?:\s+|-)(\d{4})\b",
    re.IGNORECASE,
)
PRIORITY_DATE_PATTERNS = (
    (
        80,
        re.compile(
            rf"\bexpiry\s+date(?:\s+of\s+registration)?"
            rf"(?:\s*\([^)]+\))?\s*(?::|-)?\s*(?P<date>{DATE_VALUE_PATTERN})",
            re.IGNORECASE,
        ),
    ),
    (
        75,
        re.compile(
            rf"\bexpires?(?:\s+on)?\s*(?::|-)?\s*(?P<date>{DATE_VALUE_PATTERN})",
            re.IGNORECASE,
        ),
    ),
    (
        70,
        re.compile(
            rf"\buntil\s*(?::|-)?\s*(?P<date>{DATE_VALUE_PATTERN})",
            re.IGNORECASE,
        ),
    ),
    (
        65,
        re.compile(
            rf"\bexpiry\s*(?::|-)?\s*(?P<date>{DATE_VALUE_PATTERN})",
            re.IGNORECASE,
        ),
    ),
    (
        60,
        re.compile(
            rf"\bexp\s*(?::|-)?\s*(?P<date>{DATE_VALUE_PATTERN})",
            re.IGNORECASE,
        ),
    ),
)


def _weekly_site_check_template_tag(day_key: str, row_number: int) -> str:
    """Return the placeholder tag name for one File 2 weekly checklist cell."""

    return f"{day_key}_{row_number}"
LOW_PRIORITY_KEYWORDS = ("expiry", "until", "period", "ends")
IGNORE_DATE_CONTEXT_KEYWORDS = (
    "dated",
    "created on",
    "created",
    "generation",
    "generated on",
)
CARRIER_NAME_ALIASES = {
    "abucs": "Abucs",
    "abacus bristol ltd": "Abucs",
    "abacus (bristol) ltd": "Abucs",
    "cbdu104060": "Abucs",
    "l&s waste management limited": "Abucs",
    "l & s waste management limited": "Abucs",
    "l&s waste management": "Abucs",
    "l & s waste management": "Abucs",
    "biffa limited": "Abucs",
    "biffa ltd": "Abucs",
    "biffa limited subsidiaries": "Abucs",
    "biffa ltd subsidiaries": "Abucs",
}
MONTH_NAME_MAP = {
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
}
DEFAULT_SITE_NAME = "NG Lovedean Substation"
LOVEDEAN_SITE_LATITUDE = 50.917
LOVEDEAN_SITE_LONGITUDE = -1.036
DEFAULT_WASTE_CARRIER_NAME = "Abucs"
DEFAULT_WASTE_DESCRIPTION = "Mixed Construction"
DEFAULT_EWC_CODE = "17 09 04"
DEFAULT_FOUL_WASTE_DESCRIPTION = "Cess Pit / Foul Waste"
DEFAULT_FOUL_WASTE_EWC_CODE = "20 03 04"
DEFAULT_DESTINATION_FACILITY = "Not captured from ticket PDF"
WASTE_TYPE_PATTERN = re.compile(
    r"\bwaste\s+type[:\s]*(?P<description>.+?)\s+"
    r"(?:customer|payment\s+type|total|i\s+confirm|print|sign)\b",
    re.IGNORECASE,
)
WASTE_COLLECTION_TYPE_PATTERN = re.compile(
    r"\bskip\s+type\s*(?P<value>.+?)\s+"
    r"(?:wastetruck|notes|movement\s+type|waste\s+type|customer|payment\s+type|account|total)\b",
    re.IGNORECASE,
)
EWC_CODE_PATTERN = re.compile(r"\b\d{2}\s?\d{2}\s?\d{2}\*?\b")
FOUL_WASTE_KEYWORDS = (
    "cess pit",
    "cesspit",
    "foul",
    "foul water",
    "foul waste",
    "septic",
    "septic tank",
    "sludge",
    "effluent",
    "tanker",
)
VEHICLE_REG_PRIORITY_PATTERNS = (
    re.compile(
        r"\b(?:vehicle\s+reg(?:istration)?|reg(?:istration)?\s+no\.?|vrm)\s*(?::|-)?\s*"
        r"(?P<reg>[A-Z]{2}\d{2}\s?[A-Z]{3}|[A-Z]{1,3}\s?\d{1,4}\s?[A-Z]{1,3})\b",
        re.IGNORECASE,
    ),
)
GENERIC_VEHICLE_REG_PATTERN = re.compile(
    r"\b([A-Z]{2}\d{2}\s?[A-Z]{3}|[A-Z]{1,3}\s?\d{1,4}\s?[A-Z]{1,3})\b",
    re.IGNORECASE,
)
PROJECT_NUMBER_PATTERNS = (
    re.compile(
        r"\b(?:project|job)\s*(?:number|no\.?)\s*(?::|-)?\s*(?P<value>[A-Z0-9./_-]+)",
        re.IGNORECASE,
    ),
)
TONNAGE_PATTERNS = (
    (
        90,
        re.compile(
            r"\bnet\s+weight\s*(?::|-)?\s*(?P<quantity>\d[\d,]*(?:\.\d+)?)\s*"
            r"(?P<unit>kgs?|kilograms?|tonnes?|tons?|t)\b",
            re.IGNORECASE,
        ),
    ),
    (
        80,
        re.compile(
            r"\btonnage\s*(?::|-)?\s*(?P<quantity>\d[\d,]*(?:\.\d+)?)\s*"
            r"(?P<unit>kgs?|kilograms?|tonnes?|tons?|t)\b",
            re.IGNORECASE,
        ),
    ),
    (
        70,
        re.compile(
            r"\bweight\s*(?::|-)?\s*(?P<quantity>\d[\d,]*(?:\.\d+)?)\s*"
            r"(?P<unit>kgs?|kilograms?|tonnes?|tons?|t)\b",
            re.IGNORECASE,
        ),
    ),
)
HSS_ORDER_REF_PATTERN = re.compile(r"\bH-[A-Z0-9]{6,}\b", re.IGNORECASE)
HSS_STOCK_CODE_PATTERN = re.compile(r"^\d{4,6}$")
HSS_DATE_RANGE_PATTERN = re.compile(
    r"(?P<start>\d{2}/\d{2}/\d{4})\s+(?P<end>\d{2}/\d{2}/\d{4})(?:\*[A-Za-z0-9]+)?",
    re.IGNORECASE,
)
PHONE_PATTERN = re.compile(r"\b0\d{3,4}\s?\d{3}\s?\d{3,4}\b")
EMAIL_PATTERN = re.compile(r"\b[^@\s]+@[^@\s]+\.[^@\s]+\b")
HSS_DEFAULT_PHONE = "0161 749 4090"
PLANT_PENDING_SERIAL_TEXT = ""
DEFAULT_LADDER_PERMIT_MANAGER_NAME = "Ceri Edwards"
DEFAULT_LADDER_PERMIT_MANAGER_POSITION = "Project Manager"
DEFAULT_SAFETY_MANAGER_NAME = "Ceri Edwards"
DEFAULT_SAFETY_MANAGER_POSITION = "Project Manager"
DEFAULT_SAFETY_VERSION = "1.0"
FILE_3_SAFETY_SOURCE_SUFFIXES = frozenset({".pdf", ".docx", ".doc"})
FILE_3_FILENAME_SPLIT_PATTERN = re.compile(r"[-_()]+")
FILE_3_VERSION_SEGMENT_PATTERN = re.compile(
    r"\b(?:rev(?:ision)?|ver(?:sion)?|v)\s*[:._-]?\s*(?P<value>\d+(?:\.\d+)*[a-zA-Z]?)\b",
    re.IGNORECASE,
)
FILE_3_REFERENCE_SEGMENT_PATTERN = re.compile(
    r"\b(?:ref(?:erence)?|doc(?:ument)?)\s*[:._-]?\s*(?P<value>[A-Z0-9./_-]+)\b",
    re.IGNORECASE,
)
FILE_3_COMPANY_HINT_WORDS = frozenset(
    {
        "electrical",
        "engineering",
        "services",
        "service",
        "solutions",
        "construction",
        "contractor",
        "contractors",
        "scaffolding",
        "scaffold",
        "groundworks",
        "maintenance",
        "installations",
        "installers",
        "civil",
        "civils",
        "hire",
        "plant",
        "group",
        "limited",
        "ltd",
        "llp",
        "plc",
        "waste",
    }
)
FILE_3_CONTRACTOR_BLACKLIST = frozenset(
    {
        "ng",
        "national grid",
        "lovedean",
        "waterlooville",
        "substation",
        "rams",
        "coshh",
        "method statement",
        "site",
    }
)
FILE_3_CONTRACTOR_BLACKLIST_WORDS = frozenset(
    {
        word
        for phrase in FILE_3_CONTRACTOR_BLACKLIST
        for word in re.findall(r"[a-z0-9]+", phrase.casefold())
    }
)
FILE_3_PREFERRED_COMPANY_NAMES = (
    "Bluecord",
    "West Coast",
    "Uplands",
    "Lucion",
)
FILE_3_COMPANY_ANCHOR_LABELS = (
    "company name",
    "company",
    "contractor",
)
FILE_3_TITLE_ANCHOR_LABELS = (
    "rams title",
    "activity description",
    "brief description of work",
    "title",
)
FILE_3_VERSION_ANCHOR_LABELS = (
    "rams version",
    "version",
    "rev",
)
FILE_3_SAFETY_FILENAME_STOPWORDS = frozenset(
    {
        "rams",
        "rams document",
        "risk assessment",
        "risk assessment and method statement",
        "method statement",
        "coshh",
        "coshh assessment",
        "safety data sheet",
        "data sheet",
        "sds",
        "msds",
        "final",
        "signed",
        "issue",
        "issued",
        "current",
        "copy",
    }
)
FILE_3_TEXT_FALLBACK_MAX_CHARS = 1500
FILE_3_TEXT_FALLBACK_MAX_LINE_LENGTH = 50
COSHH_KEYWORDS = (
    "coshh",
    "safety data sheet",
    "sds",
    "material safety data sheet",
)
RAMS_KEYWORDS = (
    "rams",
    "risk assessment and method statement",
    "risk assessment",
    "method statement",
)
SAFETY_REFERENCE_PATTERNS = (
    re.compile(
        r"\b(?:reference|document\s+ref(?:erence)?|ref(?:erence)?\s+no\.?|doc(?:ument)?\s+no\.?)\s*(?::|-)?\s*(?P<value>[A-Z0-9./_-]+)",
        re.IGNORECASE,
    ),
)
SAFETY_VERSION_PATTERNS = (
    re.compile(
        r"\b(?:version|revision|rev\.?|v)\s*[:._-]?\s*(?P<value>\d+(?:\.\d+)*[a-zA-Z]?)\b",
        re.IGNORECASE,
    ),
)
FILE_3_EXCLUDED_RAMS_FILENAME_MARKERS = (
    "uhsf20.1",
    "review form",
)
SAFETY_REVIEW_DATE_LABELS = (
    "review date",
    "revision date",
    "date of review",
    "approved date",
    "approval date",
    "issue date",
    "date of issue",
)
COSHH_SUBSTANCE_LABELS = (
    "product name",
    "product name(s)",
    "substance name",
    "product identifier",
    "trade name",
    "mixture name",
)
COSHH_SUPPLIER_LABELS = (
    "supplier",
    "suppliers name",
    "supplier name",
    "supplier name & address",
    "suppliers name & address",
    "manufacturer/supplier",
    "manufacturer",
    "company name",
)
COSHH_USE_LABELS = (
    "recommended use",
    "identified uses",
    "product use",
    "use of substance",
    "intended use",
)
RAMS_ACTIVITY_LABELS = (
    "activity description",
    "scope of works",
    "description of works",
    "works description",
    "activity",
    "brief description of work",
    "brief description of works",
    "job/activity/sequence",
    "task / activity",
)
FILE_3_STRONG_COSHH_TEXT_LABELS = (
    "coshh assessment",
    "safety data sheet",
    "material safety data sheet",
    "product name",
    "substance name",
    "product identifier",
    "trade name",
    "description of substance",
    "suppliers name",
    "supplier name",
    "supplier",
    "workplace exposure limits",
    "risk phrases",
    "safety phrases",
    "msds attached",
)
FILE_3_REVIEW_HOLD_KEYWORDS = (
    *COSHH_KEYWORDS,
    *RAMS_KEYWORDS,
    "review form",
    "construction phase plan",
    "uhsf15.1",
)
FILE_3_SUSPICIOUS_COMPANY_LABELS = frozenset(
    {
        "review",
        "company",
        "company name",
        "contractor",
        "contractor name",
        "supplier",
        "manufacturer",
        "main contractor",
        "subcontractor",
        "site contractor",
    }
)
FILE_3_COMMON_PERSON_FIRST_NAMES = frozenset(
    {
        "adam",
        "alex",
        "andrew",
        "ben",
        "cameron",
        "ceri",
        "chris",
        "dan",
        "david",
        "gary",
        "george",
        "jake",
        "jacob",
        "james",
        "jane",
        "john",
        "josh",
        "lee",
        "luke",
        "mark",
        "matt",
        "michael",
        "nick",
        "paul",
        "rob",
        "robert",
        "ryan",
        "sam",
        "sean",
        "steve",
        "steven",
        "tom",
    }
)


@dataclass(frozen=True)
class FiledAsset:
    """A file moved from the ingest inbox into the workspace."""

    original_path: Path
    destination_path: Path
    file_category: str
    related_doc_id: Optional[str] = None
    auto_captured_expiry_date: Optional[date] = None
    auto_captured_carrier_name: Optional[str] = None
    auto_captured_document_type: Optional[CarrierComplianceDocumentType] = None


@dataclass(frozen=True)
class SiteInductionAuditResult:
    """Summary of today's induction coverage against the attendance register."""

    audit_date: date
    site_name: Optional[str]
    workers_on_site: List[str]
    inducted_workers: List[str]
    missing_workers: List[str]
    matched_files: Dict[str, Path]

    @property
    def is_compliant(self) -> bool:
        """Return True when no worker on site is missing an induction file."""

        return not self.missing_workers


@dataclass(frozen=True)
class GeneratedLadderPermit:
    """A populated ladder permit file ready for review and signature."""

    permit: LadderPermit
    output_path: Path
    induction_file: Optional[Path] = None


@dataclass(frozen=True)
class GeneratedPermitRegisterDocument:
    """A populated File 4 permit register document ready for printing."""

    output_path: Path
    permit_count: int


@dataclass(frozen=True)
class GeneratedPlantRegisterDocument:
    """A populated File 2 plant register document ready for printing."""

    output_path: Path
    asset_count: int


@dataclass(frozen=True)
class GeneratedSafetyRegisterDocument:
    """A populated File 3 safety register document ready for printing."""

    output_path: Path
    row_count: int
    register_type: str


@dataclass(frozen=True)
class WasteKpiSheetMetadata:
    """Header values sourced from the File 1 KPI workbook."""

    workbook_path: Optional[Path]
    client_name: str
    site_address: str
    project_number: str
    manager_name: str


@dataclass(frozen=True)
class SmartScannedWasteTransferNote:
    """Best-effort fields extracted from an uploaded WTN image or PDF."""

    source_name: str
    wtn_number: str
    carrier_name: str
    vehicle_registration: str
    collection_type: str
    waste_description: str
    ticket_date: date
    quantity_tonnes: Optional[float]
    ewc_code: str
    destination_facility: str
    extracted_text: str


@dataclass(frozen=True)
class WasteTransferNoteSourceCandidate:
    """One physical waste-ticket source file evaluated for a WTN."""

    source_path: Path
    scanned_note: SmartScannedWasteTransferNote
    site_name: str
    source_mtime: float


@dataclass(frozen=True)
class WasteTransferNoteSourceConflict:
    """A WTN that has more than one physical source file on disk."""

    wtn_number: str
    site_name: str
    canonical_source: WasteTransferNoteSourceCandidate
    source_candidates: List[WasteTransferNoteSourceCandidate]

    @property
    def duplicate_count(self) -> int:
        """Return the number of competing physical source files."""

        return len(self.source_candidates)


@dataclass(frozen=True)
class LoggedWasteTransferNote:
    """A WTN file plus its persisted File 1 document."""

    waste_transfer_note: WasteTransferNoteDocument
    stored_file_path: Path
    register_document: Optional[WasteRegister] = None


@dataclass(frozen=True)
class GeneratedWasteRegisterDocument:
    """A populated File 1 waste register ready for printing."""

    output_path: Path
    row_count: int


@dataclass(frozen=True)
class File4PermitSyncResult:
    """Summary of File 4 permit records removed during a folder/database sync."""

    removed_doc_ids: List[str]
    removed_indexed_files: List[Path]

    @property
    def removed_count(self) -> int:
        """Return the number of removed ladder permit records."""

        return len(self.removed_doc_ids)


@dataclass(frozen=True)
class GeneratedSiteCheckChecklist:
    """A populated File 2 checklist document ready for printing."""

    register: SiteCheckRegister
    output_path: Path


@dataclass(frozen=True)
class GeneratedWeeklySiteCheckChecklist:
    """A populated weekly File 2 checklist document ready for printing."""

    weekly_site_check: WeeklySiteCheck
    output_path: Path


@dataclass(frozen=True)
class InductionRecord:
    """Typed File 3 induction form payload captured from the kiosk."""

    full_name: str
    company: str
    cscs_number: str
    cscs_expiry: Optional[date]
    emergency_contact: str
    emergency_tel: str
    medical: str
    asbestos_cert: bool
    erect_scaffold: bool
    cisrs_no: str
    cisrs_expiry: Optional[date]
    operate_plant: bool
    cpcs_no: str
    cpcs_expiry: Optional[date]
    client_training_desc: str
    client_training_date: Optional[date]
    client_training_expiry: Optional[date]
    first_aider: bool
    fire_warden: bool
    supervisor: bool
    smsts: bool
    signature_image_path: Path
    home_address: str = ""
    contact_number: str = ""
    occupation: str = ""
    site_name: str = DEFAULT_SITE_NAME


@dataclass(frozen=True)
class GeneratedInductionDocument:
    """A completed site induction document plus its persisted record."""

    induction_document: InductionDocument
    output_path: Path
    signature_path: Path


@dataclass(frozen=True)
class LoggedDailyAttendanceEntry:
    """One persisted UHSF16.09 sign-in/out entry with its latest signature path."""

    attendance_entry: DailyAttendanceEntryDocument
    signature_path: Path


@dataclass(frozen=True)
class SiteBroadcastContact:
    """One active on-site contact ready for emergency broadcast use."""

    individual_name: str
    contractor_name: str
    mobile_number: str
    vehicle_registration: str
    linked_induction_doc_id: str


@dataclass(frozen=True)
class MessagesDraftLaunchResult:
    """Result summary for one Messages draft launch batch."""

    draft_links: List[str]
    recipient_count: int
    chunk_count: int
    launched_successfully: bool
    error_message: str = ""


def _launch_messages_group_draft_via_gui_automation(
    mobile_numbers: Iterable[str],
    *,
    message: str = "",
) -> Tuple[bool, str]:
    """Try to open one grouped Messages draft via macOS UI automation."""

    resolved_numbers = _deduplicate_mobile_numbers(mobile_numbers)
    if len(resolved_numbers) < 2:
        return False, "A grouped Messages draft needs at least two recipients."

    applescript_source = """
on run argv
    if (count of argv) < 2 then error "Expected a message followed by one or more recipients."
    set messageBody to item 1 of argv
    set recipientList to items 2 thru -1 of argv

    tell application "Messages"
        activate
    end tell

    delay 0.6

    tell application "System Events"
        tell process "Messages"
            keystroke "n" using command down
            delay 0.8

            repeat with targetRecipient in recipientList
                set the clipboard to (targetRecipient as text)
                keystroke "v" using command down
                delay 0.2
                key code 36
                delay 0.45
            end repeat

            key code 48
            delay 0.35

            if (messageBody as text) is not "" then
                set the clipboard to (messageBody as text)
                keystroke "v" using command down
            end if
        end tell
    end tell

    return "ok"
end run
""".strip()

    with tempfile.NamedTemporaryFile(
        mode="w",
        suffix=".applescript",
        encoding="utf-8",
        delete=False,
    ) as script_file:
        script_file.write(applescript_source)
        script_path = Path(script_file.name)

    try:
        launch_result = subprocess.run(
            ["osascript", str(script_path), message.strip(), *resolved_numbers],
            capture_output=True,
            text=True,
            check=False,
        )
    finally:
        try:
            script_path.unlink()
        except FileNotFoundError:
            pass

    if launch_result.returncode == 0:
        return True, ""

    launch_error = (
        launch_result.stderr.strip()
        or launch_result.stdout.strip()
        or "Grouped Messages draft automation failed."
    )
    return False, launch_error


@dataclass(frozen=True)
class LoggedBroadcastDispatch:
    """One persisted broadcast dispatch plus the launch result used to create it."""

    dispatch_document: BroadcastDispatchDocument
    launch_result: MessagesDraftLaunchResult


@dataclass(frozen=True)
class LoggedToolboxTalkCompletion:
    """One remote UHSF16.2 completion plus its stored signature path."""

    toolbox_talk_completion: ToolboxTalkCompletionDocument
    signature_path: Path


@dataclass(frozen=True)
class SavedToolboxTalkDocument:
    """One uploaded toolbox talk source file plus its stored path."""

    toolbox_talk_document: ToolboxTalkDocument
    stored_path: Path


@dataclass(frozen=True)
class GeneratedToolboxTalkRegisterDocument:
    """A populated UHSF16.2 toolbox talk register document ready for printing."""

    output_path: Path
    row_count: int


@dataclass(frozen=True)
class GeneratedAttendanceRegisterDocument:
    """A populated File 2 attendance register document ready for printing."""

    output_path: Path
    row_count: int


@dataclass(frozen=True)
class GeneratedSiteDiaryDocument:
    """A populated UHSF15.63 diary document ready for printing."""

    output_path: Path
    contractor_count: int
    visitor_count: int
    site_diary_document: SiteDiaryDocument


@dataclass(frozen=True)
class GeneratedSiteInductionPoster:
    """A printable induction poster plus the QR target URL."""

    induction_url: str
    qr_code_png: bytes
    poster_png: bytes


@dataclass(frozen=True)
class WorkspaceDiagnosticCheck:
    """One filesystem health check for the live workspace."""

    label: str
    path: Path
    exists: bool
    expected_kind: str
    display_path: str


@dataclass(frozen=True)
class RebuiltSafetyInventoryResult:
    """Summary of one File 3 rebuild pass from filed source documents."""

    total_sources: int
    rams_records: int
    coshh_records: int
    moved_files: int
    ignored_sources: int
    moved_file_names: Tuple[str, ...] = ()
    ignored_file_names: Tuple[str, ...] = ()


def run_workspace_diagnostic() -> List[WorkspaceDiagnosticCheck]:
    """Return the live workspace health checks used by the sidebar doctor."""

    checks: List[WorkspaceDiagnosticCheck] = []

    def add_check(
        label: str,
        path: Path,
        expected_kind: str,
        *,
        display_path_override: Optional[str] = None,
    ) -> None:
        if expected_kind == "dir":
            exists = path.exists() and path.is_dir()
        else:
            exists = path.exists() and path.is_file()
        if display_path_override is not None:
            display_path = display_path_override
        else:
            try:
                display_path = os.path.relpath(str(path), str(config.PROJECT_ROOT))
                if display_path.startswith(".."):
                    raise ValueError
            except ValueError:
                display_path = str(path)
        checks.append(
            WorkspaceDiagnosticCheck(
                label=label,
                path=path,
                exists=exists,
                expected_kind=expected_kind,
                display_path=display_path,
            )
        )

    add_check("Workspace Root", config.BASE_DATA_DIR, "dir")
    add_check("Templates Folder", config.PROJECT_ROOT / "templates", "dir")
    add_check("File 2 Output", config.FILE_2_OUTPUT_DIR, "dir")
    add_check("File 3 Needs Review", config.FILE_3_REVIEW_DIR, "dir")
    add_check("File 3 Signatures", config.FILE_3_SIGNATURES_DIR, "dir")
    add_check(
        "File 3 Completed Inductions",
        config.FILE_3_COMPLETED_INDUCTIONS_DIR,
        "dir",
    )

    seen_template_paths: set[Path] = set()
    for document_type, registered_path in TemplateRegistry.TEMPLATE_PATHS.items():
        resolved_path = (
            registered_path.resolve()
            if registered_path.is_absolute()
            else (TemplateRegistry.PROJECT_ROOT / registered_path).resolve()
        )
        if resolved_path in seen_template_paths:
            continue
        seen_template_paths.add(resolved_path)
        add_check(
            f"Template · {document_type}",
            resolved_path,
            "file",
            display_path_override=str(registered_path),
        )

    add_check("Database", config.DATABASE_PATH, "file")
    return checks


def _normalise_site_history_payload(
    value: Any,
    *,
    max_entries: int = 6,
) -> Dict[str, List[str]]:
    """Return one cleaned site->recent-values mapping from persisted settings."""

    if not isinstance(value, Mapping):
        return {}

    normalised_payload: Dict[str, List[str]] = {}
    for raw_site_name, raw_entries in value.items():
        site_name = str(raw_site_name or "").strip()
        if not site_name or not isinstance(raw_entries, list):
            continue
        cleaned_entries: List[str] = []
        seen_entries: set[str] = set()
        for raw_entry in raw_entries:
            cleaned_entry = str(raw_entry or "").strip()
            if not cleaned_entry:
                continue
            dedupe_key = cleaned_entry.casefold()
            if dedupe_key in seen_entries:
                continue
            cleaned_entries.append(cleaned_entry)
            seen_entries.add(dedupe_key)
            if len(cleaned_entries) >= max_entries:
                break
        if cleaned_entries:
            normalised_payload[site_name] = cleaned_entries
    return normalised_payload


def load_app_settings() -> Dict[str, Any]:
    """Load the root-level settings used by the permanent tunnel workflow."""

    try:
        payload = json.loads(config.SETTINGS_PATH.read_text(encoding="utf-8"))
    except FileNotFoundError:
        return {
            "public_tunnel_url": "",
            "broadcast_message_history_by_site": {},
            "tbt_topic_history_by_site": {},
            "gate_access_secret": "",
        }
    except (OSError, json.JSONDecodeError, TypeError, ValueError):
        return {
            "public_tunnel_url": "",
            "broadcast_message_history_by_site": {},
            "tbt_topic_history_by_site": {},
            "gate_access_secret": "",
        }

    return {
        "public_tunnel_url": normalize_public_app_url(
            str(payload.get("public_tunnel_url") or "").strip()
        ),
        "broadcast_message_history_by_site": _normalise_site_history_payload(
            payload.get("broadcast_message_history_by_site")
        ),
        "tbt_topic_history_by_site": _normalise_site_history_payload(
            payload.get("tbt_topic_history_by_site")
        ),
        "gate_access_secret": str(payload.get("gate_access_secret") or "").strip(),
    }


def save_app_settings(
    *,
    public_tunnel_url: Optional[str] = None,
    broadcast_message_history_by_site: Optional[Mapping[str, List[str]]] = None,
    tbt_topic_history_by_site: Optional[Mapping[str, List[str]]] = None,
    gate_access_secret: Optional[str] = None,
) -> None:
    """Persist the root-level app settings used by the live manager workflows."""

    existing_settings = load_app_settings()
    resolved_public_tunnel_url = (
        normalize_public_app_url(str(public_tunnel_url).strip())
        if public_tunnel_url is not None
        else normalize_public_app_url(
            str(existing_settings.get("public_tunnel_url") or "").strip()
        )
    )
    resolved_broadcast_history = (
        _normalise_site_history_payload(broadcast_message_history_by_site)
        if broadcast_message_history_by_site is not None
        else _normalise_site_history_payload(
            existing_settings.get("broadcast_message_history_by_site")
        )
    )
    resolved_tbt_history = (
        _normalise_site_history_payload(tbt_topic_history_by_site)
        if tbt_topic_history_by_site is not None
        else _normalise_site_history_payload(
            existing_settings.get("tbt_topic_history_by_site")
        )
    )
    resolved_gate_access_secret = (
        str(gate_access_secret).strip()
        if gate_access_secret is not None
        else str(existing_settings.get("gate_access_secret") or "").strip()
    )
    config.SETTINGS_PATH.write_text(
        json.dumps(
            {
                "public_tunnel_url": resolved_public_tunnel_url,
                "broadcast_message_history_by_site": resolved_broadcast_history,
                "tbt_topic_history_by_site": resolved_tbt_history,
                "gate_access_secret": resolved_gate_access_secret,
            },
            indent=2,
            sort_keys=True,
        ),
        encoding="utf-8",
    )


def ensure_gate_access_secret() -> str:
    """Return one persistent secret used to build short-lived gate access codes."""

    existing_secret = str(load_app_settings().get("gate_access_secret") or "").strip()
    if existing_secret:
        return existing_secret

    generated_secret = secrets.token_hex(24)
    save_app_settings(gate_access_secret=generated_secret)
    return generated_secret


def _build_gate_access_payload(site_name: str, slot_index: int, slot_minutes: int) -> str:
    """Return the HMAC payload for one site-specific gate code slot."""

    normalized_site_name = re.sub(r"[^a-z0-9]+", "-", site_name.casefold()).strip("-")
    return f"{normalized_site_name}|{slot_index}|{slot_minutes}"


def build_site_gate_access_code(
    site_name: str,
    *,
    at_time: Optional[datetime] = None,
    slot_minutes: int = 30,
) -> Tuple[str, int]:
    """Return the active six-digit site gate code and minutes until it refreshes."""

    resolved_time = at_time or datetime.now()
    slot_seconds = max(1, int(slot_minutes)) * 60
    slot_index = int(resolved_time.timestamp() // slot_seconds)
    payload = _build_gate_access_payload(site_name, slot_index, max(1, int(slot_minutes)))
    digest = hmac.new(
        ensure_gate_access_secret().encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    gate_code = f"{int(digest[:12], 16) % 1_000_000:06d}"
    slot_expires_at = datetime.fromtimestamp((slot_index + 1) * slot_seconds)
    minutes_remaining = max(
        1,
        int(math.ceil((slot_expires_at - resolved_time).total_seconds() / 60.0)),
    )
    return gate_code, minutes_remaining


def validate_site_gate_access_code(
    site_name: str,
    submitted_code: str,
    *,
    at_time: Optional[datetime] = None,
    slot_minutes: int = 30,
    accepted_previous_slots: int = 1,
) -> bool:
    """Return True when the submitted gate code matches the current live site code."""

    cleaned_code = re.sub(r"\D", "", str(submitted_code or ""))
    if len(cleaned_code) != 6:
        return False

    resolved_time = at_time or datetime.now()
    slot_seconds = max(1, int(slot_minutes)) * 60
    current_slot_index = int(resolved_time.timestamp() // slot_seconds)
    max_previous_slots = max(0, int(accepted_previous_slots))

    for slot_offset in range(0, max_previous_slots + 1):
        candidate_slot_index = current_slot_index - slot_offset
        payload = _build_gate_access_payload(
            site_name,
            candidate_slot_index,
            max(1, int(slot_minutes)),
        )
        digest = hmac.new(
            ensure_gate_access_secret().encode("utf-8"),
            payload.encode("utf-8"),
            hashlib.sha256,
        ).hexdigest()
        expected_code = f"{int(digest[:12], 16) % 1_000_000:06d}"
        if hmac.compare_digest(cleaned_code, expected_code):
            return True

    return False


def detect_public_tunnel_url_from_log() -> str:
    """Return the latest public tunnel URL from tunnel.log."""

    try:
        log_text = config.TUNNEL_LOG_PATH.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return ""

    permanent_tunnel_match = re.search(
        r"https://uplands-site-induction\.omegaleague\.win(?:[A-Za-z0-9._~:/?#[\]@!$&'()*+,;=%-]*)?",
        log_text,
        flags=re.IGNORECASE,
    )
    if permanent_tunnel_match is not None:
        return permanent_tunnel_match.group(0).rstrip(".,);")

    tunnel_urls = re.findall(
        r"https://[A-Za-z0-9.-]+(?:trycloudflare\.com|omegaleague\.win)(?:[A-Za-z0-9._~:/?#[\]@!$&'()*+,;=%-]*)?",
        log_text,
        flags=re.IGNORECASE,
    )
    for tunnel_url in reversed(tunnel_urls):
        return tunnel_url.rstrip(".,);")
    return ""


def get_local_ip_address() -> str:
    """Return the current machine's local IP address for on-site QR links."""

    udp_socket = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        udp_socket.connect(("8.8.8.8", 80))
        local_ip_address = udp_socket.getsockname()[0]
    except OSError:
        local_ip_address = "127.0.0.1"
    finally:
        udp_socket.close()

    return local_ip_address or "127.0.0.1"


def lookup_uk_postcode_details(postcode: str) -> Optional[Dict[str, Any]]:
    """Return normalized postcode details from postcodes.io for Project Setup."""

    normalized_postcode = " ".join(str(postcode or "").upper().split())
    if not normalized_postcode:
        return None

    lookup_url = (
        "https://api.postcodes.io/postcodes/"
        f"{quote(normalized_postcode, safe='')}"
    )
    try:
        with urlopen(lookup_url, timeout=8) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except (HTTPError, URLError, TimeoutError, json.JSONDecodeError, OSError, ValueError):
        return None

    result = payload.get("result")
    if not isinstance(result, dict):
        return None

    latitude = result.get("latitude")
    longitude = result.get("longitude")
    returned_postcode = str(result.get("postcode") or normalized_postcode).strip()
    try:
        resolved_latitude = float(latitude)
        resolved_longitude = float(longitude)
    except (TypeError, ValueError):
        return None

    def _clean_part(value: Any) -> str:
        return str(value or "").strip()

    district = (
        _clean_part(result.get("admin_district"))
        or _clean_part(result.get("admin_county"))
        or _clean_part(result.get("region"))
    )
    locality = (
        _clean_part(result.get("parish"))
        or _clean_part(result.get("admin_ward"))
        or district
    )
    country = _clean_part(result.get("country"))

    formatted_parts: List[str] = []
    for candidate in (locality, district, returned_postcode):
        if not candidate:
            continue
        if any(existing.casefold() == candidate.casefold() for existing in formatted_parts):
            continue
        formatted_parts.append(candidate)
    formatted_address = ", ".join(formatted_parts) or returned_postcode

    return {
        "postcode": returned_postcode,
        "latitude": resolved_latitude,
        "longitude": resolved_longitude,
        "locality": locality,
        "district": district,
        "country": country,
        "formatted_address": formatted_address,
    }


def lookup_uk_postcode_coordinates(postcode: str) -> Optional[Tuple[float, float, str]]:
    """Return latitude, longitude, and normalized postcode from postcodes.io."""

    postcode_details = lookup_uk_postcode_details(postcode)
    if postcode_details is None:
        return None
    return (
        float(postcode_details["latitude"]),
        float(postcode_details["longitude"]),
        str(postcode_details["postcode"]),
    )


def get_site_induction_url(*, port: int = 8501) -> str:
    """Return the mobile sign-in URL for the current machine."""

    return build_site_induction_url(port=port)


def normalize_public_app_url(public_url: str, *, port: int = 8501) -> str:
    """Return the real app URL when a saved tunnel URL points at a component iframe."""

    cleaned_public_url = str(public_url or "").strip()
    if not cleaned_public_url:
        return ""
    if "://" not in cleaned_public_url:
        cleaned_public_url = f"https://{cleaned_public_url}"

    parsed_url = urlparse(cleaned_public_url)
    streamlit_targets = parse_qs(parsed_url.query).get("streamlitUrl", [])
    if streamlit_targets:
        return normalize_public_app_url(streamlit_targets[0], port=port)

    normalized_path = parsed_url.path or "/"
    if normalized_path.startswith("/component/"):
        normalized_path = "/"

    rebuilt_url = parsed_url._replace(
        path=normalized_path,
        params="",
        query="",
        fragment="",
    )
    normalized_url = urlunparse(rebuilt_url).strip()
    return normalized_url[:-1] if normalized_url.endswith("/") else normalized_url


def build_site_induction_url(
    *,
    public_url: str = "",
    port: int = 8501,
    kiosk_mode: bool = False,
) -> str:
    """Return the best available induction URL for manager or kiosk use."""

    cleaned_public_url = normalize_public_app_url(public_url, port=port)

    if cleaned_public_url:
        parsed_url = urlparse(cleaned_public_url)
    else:
        parsed_url = urlparse(f"http://{get_local_ip_address()}:{port}/")

    query_items = dict(parse_qsl(parsed_url.query, keep_blank_values=True))
    query_items["station"] = "induction"
    if kiosk_mode:
        query_items["mode"] = "kiosk"
    else:
        query_items.pop("mode", None)

    normalized_path = parsed_url.path or "/"
    rebuilt_url = parsed_url._replace(
        path=normalized_path,
        query=urlencode(query_items),
    )
    return urlunparse(rebuilt_url)


def build_toolbox_talk_url(
    topic: str,
    *,
    public_url: str = "",
    port: int = 8501,
) -> str:
    """Return the remote mobile UHSF16.2 signing link for one toolbox talk topic."""

    cleaned_public_url = normalize_public_app_url(public_url, port=port)

    if cleaned_public_url:
        parsed_url = urlparse(cleaned_public_url)
    else:
        parsed_url = urlparse(f"http://{get_local_ip_address()}:{port}/")

    query_items = dict(parse_qsl(parsed_url.query, keep_blank_values=True))
    query_items["station"] = "tbt"
    query_items["topic"] = topic.strip()

    normalized_path = parsed_url.path
    if normalized_path in ("", "/"):
        normalized_path = ""
    rebuilt_url = parsed_url._replace(
        path=normalized_path,
        query=urlencode(query_items),
    )
    return urlunparse(rebuilt_url)


def build_toolbox_talk_document_view_url(
    toolbox_talk_doc_id: str,
    *,
    public_url: str = "",
    port: int = 8501,
) -> str:
    """Return one public browser preview URL for an uploaded toolbox talk document."""

    resolved_doc_id = str(toolbox_talk_doc_id or "").strip()
    if not resolved_doc_id:
        return ""

    cleaned_public_url = normalize_public_app_url(public_url, port=port)
    if cleaned_public_url:
        parsed_url = urlparse(cleaned_public_url)
    else:
        parsed_url = urlparse(f"http://{get_local_ip_address()}:{port}/")

    rebuilt_url = parsed_url._replace(
        path="/gps/tbt-preview",
        query=urlencode({"doc_id": resolved_doc_id}),
    )
    return urlunparse(rebuilt_url)


def generate_site_induction_poster(
    *,
    site_name: str,
    logo_path: Optional[Path] = None,
    public_url: str = "",
    port: int = 8501,
) -> GeneratedSiteInductionPoster:
    """Build a QR code poster PNG for the mobile site induction kiosk."""

    try:
        import qrcode
        from PIL import Image, ImageDraw, ImageFilter, ImageFont, ImageOps
    except ImportError as exc:
        raise RuntimeError(
            "qrcode and Pillow are required to generate the induction poster."
        ) from exc

    induction_url = config.PERMANENT_INDUCTION_KIOSK_URL

    qr_code = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_Q,
        box_size=12,
        border=2,
    )
    qr_code.add_data(induction_url)
    qr_code.make(fit=True)
    qr_image = qr_code.make_image(fill_color="black", back_color="white").convert("RGB")

    qr_buffer = BytesIO()
    qr_image.save(qr_buffer, format="PNG")

    poster_width = 1240
    poster_height = 1754
    poster_image = Image.new("RGBA", (poster_width, poster_height), "#FFFFFF")
    heading_font = _load_poster_font(94, bold=True)
    subheading_font = _load_poster_font(48, bold=True)
    site_font = _load_poster_font(30, bold=False)
    step_number_font = _load_poster_font(44, bold=True)
    step_label_font = _load_poster_font(25, bold=True)
    body_font = _load_poster_font(22, bold=False)

    border_color = "#18181B"
    primary_text = "#111827"
    secondary_text = "#4B5563"
    accent_magenta = "#D1228E"
    accent_red = "#E11D48"
    glass_outline = "#E8AECF"

    background_overlay = Image.new("RGBA", (poster_width, poster_height), (255, 255, 255, 0))
    background_draw = ImageDraw.Draw(background_overlay)
    background_draw.ellipse(
        [(-120, 180), (440, 740)],
        fill=(242, 198, 225, 120),
    )
    background_draw.ellipse(
        [(poster_width - 500, -120), (poster_width + 140, 460)],
        fill=(214, 229, 255, 110),
    )
    background_draw.ellipse(
        [(poster_width - 420, poster_height - 620), (poster_width + 120, poster_height - 80)],
        fill=(255, 220, 226, 105),
    )
    background_overlay = background_overlay.filter(ImageFilter.GaussianBlur(radius=48))
    poster_image.alpha_composite(background_overlay)

    draw = ImageDraw.Draw(poster_image)
    draw.rounded_rectangle(
        [(18, 18), (poster_width - 18, poster_height - 18)],
        radius=26,
        outline=border_color,
        width=5,
    )
    _draw_glass_panel(
        poster_image,
        (56, 48, poster_width - 56, 232),
        radius=42,
        fill=(255, 255, 255, 228),
        outline="#E5E7EB",
        shadow=(17, 24, 39, 26),
    )
    _draw_gradient_line(
        draw,
        left=86,
        right=poster_width - 86,
        y=224,
        width=7,
        start_color=(209, 34, 142),
        end_color=(225, 29, 72),
    )

    current_y = 74
    logo_box = (252, current_y, poster_width - 252, current_y + 112)
    draw.rounded_rectangle(
        logo_box,
        radius=26,
        fill=(255, 255, 255, 255),
        outline="#F8FAFC",
        width=1,
    )
    if logo_path is not None and logo_path.exists():
        try:
            logo_image = Image.open(logo_path).convert("RGBA")
            logo_image = ImageOps.contain(logo_image, (500, 96))
            logo_x = (poster_width - logo_image.width) // 2
            logo_y = logo_box[1] + ((logo_box[3] - logo_box[1]) - logo_image.height) // 2
            poster_image.paste(logo_image, (logo_x, logo_y), logo_image)
        except OSError:
            pass
    current_y = 278

    _draw_centered_text(
        draw,
        poster_width // 2,
        current_y,
        "SITE INDUCTION",
        font=heading_font,
        fill=primary_text,
    )
    current_y += 90
    _draw_centered_text(
        draw,
        poster_width // 2,
        current_y,
        "SCAN TO SIGN IN",
        font=subheading_font,
        fill=accent_magenta,
    )
    current_y += 72
    _draw_centered_text(
        draw,
        poster_width // 2,
        current_y,
        site_name,
        font=site_font,
        fill=secondary_text,
    )
    current_y += 62

    step_box_top = current_y
    step_box_height = 182
    step_gap = 28
    step_box_width = (poster_width - 180 - (2 * step_gap)) // 3
    step_labels = [
        ("1", "SCAN\nQR", "Open your camera"),
        ("2", "COMPLETE\nFORM", "Enter your details"),
        ("3", "SIGN\nSCREEN", "Sign to finish"),
    ]
    for index, (step_number, step_label, step_caption) in enumerate(step_labels):
        step_left = 90 + index * (step_box_width + step_gap)
        step_right = step_left + step_box_width
        _draw_glass_panel(
            poster_image,
            (
                step_left,
                step_box_top,
                step_right,
                step_box_top + step_box_height,
            ),
            radius=28,
            fill=(255, 255, 255, 176),
            outline=glass_outline,
            shadow=(209, 34, 142, 30),
        )
        _draw_glass_panel(
            poster_image,
            (
                step_left + (step_box_width // 2) - 38,
                step_box_top + 18,
                step_left + (step_box_width // 2) + 38,
                step_box_top + 88,
            ),
            radius=22,
            fill=(255, 255, 255, 212),
            outline="#F5BAD7",
            shadow=(17, 24, 39, 20),
        )
        draw.text(
            (step_left + (step_box_width // 2), step_box_top + 53),
            step_number,
            fill=accent_red,
            font=step_number_font,
            anchor="mm",
        )
        _draw_centered_multiline_text(
            draw,
            step_left + (step_box_width // 2),
            step_box_top + 114,
            step_label,
            font=step_label_font,
            fill=primary_text,
            spacing=2,
        )
        _draw_centered_text(
            draw,
            step_left + (step_box_width // 2),
            step_box_top + 156,
            step_caption,
            font=body_font,
            fill=secondary_text,
        )
    current_y = step_box_top + step_box_height + 52

    qr_display_size = 680
    qr_display = ImageOps.contain(qr_image, (qr_display_size, qr_display_size))
    qr_panel = (
        (poster_width - 820) // 2,
        current_y - 6,
        (poster_width - 820) // 2 + 820,
        current_y - 6 + 820,
    )
    _draw_glass_panel(
        poster_image,
        qr_panel,
        radius=42,
        fill=(255, 255, 255, 228),
        outline="#E5E7EB",
        shadow=(17, 24, 39, 34),
    )
    qr_x = (poster_width - qr_display.width) // 2
    poster_image.paste(qr_display, (qr_x, current_y + 64))

    poster_buffer = BytesIO()
    poster_image.convert("RGB").save(poster_buffer, format="PNG")

    return GeneratedSiteInductionPoster(
        induction_url=induction_url,
        qr_code_png=qr_buffer.getvalue(),
        poster_png=poster_buffer.getvalue(),
    )


def _load_poster_font(size: int, *, bold: bool) -> Any:
    """Return a poster font with a sensible macOS fallback chain."""

    try:
        from PIL import ImageFont
    except ImportError:
        return None

    candidate_font_paths = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
    ]
    for candidate_font_path in candidate_font_paths:
        try:
            return ImageFont.truetype(candidate_font_path, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _draw_centered_text(
    draw: Any,
    x: int,
    y: int,
    text: str,
    *,
    font: Any,
    fill: str,
) -> None:
    """Draw one centered text line on the poster."""

    draw.text((x, y), text, fill=fill, font=font, anchor="ma")


def _draw_centered_multiline_text(
    draw: Any,
    x: int,
    y: int,
    text: str,
    *,
    font: Any,
    fill: str,
    spacing: int = 4,
) -> None:
    """Draw centered multiline text using its rendered bounding box."""

    bounding_box = draw.multiline_textbbox(
        (0, 0),
        text,
        font=font,
        spacing=spacing,
        align="center",
    )
    text_width = bounding_box[2] - bounding_box[0]
    text_height = bounding_box[3] - bounding_box[1]
    draw.multiline_text(
        (x - (text_width / 2), y - (text_height / 2)),
        text,
        fill=fill,
        font=font,
        spacing=spacing,
        align="center",
    )


def _draw_gradient_line(
    draw: Any,
    *,
    left: int,
    right: int,
    y: int,
    width: int,
    start_color: Tuple[int, int, int],
    end_color: Tuple[int, int, int],
) -> None:
    """Draw a simple left-to-right gradient accent line."""

    span = max(1, right - left)
    for offset in range(span):
        ratio = offset / span
        red = int(start_color[0] + (end_color[0] - start_color[0]) * ratio)
        green = int(start_color[1] + (end_color[1] - start_color[1]) * ratio)
        blue = int(start_color[2] + (end_color[2] - start_color[2]) * ratio)
        draw.line(
            [(left + offset, y), (left + offset, y)],
            fill=(red, green, blue),
            width=width,
        )


def _draw_glass_panel(
    image: Any,
    bounds: Tuple[int, int, int, int],
    *,
    radius: int,
    fill: Tuple[int, int, int, int],
    outline: str,
    shadow: Tuple[int, int, int, int],
) -> None:
    """Draw a soft glassmorphic rounded panel with a subtle shadow."""

    from PIL import Image, ImageDraw, ImageFilter

    left, top, right, bottom = bounds
    shadow_layer = Image.new("RGBA", image.size, (255, 255, 255, 0))
    shadow_draw = ImageDraw.Draw(shadow_layer)
    shadow_draw.rounded_rectangle(
        [(left + 6, top + 14), (right + 6, bottom + 16)],
        radius=radius,
        fill=shadow,
    )
    shadow_layer = shadow_layer.filter(ImageFilter.GaussianBlur(radius=18))
    image.alpha_composite(shadow_layer)

    panel_layer = Image.new("RGBA", image.size, (255, 255, 255, 0))
    panel_draw = ImageDraw.Draw(panel_layer)
    panel_draw.rounded_rectangle(
        [(left, top), (right, bottom)],
        radius=radius,
        fill=fill,
        outline=outline,
        width=2,
    )
    panel_draw.line(
        [(left + 26, top + 20), (right - 26, top + 20)],
        fill=(255, 255, 255, 168),
        width=2,
    )
    image.alpha_composite(panel_layer)


def build_site_worker_roster(
    *,
    site_name: Optional[str] = None,
    source_paths: Optional[Iterable[Path]] = None,
) -> List[SiteWorker]:
    """Build a live contractor roster from KPI backup JSON files."""

    resolved_source_paths = (
        [Path(path) for path in source_paths]
        if source_paths is not None
        else _discover_kpi_backup_json_paths()
    )
    roster_by_key: Dict[Tuple[str, str], SiteWorker] = {}

    for source_path in resolved_source_paths:
        if not source_path.exists() or not source_path.is_file():
            continue

        payload = _load_kpi_json_payload(source_path)
        payload_site_name = _extract_kpi_payload_site_name(payload)
        if (
            site_name is not None
            and payload_site_name is not None
            and payload_site_name.casefold() != site_name.casefold()
        ):
            continue

        for row in _extract_kpi_active_rows(payload):
            worker = SiteWorker.from_kpi_row(row)
            roster_key = worker.roster_key()
            existing_worker = roster_by_key.get(roster_key)
            if (
                existing_worker is None
                or worker.last_on_site_date > existing_worker.last_on_site_date
            ):
                roster_by_key[roster_key] = worker

    return sorted(
        roster_by_key.values(),
        key=lambda worker: (
            worker.company.casefold(),
            worker.worker_name.casefold(),
            worker.last_on_site_date,
        ),
    )


def get_waste_kpi_sheet_metadata(
    *,
    site_name: Optional[str] = None,
    site_address: str = "",
    fallback_project_number: str = "",
) -> WasteKpiSheetMetadata:
    """Return the best matching File 1 KPI workbook header metadata."""

    candidate_metadata: List[Tuple[int, float, WasteKpiSheetMetadata]] = []
    for workbook_path in sorted(_discover_waste_kpi_workbooks()):
        try:
            metadata = _read_waste_kpi_sheet_metadata(workbook_path)
        except RuntimeError:
            continue
        score = _score_waste_kpi_sheet_metadata(
            metadata,
            site_name=site_name,
            site_address=site_address,
        )
        try:
            modified_time = workbook_path.stat().st_mtime
        except OSError:
            modified_time = 0.0
        candidate_metadata.append((score, modified_time, metadata))

    if not candidate_metadata:
        return WasteKpiSheetMetadata(
            workbook_path=None,
            client_name="",
            site_address=site_address.strip(),
            project_number=fallback_project_number.strip(),
            manager_name="",
        )

    _, _, selected_metadata = max(
        candidate_metadata,
        key=lambda item: (item[0], item[1], item[2].workbook_path.name if item[2].workbook_path else ""),
    )
    return WasteKpiSheetMetadata(
        workbook_path=selected_metadata.workbook_path,
        client_name=selected_metadata.client_name,
        site_address=selected_metadata.site_address or site_address.strip(),
        project_number=selected_metadata.project_number or fallback_project_number.strip(),
        manager_name=selected_metadata.manager_name,
    )


def smart_scan_waste_transfer_note(
    repository: DocumentRepository,
    *,
    source_path: Path,
) -> SmartScannedWasteTransferNote:
    """Extract best-effort WTN fields from an uploaded PDF or image."""

    embedded_pdf_text = ""
    if source_path.suffix.lower() == ".pdf":
        try:
            embedded_pdf_text = "\n".join(_extract_pdf_text_pages(source_path)).strip()
        except RuntimeError:
            embedded_pdf_text = ""
    extracted_text = _extract_waste_transfer_note_text(source_path)
    normalized_text = _normalize_text(extracted_text)
    normalized_embedded_pdf_text = _normalize_text(embedded_pdf_text)

    quantity_tonnes = _extract_tonnage_from_text(normalized_text)
    if quantity_tonnes is None and source_path.suffix.lower() == ".pdf":
        try:
            quantity_tonnes = extract_tonnage_from_ticket(source_path)
        except RuntimeError:
            quantity_tonnes = None

    carrier_name = _infer_carrier_name(
        repository,
        source_path,
        pdf_text=normalized_text,
    )
    if source_path.suffix.lower() == ".pdf":
        vehicle_registration = _extract_vehicle_registration_from_pdf(
            source_path,
            normalized_text,
            normalized_embedded_pdf_text,
        )
        wtn_number = _extract_ticket_number_from_pdf(
            source_path,
            normalized_text,
        )
    else:
        vehicle_registration = _extract_vehicle_registration(normalized_text)
        wtn_number = _derive_waste_transfer_note_number(source_path, normalized_text)

    if ABUCS_PDF_PATTERN.match(source_path.name):
        duplicate_filename_match = re.fullmatch(r"(\d+)-\d+", source_path.stem)
        if (
            duplicate_filename_match is not None
            and wtn_number.strip().casefold() == source_path.stem.casefold()
        ):
            wtn_number = duplicate_filename_match.group(1)

    if ABUCS_PDF_PATTERN.match(source_path.name) and (
        not carrier_name or carrier_name.isdigit()
    ):
        carrier_name = DEFAULT_WASTE_CARRIER_NAME
    collection_type = _extract_waste_collection_type(normalized_text)
    waste_description = _extract_waste_description(normalized_text)
    ewc_code = _extract_ewc_code(normalized_text)
    ticket_date = _extract_waste_ticket_date(normalized_text) or date.today()

    return SmartScannedWasteTransferNote(
        source_name=source_path.name,
        wtn_number=wtn_number,
        carrier_name=carrier_name,
        vehicle_registration=vehicle_registration,
        collection_type=collection_type,
        waste_description=waste_description,
        ticket_date=ticket_date,
        quantity_tonnes=quantity_tonnes,
        ewc_code=ewc_code,
        destination_facility=DEFAULT_DESTINATION_FACILITY,
        extracted_text=extracted_text,
    )


def log_uploaded_waste_transfer_note(
    repository: DocumentRepository,
    *,
    upload_path: Path,
    original_filename: str,
    site_name: str,
    carrier_name: str,
    vehicle_registration: str,
    waste_description: str,
    ticket_date: date,
    quantity_tonnes: float,
    ewc_code: str,
    destination_facility: str = DEFAULT_DESTINATION_FACILITY,
    wtn_number: Optional[str] = None,
) -> LoggedWasteTransferNote:
    """Persist one uploaded WTN file and append it to the live File 1 register."""

    repository.create_schema()
    config.WASTE_DESTINATION.mkdir(parents=True, exist_ok=True)

    safe_filename = Path(original_filename).name or upload_path.name
    destination_path = _build_available_destination(
        Path(safe_filename),
        config.WASTE_DESTINATION,
    )
    stored_file_path = Path(shutil.move(str(upload_path), str(destination_path))).resolve()

    resolved_wtn_number = (wtn_number or stored_file_path.stem).strip() or stored_file_path.stem
    existing_document = _get_waste_transfer_note_document(repository, resolved_wtn_number)
    waste_transfer_note = WasteTransferNoteDocument(
        doc_id=(
            existing_document.doc_id
            if existing_document is not None
            else f"WTN-{_slugify_identifier(resolved_wtn_number)}"
        ),
        site_name=site_name.strip() or DEFAULT_SITE_NAME,
        created_at=(
            existing_document.created_at
            if existing_document is not None
            else datetime.now().replace(second=0, microsecond=0)
        ),
        status=DocumentStatus.ACTIVE,
        wtn_number=resolved_wtn_number,
        date=ticket_date,
        waste_description=waste_description.strip() or DEFAULT_WASTE_DESCRIPTION,
        ewc_code=ewc_code.strip() or DEFAULT_EWC_CODE,
        quantity_tonnes=quantity_tonnes,
        carrier_name=carrier_name.strip() or DEFAULT_WASTE_CARRIER_NAME,
        destination_facility=destination_facility.strip() or DEFAULT_DESTINATION_FACILITY,
        vehicle_registration=vehicle_registration.strip(),
    )
    repository.save(waste_transfer_note)
    repository.index_file(
        file_name=stored_file_path.name,
        file_path=stored_file_path,
        file_category="uploaded_waste_transfer_note",
        file_group=FileGroup.FILE_1,
        site_name=waste_transfer_note.site_name,
        related_doc_id=waste_transfer_note.doc_id,
    )

    register_document = _upsert_site_waste_register(
        repository,
        site_name=waste_transfer_note.site_name,
    )
    return LoggedWasteTransferNote(
        waste_transfer_note=waste_transfer_note,
        stored_file_path=stored_file_path,
        register_document=register_document,
    )


def update_logged_waste_transfer_note(
    repository: DocumentRepository,
    *,
    source_document: WasteTransferNoteDocument,
    site_name: str,
    carrier_name: str,
    vehicle_registration: str,
    waste_description: str,
    ticket_date: date,
    quantity_tonnes: float,
    ewc_code: str,
    destination_facility: str = DEFAULT_DESTINATION_FACILITY,
    tonnage_review_status: str = "",
) -> LoggedWasteTransferNote:
    """Update an already-filed WTN document from the File 1 smart-scan form."""

    repository.create_schema()
    indexed_file_path = _get_waste_transfer_note_source_path(repository, source_document)
    refreshed_waste_transfer_note = WasteTransferNoteDocument(
        doc_id=source_document.doc_id,
        site_name=site_name.strip() or source_document.site_name,
        created_at=source_document.created_at,
        status=DocumentStatus.ACTIVE,
        wtn_number=source_document.wtn_number,
        date=ticket_date,
        waste_description=waste_description.strip() or source_document.waste_description,
        ewc_code=ewc_code.strip() or source_document.ewc_code,
        quantity_tonnes=quantity_tonnes,
        carrier_name=carrier_name.strip() or source_document.carrier_name,
        destination_facility=destination_facility.strip() or source_document.destination_facility,
        vehicle_registration=vehicle_registration.strip(),
        source_file_override_path=source_document.source_file_override_path,
        canonical_source_path=source_document.canonical_source_path,
        source_conflict_candidates=source_document.source_conflict_candidates,
        tonnage_review_status=(
            tonnage_review_status.strip()
            if quantity_tonnes <= 0
            else ""
        ),
    )
    repository.save(refreshed_waste_transfer_note)
    register_document = _upsert_site_waste_register(
        repository,
        site_name=refreshed_waste_transfer_note.site_name,
    )
    return LoggedWasteTransferNote(
        waste_transfer_note=refreshed_waste_transfer_note,
        stored_file_path=indexed_file_path or Path(),
        register_document=register_document,
    )


def generate_waste_register_document(
    repository: DocumentRepository,
    *,
    site_name: str,
    client_name: str,
    site_address: str,
    manager_name: str,
) -> GeneratedWasteRegisterDocument:
    """Render the approved File 1 waste register template for one site."""

    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the printable File 1 waste register."
        ) from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError(
            "jinja2 is required to generate the printable File 1 waste register."
        ) from exc

    repository.create_schema()
    config.FILE_1_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    template_path = TemplateRegistry.resolve_template_path("waste_register")
    waste_transfer_notes = _list_site_waste_transfer_notes(
        repository,
        site_name=site_name,
    )
    context = {
        "client_name": client_name.strip(),
        "site_address": site_address.strip(),
        "manager_name": manager_name.strip(),
        "waste_entries": [
            {
                "carrier": waste_transfer_note.carrier_name,
                "date": waste_transfer_note.date.strftime("%d/%m/%y"),
                "description": _format_waste_register_description(waste_transfer_note),
                "reg_no": _format_waste_register_reference(waste_transfer_note),
            }
            for waste_transfer_note in waste_transfer_notes
        ],
    }

    output_name = Path(
        "UHSF50.0 Register of Waste Removal - "
        f"{_sanitize_filename_fragment(site_name)} - "
        f"{date.today():%Y-%m-%d}.docx"
    )
    output_path = _build_available_destination(output_name, config.FILE_1_OUTPUT_DIR)

    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "waste-register-template.docx"
        _build_patched_docxtpl_template(template_path, repaired_template_path)

        discovered_placeholders = _discover_docx_template_tags(repaired_template_path)
        missing_placeholders = sorted(
            {"client_name", "site_address", "manager_name"} - discovered_placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "Waste register template is missing required placeholders: "
                + ", ".join(missing_placeholders)
            )

        document_template = DocxTemplate(str(repaired_template_path))
        clean_jinja_environment = Environment(autoescape=False)
        document_template.render(
            context,
            jinja_env=clean_jinja_environment,
            autoescape=False,
        )
        document_template.save(output_path)
        _normalise_plant_register_table_rows(output_path)

    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="waste_register_docx",
        file_group=FileGroup.FILE_1,
        site_name=site_name,
    )
    return GeneratedWasteRegisterDocument(
        output_path=output_path,
        row_count=len(waste_transfer_notes),
    )


def file_and_index_all(repository: DocumentRepository) -> List[FiledAsset]:
    """Move supported inbox files into the workspace and index their new paths."""

    repository.create_schema()

    inbox = config.INBOX
    waste_destination = config.WASTE_DESTINATION
    carrier_docs_destination = config.CARRIER_DOCS_DESTINATION
    waste_reports_destination = config.WASTE_REPORTS_DESTINATION
    attendance_destination = config.ATTENDANCE_DESTINATION
    plant_hire_destination = config.PLANT_HIRE_REGISTER_DIR
    induction_directory = config.INDUCTION_DIR
    rams_destination = config.RAMS_DESTINATION
    coshh_destination = config.COSHH_DESTINATION
    file_3_review_directory = config.FILE_3_REVIEW_DIR
    file_3_output_directory = config.FILE_3_OUTPUT_DIR

    inbox.mkdir(parents=True, exist_ok=True)
    waste_destination.mkdir(parents=True, exist_ok=True)
    carrier_docs_destination.mkdir(parents=True, exist_ok=True)
    waste_reports_destination.mkdir(parents=True, exist_ok=True)
    attendance_destination.mkdir(parents=True, exist_ok=True)
    plant_hire_destination.mkdir(parents=True, exist_ok=True)
    induction_directory.mkdir(parents=True, exist_ok=True)
    rams_destination.mkdir(parents=True, exist_ok=True)
    coshh_destination.mkdir(parents=True, exist_ok=True)
    file_3_review_directory.mkdir(parents=True, exist_ok=True)
    file_3_output_directory.mkdir(parents=True, exist_ok=True)

    attendance_engine = IngestionEngine(repository)
    filed_assets: List[FiledAsset] = []

    for source_path in sorted(inbox.iterdir(), key=lambda path: path.name.lower()):
        if not source_path.is_file():
            continue

        if _is_carrier_compliance_pdf(source_path):
            destination_path = _move_file(source_path, carrier_docs_destination)
            extracted_expiry_date = None
            compliance_document = None
            try:
                carrier_pdf_text = " ".join(_extract_pdf_text_pages(destination_path))
            except RuntimeError:
                carrier_pdf_text = ""
            carrier_name = _infer_carrier_name(
                repository,
                destination_path,
                pdf_text=carrier_pdf_text,
            )
            carrier_document_type = _infer_carrier_document_type(destination_path)
            try:
                extracted_expiry_date = extract_expiry_date_from_pdf(destination_path)
            except RuntimeError:
                extracted_expiry_date = None

            if extracted_expiry_date is not None:
                compliance_document = _upsert_carrier_compliance_document(
                    repository,
                    carrier_name=carrier_name,
                    carrier_document_type=carrier_document_type,
                    expiry_date=extracted_expiry_date,
                    source_path=destination_path,
                )
            repository.index_file(
                file_name=destination_path.name,
                file_path=destination_path,
                file_category="carrier_doc_pdf",
                file_group=FileGroup.FILE_1,
                site_name=(
                    compliance_document.site_name
                    if compliance_document is not None
                    else None
                ),
                related_doc_id=(
                    compliance_document.doc_id
                    if compliance_document is not None
                    else None
                ),
            )
            filed_assets.append(
                FiledAsset(
                    original_path=source_path,
                    destination_path=destination_path,
                    file_category="carrier_doc_pdf",
                    related_doc_id=(
                        compliance_document.doc_id
                        if compliance_document is not None
                        else None
                    ),
                    auto_captured_expiry_date=extracted_expiry_date,
                    auto_captured_carrier_name=(
                        carrier_name if extracted_expiry_date is not None else None
                    ),
                    auto_captured_document_type=(
                        carrier_document_type
                        if extracted_expiry_date is not None
                        else None
                    ),
                )
            )
            continue

        if ABUCS_PDF_PATTERN.match(source_path.name):
            source_text = _safe_extract_pdf_text(source_path)
            if _is_plant_hire_pdf(source_path, source_text):
                destination_path = _move_file(source_path, plant_hire_destination)
                plant_assets = _upsert_plant_assets_from_pdf(
                    repository,
                    destination_path,
                    pdf_text=source_text,
                )
                file_category = _classify_plant_hire_pdf(
                    destination_path,
                    source_text,
                )
                repository.index_file(
                    file_name=destination_path.name,
                    file_path=destination_path,
                    file_category=file_category,
                    file_group=FileGroup.FILE_2,
                    site_name=(
                        plant_assets[0].site_name
                        if plant_assets
                        else _load_workspace_project_setup().get("current_site_name")
                    ),
                    related_doc_id=plant_assets[0].doc_id if len(plant_assets) == 1 else None,
                )
                filed_assets.append(
                    FiledAsset(
                        original_path=source_path,
                        destination_path=destination_path,
                        file_category=file_category,
                        related_doc_id=(
                            plant_assets[0].doc_id
                            if len(plant_assets) == 1
                            else None
                        ),
                    )
                )
                continue

            destination_path = _move_file(source_path, waste_destination)
            waste_transfer_note = _upsert_waste_transfer_note_document(
                repository,
                destination_path,
            )
            repository.index_file(
                file_name=destination_path.name,
                file_path=destination_path,
                file_category=_classify_waste_ticket_file(destination_path),
                file_group=FileGroup.FILE_1,
                site_name=(
                    waste_transfer_note.site_name
                    if waste_transfer_note is not None
                    else None
                ),
                related_doc_id=(
                    waste_transfer_note.doc_id
                    if waste_transfer_note is not None
                    else None
                ),
            )
            filed_assets.append(
                FiledAsset(
                    original_path=source_path,
                    destination_path=destination_path,
                    file_category=_classify_waste_ticket_file(destination_path),
                    related_doc_id=(
                        waste_transfer_note.doc_id
                        if waste_transfer_note is not None
                        else None
                    ),
                )
            )
            continue

        if source_path.suffix.lower() in FILE_3_SAFETY_SOURCE_SUFFIXES:
            source_text = _safe_extract_safety_source_text(source_path)
            if source_path.suffix.lower() == ".pdf" and _looks_like_waste_ticket_source(
                source_path,
                source_text,
            ):
                destination_path = _move_file(source_path, waste_destination)
                waste_transfer_note = _upsert_waste_transfer_note_document(
                    repository,
                    destination_path,
                )
                repository.index_file(
                    file_name=destination_path.name,
                    file_path=destination_path,
                    file_category="waste_ticket_pdf",
                    file_group=FileGroup.FILE_1,
                    site_name=(
                        waste_transfer_note.site_name
                        if waste_transfer_note is not None
                        else None
                    ),
                    related_doc_id=(
                        waste_transfer_note.doc_id
                        if waste_transfer_note is not None
                        else None
                    ),
                )
                filed_assets.append(
                    FiledAsset(
                        original_path=source_path,
                        destination_path=destination_path,
                        file_category="waste_ticket_pdf",
                        related_doc_id=(
                            waste_transfer_note.doc_id
                            if waste_transfer_note is not None
                            else None
                        ),
                    )
                )
                continue

            if source_path.suffix.lower() in {".doc", ".docx"} and _looks_like_waste_support_document(
                source_text
            ):
                destination_path = _move_file(source_path, waste_reports_destination)
                repository.index_file(
                    file_name=destination_path.name,
                    file_path=destination_path,
                    file_category="waste_report_word",
                    file_group=FileGroup.FILE_1,
                )
                filed_assets.append(
                    FiledAsset(
                        original_path=source_path,
                        destination_path=destination_path,
                        file_category="waste_report_word",
                    )
                )
                continue

            if _is_coshh_safety_source(source_path, source_text):
                destination_path = _move_file(source_path, coshh_destination)
                coshh_document = _upsert_coshh_document_from_source(
                    repository,
                    destination_path,
                    source_text=source_text,
                )
                repository.index_file(
                    file_name=destination_path.name,
                    file_path=destination_path,
                    file_category=_resolve_safety_file_category(
                        "coshh_pdf",
                        destination_path,
                    ),
                    file_group=FileGroup.FILE_3,
                    site_name=(
                        coshh_document.site_name
                        if coshh_document is not None
                        else _load_workspace_project_setup().get("current_site_name")
                    ),
                    related_doc_id=(
                        coshh_document.doc_id
                        if coshh_document is not None
                        else None
                    ),
                )
                filed_assets.append(
                    FiledAsset(
                        original_path=source_path,
                        destination_path=destination_path,
                        file_category=_resolve_safety_file_category(
                            "coshh_pdf",
                            destination_path,
                        ),
                        related_doc_id=(
                            coshh_document.doc_id
                            if coshh_document is not None
                            else None
                        ),
                    )
                )
                continue

            if _is_rams_safety_source(source_path, source_text):
                destination_path = _move_file(source_path, rams_destination)
                rams_document = _upsert_rams_document_from_source(
                    repository,
                    destination_path,
                    source_text=source_text,
                )
                repository.index_file(
                    file_name=destination_path.name,
                    file_path=destination_path,
                    file_category=_resolve_safety_file_category(
                        "rams_pdf",
                        destination_path,
                    ),
                    file_group=FileGroup.FILE_3,
                    site_name=(
                        rams_document.site_name
                        if rams_document is not None
                        else _load_workspace_project_setup().get("current_site_name")
                    ),
                    related_doc_id=(
                        rams_document.doc_id
                        if rams_document is not None
                        else None
                    ),
                )
                filed_assets.append(
                    FiledAsset(
                        original_path=source_path,
                        destination_path=destination_path,
                        file_category=_resolve_safety_file_category(
                            "rams_pdf",
                            destination_path,
                        ),
                        related_doc_id=(
                            rams_document.doc_id
                            if rams_document is not None
                            else None
                        ),
                    )
                )
                continue

            if source_path.suffix.lower() == ".pdf" and _is_plant_hire_pdf(source_path, source_text):
                destination_path = _move_file(source_path, plant_hire_destination)
                plant_assets = _upsert_plant_assets_from_pdf(
                    repository,
                    destination_path,
                    pdf_text=source_text,
                )
                file_category = _classify_plant_hire_pdf(
                    destination_path,
                    source_text,
                )
                repository.index_file(
                    file_name=destination_path.name,
                    file_path=destination_path,
                    file_category=file_category,
                    file_group=FileGroup.FILE_2,
                    site_name=(
                        plant_assets[0].site_name
                        if plant_assets
                        else _load_workspace_project_setup().get("current_site_name")
                    ),
                    related_doc_id=plant_assets[0].doc_id if len(plant_assets) == 1 else None,
                )
                filed_assets.append(
                    FiledAsset(
                        original_path=source_path,
                        destination_path=destination_path,
                        file_category=file_category,
                        related_doc_id=plant_assets[0].doc_id if len(plant_assets) == 1 else None,
                    )
                )
                continue

            if _looks_like_file_3_review_hold_source(source_path, source_text):
                destination_path = _move_file(source_path, file_3_review_directory)
                repository.index_file(
                    file_name=destination_path.name,
                    file_path=destination_path,
                    file_category=_resolve_safety_file_category(
                        "safety_review_pdf",
                        destination_path,
                    ),
                    file_group=FileGroup.FILE_3,
                    site_name=_load_workspace_project_setup().get("current_site_name"),
                )
                filed_assets.append(
                    FiledAsset(
                        original_path=source_path,
                        destination_path=destination_path,
                        file_category=_resolve_safety_file_category(
                            "safety_review_pdf",
                            destination_path,
                        ),
                    )
                )
                continue

        if source_path.suffix.lower() in {".xls", ".xlsx"}:
            destination_path = _move_file(source_path, waste_reports_destination)
            repository.index_file(
                file_name=destination_path.name,
                file_path=destination_path,
                file_category="waste_report_excel",
                file_group=FileGroup.FILE_1,
            )
            filed_assets.append(
                FiledAsset(
                    original_path=source_path,
                    destination_path=destination_path,
                    file_category="waste_report_excel",
                )
            )
            continue

        if source_path.suffix.lower() == ".json":
            destination_path = _move_file(source_path, attendance_destination)
            register = attendance_engine.ingest_site_attendance_json(destination_path)
            repository.index_file(
                file_name=destination_path.name,
                file_path=destination_path,
                file_category="kpi_json",
                file_group=FileGroup.FILE_2,
                site_name=register.site_name,
                related_doc_id=register.doc_id,
            )
            filed_assets.append(
                FiledAsset(
                    original_path=source_path,
                    destination_path=destination_path,
                    file_category="kpi_json",
                    related_doc_id=register.doc_id,
                )
            )

    _sync_existing_waste_transfer_notes(repository, waste_destination)
    _sync_existing_plant_hire_pdfs(repository, plant_hire_destination)
    _sync_existing_safety_sources(
        repository,
        destination_directory=rams_destination,
        detector=_is_rams_safety_source,
        upsert_document=_upsert_rams_document_from_source,
        file_category="rams_pdf",
    )
    _sync_existing_safety_sources(
        repository,
        destination_directory=coshh_destination,
        detector=_is_coshh_safety_source,
        upsert_document=_upsert_coshh_document_from_source,
        file_category="coshh_pdf",
    )
    return filed_assets


def _load_workspace_project_setup() -> Dict[str, str]:
    """Load the persisted project setup used by sync-time document intake."""

    project_setup_path = config.BASE_DATA_DIR / "project_setup.json"
    try:
        payload = json.loads(project_setup_path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        return {}
    except (OSError, json.JSONDecodeError, TypeError, ValueError):
        return {}

    if not isinstance(payload, dict):
        return {}

    return {
        "current_site_name": str(payload.get("current_site_name") or "").strip(),
        "job_number": str(payload.get("job_number") or "").strip(),
        "site_address": str(payload.get("site_address") or "").strip(),
        "client_name": str(payload.get("client_name") or "").strip(),
    }


def _default_hired_by_for_project(project_setup: Mapping[str, str]) -> str:
    """Return the default hired-by value for the current project."""

    searchable_text = " ".join(
        project_setup.get(key, "")
        for key in ("current_site_name", "job_number", "site_address", "client_name")
    ).casefold()
    if "archer" in searchable_text:
        return "A. Archer Electrical"
    if "tde" in searchable_text:
        return "TDE"
    if any(alias in searchable_text for alias in ("uplands", "url", "uplands retail")):
        return "URL (Uplands)"
    return "URL (Uplands)"


def _is_coshh_safety_source(source_path: Path, source_text: str = "") -> bool:
    """Return True when one safety source belongs in the File 3 COSHH inventory."""

    lowered_name = source_path.stem.casefold()
    if any(marker in lowered_name for marker in FILE_3_EXCLUDED_RAMS_FILENAME_MARKERS):
        return False

    lowered_text = source_text.casefold()
    if "construction phase plan" in lowered_text or "uhsf15.1" in lowered_name:
        return False
    if any(_file_3_contains_keyword(lowered_name, keyword) for keyword in COSHH_KEYWORDS):
        return True
    if any(_file_3_contains_keyword(lowered_name, keyword) for keyword in RAMS_KEYWORDS):
        return False

    fallback_text = _build_file_3_fallback_text(source_text).casefold()
    strong_signal_count = sum(
        1
        for label in FILE_3_STRONG_COSHH_TEXT_LABELS
        if label in fallback_text or label in lowered_text
    )
    if strong_signal_count < 3 and not any(
        _file_3_contains_keyword(lowered_text, keyword) for keyword in COSHH_KEYWORDS
    ):
        return False
    if _is_rams_safety_source(source_path, source_text):
        return False
    return True


def _is_rams_safety_source(source_path: Path, source_text: str = "") -> bool:
    """Return True when one safety source belongs in the File 3 RAMS inventory."""

    lowered_text = f"{source_path.name} {source_text}".casefold()
    if "register" in lowered_text:
        return False
    lowered_name = source_path.stem.casefold()
    if any(marker in lowered_name for marker in FILE_3_EXCLUDED_RAMS_FILENAME_MARKERS):
        return False
    if "construction phase plan" in lowered_text or "uhsf15.1" in lowered_name:
        return False

    fallback_text = _build_file_3_fallback_text(source_text).casefold()
    strong_coshh_signal_count = sum(
        1
        for label in FILE_3_STRONG_COSHH_TEXT_LABELS
        if label in fallback_text or label in lowered_text
    )
    if any(_file_3_contains_keyword(lowered_name, keyword) for keyword in RAMS_KEYWORDS):
        return True

    explicit_rams_signal = any(
        _file_3_contains_keyword(lowered_text, keyword)
        or _file_3_contains_keyword(fallback_text, keyword)
        for keyword in RAMS_KEYWORDS
    )
    rams_layout_signal = any(
        label in fallback_text
        for label in (
            "brief description of work",
            "job/activity/sequence",
            "hazards and control procedures",
            "combined risk assessment",
        )
    )
    if not explicit_rams_signal and not rams_layout_signal:
        return False
    if strong_coshh_signal_count >= 3 and not any(keyword in lowered_name for keyword in RAMS_KEYWORDS):
        return False
    return True


def _is_coshh_pdf(source_path: Path, pdf_text: str = "") -> bool:
    """Compatibility wrapper for the original PDF-only COSHH detector."""

    return _is_coshh_safety_source(source_path, pdf_text)


def _is_rams_pdf(source_path: Path, pdf_text: str = "") -> bool:
    """Compatibility wrapper for the original PDF-only RAMS detector."""

    return _is_rams_safety_source(source_path, pdf_text)


def _build_safety_doc_id(prefix: str, site_name: str, source_stem: str) -> str:
    """Return a stable File 3 document id for one synced safety PDF."""

    return (
        f"{prefix}-{_slugify_identifier(site_name)}-"
        f"{_slugify_identifier(source_stem)}"
    )


def _extract_text_after_labels(text: str, labels: Iterable[str]) -> str:
    """Return the first value found after one of the supplied line labels."""

    normalized_lines = [
        _normalize_text(line)
        for line in text.splitlines()
        if _normalize_text(line)
    ]
    labels_list = sorted(
        [label for label in labels if label],
        key=len,
        reverse=True,
    )

    for index, line in enumerate(normalized_lines):
        lowered_line = line.casefold()
        for label in labels_list:
            lowered_label = label.casefold()
            if lowered_label not in lowered_line:
                continue

            same_line_match = re.search(
                rf"{re.escape(label)}\s*(?::|-)?\s*(?P<value>.+)$",
                line,
                re.IGNORECASE,
            )
            if same_line_match is not None:
                candidate_value = _clean_safety_value(same_line_match.group("value"))
                if (
                    candidate_value
                    and candidate_value.casefold() != lowered_label
                    and re.fullmatch(
                        r"(?:\(?s\)?|&\s*address:?|address:?|name\s*&?)",
                        candidate_value.casefold(),
                    )
                    is None
                ):
                    return candidate_value

            if index + 1 < len(normalized_lines):
                next_index = index + 1
                next_line = _clean_safety_value(normalized_lines[next_index])
                while next_index + 1 < len(normalized_lines) and (
                    next_line.startswith("&")
                    or (len(next_line) <= 20 and next_line.endswith(":"))
                ):
                    next_index += 1
                    next_line = _clean_safety_value(normalized_lines[next_index])
                if next_line and not any(
                    other_label.casefold() in next_line.casefold()
                    for other_label in labels_list
                ):
                    return next_line

    normalized_text = _normalize_text(text)
    for label in labels_list:
        fallback_match = re.search(
            rf"{re.escape(label)}\s*(?::|-)?\s*(?P<value>[^:]+?)(?=\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*:|$)",
            normalized_text,
            re.IGNORECASE,
        )
        if fallback_match is None:
            continue
        candidate_value = _clean_safety_value(fallback_match.group("value"))
        if candidate_value:
            return candidate_value
    return ""


def _file_3_contains_keyword(text: str, keyword: str) -> bool:
    """Return True when one keyword appears as a real phrase inside File 3 text."""

    lowered_text = text.casefold()
    lowered_keyword = keyword.casefold()
    if " " in lowered_keyword:
        return lowered_keyword in lowered_text
    return re.search(rf"\b{re.escape(lowered_keyword)}\b", lowered_text) is not None


def _clean_safety_value(value: str) -> str:
    """Normalize one extracted File 3 field value."""

    cleaned_value = re.sub(r"\s+\|\s+", " ", value)
    cleaned_value = re.sub(r"\|+", " ", cleaned_value)
    cleaned_value = re.sub(r"\s+", " ", cleaned_value).strip(" :;|-")
    cleaned_value = _collapse_repeated_safety_phrase(cleaned_value)
    return cleaned_value


def _collapse_repeated_safety_phrase(value: str) -> str:
    """Collapse obvious OCR-style repeated phrases into a single clean phrase."""

    words = value.split()
    if len(words) < 6:
        return value

    max_chunk_size = min(12, len(words) // 2)
    for chunk_size in range(2, max_chunk_size + 1):
        chunk_words = words[:chunk_size]
        repeated = chunk_words * max(2, len(words) // chunk_size)
        if repeated[: len(words)] == words and len(words) >= chunk_size * 2:
            return " ".join(chunk_words)

    if len(words) >= 9:
        first_half = words[: len(words) // 2]
        second_half = words[len(words) // 2 :]
        if first_half[: len(second_half)] == second_half[: len(first_half)]:
            return " ".join(first_half)
    return value


def _split_file_3_filename_segments(source_path: Path) -> List[str]:
    """Return cleaned filename segments split on hyphens and underscores."""

    return [
        cleaned_segment
        for cleaned_segment in (
            _clean_safety_value(segment)
            for segment in FILE_3_FILENAME_SPLIT_PATTERN.split(source_path.stem)
        )
        if cleaned_segment
    ]


def _file_3_is_version_segment(segment: str) -> bool:
    """Return True when one filename segment primarily encodes a version."""

    if FILE_3_VERSION_SEGMENT_PATTERN.search(segment) is not None:
        return True
    return segment.casefold() in {"rev", "revision", "ver", "version", "v"}


def _file_3_find_version_from_segments(segments: List[str]) -> str:
    """Return one version value from a structured File 3 filename."""

    for index, segment in enumerate(segments):
        match = FILE_3_VERSION_SEGMENT_PATTERN.search(segment)
        if match is not None:
            return _clean_safety_value(match.group("value"))
        if segment.casefold() in {"rev", "revision", "ver", "version", "v"} and index + 1 < len(segments):
            next_segment = _clean_safety_value(segments[index + 1])
            if re.fullmatch(r"\d+(?:\.\d+)*", next_segment):
                return next_segment
    return ""


def _file_3_find_reference_from_segments(segments: List[str]) -> str:
    """Return one reference token from a structured File 3 filename."""

    for index, segment in enumerate(segments):
        match = FILE_3_REFERENCE_SEGMENT_PATTERN.search(segment)
        if match is not None:
            return _clean_safety_value(match.group("value"))
        if segment.casefold() in {"ref", "reference", "doc", "document"} and index + 1 < len(segments):
            next_segment = _clean_safety_value(segments[index + 1])
            if next_segment:
                return next_segment
    return ""


def _file_3_is_safety_label_segment(segment: str) -> bool:
    """Return True when one filename segment is only a safety-document label."""

    lowered_segment = segment.casefold()
    return lowered_segment in FILE_3_SAFETY_FILENAME_STOPWORDS or any(
        lowered_segment == keyword
        for keyword in (*COSHH_KEYWORDS, *RAMS_KEYWORDS)
    )


def _strip_file_3_safety_markers(value: str) -> str:
    """Remove File 3 safety-label text and version markers from one string."""

    stripped_value = value
    stripped_value = FILE_3_VERSION_SEGMENT_PATTERN.sub(" ", stripped_value)
    keywords_to_strip = sorted(
        {
            *FILE_3_SAFETY_FILENAME_STOPWORDS,
            *COSHH_KEYWORDS,
            *RAMS_KEYWORDS,
        },
        key=len,
        reverse=True,
    )
    for keyword in keywords_to_strip:
        stripped_value = re.sub(
            rf"\b{re.escape(keyword)}\b",
            " ",
            stripped_value,
            flags=re.IGNORECASE,
        )
    return _clean_safety_value(stripped_value)


def _file_3_safety_label_index(segments: List[str]) -> Optional[int]:
    """Return the index of the first RAMS/COSHH label-like filename segment."""

    for index, segment in enumerate(segments):
        lowered_segment = segment.casefold()
        if _file_3_is_safety_label_segment(segment):
            return index
        if any(keyword in lowered_segment for keyword in COSHH_KEYWORDS + RAMS_KEYWORDS):
            return index
    return None


def _file_3_candidate_company_names(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> List[str]:
    """Return known contractor/company names that can be matched in File 3 filenames."""

    candidate_names = {
        worker.company.strip()
        for worker in build_site_worker_roster(site_name=site_name)
        if worker.company.strip()
    }
    candidate_names.update(
        document.contractor_name.strip()
        for document in repository.list_documents(document_type=RAMSDocument.document_type)
        if isinstance(document, RAMSDocument)
        and document.site_name.casefold() == site_name.casefold()
        and document.contractor_name.strip()
    )
    candidate_names.update(
        document.contractor_name.strip()
        for document in repository.list_documents(document_type=COSHHDocument.document_type)
        if isinstance(document, COSHHDocument)
        and document.site_name.casefold() == site_name.casefold()
        and document.contractor_name.strip()
    )
    candidate_names.update(FILE_3_PREFERRED_COMPANY_NAMES)
    return sorted(candidate_names, key=len, reverse=True)


def _file_3_is_blacklisted_contractor_value(value: str) -> bool:
    """Return True when one filename segment is only site/client noise."""

    normalized_value = _clean_safety_value(value).casefold()
    if not normalized_value:
        return True
    if normalized_value in FILE_3_CONTRACTOR_BLACKLIST:
        return True

    value_words = {
        word
        for word in re.findall(r"[a-z0-9]+", normalized_value)
        if word
    }
    return bool(value_words) and value_words.issubset(FILE_3_CONTRACTOR_BLACKLIST_WORDS)


def _file_3_looks_like_company_segment(segment: str) -> bool:
    """Return True when one filename segment looks like a contractor/company name."""

    if _file_3_is_blacklisted_contractor_value(segment):
        return False
    lowered_segment = segment.casefold()
    words = re.findall(r"[a-z0-9]+", lowered_segment)
    if not words:
        return False
    if any(word in FILE_3_COMPANY_HINT_WORDS for word in words):
        return True
    compact_segment = re.sub(r"[^A-Z0-9]+", "", segment.upper())
    return compact_segment.isalpha() and 2 <= len(compact_segment) <= 6


def _file_3_normalize_company_match_text(value: str) -> str:
    """Return a punctuation-light string for tolerant company-name matching."""

    return re.sub(r"[^a-z0-9]+", " ", value.casefold()).strip()


def _file_3_matches_company_name(candidate_name: str, haystack: str) -> bool:
    """Return True when one company name appears in a tolerant text view."""

    normalized_candidate = _file_3_normalize_company_match_text(candidate_name)
    normalized_haystack = _file_3_normalize_company_match_text(haystack)
    if not normalized_candidate or not normalized_haystack:
        return False
    if f" {normalized_candidate} " in f" {normalized_haystack} ":
        return True

    trimmed_candidate = re.sub(
        r"\b(?:ltd|limited|llp|plc|group|uk)\b",
        " ",
        normalized_candidate,
    )
    trimmed_candidate = re.sub(r"\s+", " ", trimmed_candidate).strip()
    return bool(trimmed_candidate) and (
        f" {trimmed_candidate} " in f" {normalized_haystack} "
    )


def _file_3_looks_like_person_name(value: str) -> bool:
    """Return True when one extracted value looks more like a person than a company."""

    cleaned_value = _clean_safety_value(value)
    if not cleaned_value or any(character.isdigit() for character in cleaned_value):
        return False
    words = cleaned_value.split()
    if len(words) not in {2, 3}:
        return False
    if words[0].casefold() not in FILE_3_COMMON_PERSON_FIRST_NAMES:
        return False
    return all(word[:1].isupper() and word[1:].islower() for word in words if word)


def _file_3_is_suspicious_company_value(value: str) -> bool:
    """Return True when one extracted company value looks like OCR or label noise."""

    cleaned_value = _clean_safety_value(value)
    if not cleaned_value:
        return True
    if _file_3_is_blacklisted_contractor_value(cleaned_value):
        return True
    if _file_3_looks_like_person_name(cleaned_value):
        return True

    lowered_value = cleaned_value.casefold()
    if lowered_value in FILE_3_SUSPICIOUS_COMPANY_LABELS:
        return True
    if any(token in cleaned_value for token in {"@", "http://", "https://"}):
        return True
    if cleaned_value.startswith(("/", "|", ".")):
        return True
    if re.match(r"^\d+(?:\.\d+)?\s", cleaned_value) is not None:
        return True
    if "main contractor" in lowered_value or "signature" in lowered_value:
        return True

    value_words = re.findall(r"[a-z0-9]+", lowered_value)
    if not value_words:
        return True
    if len(value_words) == 1:
        single_word = value_words[0]
        if (
            len(single_word) <= 2
            or single_word.isdigit()
            or re.fullmatch(r"(?:ec|ms|w)\d+[a-z]*", single_word) is not None
        ):
            return True
    return False


def _file_3_is_suspicious_activity_value(
    value: str,
    *,
    company_name: str = "",
    reference: str = "",
) -> bool:
    """Return True when one extracted activity/title looks like noise."""

    cleaned_value = _clean_safety_value(value)
    if not cleaned_value:
        return True
    if _file_3_is_blacklisted_contractor_value(cleaned_value):
        return True
    if _file_3_looks_like_person_name(cleaned_value):
        return True

    lowered_value = cleaned_value.casefold()
    if any(token in cleaned_value for token in {"@", "http://", "https://"}):
        return True
    if lowered_value in {"review", "review form", "rams", "method statement"}:
        return True
    if "signature" in lowered_value or lowered_value.startswith("/sequence"):
        return True
    if company_name and lowered_value == company_name.casefold():
        return True
    if reference and lowered_value == reference.casefold():
        return True

    value_words = re.findall(r"[a-z0-9]+", lowered_value)
    if not value_words:
        return True
    if len(value_words) == 1:
        single_word = value_words[0]
        if (
            single_word.isdigit()
            or re.fullmatch(r"(?:ec|ms|w)\d+[a-z]*", single_word) is not None
        ):
            return True
    return False


def _extract_file_3_company_from_text(source_text: str) -> str:
    """Return a best-effort company line from the raw File 3 text body."""

    preferred_lines: List[str] = []
    for raw_line in source_text[:1200].splitlines():
        cleaned_line = _clean_safety_value(_normalize_text(raw_line))
        if not cleaned_line or len(cleaned_line) > 80:
            continue
        if _file_3_is_suspicious_company_value(cleaned_line):
            continue
        if _file_3_looks_like_company_segment(cleaned_line):
            preferred_lines.append(cleaned_line)

    if preferred_lines:
        return preferred_lines[0]
    return ""


def _extract_file_3_company_from_filename(
    source_path: Path,
    *,
    candidate_names: Iterable[str],
) -> str:
    """Return the contractor/company inferred from the File 3 filename."""

    lowered_stem = source_path.stem.casefold()
    matched_names = [
        candidate_name
        for candidate_name in candidate_names
        if (
            candidate_name
            and not _file_3_is_blacklisted_contractor_value(candidate_name)
            and _file_3_matches_company_name(candidate_name, lowered_stem)
        )
    ]
    if matched_names:
        return max(matched_names, key=len)

    segments = _split_file_3_filename_segments(source_path)
    for segment in segments:
        stripped_segment = _strip_file_3_safety_markers(segment)
        if (
            not stripped_segment
            or _file_3_is_blacklisted_contractor_value(stripped_segment)
            or _file_3_is_suspicious_company_value(stripped_segment)
        ):
            continue
        stripped_matches = [
            candidate_name
            for candidate_name in candidate_names
            if candidate_name and _file_3_matches_company_name(candidate_name, stripped_segment)
        ]
        if stripped_matches:
            return max(stripped_matches, key=len)
        if _file_3_looks_like_company_segment(stripped_segment):
            return stripped_segment

    safety_label_index = _file_3_safety_label_index(segments)
    if safety_label_index is not None and safety_label_index > 0:
        company_segments = [
            segment
            for segment in segments[:safety_label_index]
            if not _file_3_is_version_segment(segment)
            and not _file_3_is_blacklisted_contractor_value(segment)
            and not _file_3_is_suspicious_company_value(segment)
        ]
        if company_segments:
            return _clean_safety_value(" ".join(company_segments))

        for segment in segments[:safety_label_index]:
            if _file_3_is_version_segment(segment):
                continue
            if (
                _file_3_is_blacklisted_contractor_value(segment)
                or _file_3_is_suspicious_company_value(segment)
            ):
                continue
            return segment
        return ""

    for segment in segments:
        if _file_3_is_version_segment(segment) or _file_3_is_safety_label_segment(segment):
            continue
        if (
            _file_3_is_blacklisted_contractor_value(segment)
            or _file_3_is_suspicious_company_value(segment)
        ):
            continue
        if _file_3_looks_like_company_segment(segment):
            return segment
    if safety_label_index is None:
        for segment in segments:
            if _file_3_is_version_segment(segment) or _file_3_is_safety_label_segment(segment):
                continue
            if (
                _file_3_is_blacklisted_contractor_value(segment)
                or _file_3_is_suspicious_company_value(segment)
            ):
                continue
            return segment
    return ""


def _extract_file_3_title_from_filename(
    source_path: Path,
    *,
    company_name: str = "",
) -> str:
    """Return the activity/substance title inferred from the File 3 filename."""

    segments = _split_file_3_filename_segments(source_path)
    if not segments:
        return ""

    safety_label_index = _file_3_safety_label_index(segments)
    title_segments = (
        segments[safety_label_index + 1 :]
        if safety_label_index is not None and safety_label_index + 1 < len(segments)
        else segments
    )
    filtered_segments: List[str] = []
    lowered_company_name = company_name.casefold()
    for segment in title_segments:
        lowered_segment = segment.casefold()
        if _file_3_is_version_segment(segment) or _file_3_is_safety_label_segment(segment):
            continue
        if _file_3_is_blacklisted_contractor_value(segment):
            continue
        stripped_segment = _strip_file_3_safety_markers(segment)
        if lowered_company_name and lowered_company_name in lowered_segment:
            continue
        if (
            safety_label_index is None
            and stripped_segment
            and stripped_segment.casefold() != lowered_segment
        ):
            if not _file_3_is_blacklisted_contractor_value(stripped_segment):
                filtered_segments.append(stripped_segment)
            continue
        filtered_segments.append(segment)

    if not filtered_segments and safety_label_index is None and segments:
        filtered_segments = [
            segment
            for segment in segments[1:]
            if not _file_3_is_version_segment(segment)
            and not _file_3_is_safety_label_segment(segment)
            and not _file_3_is_blacklisted_contractor_value(segment)
        ]

    if not filtered_segments and segments:
        stripped_segments: List[str] = []
        for segment in segments:
            stripped_segment = _strip_file_3_safety_markers(segment)
            if not stripped_segment:
                continue
            if lowered_company_name and lowered_company_name in stripped_segment.casefold():
                stripped_segment = re.sub(
                    re.escape(company_name),
                    " ",
                    stripped_segment,
                    flags=re.IGNORECASE,
                )
                stripped_segment = _clean_safety_value(stripped_segment)
            if not stripped_segment or _file_3_is_blacklisted_contractor_value(stripped_segment):
                continue
            stripped_segments.append(stripped_segment)
        filtered_segments = stripped_segments

    if not filtered_segments:
        return ""
    while filtered_segments and re.fullmatch(r"\d+[a-zA-Z]*", filtered_segments[0]) is not None:
        filtered_segments.pop(0)
    if not filtered_segments:
        return ""
    return _clean_safety_value(" ".join(filtered_segments))


def _build_file_3_fallback_text(source_text: str) -> str:
    """Return a low-noise fallback text window for File 3 metadata parsing."""

    trimmed_text = source_text[:FILE_3_TEXT_FALLBACK_MAX_CHARS]
    fallback_lines: List[str] = []
    for raw_line in trimmed_text.splitlines():
        normalized_line = _normalize_text(raw_line)
        if not normalized_line:
            continue
        if len(normalized_line) > FILE_3_TEXT_FALLBACK_MAX_LINE_LENGTH:
            label_value_match = re.match(r"^[^:]{1,30}:\s*(?P<value>.+)$", normalized_line)
            if (
                label_value_match is None
                or len(_clean_safety_value(label_value_match.group("value")))
                > FILE_3_TEXT_FALLBACK_MAX_LINE_LENGTH
            ):
                continue
        fallback_lines.append(normalized_line)

    if not fallback_lines:
        for raw_sentence in re.split(r"(?<=[.;:])\s+", _normalize_text(trimmed_text)):
            normalized_sentence = _clean_safety_value(raw_sentence)
            if (
                normalized_sentence
                and len(normalized_sentence) <= FILE_3_TEXT_FALLBACK_MAX_LINE_LENGTH
            ):
                fallback_lines.append(normalized_sentence)
    return "\n".join(fallback_lines)


def _extract_file_3_anchor_value(
    source_text: str,
    labels: Iterable[str],
    *,
    max_length: Optional[int] = None,
    value_pattern: Optional[str] = None,
) -> str:
    """Return one explicit File 3 key-value anchor match from raw document text."""

    candidate_value = _extract_text_after_labels(source_text, labels)
    if not candidate_value:
        return ""
    cleaned_candidate = _clean_safety_value(candidate_value)
    if not cleaned_candidate:
        return ""
    if max_length is not None and len(cleaned_candidate) > max_length:
        return ""
    if value_pattern is not None and re.fullmatch(value_pattern, cleaned_candidate) is None:
        return ""
    return cleaned_candidate


def _extract_safety_reference(
    pdf_text: str,
    fallback: str,
    *,
    source_path: Optional[Path] = None,
) -> str:
    """Return a best-effort File 3 document reference."""

    if source_path is not None:
        filename_reference = _file_3_find_reference_from_segments(
            _split_file_3_filename_segments(source_path)
        )
        if filename_reference:
            return filename_reference

    normalized_text = _normalize_text(pdf_text)
    for pattern in SAFETY_REFERENCE_PATTERNS:
        match = pattern.search(normalized_text)
        if match is not None:
            return _clean_safety_value(match.group("value"))
    return fallback


def _extract_safety_version(
    pdf_text: str,
    *,
    source_path: Optional[Path] = None,
) -> str:
    """Return a best-effort safety document version."""

    anchored_version = _extract_file_3_anchor_value(
        pdf_text,
        FILE_3_VERSION_ANCHOR_LABELS,
        max_length=12,
        value_pattern=r"\d+(?:\.\d+)*[a-zA-Z]?",
    )
    if anchored_version:
        return anchored_version

    if source_path is not None:
        filename_version = _file_3_find_version_from_segments(
            _split_file_3_filename_segments(source_path)
        )
        if filename_version:
            return filename_version

    normalized_text = _build_file_3_fallback_text(pdf_text)
    for pattern in SAFETY_VERSION_PATTERNS:
        match = pattern.search(normalized_text)
        if match is not None:
            candidate_value = _clean_safety_value(match.group("value"))
            if re.fullmatch(r"\d+(?:\.\d+)*[a-zA-Z]?", candidate_value):
                return candidate_value
    return DEFAULT_SAFETY_VERSION


def _extract_safety_review_date(
    pdf_text: str,
    *,
    fallback_date: Optional[date] = None,
) -> date:
    """Return the most relevant review-style date from one safety PDF."""

    normalized_text = _normalize_text(pdf_text)
    for label in SAFETY_REVIEW_DATE_LABELS:
        match = re.search(
            rf"{re.escape(label)}\s*(?::|-)?\s*(?P<date>{DATE_VALUE_PATTERN})",
            normalized_text,
            re.IGNORECASE,
        )
        if match is None:
            continue
        parsed_date = _parse_date_string(match.group("date"))
        if parsed_date is not None:
            return parsed_date

    extracted_dates = [candidate[0] for candidate in _extract_dates_with_positions(normalized_text)]
    if extracted_dates:
        return max(extracted_dates)
    return fallback_date or date.today()


def _extract_coshh_substance_name(pdf_text: str, source_path: Path) -> str:
    """Return the substance name from a COSHH / SDS PDF."""

    filename_substance = _extract_file_3_title_from_filename(source_path)
    extracted_value = _extract_text_after_labels(
        _build_file_3_fallback_text(pdf_text),
        COSHH_SUBSTANCE_LABELS,
    )
    if extracted_value and (
        not filename_substance or len(filename_substance.split()) <= 1
    ) and not _file_3_is_suspicious_activity_value(extracted_value):
        return extracted_value
    if filename_substance and not re.search(
        r"\b(?:coshh|sds|assessment|sheet)\b",
        filename_substance,
        re.IGNORECASE,
    ) and not _file_3_is_suspicious_activity_value(filename_substance):
        return filename_substance
    if extracted_value and not _file_3_is_suspicious_activity_value(extracted_value):
        return extracted_value

    fallback = re.sub(
        r"(?i)\b(?:coshh|safety\s+data\s+sheet|sds)\b",
        "",
        source_path.stem.replace("_", " ").replace("-", " "),
    )
    cleaned_fallback = _clean_safety_value(fallback)
    if not _file_3_is_suspicious_activity_value(cleaned_fallback):
        return cleaned_fallback
    return source_path.stem


def _extract_coshh_supplier(pdf_text: str, source_path: Optional[Path] = None) -> str:
    """Return the supplier / manufacturer from a COSHH / SDS PDF."""

    if source_path is not None:
        filename_segments = _split_file_3_filename_segments(source_path)
        safety_label_index = _file_3_safety_label_index(filename_segments)
        if safety_label_index is not None and safety_label_index > 0:
            supplier_candidate = _clean_safety_value(
                " ".join(filename_segments[:safety_label_index])
            )
            if supplier_candidate and not _file_3_is_suspicious_company_value(
                supplier_candidate
            ):
                return supplier_candidate

    extracted_value = _extract_text_after_labels(
        _build_file_3_fallback_text(pdf_text),
        COSHH_SUPPLIER_LABELS,
    )
    if extracted_value and not _file_3_is_suspicious_company_value(extracted_value):
        return extracted_value

    text_company = _extract_file_3_company_from_text(pdf_text)
    if text_company:
        return text_company
    return "Unknown Supplier"


def _extract_coshh_use(pdf_text: str) -> str:
    """Return the intended use text from a COSHH / SDS PDF."""

    extracted_value = _extract_text_after_labels(pdf_text, COSHH_USE_LABELS)
    return extracted_value or "General site use"


def _extract_coshh_hazard_keywords(pdf_text: str) -> List[str]:
    """Return a lightweight hazard summary from one COSHH / SDS PDF."""

    lowered_text = pdf_text.casefold()
    detected_hazards: List[str] = []
    for keyword, label in (
        ("flammable", "Flammable"),
        ("corrosive", "Corrosive"),
        ("irritant", "Irritant"),
        ("toxic", "Toxic"),
        ("oxidising", "Oxidising"),
        ("health hazard", "Health Hazard"),
        ("compressed gas", "Compressed Gas"),
        ("environment", "Environmental"),
    ):
        if keyword in lowered_text and label not in detected_hazards:
            detected_hazards.append(label)
    return detected_hazards


def _extract_rams_activity_description(
    pdf_text: str,
    source_path: Path,
    *,
    company_name: str = "",
) -> str:
    """Return the work activity from one RAMS PDF."""

    anchored_activity = _extract_file_3_anchor_value(
        pdf_text,
        FILE_3_TITLE_ANCHOR_LABELS,
        max_length=160,
    )
    if anchored_activity and not _file_3_is_suspicious_activity_value(
        anchored_activity,
        company_name=company_name,
    ):
        return anchored_activity

    filename_activity = _extract_file_3_title_from_filename(
        source_path,
        company_name=company_name,
    )
    if filename_activity and not _file_3_is_suspicious_activity_value(
        filename_activity,
        company_name=company_name,
    ):
        return filename_activity

    extracted_value = _extract_text_after_labels(
        _build_file_3_fallback_text(pdf_text),
        RAMS_ACTIVITY_LABELS,
    )
    if extracted_value and not _file_3_is_suspicious_activity_value(
        extracted_value,
        company_name=company_name,
    ):
        return extracted_value

    fallback = re.sub(
        r"(?i)\b(?:rams|risk\s+assessment(?:\s+and)?|method\s+statement)\b",
        "",
        source_path.stem.replace("_", " ").replace("-", " "),
    )
    cleaned_fallback = _clean_safety_value(fallback)
    if _file_3_is_suspicious_activity_value(
        cleaned_fallback,
        company_name=company_name,
    ):
        return "RAMS Document"
    return cleaned_fallback or "RAMS Document"


def _guess_file_3_contractor_name(
    repository: DocumentRepository,
    *,
    site_name: str,
    pdf_text: str,
    source_path: Path,
    fallback: str = "",
) -> str:
    """Return the most likely contractor name for one File 3 safety PDF."""

    anchored_company = _extract_file_3_anchor_value(
        pdf_text,
        FILE_3_COMPANY_ANCHOR_LABELS,
        max_length=120,
    )
    if anchored_company and not _file_3_is_suspicious_company_value(anchored_company):
        return anchored_company

    candidate_names = _file_3_candidate_company_names(
        repository,
        site_name=site_name,
    )
    filename_company = _extract_file_3_company_from_filename(
        source_path,
        candidate_names=candidate_names,
    )
    if filename_company:
        return filename_company

    search_text = pdf_text
    matched_names = [
        candidate_name
        for candidate_name in candidate_names
        if (
            _file_3_matches_company_name(candidate_name, search_text)
            and not _file_3_is_blacklisted_contractor_value(candidate_name)
        )
    ]
    if matched_names:
        return max(matched_names, key=len)

    text_company = _extract_file_3_company_from_text(pdf_text)
    if text_company:
        return text_company

    if fallback.strip() and not _file_3_is_suspicious_company_value(fallback.strip()):
        return fallback.strip()

    filename_tokens = _split_file_3_filename_segments(source_path)
    cleaned_filename_tokens = [
        cleaned_token
        for cleaned_token in (
            _strip_file_3_safety_markers(token)
            for token in filename_tokens
        )
        if cleaned_token and not _file_3_is_blacklisted_contractor_value(cleaned_token)
    ]
    if cleaned_filename_tokens:
        return _clean_safety_value(" ".join(cleaned_filename_tokens[:2]))
    return "Site Contractor"


def _upsert_coshh_document_from_source(
    repository: DocumentRepository,
    source_path: Path,
    *,
    source_text: Optional[str] = None,
    site_name_override: Optional[str] = None,
) -> Optional[COSHHDocument]:
    """Create or update one File 3 COSHH record from a synced safety source."""

    resolved_source_text = (
        source_text if source_text is not None else _safe_extract_safety_source_text(source_path)
    )
    if not _is_coshh_safety_source(source_path, resolved_source_text):
        return None

    project_setup = _load_workspace_project_setup()
    site_name = (
        site_name_override
        or project_setup.get("current_site_name")
        or _infer_default_site_name(repository)
    )
    supplier_name = _extract_coshh_supplier(resolved_source_text, source_path)
    contractor_name = _guess_file_3_contractor_name(
        repository,
        site_name=site_name,
        pdf_text=resolved_source_text,
        source_path=source_path,
        fallback=supplier_name,
    )
    coshh_document = COSHHDocument(
        doc_id=_build_safety_doc_id("COSHH", site_name, source_path.stem),
        site_name=site_name,
        created_at=datetime.now(),
        status=DocumentStatus.ACTIVE,
        contractor_name=contractor_name,
        substance_name=_extract_coshh_substance_name(resolved_source_text, source_path),
        hazard_pictograms=_extract_coshh_hazard_keywords(resolved_source_text),
        ppe_required=[],
        emergency_first_aid="Refer to the Safety Data Sheet.",
        reference=_extract_safety_reference(
            resolved_source_text,
            source_path.stem,
            source_path=source_path,
        ),
        version=_extract_safety_version(
            resolved_source_text,
            source_path=source_path,
        ),
        manufacturer=supplier_name,
        review_date=_extract_safety_review_date(resolved_source_text),
        supplier_name=supplier_name,
        intended_use=_extract_coshh_use(resolved_source_text),
        assessor_name=DEFAULT_SAFETY_MANAGER_NAME,
        manager_name=DEFAULT_SAFETY_MANAGER_NAME,
        manager_position=DEFAULT_SAFETY_MANAGER_POSITION,
    )
    repository.save(coshh_document)
    return coshh_document


def _upsert_coshh_document_from_pdf(
    repository: DocumentRepository,
    pdf_path: Path,
    *,
    pdf_text: Optional[str] = None,
) -> Optional[COSHHDocument]:
    """Compatibility wrapper for the original PDF-only COSHH upsert path."""

    return _upsert_coshh_document_from_source(
        repository,
        pdf_path,
        source_text=pdf_text,
    )


def _upsert_rams_document_from_source(
    repository: DocumentRepository,
    source_path: Path,
    *,
    source_text: Optional[str] = None,
    site_name_override: Optional[str] = None,
) -> Optional[RAMSDocument]:
    """Create or update one File 3 RAMS record from a synced safety source."""

    resolved_source_text = (
        source_text if source_text is not None else _safe_extract_safety_source_text(source_path)
    )
    if not _is_rams_safety_source(source_path, resolved_source_text):
        return None

    project_setup = _load_workspace_project_setup()
    site_name = (
        site_name_override
        or project_setup.get("current_site_name")
        or _infer_default_site_name(repository)
    )
    contractor_name = _guess_file_3_contractor_name(
        repository,
        site_name=site_name,
        pdf_text=resolved_source_text,
        source_path=source_path,
    )
    rams_document = RAMSDocument(
        doc_id=_build_safety_doc_id("RAMS", site_name, source_path.stem),
        site_name=site_name,
        created_at=datetime.now(),
        status=DocumentStatus.ACTIVE,
        contractor_name=contractor_name,
        activity_description=_extract_rams_activity_description(
            resolved_source_text,
            source_path,
            company_name=contractor_name,
        ),
        approval_date=_extract_safety_review_date(resolved_source_text),
        reference=_extract_safety_reference(
            resolved_source_text,
            source_path.stem,
            source_path=source_path,
        ),
        version=_extract_safety_version(
            resolved_source_text,
            source_path=source_path,
        ),
        manufacturer="",
        review_date=_extract_safety_review_date(resolved_source_text),
        assessor_name=DEFAULT_SAFETY_MANAGER_NAME,
        manager_name=DEFAULT_SAFETY_MANAGER_NAME,
        manager_position=DEFAULT_SAFETY_MANAGER_POSITION,
    )
    repository.save(rams_document)
    return rams_document


def _upsert_rams_document_from_pdf(
    repository: DocumentRepository,
    pdf_path: Path,
    *,
    pdf_text: Optional[str] = None,
) -> Optional[RAMSDocument]:
    """Compatibility wrapper for the original PDF-only RAMS upsert path."""

    return _upsert_rams_document_from_source(
        repository,
        pdf_path,
        source_text=pdf_text,
    )


def _sync_existing_safety_sources(
    repository: DocumentRepository,
    *,
    destination_directory: Path,
    detector: Callable[[Path, str], bool],
    upsert_document: Callable[..., Optional[BaseDocument]],
    file_category: str,
) -> None:
    """Backfill File 3 safety documents from already-filed PDFs and Word files."""

    for source_path in sorted(destination_directory.iterdir(), key=lambda path: path.name.lower()):
        if (
            not source_path.is_file()
            or source_path.suffix.lower() not in FILE_3_SAFETY_SOURCE_SUFFIXES
        ):
            continue
        source_text = _safe_extract_safety_source_text(source_path)
        if not detector(source_path, source_text):
            continue
        safety_document = upsert_document(
            repository,
            source_path,
            source_text=source_text,
        )
        repository.index_file(
            file_name=source_path.name,
            file_path=source_path,
            file_category=_resolve_safety_file_category(file_category, source_path),
            file_group=FileGroup.FILE_3,
            site_name=(
                safety_document.site_name
                if safety_document is not None
                else _load_workspace_project_setup().get("current_site_name")
            ),
            related_doc_id=(
                safety_document.doc_id
                if safety_document is not None
                else None
            ),
        )


def _looks_like_file_3_review_hold_source(source_path: Path, source_text: str = "") -> bool:
    """Return True when one source looks safety-related but should be held out of registers."""

    lowered_text = f"{source_path.name} {source_text}".casefold()
    return any(keyword in lowered_text for keyword in FILE_3_REVIEW_HOLD_KEYWORDS)


def rebuild_file_3_safety_inventory(
    repository: DocumentRepository,
    *,
    site_name: Optional[str] = None,
) -> RebuiltSafetyInventoryResult:
    """Rebuild File 3 RAMS/COSHH records from the filed source documents."""

    repository.create_schema()

    active_site_name = (
        site_name
        or _load_workspace_project_setup().get("current_site_name")
        or _infer_default_site_name(repository)
    )
    for directory in (
        config.RAMS_DESTINATION,
        config.COSHH_DESTINATION,
        config.FILE_3_REVIEW_DIR,
    ):
        directory.mkdir(parents=True, exist_ok=True)

    for document_type in (RAMSDocument.document_type, COSHHDocument.document_type):
        for document in repository.list_documents(
            document_type=document_type,
            site_name=active_site_name,
        ):
            repository.delete_document(document.doc_id)

    source_paths = sorted(
        {
            source_path.resolve()
            for destination_directory in (
                config.RAMS_DESTINATION,
                config.COSHH_DESTINATION,
                config.FILE_3_REVIEW_DIR,
            )
            for source_path in destination_directory.iterdir()
            if source_path.is_file()
            and source_path.suffix.lower() in FILE_3_SAFETY_SOURCE_SUFFIXES
        },
        key=lambda path: (path.parent.name.casefold(), path.name.casefold()),
    )

    moved_file_names: List[str] = []
    ignored_file_names: List[str] = []
    rams_records = 0
    coshh_records = 0
    moved_files = 0

    for source_path in source_paths:
        source_text = _safe_extract_safety_source_text(source_path)

        if _is_coshh_safety_source(source_path, source_text):
            destination_directory = config.COSHH_DESTINATION
            base_category = "coshh_pdf"
            created_document = _upsert_coshh_document_from_source(
                repository,
                source_path,
                source_text=source_text,
                site_name_override=active_site_name,
            )
            coshh_records += 1
        elif _is_rams_safety_source(source_path, source_text):
            destination_directory = config.RAMS_DESTINATION
            base_category = "rams_pdf"
            created_document = _upsert_rams_document_from_source(
                repository,
                source_path,
                source_text=source_text,
                site_name_override=active_site_name,
            )
            rams_records += 1
        else:
            destination_directory = config.FILE_3_REVIEW_DIR
            base_category = "safety_review_pdf"
            created_document = None
            ignored_file_names.append(source_path.name)

        target_path = source_path
        if source_path.parent != destination_directory:
            target_path = _move_file(source_path, destination_directory)
            moved_files += 1
            moved_file_names.append(target_path.name)

        if created_document is not None and target_path != source_path:
            created_document = (
                _upsert_coshh_document_from_source(
                    repository,
                    target_path,
                    source_text=source_text,
                    site_name_override=active_site_name,
                )
                if isinstance(created_document, COSHHDocument)
                else _upsert_rams_document_from_source(
                    repository,
                    target_path,
                    source_text=source_text,
                    site_name_override=active_site_name,
                )
            )

        repository.index_file(
            file_name=target_path.name,
            file_path=target_path,
            file_category=_resolve_safety_file_category(base_category, target_path),
            file_group=FileGroup.FILE_3,
            site_name=active_site_name,
            related_doc_id=created_document.doc_id if created_document is not None else None,
        )

    return RebuiltSafetyInventoryResult(
        total_sources=len(source_paths),
        rams_records=rams_records,
        coshh_records=coshh_records,
        moved_files=moved_files,
        ignored_sources=len(ignored_file_names),
        moved_file_names=tuple(sorted(moved_file_names, key=str.casefold)),
        ignored_file_names=tuple(sorted(ignored_file_names, key=str.casefold)),
    )


def park_file_3_document_for_review(
    repository: DocumentRepository,
    doc_id: str,
) -> Tuple[Path, ...]:
    """Move one live File 3 source into Needs Review and remove its register row."""

    repository.create_schema()
    document = repository.get(doc_id)
    if not isinstance(document, (RAMSDocument, COSHHDocument)):
        raise ValueError("Only RAMS and COSHH documents can be parked for File 3 review.")

    indexed_files = repository.list_indexed_files(related_doc_id=doc_id)
    moved_paths: List[Path] = []
    config.FILE_3_REVIEW_DIR.mkdir(parents=True, exist_ok=True)

    for indexed_file in indexed_files:
        current_path = indexed_file.file_path
        if not current_path.exists():
            continue
        target_path = current_path
        if current_path.parent != config.FILE_3_REVIEW_DIR:
            target_path = _move_file(current_path, config.FILE_3_REVIEW_DIR)
        moved_paths.append(target_path)

    repository.delete_document(doc_id)

    for moved_path in moved_paths:
        repository.index_file(
            file_name=moved_path.name,
            file_path=moved_path,
            file_category=_resolve_safety_file_category("safety_review_pdf", moved_path),
            file_group=FileGroup.FILE_3,
            site_name=document.site_name,
            related_doc_id=None,
        )

    return tuple(moved_paths)


def _safe_extract_safety_source_text(source_path: Path) -> str:
    """Return the best-effort raw text for one safety source file."""

    suffix = source_path.suffix.lower()
    if suffix == ".pdf":
        return _safe_extract_pdf_text(source_path)
    if suffix in {".docx", ".doc"}:
        return _safe_extract_word_text(source_path)
    return ""


def _safe_extract_word_text(word_path: Path) -> str:
    """Return raw text from a Word document without raising on malformed files."""

    try:
        if word_path.suffix.lower() == ".docx":
            return _extract_docx_text(word_path)
        if word_path.suffix.lower() == ".doc":
            with tempfile.TemporaryDirectory(prefix="uplands-doc-convert-") as temp_dir:
                converted_docx_path = Path(temp_dir) / f"{word_path.stem}.docx"
                _convert_legacy_word_document_to_docx(word_path, converted_docx_path)
                return _extract_docx_text(converted_docx_path)
    except (OSError, RuntimeError, ValueError, zipfile.BadZipFile):
        return ""
    return ""


def _extract_docx_text(docx_path: Path) -> str:
    """Return a flattened text view of one DOCX, including tables and headers."""

    document = Document(docx_path)
    extracted_segments: List[str] = []
    _extend_docx_container_text(extracted_segments, document)
    for section in document.sections:
        _extend_docx_container_text(extracted_segments, section.header)
        _extend_docx_container_text(extracted_segments, section.footer)
    return "\n".join(segment for segment in extracted_segments if segment).strip()


def _extend_docx_container_text(extracted_segments: List[str], container: Any) -> None:
    """Append cleaned text from one python-docx container into the output list."""

    for paragraph in getattr(container, "paragraphs", []):
        paragraph_text = _normalize_text(paragraph.text)
        if paragraph_text:
            extracted_segments.append(paragraph_text)

    for table in getattr(container, "tables", []):
        for row in table.rows:
            row_segments = [
                _normalize_text(cell.text)
                for cell in row.cells
                if _normalize_text(cell.text)
            ]
            if row_segments:
                extracted_segments.append(" | ".join(row_segments))


def _convert_legacy_word_document_to_docx(doc_path: Path, destination_path: Path) -> None:
    """Convert one legacy .doc file into a temporary .docx via macOS textutil."""

    completed_process = subprocess.run(
        [
            "textutil",
            "-convert",
            "docx",
            "-output",
            str(destination_path),
            str(doc_path),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if completed_process.returncode != 0 or not destination_path.exists():
        raise RuntimeError(
            "Failed to convert legacy Word document for File 3 safety intake."
        )


def _resolve_safety_file_category(base_category: str, source_path: Path) -> str:
    """Return the indexed file-category for one safety source extension."""

    suffix = source_path.suffix.lower().lstrip(".")
    if not suffix or suffix == "pdf":
        return base_category
    if base_category.endswith("_pdf"):
        return f"{base_category[:-4]}_{suffix}"
    return f"{base_category}_{suffix}"


def _safe_extract_pdf_text(pdf_path: Path) -> str:
    """Return embedded PDF text without raising on malformed source files."""

    try:
        return "\n".join(_extract_pdf_text_pages(pdf_path)).strip()
    except RuntimeError:
        return ""


def _is_plant_hire_pdf(source_path: Path, pdf_text: str = "") -> bool:
    """Return True when one PDF belongs in the File 2 plant register folder."""

    if source_path.suffix.lower() != ".pdf":
        return False

    lowered_name = source_path.name.casefold()
    lowered_text = pdf_text.casefold()
    return any(
        (
            lowered_name.startswith("contract-h-"),
            lowered_name.startswith("quote-h-"),
            "order confirmation" in lowered_text,
            "proof of delivery" in lowered_text,
            "proof of collection" in lowered_text,
            "hss proservice" in lowered_text,
            "prohire" in lowered_text,
            "the hire service company" in lowered_text,
            "mep hire" in lowered_text,
        )
    )


def _is_hss_order_confirmation_pdf(source_path: Path, pdf_text: str) -> bool:
    """Return True when one PDF is an HSS order confirmation."""

    lowered_name = source_path.name.casefold()
    lowered_text = pdf_text.casefold()
    if lowered_name.startswith("quote-h-"):
        return False
    return (
        lowered_name.startswith("contract-h-")
        or ("order confirmation" in lowered_text and HSS_ORDER_REF_PATTERN.search(pdf_text) is not None)
    )


def _is_plant_collection_note_pdf(source_path: Path, pdf_text: str) -> bool:
    """Return True when one PDF is a plant proof-of-collection / off-hire note."""

    lowered_name = source_path.name.casefold()
    lowered_text = pdf_text.casefold()
    return any(
        (
            "proof of collection" in lowered_text,
            "off hire:" in lowered_text and "collection:" in lowered_text,
            lowered_name.startswith("proof-of-collection"),
        )
    )


def _classify_plant_hire_pdf(source_path: Path, pdf_text: str) -> str:
    """Return a stable file-category label for one plant PDF."""

    lowered_name = source_path.name.casefold()
    lowered_text = pdf_text.casefold()
    if lowered_name.startswith("quote-h-") or "quote reference" in lowered_text:
        return "plant_hire_quote_pdf"
    if _is_hss_order_confirmation_pdf(source_path, pdf_text):
        return "plant_hire_order_pdf"
    if _is_plant_collection_note_pdf(source_path, pdf_text):
        return "plant_hire_collection_pdf"
    if "proof of delivery" in lowered_text:
        return "plant_hire_delivery_pdf"
    return "plant_hire_pdf"


def _tokenize_plant_pdf_text(pdf_text: str) -> List[str]:
    """Split PDF text into cleaned tokens for table-like hire documents."""

    return [
        token.strip()
        for token in re.split(r"[|\n\r]+", pdf_text)
        if token and token.strip()
    ]


def _looks_like_purchase_order(token: str) -> bool:
    """Return True when one token resembles a PO/reference field."""

    cleaned_token = token.strip()
    if not cleaned_token or len(cleaned_token) > 24:
        return False
    if EMAIL_PATTERN.search(cleaned_token) or PHONE_PATTERN.search(cleaned_token):
        return False
    if re.search(r"\s", cleaned_token):
        return False
    return any(character.isdigit() for character in cleaned_token)


def _looks_like_money(token: str) -> bool:
    """Return True when one token looks like a GBP price cell."""

    return bool(re.fullmatch(r"£?\d+(?:\.\d+)?", token.strip()))


def _clean_plant_description(description: str) -> str:
    """Normalize OCR quirks from plant item descriptions."""

    cleaned_description = " ".join(description.split())
    cleaned_description = re.sub(r"\(\s*UP\s*$", "", cleaned_description, flags=re.IGNORECASE)
    cleaned_description = cleaned_description.rstrip(" -(")
    return cleaned_description


def _plant_description_match_key(description: str) -> str:
    """Return a stable comparison key for one plant description."""

    cleaned_description = _clean_plant_description(description)
    cleaned_description = re.sub(r"\s*\(x\d+\)\s*$", "", cleaned_description, flags=re.IGNORECASE)
    return cleaned_description.casefold()


def _extract_plant_description_quantity(description: str) -> int:
    """Return the quantity suffix embedded in one plant description."""

    match = re.search(r"\(x(\d+)\)\s*$", description, flags=re.IGNORECASE)
    return int(match.group(1)) if match is not None else 1


def _extract_hss_company_phone(pdf_text: str) -> str:
    """Return the best HSS contact phone from one order confirmation."""

    for match in PHONE_PATTERN.findall(pdf_text):
        cleaned_phone = " ".join(match.split())
        if cleaned_phone.startswith("07"):
            continue
        return cleaned_phone
    return HSS_DEFAULT_PHONE


def _extract_hss_purchase_order(tokens: List[str], order_ref: str) -> str:
    """Return the purchase-order token that precedes one HSS order ref."""

    try:
        order_index = next(
            index
            for index, token in enumerate(tokens)
            if order_ref.casefold() in token.casefold()
        )
    except StopIteration:
        return ""

    for offset in range(1, 6):
        candidate_index = order_index - offset
        if candidate_index < 0:
            break
        candidate_token = tokens[candidate_index]
        if _looks_like_purchase_order(candidate_token):
            return candidate_token
    return ""


def _extract_hss_customer_name(tokens: List[str], order_ref: str) -> str:
    """Return the best available customer/hirer label from one HSS order."""

    try:
        order_index = next(
            index
            for index, token in enumerate(tokens)
            if order_ref.casefold() in token.casefold()
        )
    except StopIteration:
        return ""

    for offset in range(1, 8):
        candidate_index = order_index - offset
        if candidate_index < 0:
            break
        candidate_token = " ".join(tokens[candidate_index].split()).strip()
        if not candidate_token:
            continue
        if EMAIL_PATTERN.search(candidate_token) or PHONE_PATTERN.search(candidate_token):
            continue
        if HSS_ORDER_REF_PATTERN.search(candidate_token):
            continue
        if _looks_like_purchase_order(candidate_token):
            continue
        if not re.search(r"[A-Za-z]", candidate_token):
            continue
        return candidate_token
    return ""


def _normalise_plant_hired_by_label(value: str) -> str:
    """Normalize common plant hirer labels into a cleaner register value."""

    cleaned_value = " ".join(value.split()).strip()
    if not cleaned_value:
        return ""

    lowered_value = cleaned_value.casefold()
    if any(alias in lowered_value for alias in ("uplands", "url", "uplands retail")):
        return "URL (Uplands)"
    if "archer" in lowered_value:
        return "A. Archer Electrical"
    if "tde" in lowered_value:
        return "TDE"
    return cleaned_value


def _is_pending_plant_inspection_value(value: str) -> bool:
    """Return True when one plant inspection value is only a placeholder."""

    return is_pending_plant_inspection_reference(value)


def _extract_hss_product_lines(pdf_text: str) -> List[Dict[str, Any]]:
    """Return one parsed plant entry per product line on an HSS contract PDF."""

    tokens = _tokenize_plant_pdf_text(pdf_text)
    product_lines: List[Dict[str, Any]] = []
    seen_keys: set[Tuple[str, str, str, int]] = set()
    company_phone = _extract_hss_company_phone(pdf_text)

    for index, token in enumerate(tokens):
        date_match = HSS_DATE_RANGE_PATTERN.search(token)
        if date_match is None or index < 1:
            continue

        stock_token = ""
        description = ""
        previous_token = tokens[index - 1].strip()
        previous_stock_token = previous_token if HSS_STOCK_CODE_PATTERN.fullmatch(previous_token) else ""
        if previous_stock_token and index >= 2:
            stock_token = previous_stock_token
            description = _clean_plant_description(tokens[index - 2])
        else:
            merged_match = re.match(r"(?P<description>.+?)\s+(?P<stock>\d{4,6})$", previous_token)
            if merged_match is None:
                continue
            stock_token = merged_match.group("stock")
            description = _clean_plant_description(merged_match.group("description"))
        if len(description) < 4 or not re.search(r"[A-Z]", description, flags=re.IGNORECASE):
            continue

        quantity = 1
        if (
            index + 2 < len(tokens)
            and _looks_like_money(tokens[index + 1])
            and re.fullmatch(r"\d+", tokens[index + 2].strip())
        ):
            quantity = int(tokens[index + 2].strip())

        display_description = (
            f"{description} (x{quantity})"
            if quantity > 1
            else description
        )
        dedupe_key = (
            display_description.casefold(),
            stock_token,
            date_match.group("start"),
            quantity,
        )
        if dedupe_key in seen_keys:
            continue
        seen_keys.add(dedupe_key)

        product_lines.append(
            {
                "description": display_description,
                "stock_code": stock_token,
                "company": "HSS",
                "phone": company_phone,
                "on_hire": _coerce_date_or_none(date_match.group("start")) or date.today(),
            }
        )

    return product_lines


def _parse_hss_order_confirmation(pdf_text: str) -> Dict[str, Any]:
    """Extract HSS order-level metadata and line items from one PDF."""

    order_ref_match = HSS_ORDER_REF_PATTERN.search(pdf_text)
    if order_ref_match is None:
        return {}

    tokens = _tokenize_plant_pdf_text(pdf_text)
    order_ref = order_ref_match.group(0).upper()
    return {
        "order_ref": order_ref,
        "customer_name": _extract_hss_customer_name(tokens, order_ref),
        "purchase_order": _extract_hss_purchase_order(tokens, order_ref),
        "product_lines": _extract_hss_product_lines(pdf_text),
    }


def _extract_collection_note_header_value(lines: List[str], label: str) -> str:
    """Return one label value from a proof-of-collection PDF line list."""

    label_prefix = f"{label.casefold()}:"
    for index, line in enumerate(lines):
        lowered_line = line.casefold()
        if lowered_line == label_prefix and index + 1 < len(lines):
            return lines[index + 1].strip()
        if lowered_line.startswith(label_prefix):
            return line.split(":", 1)[1].strip()
    return ""


def _looks_like_plant_collection_serial(token: str) -> bool:
    """Return True when one token looks like an item-level returned serial/reference."""

    cleaned_token = token.strip().upper()
    if not cleaned_token:
        return False
    if re.fullmatch(r"\d+", cleaned_token):
        return False
    if PHONE_PATTERN.search(cleaned_token) or EMAIL_PATTERN.search(cleaned_token):
        return False
    return bool(re.fullmatch(r"[A-Z0-9-]{4,}", cleaned_token)) and any(
        character.isalpha() for character in cleaned_token
    )


def _parse_plant_collection_product_lines(pdf_text: str) -> List[Dict[str, Any]]:
    """Return aggregated returned line items from one proof-of-collection PDF."""

    lines = [
        " ".join(line.split()).strip()
        for line in pdf_text.splitlines()
        if line and line.strip()
    ]
    ignored_description_lines = {
        "comm code",
        "description",
        "e/code",
        "o/s qty",
        "advised qty",
        "actual qty",
        "damage qty",
        "dirty qty",
    }
    stop_prefixes = (
        "collection slot:",
        "driver instruction:",
        "customer name:",
        "damage waiver option",
        "returned equipment is subject",
        "customer signature",
        "job photo",
        "gps location",
        "gps latitude:",
        "gps longitude:",
        "click here for google map",
        "driver name:",
        "vehicle registration:",
        "date:",
    )

    aggregated: Dict[Tuple[str, str], Dict[str, Any]] = {}
    index = 0
    while index + 1 < len(lines):
        stock_code = lines[index]
        description = lines[index + 1]
        description_key = description.casefold()
        if not re.fullmatch(r"\d{4,6}", stock_code):
            index += 1
            continue
        if description_key in ignored_description_lines or not re.search(r"[A-Za-z]", description):
            index += 1
            continue

        next_index = index + 2
        quantities: List[int] = []
        serials: List[str] = []
        while next_index < len(lines):
            token = lines[next_index]
            lowered_token = token.casefold()
            if any(lowered_token.startswith(prefix) for prefix in stop_prefixes):
                break
            if (
                re.fullmatch(r"\d{4,6}", token)
                and next_index + 1 < len(lines)
                and re.search(r"[A-Za-z]", lines[next_index + 1])
                and lines[next_index + 1].casefold() not in ignored_description_lines
            ):
                break
            if re.fullmatch(r"\d+", token):
                quantities.append(int(token))
            elif _looks_like_plant_collection_serial(token):
                serials.append(token)
            next_index += 1

        match_key = (stock_code, _plant_description_match_key(description))
        existing_entry = aggregated.setdefault(
            match_key,
            {
                "stock_code": stock_code,
                "description": _clean_plant_description(description),
                "actual_quantity": 0,
                "serials": [],
            },
        )
        if len(quantities) >= 3:
            existing_entry["actual_quantity"] = max(
                int(existing_entry["actual_quantity"]),
                quantities[2],
            )
        existing_entry["serials"].extend(serials)
        index = next_index

    product_lines: List[Dict[str, Any]] = []
    for entry in aggregated.values():
        unique_serials = list(dict.fromkeys(entry["serials"]))
        actual_quantity = int(entry["actual_quantity"]) or len(unique_serials) or 1
        product_lines.append(
            {
                "stock_code": str(entry["stock_code"]),
                "description": str(entry["description"]),
                "actual_quantity": actual_quantity,
                "serials": unique_serials,
            }
        )
    return product_lines


def _parse_plant_collection_note(pdf_text: str) -> Dict[str, Any]:
    """Extract collection-level metadata and returned line items from one PDF."""

    lines = [
        " ".join(line.split()).strip()
        for line in pdf_text.splitlines()
        if line and line.strip()
    ]
    return {
        "contract_ref": _extract_collection_note_header_value(lines, "Contract"),
        "customer_name": _extract_collection_note_header_value(lines, "Customer Name"),
        "on_hire": _coerce_date_or_none(_extract_collection_note_header_value(lines, "On Hire")),
        "off_hire": _coerce_date_or_none(_extract_collection_note_header_value(lines, "Off Hire")),
        "collection_date": _coerce_date_or_none(_extract_collection_note_header_value(lines, "Collection")),
        "product_lines": _parse_plant_collection_product_lines(pdf_text),
    }


def _coerce_date_or_none(value: str) -> Optional[date]:
    """Parse one date string when present, returning None on failure."""

    try:
        return date.fromisoformat(value)
    except ValueError:
        try:
            return datetime.strptime(value, "%d/%m/%Y").date()
        except ValueError:
            return None


def _build_plant_asset_doc_id(site_name: str, order_ref: str, line_number: int) -> str:
    """Return a stable doc_id for one scanned plant product line."""

    return (
        f"PLANT-{_slugify_identifier(site_name)}-"
        f"{_slugify_identifier(order_ref)}-{line_number:03d}"
    )


def _extract_plant_hire_sequence(hire_num: str) -> int:
    """Return the trailing numeric sequence from one plant hire number."""

    match = re.search(r"(\d+)$", hire_num)
    if match is None:
        return 0
    return int(match.group(1))


def _format_plant_hire_number(job_number: str, sequence: int) -> str:
    """Return the printable hire number for one plant asset."""

    if job_number.strip():
        return f"{job_number.strip()}-{sequence:02d}"
    return f"{sequence:02d}"


def _next_plant_hire_sequence(repository: DocumentRepository, site_name: str) -> int:
    """Return the next available File 2 plant hire sequence for one site."""

    existing_sequences = [
        _extract_plant_hire_sequence(document.hire_num)
        for document in repository.list_documents(
            document_type=PlantAssetDocument.document_type,
            site_name=site_name,
        )
        if isinstance(document, PlantAssetDocument)
        and document.status != DocumentStatus.ARCHIVED
    ]
    return (max(existing_sequences) if existing_sequences else 0) + 1


def _get_plant_asset_document(
    repository: DocumentRepository,
    doc_id: str,
) -> Optional[PlantAssetDocument]:
    """Return one plant asset by doc_id when it already exists."""

    try:
        document = repository.get(doc_id)
    except DocumentNotFoundError:
        return None
    return document if isinstance(document, PlantAssetDocument) else None


def _upsert_plant_assets_from_pdf(
    repository: DocumentRepository,
    pdf_path: Path,
    *,
    pdf_text: Optional[str] = None,
) -> List[PlantAssetDocument]:
    """Create or update pending plant assets from one synced contract PDF."""

    resolved_pdf_text = pdf_text if pdf_text is not None else _safe_extract_pdf_text(pdf_path)
    if _is_plant_collection_note_pdf(pdf_path, resolved_pdf_text):
        return _apply_plant_collection_note_pdf(
            repository,
            pdf_path,
            pdf_text=resolved_pdf_text,
        )
    if not _is_hss_order_confirmation_pdf(pdf_path, resolved_pdf_text):
        return []

    parsed_contract = _parse_hss_order_confirmation(resolved_pdf_text)
    order_ref = str(parsed_contract.get("order_ref") or "").strip()
    product_lines = list(parsed_contract.get("product_lines") or [])
    if not order_ref or not product_lines:
        return []

    project_setup = _load_workspace_project_setup()
    site_name = (
        project_setup.get("current_site_name")
        or _infer_default_site_name(repository)
    )
    job_number = project_setup.get("job_number", "")
    default_hired_by = _default_hired_by_for_project(project_setup)
    parsed_hired_by = _normalise_plant_hired_by_label(
        str(parsed_contract.get("customer_name") or "")
    )
    hired_by = parsed_hired_by or default_hired_by
    purchase_order = str(parsed_contract.get("purchase_order") or "").strip()
    created_at = datetime.now()
    next_sequence = _next_plant_hire_sequence(repository, site_name)
    plant_assets: List[PlantAssetDocument] = []

    for line_number, product_line in enumerate(product_lines, start=1):
        doc_id = _build_plant_asset_doc_id(site_name, order_ref, line_number)
        existing_document = _get_plant_asset_document(repository, doc_id)
        hire_num = (
            existing_document.hire_num
            if existing_document is not None
            else _format_plant_hire_number(job_number, next_sequence)
        )
        if existing_document is None:
            next_sequence += 1

        existing_hired_by = (
            _normalise_plant_hired_by_label(existing_document.hired_by)
            if existing_document is not None
            else ""
        )
        if parsed_hired_by and existing_hired_by in {"", "TDE", default_hired_by}:
            resolved_hired_by = parsed_hired_by
        else:
            resolved_hired_by = existing_hired_by or hired_by

        existing_inspection = (
            str(existing_document.inspection or "")
            if existing_document is not None
            else ""
        )
        resolved_inspection = (
            existing_inspection
            if existing_inspection and not _is_pending_plant_inspection_value(existing_inspection)
            else PLANT_PENDING_INSPECTION_TEXT
        )
        existing_inspection_type = (
            existing_document.inspection_type
            if existing_document is not None
            else ""
        )
        resolved_inspection_type = (
            existing_inspection_type
            if str(existing_inspection_type or "").strip()
            else infer_plant_inspection_type(str(product_line["description"]))
        )
        resolved_stock_code = (
            str(product_line.get("stock_code") or "").strip()
            or (
                existing_document.stock_code
                if existing_document is not None
                else ""
            )
        )
        existing_serial = (
            str(existing_document.serial or "")
            if existing_document is not None
            else ""
        )
        resolved_serial = (
            existing_serial
            if existing_serial
            else (
                resolved_stock_code
                if resolved_stock_code
                else PLANT_PENDING_SERIAL_TEXT
            )
        )

        plant_asset = PlantAssetDocument(
            doc_id=doc_id,
            site_name=site_name,
            created_at=(
                existing_document.created_at
                if existing_document is not None
                else created_at
            ),
            status=(
                DocumentStatus.ARCHIVED
                if existing_document is not None
                and existing_document.status == DocumentStatus.ARCHIVED
                else DocumentStatus.ACTIVE
            ),
            hire_num=hire_num,
            description=str(product_line["description"]),
            company=str(product_line["company"]),
            phone=str(product_line["phone"]),
            on_hire=product_line["on_hire"],
            off_hire=(
                existing_document.off_hire
                if existing_document is not None
                else None
            ),
            hired_by=resolved_hired_by,
            serial=resolved_serial,
            stock_code=resolved_stock_code,
            inspection_type=resolved_inspection_type,
            inspection=resolved_inspection,
            source_reference=order_ref,
            purchase_order=(
                existing_document.purchase_order
                if existing_document is not None and existing_document.purchase_order
                else purchase_order
            ),
        )
        repository.save(plant_asset)
        plant_assets.append(plant_asset)

    return plant_assets


def _apply_plant_collection_note_pdf(
    repository: DocumentRepository,
    pdf_path: Path,
    *,
    pdf_text: Optional[str] = None,
) -> List[PlantAssetDocument]:
    """Archive matching plant assets when a proof-of-collection PDF is filed."""

    resolved_pdf_text = pdf_text if pdf_text is not None else _safe_extract_pdf_text(pdf_path)
    parsed_collection = _parse_plant_collection_note(resolved_pdf_text)
    product_lines = list(parsed_collection.get("product_lines") or [])
    if not product_lines:
        return []

    project_setup = _load_workspace_project_setup()
    site_name = (
        project_setup.get("current_site_name")
        or _infer_default_site_name(repository)
    )
    active_assets = [
        document
        for document in repository.list_documents(
            document_type=PlantAssetDocument.document_type,
            site_name=site_name,
        )
        if isinstance(document, PlantAssetDocument)
        and document.status != DocumentStatus.ARCHIVED
    ]
    archived_assets: List[PlantAssetDocument] = []
    off_hire_date = (
        parsed_collection.get("off_hire")
        or parsed_collection.get("collection_date")
    )

    for product_line in product_lines:
        product_stock_code = str(product_line.get("stock_code") or "").strip()
        product_description_key = _plant_description_match_key(
            str(product_line.get("description") or "")
        )
        quantity_to_archive = max(int(product_line.get("actual_quantity") or 0), 1)
        returned_serials = list(dict.fromkeys(product_line.get("serials") or []))

        candidate_assets = sorted(
            (
                asset
                for asset in active_assets
                if (
                    (product_stock_code and asset.stock_code == product_stock_code)
                    or _plant_description_match_key(asset.description) == product_description_key
                )
            ),
            key=lambda asset: (
                asset.stock_code != product_stock_code,
                _plant_description_match_key(asset.description) != product_description_key,
                asset.on_hire,
                asset.hire_num,
            ),
        )
        if not candidate_assets:
            continue

        exact_quantity_asset = next(
            (
                asset
                for asset in candidate_assets
                if _extract_plant_description_quantity(asset.description) == quantity_to_archive
            ),
            None,
        )
        if exact_quantity_asset is not None:
            assets_to_archive = [exact_quantity_asset]
        else:
            assets_to_archive = []
            running_quantity = 0
            for candidate_asset in candidate_assets:
                candidate_quantity = _extract_plant_description_quantity(candidate_asset.description)
                if running_quantity + candidate_quantity > quantity_to_archive:
                    if running_quantity == 0:
                        assets_to_archive = []
                    break
                assets_to_archive.append(candidate_asset)
                running_quantity += candidate_quantity
                if running_quantity == quantity_to_archive:
                    break
            if running_quantity != quantity_to_archive:
                continue

        for asset_index, candidate_asset in enumerate(assets_to_archive):
            if candidate_asset not in active_assets:
                continue
            updated_serial = candidate_asset.serial
            asset_quantity = _extract_plant_description_quantity(candidate_asset.description)
            if returned_serials:
                if asset_quantity == len(returned_serials):
                    updated_serial = ", ".join(returned_serials)
                elif asset_quantity == 1 and asset_index < len(returned_serials):
                    updated_serial = returned_serials[asset_index]

            archived_asset = replace(
                candidate_asset,
                status=DocumentStatus.ARCHIVED,
                serial=updated_serial,
                off_hire=off_hire_date,
            )
            repository.save(archived_asset)
            archived_assets.append(archived_asset)
            active_assets.remove(candidate_asset)

    return archived_assets


def _sync_existing_plant_hire_pdfs(
    repository: DocumentRepository,
    plant_hire_destination: Path,
) -> None:
    """Backfill plant assets from already-filed plant hire PDFs."""

    pdf_paths = [
        pdf_path
        for pdf_path in sorted(plant_hire_destination.iterdir(), key=lambda path: path.name.lower())
        if pdf_path.is_file() and pdf_path.suffix.lower() == ".pdf"
    ]

    for pdf_path in pdf_paths:
        pdf_text = _safe_extract_pdf_text(pdf_path)
        if not _is_plant_hire_pdf(pdf_path, pdf_text):
            continue
        if _is_plant_collection_note_pdf(pdf_path, pdf_text):
            continue
        plant_assets = _upsert_plant_assets_from_pdf(
            repository,
            pdf_path,
            pdf_text=pdf_text,
        )
        repository.index_file(
            file_name=pdf_path.name,
            file_path=pdf_path,
            file_category=_classify_plant_hire_pdf(pdf_path, pdf_text),
            file_group=FileGroup.FILE_2,
            site_name=(
                plant_assets[0].site_name
                if plant_assets
                else _load_workspace_project_setup().get("current_site_name")
            ),
            related_doc_id=plant_assets[0].doc_id if len(plant_assets) == 1 else None,
        )

    for pdf_path in pdf_paths:
        if not pdf_path.is_file() or pdf_path.suffix.lower() != ".pdf":
            continue
        pdf_text = _safe_extract_pdf_text(pdf_path)
        if not _is_plant_hire_pdf(pdf_path, pdf_text):
            continue
        if not _is_plant_collection_note_pdf(pdf_path, pdf_text):
            continue
        plant_assets = _upsert_plant_assets_from_pdf(
            repository,
            pdf_path,
            pdf_text=pdf_text,
        )
        repository.index_file(
            file_name=pdf_path.name,
            file_path=pdf_path,
            file_category=_classify_plant_hire_pdf(pdf_path, pdf_text),
            file_group=FileGroup.FILE_2,
            site_name=(
                plant_assets[0].site_name
                if plant_assets
                else _load_workspace_project_setup().get("current_site_name")
            ),
            related_doc_id=plant_assets[0].doc_id if len(plant_assets) == 1 else None,
        )


def _discover_kpi_backup_json_paths() -> List[Path]:
    """Return unique KPI backup JSON paths from the project root and File 2."""

    candidate_paths = list(config.ATTENDANCE_DESTINATION.glob("site-kpi-backup*.json"))
    candidate_paths.extend(config.PROJECT_ROOT.glob("site-kpi-backup*.json"))

    unique_paths: Dict[Path, Path] = {}
    for candidate_path in candidate_paths:
        unique_paths[candidate_path.resolve()] = candidate_path.resolve()
    return sorted(unique_paths.values())


def _load_kpi_json_payload(json_path: Path) -> Any:
    """Load one KPI backup JSON payload from disk."""

    with json_path.open("r", encoding="utf-8") as file_handle:
        return json.load(file_handle)


def _extract_kpi_payload_site_name(payload: Any) -> Optional[str]:
    """Return the site name embedded in one KPI export when available."""

    if not isinstance(payload, Mapping):
        return None
    settings = payload.get("settings")
    if not isinstance(settings, Mapping):
        return None
    candidate = settings.get("siteName")
    if not isinstance(candidate, str):
        return None
    cleaned_candidate = candidate.strip()
    return cleaned_candidate or None


def _extract_kpi_active_rows(payload: Any) -> Iterable[Mapping[str, Any]]:
    """Yield KPI rows from all active extracted row arrays."""

    if not isinstance(payload, Mapping):
        return []

    extracted_rows = payload.get("extractedRows")
    if not isinstance(extracted_rows, Mapping):
        return []

    rows: List[Mapping[str, Any]] = []
    for candidate in extracted_rows.values():
        if not isinstance(candidate, list):
            continue
        for row in candidate:
            if isinstance(row, Mapping):
                rows.append(row)
    return rows


def _discover_docx_template_tags(template_path: Path) -> FrozenSet[str]:
    """Return placeholder tags found anywhere inside a DOCX package."""

    placeholders = set()
    with zipfile.ZipFile(template_path) as archive:
        for member_name in archive.namelist():
            if not member_name.startswith("word/") or not member_name.endswith(".xml"):
                continue
            xml_text = archive.read(member_name).decode("utf-8", errors="ignore")
            xml_text = re.sub(
                r"</w:t>\s*</w:r>\s*<w:r[^>]*>\s*<w:t[^>]*>",
                "",
                xml_text,
            )
            xml_text = re.sub(r"</w:t>\s*<w:t[^>]*>", "", xml_text)
            text_content = re.sub(r"<[^>]+>", " ", xml_text)
            placeholders.update(TEMPLATE_TAG_PATTERN.findall(text_content))
    return frozenset(placeholders)


@lru_cache(maxsize=8)
def _get_valid_template_tags_for_path(template_path_str: str) -> FrozenSet[str]:
    """Return valid placeholder tags for one specific weekly template path."""

    template_path = Path(template_path_str)
    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "weekly-site-check-template.docx"
        _repair_weekly_site_check_template(template_path, repaired_template_path)
        return _discover_docx_template_tags(repaired_template_path)


def get_valid_template_tags() -> FrozenSet[str]:
    """Return the valid placeholder tags present in the approved File 2 template."""

    template_path = TemplateRegistry.resolve_template_path(WeeklySiteCheck.document_type)
    return _get_valid_template_tags_for_path(str(template_path.resolve()))


get_valid_template_tags.cache_clear = _get_valid_template_tags_for_path.cache_clear  # type: ignore[attr-defined]


@lru_cache(maxsize=8)
def _get_weekly_site_check_row_definitions_for_path(
    template_path_str: str,
) -> Tuple[WeeklySiteCheckRowDefinition, ...]:
    """Return row definitions for one specific weekly template path."""

    template_path = Path(template_path_str)
    document = Document(template_path)
    if not document.tables:
        raise ValueError(f"Weekly site check template has no tables: {template_path}")

    template_table = document.tables[0]
    if len(template_table.rows) < 32:
        raise ValueError(
            "Weekly site check template must contain at least 32 rows including the header."
        )

    row_definitions: List[WeeklySiteCheckRowDefinition] = []
    for row_number in range(1, 32):
        row = template_table.rows[row_number]
        section = " ".join(row.cells[0].text.split())
        prompt = " ".join(row.cells[1].text.split())
        row_definitions.append(
            WeeklySiteCheckRowDefinition(
                row_number=row_number,
                section=section,
                prompt=prompt,
                frequency=get_weekly_site_check_frequency_for_row(row_number),
            )
        )
    return tuple(row_definitions)


def get_weekly_site_check_row_definitions() -> Tuple[WeeklySiteCheckRowDefinition, ...]:
    """Return the 31 row definitions extracted from the approved File 2 template."""

    template_path = TemplateRegistry.resolve_template_path(WeeklySiteCheck.document_type)
    return _get_weekly_site_check_row_definitions_for_path(str(template_path.resolve()))


get_weekly_site_check_row_definitions.cache_clear = (  # type: ignore[attr-defined]
    _get_weekly_site_check_row_definitions_for_path.cache_clear
)


def _repair_weekly_site_check_template(
    source_path: Path,
    destination_path: Path,
) -> None:
    """Copy the official template and repair known malformed tokens in the copy only."""

    destination_path.write_bytes(source_path.read_bytes())
    document = Document(destination_path)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "{{initials_wed}}}" in paragraph.text:
                        _replace_paragraph_text(
                            paragraph,
                            "{{initials_wed}}}",
                            "{{initials_wed}}",
                        )
                    if "{{initials_wed}" in paragraph.text:
                        _replace_paragraph_text(
                            paragraph,
                            "{{initials_wed}",
                            "{{initials_wed}}",
                        )
    document.save(destination_path)


def _build_patched_docxtpl_template(
    source_path: Path,
    destination_path: Path,
) -> None:
    """Copy a DOCX template and collapse split Jinja tags into contiguous XML."""

    from docxtpl import DocxTemplate

    template_document = DocxTemplate(str(source_path))
    with zipfile.ZipFile(source_path) as source_archive:
        with zipfile.ZipFile(destination_path, "w") as destination_archive:
            for archive_member in source_archive.infolist():
                member_bytes = source_archive.read(archive_member.filename)
                if (
                    archive_member.filename.startswith("word/")
                    and archive_member.filename.endswith(".xml")
                ):
                    try:
                        member_xml = member_bytes.decode("utf-8")
                    except UnicodeDecodeError:
                        destination_archive.writestr(archive_member, member_bytes)
                        continue

                    if any(token in member_xml for token in ("{{", "{%", "{#")):
                        member_xml = _rewrite_inline_table_row_loops(member_xml)
                        member_xml = _normalise_docxtpl_structural_tags(member_xml)
                        member_bytes = template_document.patch_xml(member_xml).encode("utf-8")

                destination_archive.writestr(archive_member, member_bytes)


def _rewrite_inline_table_row_loops(member_xml: str) -> str:
    """Rewrite one-row {% tr for ... %} ... {% tr endfor %} loops into normal Jinja."""

    row_pattern = re.compile(r"(<w:tr\b[^>]*>.*?</w:tr>)", flags=re.DOTALL)

    def rewrite_row(match: re.Match[str]) -> str:
        row_xml = match.group(1)
        start_match = re.search(
            r"\{%\s*tr\s+(for\s+[^%]+?)\s*%\}",
            row_xml,
            flags=re.IGNORECASE,
        )
        end_match = re.search(
            r"\{%\s*(?:tr\s+endfor|endtr)\s*%\}",
            row_xml,
            flags=re.IGNORECASE,
        )
        if start_match is None or end_match is None:
            return row_xml

        cleaned_row = re.sub(
            r"\{%\s*tr\s+for\s+[^%]+?\s*%\}\s*",
            "",
            row_xml,
            count=1,
            flags=re.IGNORECASE,
        )
        cleaned_row = re.sub(
            r"\s*\{%\s*(?:tr\s+endfor|endtr)\s*%\}",
            "",
            cleaned_row,
            count=1,
            flags=re.IGNORECASE,
        )
        return "{% " + start_match.group(1).strip() + " %}" + cleaned_row + "{% endfor %}"

    return row_pattern.sub(rewrite_row, member_xml)


def _normalise_docxtpl_structural_tags(member_xml: str) -> str:
    """Normalize docxtpl structural tags so patch_xml can recognize them."""

    return re.sub(
        r"(\{[%#]{1,2})\s*(tr|tc|p|r)\s+",
        r"\1\2 ",
        member_xml,
        flags=re.IGNORECASE,
    )


def _replace_paragraph_text(paragraph: object, old_text: str, new_text: str) -> None:
    """Replace text in a paragraph even when Word has split it into multiple runs."""

    runs = getattr(paragraph, "runs", [])
    if runs:
        original_text = "".join(run.text for run in runs)
    else:
        original_text = getattr(paragraph, "text", "")

    if not original_text or old_text not in original_text:
        return

    replaced_text = original_text.replace(old_text, new_text)
    if not runs:
        paragraph.add_run(replaced_text)
        return

    runs[0].text = replaced_text
    for run in runs[1:]:
        run.text = ""


def _stamp_weekly_site_check_header(
    output_path: Path,
    *,
    site_name: str,
    week_commencing: date,
) -> None:
    """Stamp the active project name into the merged top cell of the rendered checklist."""

    document = Document(output_path)
    if not document.tables:
        document.save(output_path)
        return

    header_cell = document.tables[0].cell(0, 1)
    header_cell.text = f"{site_name}\n{week_commencing.strftime('%d/%m/%Y')}"
    document.save(output_path)


def _build_export_ready_weekly_site_check(
    weekly_site_check: WeeklySiteCheck,
) -> WeeklySiteCheck:
    """Return a copy with signed blank days carried forward from the previous day."""

    export_ready_check = deepcopy(weekly_site_check)
    row_definitions = get_weekly_site_check_row_definitions()
    valid_template_tags = set(get_valid_template_tags())

    for day_index in range(1, len(SITE_CHECK_WEEKDAY_KEYS)):
        day_key = SITE_CHECK_WEEKDAY_KEYS[day_index]
        previous_day_key = SITE_CHECK_WEEKDAY_KEYS[day_index - 1]
        if not (
            export_ready_check.daily_initials.get(day_key, "").strip()
            or export_ready_check.daily_time_markers.get(day_key, "").strip()
        ):
            continue

        target_row_definitions = [
            row_definition
            for row_definition in row_definitions
            if row_definition.supports_day_key(day_key)
            and _weekly_site_check_template_tag(day_key, row_definition.row_number)
            in valid_template_tags
        ]
        if not target_row_definitions:
            continue

        current_values = [
            export_ready_check.get_row_state(row_definition.row_number).get_value(day_key)
            for row_definition in target_row_definitions
        ]
        if any(value is not None for value in current_values):
            continue

        previous_values = [
            export_ready_check.get_row_state(row_definition.row_number).get_value(
                previous_day_key
            )
            for row_definition in target_row_definitions
        ]
        if not any(value is not None for value in previous_values):
            continue

        for row_definition, previous_value in zip(
            target_row_definitions,
            previous_values,
        ):
            export_ready_check.get_row_state(row_definition.row_number).set_value(
                day_key,
                previous_value,
            )

    return export_ready_check


def _normalise_weekly_site_check_signoff_rows(output_path: Path) -> None:
    """Normalize the initials/time sign-off rows so all day cells render identically."""

    from docx.shared import Pt

    document = Document(output_path)
    if not document.tables:
        document.save(output_path)
        return

    table = document.tables[0]
    if len(table.rows) <= 34 or len(table.columns) <= 9:
        document.save(output_path)
        return

    for row_index in (33, 34):
        for column_index in range(3, 10):
            cell = table.cell(row_index, column_index)
            normalized_text = " ".join(cell.text.split())
            cell.text = normalized_text
            paragraph = cell.paragraphs[0]
            paragraph.alignment = 1
            if not paragraph.runs:
                paragraph.add_run(normalized_text)
            run = paragraph.runs[0]
            run.text = normalized_text
            run.font.name = "Arial Narrow"
            run.font.size = Pt(5)
            run.bold = False
            run.italic = False
            for extra_run in paragraph.runs[1:]:
                extra_run.text = ""

    document.save(output_path)


def _normalise_plant_register_table_rows(output_path: Path) -> None:
    """Normalize the rendered plant register rows so data text matches the template size."""

    from docx.shared import Pt

    document = Document(output_path)
    target_table = None
    for table in document.tables:
        if not table.rows or len(table.columns) < 9:
            continue
        header_values = [table.cell(0, index).text.strip() for index in range(3)]
        if header_values == ["Hire Number", "Plant description", "Hire Company"]:
            target_table = table
            break

    if target_table is None or len(target_table.rows) <= 1:
        document.save(output_path)
        return

    for row_index in range(1, len(target_table.rows)):
        for column_index in range(len(target_table.columns)):
            cell = target_table.cell(row_index, column_index)
            for paragraph in cell.paragraphs:
                if not paragraph.runs and paragraph.text:
                    paragraph.add_run(paragraph.text)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.bold = False
                    run.italic = False

    document.save(output_path)


def _normalise_permit_register_table_rows(output_path: Path) -> None:
    """Normalize the rendered permit register rows so the body matches the template scale."""

    from docx.shared import Pt

    document = Document(output_path)
    target_table = None
    for table in document.tables:
        if not table.rows or len(table.columns) < 7:
            continue
        header_values = [table.cell(0, index).text.strip() for index in range(3)]
        if header_values in (
            [
                "Permit Reference Number",
                "Date (dd/mm/yy)",
                "Permit\nType \n(HW, RW Etc)",
            ],
            ["Ref", "Date", "Type"],
        ):
            target_table = table
            break

    if target_table is None or len(target_table.rows) <= 1:
        document.save(output_path)
        return

    for row_index in range(1, len(target_table.rows)):
        for column_index in range(len(target_table.columns)):
            cell = target_table.cell(row_index, column_index)
            for paragraph in cell.paragraphs:
                if not paragraph.runs and paragraph.text:
                    paragraph.add_run(paragraph.text)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.bold = False
                    run.italic = False

    document.save(output_path)


def create_weekly_site_check_checklist_draft(
    repository: DocumentRepository,
    *,
    weekly_site_check: WeeklySiteCheck,
) -> GeneratedWeeklySiteCheckChecklist:
    """Render the approved tagged File 2 template into outputs/FILE_2_Checklists."""

    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the printable File 2 checklist."
        ) from exc

    repository.create_schema()
    output_directory = config.FILE_2_CHECKLIST_OUTPUT_DIR
    output_directory.mkdir(parents=True, exist_ok=True)
    export_ready_weekly_site_check = _build_export_ready_weekly_site_check(
        weekly_site_check
    )

    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "weekly-site-check-template.docx"
        patched_template_path = Path(temp_dir) / "weekly-site-check-template-patched.docx"
        template_path = TemplateRegistry.resolve_template_path(
            WeeklySiteCheck.document_type
        )
        _repair_weekly_site_check_template(template_path, repaired_template_path)
        _build_patched_docxtpl_template(repaired_template_path, patched_template_path)

        discovered_placeholders = _discover_docx_template_tags(patched_template_path)
        missing_placeholders = sorted(
            {"week_commencing", "checked_by"} - discovered_placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "Weekly site check template is missing required placeholders: "
                + ", ".join(missing_placeholders)
            )

        output_path = (
            output_directory
            / (
                f"{export_ready_weekly_site_check.week_commencing.strftime('%Y%m%d')}-"
                f"{export_ready_weekly_site_check.doc_id}.docx"
            )
        )
        document_template = DocxTemplate(str(patched_template_path))
        document_template.render(
            export_ready_weekly_site_check.to_template_context(),
            autoescape=False,
        )
        document_template.save(output_path)

    _stamp_weekly_site_check_header(
        output_path,
        site_name=export_ready_weekly_site_check.site_name,
        week_commencing=export_ready_weekly_site_check.week_commencing,
    )
    _normalise_weekly_site_check_signoff_rows(output_path)
    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="weekly_site_check_docx",
        file_group=FileGroup.FILE_2,
        site_name=export_ready_weekly_site_check.site_name,
        related_doc_id=export_ready_weekly_site_check.doc_id,
    )
    return GeneratedWeeklySiteCheckChecklist(
        weekly_site_check=export_ready_weekly_site_check,
        output_path=output_path,
    )


def create_site_check_checklist_draft(
    repository: DocumentRepository,
    *,
    register: SiteCheckRegister,
) -> GeneratedSiteCheckChecklist:
    """Render the approved File 2 checklist template into FILE_2_Output."""

    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the printable File 2 checklist."
        ) from exc

    repository.create_schema()
    output_directory = config.FILE_2_OUTPUT_DIR
    output_directory.mkdir(parents=True, exist_ok=True)

    template_manager = TemplateManager(register)
    template_manager.validate_template()

    output_path = (
        output_directory
        / f"{register.week_commencing.strftime('%Y%m%d')}-{register.doc_id}.docx"
    )
    document_template = DocxTemplate(str(template_manager.template_path))
    document_template.render(register.to_template_context(), autoescape=False)
    document_template.save(output_path)

    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="site_check_docx",
        file_group=FileGroup.FILE_2,
        site_name=register.site_name,
        related_doc_id=register.doc_id,
    )

    return GeneratedSiteCheckChecklist(
        register=register,
        output_path=output_path,
    )


def create_site_induction_document(
    repository: DocumentRepository,
    *,
    site_name: str,
    full_name: str,
    home_address: str,
    contact_number: str,
    company: str,
    occupation: str,
    emergency_contact: str,
    emergency_tel: str,
    medical: str,
    cscs_number: str,
    cscs_expiry: Optional[date] = None,
    asbestos_cert: bool = False,
    erect_scaffold: bool = False,
    cisrs_no: str = "",
    cisrs_expiry: Optional[date] = None,
    operate_plant: bool = False,
    cpcs_no: str = "",
    cpcs_expiry: Optional[date] = None,
    client_training_desc: str = "",
    client_training_date: Optional[date] = None,
    client_training_expiry: Optional[date] = None,
    first_aider: bool,
    fire_warden: bool,
    supervisor: bool,
    smsts: bool,
    competency_expiry_date: Optional[date] = None,
    competency_files: Optional[List[Mapping[str, Any]]] = None,
    signature_image_data: Any,
    linked_rams_doc_id: str = "",
) -> GeneratedInductionDocument:
    """Render one signed induction document and persist the logged induction record."""

    try:
        from docxtpl import DocxTemplate, InlineImage
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the completed induction document."
        ) from exc
    try:
        from docx.shared import Mm
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to size the induction signature image."
        ) from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError(
            "jinja2 is required to generate the completed induction document."
        ) from exc

    created_at = datetime.now().replace(second=0, microsecond=0)
    cleaned_full_name = full_name.strip()
    cleaned_company = company.strip()
    cleaned_home_address = home_address.strip()
    cleaned_contact_number = contact_number.strip()
    if not cleaned_full_name:
        raise ValidationError("Full Name is required.")
    if not cleaned_company:
        raise ValidationError("Company is required.")
    if not cleaned_home_address:
        raise ValidationError("Home Address is required.")
    if not cleaned_contact_number:
        raise ValidationError("Contact Number is required.")
    manual_handling_uploaded = any(
        str(uploaded_file.get("label", "") or "").strip().casefold()
        == "manual handling certificate"
        for uploaded_file in (competency_files or [])
    )
    if not manual_handling_uploaded:
        raise ValidationError("Manual Handling Certificate upload is required.")

    repository.create_schema()
    config.FILE_3_COMPETENCY_CARDS_DIR.mkdir(parents=True, exist_ok=True)
    config.FILE_3_SIGNATURES_DIR.mkdir(parents=True, exist_ok=True)
    config.FILE_3_COMPLETED_INDUCTIONS_DIR.mkdir(parents=True, exist_ok=True)

    saved_competency_card_paths = _save_induction_competency_cards(
        competency_files=competency_files or [],
        full_name=cleaned_full_name,
        company=cleaned_company,
        created_at=created_at,
    )
    signature_path = _save_induction_signature_image(
        signature_image_data=signature_image_data,
        full_name=cleaned_full_name,
        company=cleaned_company,
        created_at=created_at,
    )
    induction_document = InductionDocument(
        doc_id=_build_induction_doc_id(created_at, cleaned_full_name),
        site_name=site_name.strip() or DEFAULT_SITE_NAME,
        created_at=created_at,
        status=DocumentStatus.ACTIVE,
        contractor_name=cleaned_company,
        individual_name=cleaned_full_name,
        linked_rams_doc_id=linked_rams_doc_id.strip(),
        home_address=cleaned_home_address,
        contact_number=cleaned_contact_number,
        occupation=occupation.strip(),
        emergency_contact=emergency_contact.strip(),
        emergency_tel=emergency_tel.strip(),
        medical=medical.strip(),
        cscs_number=cscs_number.strip(),
        cscs_expiry=cscs_expiry,
        asbestos_cert=bool(asbestos_cert),
        erect_scaffold=bool(erect_scaffold),
        cisrs_no=cisrs_no.strip(),
        cisrs_expiry=cisrs_expiry,
        operate_plant=bool(operate_plant),
        cpcs_no=cpcs_no.strip(),
        cpcs_expiry=cpcs_expiry,
        client_training_desc=client_training_desc.strip(),
        client_training_date=client_training_date,
        client_training_expiry=client_training_expiry,
        first_aider=bool(first_aider),
        fire_warden=bool(fire_warden),
        supervisor=bool(supervisor),
        smsts=bool(smsts),
        competency_expiry_date=competency_expiry_date,
        competency_card_paths=",".join(str(path) for path in saved_competency_card_paths),
        signature_image_path=str(signature_path),
    )

    template_path = TemplateRegistry.resolve_template_path("site_induction")
    output_path = _build_available_destination(
        Path(
            "Induction - "
            f"{_sanitize_filename_fragment(cleaned_full_name)} - "
            f"{_sanitize_filename_fragment(cleaned_company)} - "
            f"{created_at:%Y-%m-%d}.docx"
        ),
        config.FILE_3_COMPLETED_INDUCTIONS_DIR,
    )

    _render_site_induction_docx(
        induction_document,
        output_path=output_path,
        signature_path=signature_path,
    )

    induction_document.completed_document_path = str(output_path)
    repository.save(induction_document)
    repository.index_file(
        file_name=signature_path.name,
        file_path=signature_path,
        file_category="induction_signature_png",
        file_group=FileGroup.FILE_3,
        site_name=induction_document.site_name,
        related_doc_id=induction_document.doc_id,
    )
    for competency_card_path in saved_competency_card_paths:
        repository.index_file(
            file_name=competency_card_path.name,
            file_path=competency_card_path,
            file_category="competency_card_upload",
            file_group=FileGroup.FILE_3,
            site_name=induction_document.site_name,
            related_doc_id=induction_document.doc_id,
        )
    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="completed_induction_docx",
        file_group=FileGroup.FILE_3,
        site_name=induction_document.site_name,
        related_doc_id=induction_document.doc_id,
    )
    return GeneratedInductionDocument(
        induction_document=induction_document,
        output_path=output_path,
        signature_path=signature_path,
    )


def _render_site_induction_docx(
    induction_document: InductionDocument,
    *,
    output_path: Path,
    signature_path: Path,
) -> None:
    """Render one completed site induction DOCX from the official template."""

    try:
        from docxtpl import DocxTemplate, InlineImage
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the completed induction document."
        ) from exc
    try:
        from docx.shared import Mm
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to size the induction signature image."
        ) from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError(
            "jinja2 is required to generate the completed induction document."
        ) from exc

    template_path = TemplateRegistry.resolve_template_path("site_induction")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "site-induction-template.docx"
        _build_patched_docxtpl_template(template_path, repaired_template_path)

        discovered_placeholders = _discover_docx_template_tags(repaired_template_path)
        missing_placeholders = sorted(
            induction_document.required_template_placeholders - discovered_placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "Site induction template is missing required placeholders: "
                + ", ".join(missing_placeholders)
            )

        document_template = DocxTemplate(str(repaired_template_path))
        template_context = induction_document.to_template_context()
        template_context.update(
            {
                "date": induction_document.created_at.strftime("%d/%m/%Y"),
                "induction_date": induction_document.created_at.strftime("%d/%m/%Y"),
                "today_date": induction_document.created_at.strftime("%d/%m/%Y"),
                "home_address": induction_document.home_address,
                "company": induction_document.contractor_name,
                "cscs_expiry": (
                    induction_document.cscs_expiry.strftime("%d/%m/%Y")
                    if induction_document.cscs_expiry is not None
                    else ""
                ),
                "asbestos_cert": "✔" if induction_document.asbestos_cert else "",
                "erect_scaffold": "Yes" if induction_document.erect_scaffold else "No",
                "cisrs_no": induction_document.cisrs_no,
                "cisrs_expiry": (
                    induction_document.cisrs_expiry.strftime("%d/%m/%Y")
                    if induction_document.cisrs_expiry is not None
                    else ""
                ),
                "operate_plant": "Yes" if induction_document.operate_plant else "No",
                "cpcs_no": induction_document.cpcs_no,
                "cpcs_expiry": (
                    induction_document.cpcs_expiry.strftime("%d/%m/%Y")
                    if induction_document.cpcs_expiry is not None
                    else ""
                ),
                "client_training_desc": induction_document.client_training_desc,
                "client_training_date": (
                    induction_document.client_training_date.strftime("%d/%m/%Y")
                    if induction_document.client_training_date is not None
                    else ""
                ),
                "client_training_expiry": (
                    induction_document.client_training_expiry.strftime("%d/%m/%Y")
                    if induction_document.client_training_expiry is not None
                    else ""
                ),
                "competency_expiry_date": (
                    induction_document.competency_expiry_date.strftime("%d/%m/%Y")
                    if induction_document.competency_expiry_date is not None
                    else ""
                ),
                "inductor_name_date": "Ceri Edwards",
                "inductor_title": "Site Manager",
                "signature_image": InlineImage(
                    document_template,
                    str(signature_path),
                    height=Mm(20),
                ),
            }
        )
        document_template.render(
            template_context,
            jinja_env=Environment(autoescape=False),
            autoescape=False,
        )
        document_template.save(output_path)


def update_site_induction_document(
    repository: DocumentRepository,
    *,
    induction_doc_id: str,
    full_name: str,
    home_address: str,
    contact_number: str,
    company: str,
    occupation: str,
    emergency_contact: str,
    emergency_tel: str,
    medical: str,
    cscs_number: str,
    cscs_expiry: Optional[date] = None,
    asbestos_cert: bool = False,
    erect_scaffold: bool = False,
    cisrs_no: str = "",
    cisrs_expiry: Optional[date] = None,
    operate_plant: bool = False,
    cpcs_no: str = "",
    cpcs_expiry: Optional[date] = None,
    client_training_desc: str = "",
    client_training_date: Optional[date] = None,
    client_training_expiry: Optional[date] = None,
    first_aider: bool = False,
    fire_warden: bool = False,
    supervisor: bool = False,
    smsts: bool = False,
    competency_expiry_date: Optional[date] = None,
) -> GeneratedInductionDocument:
    """Update one saved induction record and regenerate its completed DOCX."""

    repository.create_schema()
    existing_document = repository.get(induction_doc_id)
    if not isinstance(existing_document, InductionDocument):
        raise ValidationError("Selected induction record is not a UHSF16.01 entry.")

    updated_induction_document = replace(
        existing_document,
        contractor_name=company.strip(),
        individual_name=full_name.strip(),
        home_address=home_address.strip(),
        contact_number=contact_number.strip(),
        occupation=occupation.strip(),
        emergency_contact=emergency_contact.strip(),
        emergency_tel=emergency_tel.strip(),
        medical=medical.strip(),
        cscs_number=cscs_number.strip(),
        cscs_expiry=cscs_expiry,
        asbestos_cert=bool(asbestos_cert),
        erect_scaffold=bool(erect_scaffold),
        cisrs_no=cisrs_no.strip(),
        cisrs_expiry=cisrs_expiry,
        operate_plant=bool(operate_plant),
        cpcs_no=cpcs_no.strip(),
        cpcs_expiry=cpcs_expiry,
        client_training_desc=client_training_desc.strip(),
        client_training_date=client_training_date,
        client_training_expiry=client_training_expiry,
        first_aider=bool(first_aider),
        fire_warden=bool(fire_warden),
        supervisor=bool(supervisor),
        smsts=bool(smsts),
        competency_expiry_date=competency_expiry_date,
    )

    if not updated_induction_document.individual_name:
        raise ValidationError("Full Name is required.")
    if not updated_induction_document.contractor_name:
        raise ValidationError("Company is required.")
    if not updated_induction_document.home_address:
        raise ValidationError("Home Address is required.")
    if not updated_induction_document.contact_number:
        raise ValidationError("Contact Number is required.")

    signature_path = Path(updated_induction_document.signature_image_path).expanduser()
    if not updated_induction_document.signature_image_path or not signature_path.exists():
        raise ValidationError(
            "The saved induction signature file is missing, so the induction document cannot be regenerated."
        )

    output_path = (
        Path(updated_induction_document.completed_document_path).expanduser()
        if updated_induction_document.completed_document_path
        else _build_available_destination(
            Path(
                "Induction - "
                f"{_sanitize_filename_fragment(updated_induction_document.individual_name)} - "
                f"{_sanitize_filename_fragment(updated_induction_document.contractor_name)} - "
                f"{updated_induction_document.created_at:%Y-%m-%d}.docx"
            ),
            config.FILE_3_COMPLETED_INDUCTIONS_DIR,
        )
    )
    _render_site_induction_docx(
        updated_induction_document,
        output_path=output_path,
        signature_path=signature_path,
    )

    updated_induction_document.completed_document_path = str(output_path)
    repository.save(updated_induction_document)
    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="completed_induction_docx",
        file_group=FileGroup.FILE_3,
        site_name=updated_induction_document.site_name,
        related_doc_id=updated_induction_document.doc_id,
    )
    repository.index_file(
        file_name=signature_path.name,
        file_path=signature_path,
        file_category="induction_signature_png",
        file_group=FileGroup.FILE_3,
        site_name=updated_induction_document.site_name,
        related_doc_id=updated_induction_document.doc_id,
    )
    for path_text in updated_induction_document.competency_card_paths.split(","):
        if not path_text.strip():
            continue
        competency_path = Path(path_text.strip()).expanduser()
        if not competency_path.exists():
            continue
        repository.index_file(
            file_name=competency_path.name,
            file_path=competency_path,
            file_category="competency_card_upload",
            file_group=FileGroup.FILE_3,
            site_name=updated_induction_document.site_name,
            related_doc_id=updated_induction_document.doc_id,
        )
    return GeneratedInductionDocument(
        induction_document=updated_induction_document,
        output_path=output_path,
        signature_path=signature_path,
    )


def add_site_induction_evidence_files(
    repository: DocumentRepository,
    *,
    induction_doc_id: str,
    competency_files: List[Mapping[str, Any]],
) -> InductionDocument:
    """Append extra evidence files onto one saved induction record."""

    repository.create_schema()
    existing_document = repository.get(induction_doc_id)
    if not isinstance(existing_document, InductionDocument):
        raise ValidationError("Selected induction record is not a UHSF16.01 entry.")

    cleaned_competency_files = [
        uploaded_file
        for uploaded_file in competency_files
        if str(uploaded_file.get("name", "") or "").strip()
        and uploaded_file.get("bytes", b"")
    ]
    if not cleaned_competency_files:
        raise ValidationError("Select at least one evidence file to add.")

    created_at = datetime.now().replace(second=0, microsecond=0)
    saved_competency_card_paths = _save_induction_competency_cards(
        competency_files=cleaned_competency_files,
        full_name=existing_document.individual_name,
        company=existing_document.contractor_name,
        created_at=created_at,
    )
    if not saved_competency_card_paths:
        raise ValidationError("No valid evidence files were uploaded.")

    existing_paths = [
        path_text.strip()
        for path_text in existing_document.competency_card_paths.split(",")
        if path_text.strip()
    ]
    updated_induction_document = replace(
        existing_document,
        competency_card_paths=",".join(
            [*existing_paths, *[str(path) for path in saved_competency_card_paths]]
        ),
    )
    repository.save(updated_induction_document)
    for competency_card_path in saved_competency_card_paths:
        repository.index_file(
            file_name=competency_card_path.name,
            file_path=competency_card_path,
            file_category="competency_card_upload",
            file_group=FileGroup.FILE_3,
            site_name=updated_induction_document.site_name,
            related_doc_id=updated_induction_document.doc_id,
        )
    return updated_induction_document


def list_daily_attendance_entries(
    repository: DocumentRepository,
    *,
    site_name: str,
    on_date: Optional[date] = None,
    active_only: bool = False,
) -> List[DailyAttendanceEntryDocument]:
    """Return live UHSF16.09 attendance entries for one site."""

    repository.create_schema()
    resolved_date = on_date
    entries = [
        document
        for document in repository.list_documents(
            document_type=DailyAttendanceEntryDocument.document_type,
            site_name=site_name.strip() or DEFAULT_SITE_NAME,
        )
        if isinstance(document, DailyAttendanceEntryDocument)
    ]
    if resolved_date is not None:
        entries = [
            entry for entry in entries if entry.attendance_date == resolved_date
        ]
    if active_only:
        entries = [entry for entry in entries if entry.is_on_site]
    return sorted(entries, key=lambda entry: entry.time_in, reverse=True)


def build_live_site_broadcast_contacts(
    repository: DocumentRepository,
    *,
    site_name: str,
    on_date: Optional[date] = None,
) -> List[SiteBroadcastContact]:
    """Return the current on-site mobile contacts sourced from live attendance."""

    active_attendance_entries = list_daily_attendance_entries(
        repository,
        site_name=site_name,
        on_date=on_date or date.today(),
        active_only=True,
    )
    contacts: List[SiteBroadcastContact] = []
    seen_mobile_numbers: set[str] = set()
    for attendance_entry in sorted(
        active_attendance_entries,
        key=lambda attendance_entry: (
            attendance_entry.individual_name.casefold(),
            attendance_entry.contractor_name.casefold(),
            attendance_entry.time_in,
        ),
    ):
        if not attendance_entry.linked_induction_doc_id:
            continue
        try:
            linked_document = repository.get(attendance_entry.linked_induction_doc_id)
        except (DocumentNotFoundError, ValueError):
            continue
        if not isinstance(linked_document, InductionDocument):
            continue

        mobile_number = _normalise_uk_mobile_number(linked_document.contact_number)
        if not mobile_number or mobile_number in seen_mobile_numbers:
            continue

        contacts.append(
            SiteBroadcastContact(
                individual_name=attendance_entry.individual_name,
                contractor_name=attendance_entry.contractor_name,
                mobile_number=mobile_number,
                vehicle_registration=attendance_entry.vehicle_registration,
                linked_induction_doc_id=attendance_entry.linked_induction_doc_id,
            )
        )
        seen_mobile_numbers.add(mobile_number)
    return contacts


def build_pending_toolbox_talk_contacts(
    repository: DocumentRepository,
    *,
    site_name: str,
    topic: str,
    on_date: Optional[date] = None,
) -> List[SiteBroadcastContact]:
    """Return active live contacts who still need to sign one toolbox talk."""

    live_contacts = build_live_site_broadcast_contacts(
        repository,
        site_name=site_name,
        on_date=on_date or date.today(),
    )
    if not topic.strip():
        return live_contacts

    completions = list_toolbox_talk_completions(
        repository,
        site_name=site_name,
        topic=topic,
    )
    signed_induction_ids = {
        completion.linked_induction_doc_id
        for completion in completions
        if completion.linked_induction_doc_id
    }
    signed_people = {
        (
            completion.individual_name.casefold(),
            completion.contractor_name.casefold(),
        )
        for completion in completions
    }
    return [
        contact
        for contact in live_contacts
        if (
            contact.linked_induction_doc_id not in signed_induction_ids
            and (
                contact.individual_name.casefold(),
                contact.contractor_name.casefold(),
            )
            not in signed_people
        )
    ]


def _deduplicate_mobile_numbers(mobile_numbers: Iterable[str]) -> List[str]:
    """Return cleaned broadcast numbers without duplicates."""

    resolved_numbers: List[str] = []
    seen_mobile_numbers: set[str] = set()
    for mobile_number in mobile_numbers:
        cleaned_mobile_number = str(mobile_number or "").strip()
        if not cleaned_mobile_number or cleaned_mobile_number in seen_mobile_numbers:
            continue
        resolved_numbers.append(cleaned_mobile_number)
        seen_mobile_numbers.add(cleaned_mobile_number)
    return resolved_numbers


def _build_sms_link(
    mobile_numbers: Iterable[str],
    *,
    message: str = "",
) -> str:
    """Return one sms: deep link for a resolved recipient chunk."""

    resolved_numbers = _deduplicate_mobile_numbers(mobile_numbers)
    if not resolved_numbers:
        return ""

    sms_link = f"sms:{','.join(resolved_numbers)}"
    trimmed_message = message.strip()
    if trimmed_message:
        sms_link += f"&body={quote(trimmed_message)}"
    return sms_link


def build_site_alert_sms_links(
    mobile_numbers: Iterable[str],
    *,
    message: str = "",
    max_recipients_per_chunk: int = 24,
    max_url_length: int = 1800,
) -> List[str]:
    """Return one or more Messages-ready sms: links for the supplied recipients."""

    if max_recipients_per_chunk < 1:
        raise ValueError("max_recipients_per_chunk must be at least 1.")
    if max_url_length < 32:
        raise ValueError("max_url_length must be at least 32.")

    resolved_numbers = _deduplicate_mobile_numbers(mobile_numbers)
    if not resolved_numbers:
        return []

    trimmed_message = message.strip()
    draft_links: List[str] = []
    current_chunk: List[str] = []
    for mobile_number in resolved_numbers:
        candidate_chunk = current_chunk + [mobile_number]
        candidate_link = _build_sms_link(candidate_chunk, message=trimmed_message)
        if (
            current_chunk
            and (
                len(candidate_chunk) > max_recipients_per_chunk
                or len(candidate_link) > max_url_length
            )
        ):
            draft_links.append(_build_sms_link(current_chunk, message=trimmed_message))
            current_chunk = [mobile_number]
            continue
        current_chunk = candidate_chunk

    if current_chunk:
        draft_links.append(_build_sms_link(current_chunk, message=trimmed_message))
    return draft_links


def build_site_alert_sms_link(
    mobile_numbers: Iterable[str],
    *,
    message: str = "",
) -> str:
    """Return a clickable sms: link for the supplied recipients."""

    sms_links = build_site_alert_sms_links(
        mobile_numbers,
        message=message,
    )
    return sms_links[0] if sms_links else ""


def build_toolbox_talk_sms_message(
    topic: str,
    tbt_link: str,
) -> str:
    """Return the default delivery message for one remote toolbox talk."""

    resolved_topic = topic.strip()
    resolved_link = tbt_link.strip()
    return (
        f"Toolbox Talk: {resolved_topic}. "
        "Please click this link to read the document and sign the register: "
        f"{resolved_link}"
    ).strip()


def launch_messages_sms_broadcast(
    mobile_numbers: Iterable[str],
    *,
    message: str = "",
    max_recipients_per_chunk: int = 24,
    max_url_length: int = 1800,
) -> MessagesDraftLaunchResult:
    """Open one ready-to-send Messages draft per recipient.

    macOS Messages has proved unreliable when opened with one ``sms:`` deep link
    containing multiple recipients. In practice that can silently drop all but
    the first number. For gate-critical broadcasts we prefer a slightly noisier
    but deterministic path: open one draft per resolved recipient.
    """

    resolved_numbers = _deduplicate_mobile_numbers(mobile_numbers)
    if not resolved_numbers:
        return MessagesDraftLaunchResult(
            draft_links=[],
            recipient_count=0,
            chunk_count=0,
            launched_successfully=False,
            error_message="No valid mobile numbers were supplied.",
        )

    # Keep one-recipient drafts for launch reliability inside Messages.
    draft_links = [
        _build_sms_link([mobile_number], message=message)
        for mobile_number in resolved_numbers
    ]
    if not draft_links:
        return MessagesDraftLaunchResult(
            draft_links=[],
            recipient_count=0,
            chunk_count=0,
            launched_successfully=False,
            error_message="Unable to build any SMS draft links.",
        )

    app_lookup = subprocess.run(
        ["open", "-Ra", "Messages"],
        capture_output=True,
        text=True,
        check=False,
    )
    if app_lookup.returncode != 0:
        return MessagesDraftLaunchResult(
            draft_links=draft_links,
            recipient_count=len(resolved_numbers),
            chunk_count=len(draft_links),
            launched_successfully=False,
            error_message=(
                app_lookup.stderr.strip()
                or app_lookup.stdout.strip()
                or "Messages is not available on this Mac."
            ),
        )

    if len(resolved_numbers) > 1:
        grouped_launch_successful, grouped_launch_error = (
            _launch_messages_group_draft_via_gui_automation(
                resolved_numbers,
                message=message,
            )
        )
        if grouped_launch_successful:
            return MessagesDraftLaunchResult(
                draft_links=[_build_sms_link(resolved_numbers, message=message)],
                recipient_count=len(resolved_numbers),
                chunk_count=1,
                launched_successfully=True,
                error_message="Opened one grouped Messages draft for the live audience.",
            )

    launch_error_message = ""
    launched_successfully = True
    for draft_link_index, draft_link in enumerate(draft_links):
        launch_result = subprocess.run(
            ["open", "-a", "Messages", draft_link],
            capture_output=True,
            text=True,
            check=False,
        )
        if launch_result.returncode != 0:
            launched_successfully = False
            launch_error_message = (
                launch_result.stderr.strip()
                or launch_result.stdout.strip()
                or "The Messages draft could not be opened."
            )
            break
        if draft_link_index < len(draft_links) - 1:
            time_module.sleep(0.25)

    if launched_successfully and len(resolved_numbers) > 1 and grouped_launch_error:
        launch_error_message = (
            "Messages group compose was unavailable on this Mac, so the app opened "
            "one draft per recipient instead. "
            f"{grouped_launch_error}"
        )

    return MessagesDraftLaunchResult(
        draft_links=draft_links,
        recipient_count=len(resolved_numbers),
        chunk_count=len(draft_links),
        launched_successfully=launched_successfully,
        error_message=launch_error_message,
    )


def log_broadcast_dispatch(
    repository: DocumentRepository,
    *,
    site_name: str,
    dispatch_kind: str,
    audience_label: str,
    subject: str,
    message_body: str,
    recipient_numbers: Iterable[str],
    recipient_names: Iterable[str],
    launch_result: MessagesDraftLaunchResult,
    topic: str = "",
) -> LoggedBroadcastDispatch:
    """Persist one broadcast/TBT delivery batch to the workspace database."""

    repository.create_schema()
    dispatched_at = datetime.now()
    dispatch_document = BroadcastDispatchDocument(
        doc_id=(
            f"DISPATCH-{dispatched_at:%Y%m%d%H%M%S%f}-"
            f"{_slugify_identifier(subject)}"
        ),
        site_name=site_name.strip() or DEFAULT_SITE_NAME,
        created_at=dispatched_at,
        status=DocumentStatus.ACTIVE,
        dispatch_kind=dispatch_kind.strip().lower(),
        channel="messages",
        audience_label=audience_label.strip() or "Live fire roll",
        subject=subject.strip(),
        message_body=message_body,
        recipient_numbers=_deduplicate_mobile_numbers(recipient_numbers),
        recipient_names=[
            str(name).strip()
            for name in recipient_names
            if str(name).strip()
        ],
        topic=topic.strip(),
        dispatched_at=dispatched_at,
        launch_mode="messages_draft",
        launched_successfully=launch_result.launched_successfully,
        chunk_count=launch_result.chunk_count,
    )
    repository.save(dispatch_document)
    return LoggedBroadcastDispatch(
        dispatch_document=dispatch_document,
        launch_result=launch_result,
    )


def list_broadcast_dispatches(
    repository: DocumentRepository,
    *,
    site_name: str,
    topic: str = "",
    dispatch_kind: str = "",
) -> List[BroadcastDispatchDocument]:
    """Return prior broadcast/TBT deliveries for one site."""

    repository.create_schema()
    dispatches = [
        document
        for document in repository.list_documents(
            document_type=BroadcastDispatchDocument.document_type,
            site_name=site_name.strip() or DEFAULT_SITE_NAME,
        )
        if isinstance(document, BroadcastDispatchDocument)
    ]
    if topic.strip():
        lowered_topic = topic.strip().casefold()
        dispatches = [
            dispatch
            for dispatch in dispatches
            if dispatch.topic.casefold() == lowered_topic
        ]
    if dispatch_kind.strip():
        lowered_dispatch_kind = dispatch_kind.strip().casefold()
        dispatches = [
            dispatch
            for dispatch in dispatches
            if dispatch.dispatch_kind.casefold() == lowered_dispatch_kind
        ]
    return sorted(
        dispatches,
        key=lambda dispatch: (dispatch.dispatched_at, dispatch.doc_id),
        reverse=True,
    )


def calculate_haversine_distance_meters(
    latitude_a: float,
    longitude_a: float,
    latitude_b: float,
    longitude_b: float,
) -> float:
    """Return the great-circle distance between two GPS coordinates in meters."""

    earth_radius_meters = 6_371_000.0
    latitude_a_radians = math.radians(float(latitude_a))
    latitude_b_radians = math.radians(float(latitude_b))
    latitude_delta_radians = math.radians(float(latitude_b) - float(latitude_a))
    longitude_delta_radians = math.radians(float(longitude_b) - float(longitude_a))

    haversine_value = (
        math.sin(latitude_delta_radians / 2) ** 2
        + math.cos(latitude_a_radians)
        * math.cos(latitude_b_radians)
        * math.sin(longitude_delta_radians / 2) ** 2
    )
    return 2 * earth_radius_meters * math.atan2(
        math.sqrt(haversine_value),
        math.sqrt(1 - haversine_value),
    )


def list_toolbox_talk_completions(
    repository: DocumentRepository,
    *,
    site_name: str,
    topic: str = "",
    on_date: Optional[date] = None,
) -> List[ToolboxTalkCompletionDocument]:
    """Return stored UHSF16.2 completions for one site and optional topic."""

    repository.create_schema()
    completions = [
        document
        for document in repository.list_documents(
            document_type=ToolboxTalkCompletionDocument.document_type,
            site_name=site_name.strip() or DEFAULT_SITE_NAME,
        )
        if isinstance(document, ToolboxTalkCompletionDocument)
    ]
    if topic.strip():
        lowered_topic = topic.strip().casefold()
        completions = [
            completion
            for completion in completions
            if completion.topic.casefold() == lowered_topic
        ]
    if on_date is not None:
        completions = [
            completion
            for completion in completions
            if completion.completed_at.date() == on_date
        ]
    return sorted(
        completions,
        key=lambda completion: (completion.completed_at, completion.doc_id),
        reverse=True,
    )


def list_toolbox_talk_documents(
    repository: DocumentRepository,
    *,
    site_name: str,
    topic: str = "",
) -> List[ToolboxTalkDocument]:
    """Return uploaded toolbox talk source documents for one site and optional topic."""

    repository.create_schema()
    documents = [
        document
        for document in repository.list_documents(
            document_type=ToolboxTalkDocument.document_type,
            site_name=site_name.strip() or DEFAULT_SITE_NAME,
        )
        if isinstance(document, ToolboxTalkDocument)
    ]
    if topic.strip():
        lowered_topic = topic.strip().casefold()
        documents = [
            document
            for document in documents
            if document.topic.casefold() == lowered_topic
        ]
    return sorted(
        documents,
        key=lambda document: (document.created_at, document.doc_id),
        reverse=True,
    )


def get_latest_toolbox_talk_document(
    repository: DocumentRepository,
    *,
    site_name: str,
    topic: str,
) -> Optional[ToolboxTalkDocument]:
    """Return the most recent uploaded toolbox talk source document for one topic."""

    documents = list_toolbox_talk_documents(
        repository,
        site_name=site_name,
        topic=topic,
    )
    return documents[0] if documents else None


def save_toolbox_talk_document(
    repository: DocumentRepository,
    *,
    site_name: str,
    topic: str,
    uploaded_file_name: str,
    uploaded_file_bytes: bytes,
) -> SavedToolboxTalkDocument:
    """Persist one uploaded toolbox talk source document inside File 2."""

    repository.create_schema()
    resolved_topic = topic.strip()
    if not resolved_topic:
        raise ValidationError("Toolbox talk topic must not be blank.")
    cleaned_file_name = Path(uploaded_file_name or "").name
    if not cleaned_file_name:
        raise ValidationError("Upload a toolbox talk document before creating the link.")
    if not uploaded_file_bytes:
        raise ValidationError("The uploaded toolbox talk document is empty.")

    created_at = datetime.now()
    suffix = Path(cleaned_file_name).suffix or ".bin"
    stored_path = _build_available_destination(
        Path(
            f"{created_at:%Y%m%d-%H%M%S}-"
            f"{_sanitize_filename_fragment(resolved_topic)}"
            f"{suffix}"
        ),
        config.FILE_2_TBT_ACTIVE_DOCS_DIR,
    )
    config.FILE_2_TBT_ACTIVE_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    stored_path.write_bytes(uploaded_file_bytes)

    toolbox_talk_document = ToolboxTalkDocument(
        doc_id=(
            f"TBTDOC-{created_at:%Y%m%d%H%M%S%f}-"
            f"{_slugify_identifier(resolved_topic)}"
        ),
        site_name=site_name.strip() or DEFAULT_SITE_NAME,
        created_at=created_at,
        status=DocumentStatus.ACTIVE,
        topic=resolved_topic,
        original_file_name=cleaned_file_name,
        stored_file_path=str(stored_path),
    )
    repository.save(toolbox_talk_document)
    repository.index_file(
        file_name=stored_path.name,
        file_path=stored_path,
        file_category="toolbox_talk_source_doc",
        file_group=FileGroup.FILE_2,
        site_name=toolbox_talk_document.site_name,
        related_doc_id=toolbox_talk_document.doc_id,
    )
    return SavedToolboxTalkDocument(
        toolbox_talk_document=toolbox_talk_document,
        stored_path=stored_path,
    )


def log_toolbox_talk_completion(
    repository: DocumentRepository,
    *,
    site_name: str,
    topic: str,
    attendance_entry: DailyAttendanceEntryDocument,
    signature_image_data: Any,
    document_read_confirmed: bool,
) -> LoggedToolboxTalkCompletion:
    """Persist one remote toolbox talk sign-off from the active fire roll."""

    repository.create_schema()
    resolved_topic = topic.strip()
    if not resolved_topic:
        raise ValidationError("Toolbox talk topic must not be blank.")
    if not document_read_confirmed:
        raise ValidationError(
            "Confirm that you have read and understood the toolbox talk before signing."
        )
    if not attendance_entry.is_on_site:
        raise ValidationError("Only operatives currently on site can sign this toolbox talk.")

    existing_completions = list_toolbox_talk_completions(
        repository,
        site_name=site_name.strip() or DEFAULT_SITE_NAME,
        topic=resolved_topic,
        on_date=date.today(),
    )
    for existing_completion in existing_completions:
        if (
            existing_completion.linked_induction_doc_id
            and existing_completion.linked_induction_doc_id
            == attendance_entry.linked_induction_doc_id
        ) or (
            existing_completion.individual_name.casefold()
            == attendance_entry.individual_name.casefold()
            and existing_completion.contractor_name.casefold()
            == attendance_entry.contractor_name.casefold()
        ):
            raise ValidationError(
                f"{attendance_entry.individual_name} has already signed this toolbox talk today."
            )

    completed_at = datetime.now()
    signature_path = _save_toolbox_talk_signature_image(
        signature_image_data=signature_image_data,
        full_name=attendance_entry.individual_name,
        created_at=completed_at,
        topic=resolved_topic,
    )
    completion_document = ToolboxTalkCompletionDocument(
        doc_id=(
            f"TBT-{completed_at:%Y%m%d%H%M%S%f}-"
            f"{_slugify_identifier(resolved_topic)}-"
            f"{_slugify_identifier(attendance_entry.individual_name)}"
        ),
        site_name=site_name.strip() or DEFAULT_SITE_NAME,
        created_at=completed_at,
        status=DocumentStatus.ACTIVE,
        topic=resolved_topic,
        linked_induction_doc_id=attendance_entry.linked_induction_doc_id,
        individual_name=attendance_entry.individual_name,
        contractor_name=attendance_entry.contractor_name,
        completed_at=completed_at,
        signature_image_path=str(signature_path),
        document_read_confirmed=document_read_confirmed,
    )
    repository.save(completion_document)
    repository.index_file(
        file_name=signature_path.name,
        file_path=signature_path,
        file_category="toolbox_talk_signature_png",
        file_group=FileGroup.FILE_2,
        site_name=completion_document.site_name,
        related_doc_id=completion_document.doc_id,
    )
    return LoggedToolboxTalkCompletion(
        toolbox_talk_completion=completion_document,
        signature_path=signature_path,
    )


def generate_toolbox_talk_register_document(
    repository: DocumentRepository,
    *,
    site_name: str,
    topic: str,
) -> GeneratedToolboxTalkRegisterDocument:
    """Render the approved UHSF16.2 toolbox talk register for one topic."""

    repository.create_schema()
    template_path = TemplateRegistry.resolve_template_path("toolbox_talk_register")
    if not template_path.exists():
        raise TemplateValidationError(
            f"Toolbox talk template is missing: {template_path}"
        )

    completions = list_toolbox_talk_completions(
        repository,
        site_name=site_name,
        topic=topic,
    )
    export_rows = sorted(
        completions,
        key=lambda completion: (completion.completed_at, completion.doc_id),
    )
    config.FILE_2_TBT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = _build_available_destination(
        Path(
            "UHSF16.2 Toolbox Talk Register - "
            f"{_sanitize_filename_fragment(topic)} - "
            f"{date.today():%Y-%m-%d}.docx"
        ),
        config.FILE_2_TBT_OUTPUT_DIR,
    )

    document = Document(template_path)
    _replace_docx_token_everywhere(document, "{{ topic }}", topic.strip())
    _populate_toolbox_talk_rows(document, export_rows)
    document.save(output_path)

    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="toolbox_talk_register_docx",
        file_group=FileGroup.FILE_2,
        site_name=site_name,
    )
    return GeneratedToolboxTalkRegisterDocument(
        output_path=output_path,
        row_count=len(export_rows),
    )


def create_daily_attendance_sign_in(
    repository: DocumentRepository,
    *,
    site_name: str,
    induction_document: InductionDocument,
    vehicle_registration: str,
    distance_travelled: str,
    signature_image_data: Any,
    gate_verification_method: str = "",
    gate_verification_note: str = "",
    geofence_distance_meters: Optional[float] = None,
) -> LoggedDailyAttendanceEntry:
    """Create one live daily sign-in record from an existing induction record."""

    repository.create_schema()
    created_at = datetime.now().replace(second=0, microsecond=0)
    resolved_site_name = site_name.strip() or DEFAULT_SITE_NAME
    if induction_document.status == DocumentStatus.ARCHIVED:
        raise ValidationError(
            f"{induction_document.individual_name} is archived and cannot be signed in."
        )

    active_entries = list_daily_attendance_entries(
        repository,
        site_name=resolved_site_name,
        on_date=created_at.date(),
        active_only=True,
    )
    for active_entry in active_entries:
        same_induction = (
            induction_document.doc_id
            and active_entry.linked_induction_doc_id == induction_document.doc_id
        )
        same_person = (
            active_entry.individual_name.casefold()
            == induction_document.individual_name.casefold()
            and active_entry.contractor_name.casefold()
            == induction_document.contractor_name.casefold()
        )
        if same_induction or same_person:
            raise ValidationError(
                f"{induction_document.individual_name} is already signed in for today."
            )

    config.FILE_2_ATTENDANCE_SIGNATURES_DIR.mkdir(parents=True, exist_ok=True)
    signature_path = _save_attendance_signature_image(
        signature_image_data=signature_image_data,
        full_name=induction_document.individual_name,
        created_at=created_at,
        action="sign-in",
    )

    attendance_entry = DailyAttendanceEntryDocument(
        doc_id=_build_daily_attendance_doc_id(created_at, induction_document.individual_name),
        site_name=resolved_site_name,
        created_at=created_at,
        status=DocumentStatus.ACTIVE,
        linked_induction_doc_id=induction_document.doc_id,
        individual_name=induction_document.individual_name,
        contractor_name=induction_document.contractor_name,
        vehicle_registration=vehicle_registration.strip().upper(),
        distance_travelled=distance_travelled.strip(),
        gate_verification_method=gate_verification_method.strip(),
        gate_verification_note=gate_verification_note.strip(),
        geofence_distance_meters=geofence_distance_meters,
        time_in=created_at,
        sign_in_signature_path=str(signature_path),
    )

    repository.save(attendance_entry)
    repository.index_file(
        file_name=signature_path.name,
        file_path=signature_path,
        file_category="attendance_sign_in_signature_png",
        file_group=FileGroup.FILE_2,
        site_name=attendance_entry.site_name,
        related_doc_id=attendance_entry.doc_id,
    )
    return LoggedDailyAttendanceEntry(
        attendance_entry=attendance_entry,
        signature_path=signature_path,
    )


def complete_daily_attendance_sign_out(
    repository: DocumentRepository,
    *,
    attendance_doc_id: str,
    signature_image_data: Any,
) -> LoggedDailyAttendanceEntry:
    """Complete one live daily sign-out record and calculate hours worked."""

    repository.create_schema()
    attendance_entry = repository.get(attendance_doc_id)
    if not isinstance(attendance_entry, DailyAttendanceEntryDocument):
        raise ValidationError("Selected attendance record is not a UHSF16.09 entry.")
    if not attendance_entry.is_on_site:
        raise ValidationError(
            f"{attendance_entry.individual_name} is not currently signed in."
        )

    completed_at = datetime.now().replace(second=0, microsecond=0)
    signature_path = _save_attendance_signature_image(
        signature_image_data=signature_image_data,
        full_name=attendance_entry.individual_name,
        created_at=completed_at,
        action="sign-out",
    )
    worked_duration_hours = round(
        max((completed_at - attendance_entry.time_in).total_seconds(), 0.0) / 3600.0,
        2,
    )

    attendance_entry.time_out = completed_at
    attendance_entry.hours_worked = worked_duration_hours
    attendance_entry.sign_out_signature_path = str(signature_path)
    attendance_entry.status = DocumentStatus.ARCHIVED
    repository.save(attendance_entry)
    repository.index_file(
        file_name=signature_path.name,
        file_path=signature_path,
        file_category="attendance_sign_out_signature_png",
        file_group=FileGroup.FILE_2,
        site_name=attendance_entry.site_name,
        related_doc_id=attendance_entry.doc_id,
    )
    return LoggedDailyAttendanceEntry(
        attendance_entry=attendance_entry,
        signature_path=signature_path,
    )


def _delete_attendance_file_if_present(
    repository: DocumentRepository,
    file_path_text: str,
) -> None:
    """Delete one attendance-linked file and its index row when it still exists."""

    resolved_path_text = str(file_path_text or "").strip()
    if not resolved_path_text:
        return

    resolved_path = Path(resolved_path_text).expanduser().resolve()
    try:
        repository.delete_indexed_file(resolved_path)
    except Exception:
        pass

    try:
        if resolved_path.exists():
            resolved_path.unlink()
    except OSError:
        pass


def update_daily_attendance_entry(
    repository: DocumentRepository,
    *,
    attendance_doc_id: str,
    contractor_name: str,
    vehicle_registration: str = "",
    distance_travelled: str = "",
    gate_verification_method: str = "",
    gate_verification_note: str = "",
    time_in: datetime,
    time_out: Optional[datetime] = None,
) -> DailyAttendanceEntryDocument:
    """Apply a manager correction to one saved UHSF16.09 attendance entry."""

    repository.create_schema()
    attendance_entry = repository.get(attendance_doc_id)
    if not isinstance(attendance_entry, DailyAttendanceEntryDocument):
        raise ValidationError("Selected attendance record is not a UHSF16.09 entry.")

    resolved_contractor_name = contractor_name.strip()
    if not resolved_contractor_name:
        raise ValidationError("Company / contractor name is required.")

    resolved_time_in = time_in.replace(second=0, microsecond=0)
    resolved_time_out = (
        time_out.replace(second=0, microsecond=0)
        if time_out is not None
        else None
    )
    if resolved_time_out is not None and resolved_time_out < resolved_time_in:
        raise ValidationError("Time out must be on or after time in.")

    attendance_entry.contractor_name = resolved_contractor_name
    attendance_entry.vehicle_registration = vehicle_registration.strip().upper()
    attendance_entry.distance_travelled = distance_travelled.strip()
    attendance_entry.gate_verification_method = gate_verification_method.strip()
    attendance_entry.gate_verification_note = gate_verification_note.strip()
    attendance_entry.time_in = resolved_time_in

    if resolved_time_out is None:
        if attendance_entry.sign_out_signature_path:
            _delete_attendance_file_if_present(
                repository,
                attendance_entry.sign_out_signature_path,
            )
        attendance_entry.time_out = None
        attendance_entry.hours_worked = None
        attendance_entry.sign_out_signature_path = ""
        attendance_entry.status = DocumentStatus.ACTIVE
    else:
        attendance_entry.time_out = resolved_time_out
        attendance_entry.hours_worked = round(
            max((resolved_time_out - resolved_time_in).total_seconds(), 0.0) / 3600.0,
            2,
        )
        attendance_entry.status = DocumentStatus.ARCHIVED

    repository.save(attendance_entry)
    return attendance_entry


def get_daily_contractor_headcount(
    repository: DocumentRepository,
    site_name: str,
    target_date: date,
) -> List[Dict[str, Any]]:
    """Return contractor day/night counts from the live UHSF16.09 register."""

    attendance_entries = list_daily_attendance_entries(
        repository,
        site_name=site_name,
        on_date=target_date,
        active_only=True,
    )
    contractor_counts: Dict[str, int] = {}
    for attendance_entry in attendance_entries:
        contractor_name = attendance_entry.contractor_name.strip() or "Unknown"
        contractor_counts[contractor_name] = contractor_counts.get(contractor_name, 0) + 1
    return [
        {
            "company": contractor_name,
            "days": count,
            "nights": 0,
        }
        for contractor_name, count in sorted(
            contractor_counts.items(),
            key=lambda item: item[0].casefold(),
        )
    ]


def _build_site_diary_doc_id(site_name: str, target_date: date) -> str:
    """Return a stable document id for one UHSF15.63 diary date."""

    return (
        "site-diary-"
        f"{_slugify_identifier(site_name)}-"
        f"{target_date:%Y%m%d}"
    )


def generate_site_diary_document(
    repository: DocumentRepository,
    *,
    site_diary_document: SiteDiaryDocument,
) -> GeneratedSiteDiaryDocument:
    """Render and store the approved UHSF15.63 daily site diary."""

    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the printable UHSF15.63 daily site diary."
        ) from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError(
            "jinja2 is required to generate the printable UHSF15.63 daily site diary."
        ) from exc

    repository.create_schema()
    config.FILE_2_DIARY_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    template_path = TemplateRegistry.resolve_template_path("site_diary")
    if not template_path.exists():
        raise TemplateValidationError(
            f"Daily site diary template is missing: {template_path}"
        )

    output_name = Path(
        "UHSF15.63 Daily Site Diary - "
        f"{_sanitize_filename_fragment(site_diary_document.site_name)} - "
        f"{site_diary_document.date:%Y-%m-%d}.docx"
    )
    output_path = _build_available_destination(
        output_name,
        config.FILE_2_DIARY_OUTPUT_DIR,
    )

    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "site-diary-template.docx"
        _build_patched_docxtpl_template(template_path, repaired_template_path)

        discovered_placeholders = (
            _discover_docx_template_tags(template_path)
            | _discover_docx_template_tags(repaired_template_path)
        )
        missing_placeholders = sorted(
            SiteDiaryDocument.required_template_placeholders - discovered_placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "Daily site diary template is missing required placeholders: "
                + ", ".join(missing_placeholders)
            )

        document_template = DocxTemplate(str(repaired_template_path))
        clean_jinja_environment = Environment(autoescape=False)
        document_template.render(
            site_diary_document.to_template_context(),
            jinja_env=clean_jinja_environment,
            autoescape=False,
        )
        document_template.save(output_path)

    saved_document = replace(
        site_diary_document,
        generated_document_path=str(output_path),
    )
    repository.save(saved_document)
    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="site_diary_docx",
        file_group=FileGroup.FILE_2,
        site_name=saved_document.site_name,
        related_doc_id=saved_document.doc_id,
    )
    return GeneratedSiteDiaryDocument(
        output_path=output_path,
        contractor_count=len(saved_document.contractors),
        visitor_count=len(saved_document.visitors),
        site_diary_document=saved_document,
    )


def generate_attendance_register_document(
    repository: DocumentRepository,
    *,
    site_name: str,
    on_date: Optional[date] = None,
) -> GeneratedAttendanceRegisterDocument:
    """Render the approved File 2 attendance register into the attendance folder."""

    report_date = on_date or date.today()
    repository.create_schema()
    config.FILE_2_ATTENDANCE_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    template_path = TemplateRegistry.resolve_template_path("attendance_register")
    attendance_entries = list_daily_attendance_entries(
        repository,
        site_name=site_name,
        on_date=report_date,
        active_only=False,
    )
    printable_entries = sorted(
        attendance_entries,
        key=lambda attendance_entry: attendance_entry.time_in,
    )[:17]
    induction_lookup = _build_attendance_induction_lookup(repository, printable_entries)

    output_name = Path(
        "UHSF16.09 Site Attendance Register - "
        f"{_sanitize_filename_fragment(site_name)} - "
        f"{report_date:%Y-%m-%d}.docx"
    )
    output_path = _build_available_destination(
        output_name,
        config.FILE_2_ATTENDANCE_OUTPUT_DIR,
    )

    document = Document(template_path)
    _validate_attendance_register_template(document)
    _populate_attendance_register_rows(
        document,
        attendance_entries=printable_entries,
        induction_lookup=induction_lookup,
    )
    _populate_attendance_register_summary(
        document,
        attendance_entries=printable_entries,
    )
    document.save(output_path)

    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="attendance_register_docx",
        file_group=FileGroup.FILE_2,
        site_name=site_name,
    )
    return GeneratedAttendanceRegisterDocument(
        output_path=output_path,
        row_count=len(printable_entries),
    )


def generate_plant_register_document(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> GeneratedPlantRegisterDocument:
    """Render the approved File 2 plant register template for one site."""

    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the printable File 2 plant register."
        ) from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError(
            "jinja2 is required to generate the printable File 2 plant register."
        ) from exc

    repository.create_schema()
    config.PLANT_HIRE_REGISTER_DIR.mkdir(parents=True, exist_ok=True)

    template_path = TemplateRegistry.resolve_template_path("plant_register")
    plant_assets = _list_site_plant_assets(repository, site_name=site_name)
    context = {
        "plant_assets": [
            {
                "hire_num": plant_asset.hire_num,
                "description": plant_asset.description,
                "company": plant_asset.company,
                "phone": plant_asset.phone,
                "on_hire": plant_asset.on_hire.strftime("%d/%m/%y"),
                "hired_by": plant_asset.hired_by,
                "serial": plant_asset.serial or plant_asset.stock_code,
                "inspection": format_plant_inspection_reference(
                    plant_asset.inspection_type,
                    plant_asset.inspection,
                ),
                "in_file": "Yes" if plant_asset.source_reference else "",
            }
            for plant_asset in plant_assets
        ]
    }

    output_name = Path(
        "UHSF18.32 Plant Hire Site Register - "
        f"{_sanitize_filename_fragment(site_name)} - "
        f"{date.today():%Y-%m-%d}.docx"
    )
    output_path = _build_available_destination(
        output_name,
        config.PLANT_HIRE_REGISTER_DIR,
    )

    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "plant-register-template.docx"
        _build_patched_docxtpl_template(template_path, repaired_template_path)

        discovered_placeholders = _discover_docx_template_tags(repaired_template_path)
        missing_placeholders = sorted(
            {
                "p.hire_num",
                "p.description",
                "p.company",
                "p.phone",
                "p.on_hire",
                "p.hired_by",
                "p.serial",
                "p.inspection",
                "p.in_file",
            }
            - discovered_placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "Plant register template is missing required placeholders: "
                + ", ".join(missing_placeholders)
            )

        document_template = DocxTemplate(str(repaired_template_path))
        clean_jinja_environment = Environment(autoescape=False)
        document_template.render(
            context,
            jinja_env=clean_jinja_environment,
            autoescape=False,
        )
        document_template.save(output_path)
        _normalise_plant_register_table_rows(output_path)

    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="plant_register_docx",
        file_group=FileGroup.FILE_2,
        site_name=site_name,
    )
    return GeneratedPlantRegisterDocument(
        output_path=output_path,
        asset_count=len(plant_assets),
    )


def _build_attendance_induction_lookup(
    repository: DocumentRepository,
    attendance_entries: List[DailyAttendanceEntryDocument],
) -> Dict[str, InductionDocument]:
    """Return linked induction records used to backfill contact details."""

    induction_lookup: Dict[str, InductionDocument] = {}
    for attendance_entry in attendance_entries:
        linked_doc_id = attendance_entry.linked_induction_doc_id
        if not linked_doc_id or linked_doc_id in induction_lookup:
            continue
        try:
            linked_document = repository.get(linked_doc_id)
        except (DocumentNotFoundError, ValueError):
            continue
        if isinstance(linked_document, InductionDocument):
            induction_lookup[linked_doc_id] = linked_document
    return induction_lookup


def _validate_attendance_register_template(document: Document) -> None:
    """Check the approved attendance template still uses the fixed Lovedean layout."""

    if len(document.tables) < 4:
        raise TemplateValidationError(
            "Attendance register template must contain the attendance and summary tables."
        )

    attendance_table = document.tables[1]
    if len(attendance_table.rows) < 18 or len(attendance_table.columns) < 11:
        raise TemplateValidationError(
            "Attendance register template must contain 17 printable attendance rows."
        )

    expected_headers = (
        "Date",
        "Name",
        "Company",
        "Phone No",
        "Distance Travelled",
        "Vehicle Reg",
        "Time In",
        "Signature",
        "Time Out",
        "Hours Worked",
        "Signature",
    )
    actual_headers = tuple(
        re.sub(r"\s+", " ", cell.text.replace("\n", " ").strip())
        for cell in attendance_table.rows[0].cells[:11]
    )
    if actual_headers != expected_headers:
        raise TemplateValidationError(
            "Attendance register template headers no longer match the approved UHSF16.09 layout."
        )

    summary_table = document.tables[3]
    if len(summary_table.rows) < 2 or len(summary_table.columns) < 5:
        raise TemplateValidationError(
            "Attendance register template summary table is missing required totals cells."
        )


def _populate_attendance_register_rows(
    document: Document,
    *,
    attendance_entries: List[DailyAttendanceEntryDocument],
    induction_lookup: Dict[str, InductionDocument],
) -> None:
    """Fill the fixed-row UHSF16.09 table from today's live attendance entries."""

    attendance_table = document.tables[1]
    for row_index in range(1, 18):
        row = attendance_table.rows[row_index]
        if row_index - 1 >= len(attendance_entries):
            for cell in row.cells[:11]:
                _set_docx_cell_text(cell, "")
            continue

        attendance_entry = attendance_entries[row_index - 1]
        linked_induction = induction_lookup.get(attendance_entry.linked_induction_doc_id)
        row_values = (
            attendance_entry.attendance_date.strftime("%d/%m/%Y"),
            attendance_entry.individual_name,
            attendance_entry.contractor_name,
            linked_induction.contact_number if linked_induction is not None else "",
            attendance_entry.distance_travelled,
            attendance_entry.vehicle_registration,
            attendance_entry.time_in.strftime("%H:%M"),
            "",
            attendance_entry.time_out.strftime("%H:%M")
            if attendance_entry.time_out is not None
            else "",
            f"{attendance_entry.hours_worked:.2f}"
            if attendance_entry.hours_worked is not None
            else "",
            "",
        )
        for cell, value in zip(row.cells[:11], row_values):
            _set_docx_cell_text(cell, value)
        _insert_signature_into_cell(row.cells[7], attendance_entry.sign_in_signature_path)
        _insert_signature_into_cell(row.cells[10], attendance_entry.sign_out_signature_path)


def _populate_attendance_register_summary(
    document: Document,
    *,
    attendance_entries: List[DailyAttendanceEntryDocument],
) -> None:
    """Fill the summary totals for Uplands and subcontractor rows."""

    summary_table = document.tables[3]
    subcontractor_entries = [
        attendance_entry
        for attendance_entry in attendance_entries
        if not attendance_entry.is_uplands_employee
    ]
    uplands_entries = [
        attendance_entry
        for attendance_entry in attendance_entries
        if attendance_entry.is_uplands_employee
    ]

    _set_docx_cell_text(summary_table.rows[0].cells[1], str(len(subcontractor_entries)))
    _set_docx_cell_text(summary_table.rows[0].cells[4], str(len(uplands_entries)))
    _set_docx_cell_text(
        summary_table.rows[1].cells[1],
        _format_attendance_hours_total(subcontractor_entries),
    )
    _set_docx_cell_text(
        summary_table.rows[1].cells[4],
        _format_attendance_hours_total(uplands_entries),
    )


def _format_attendance_hours_total(
    attendance_entries: List[DailyAttendanceEntryDocument],
) -> str:
    """Return the printed total hours for one attendance grouping."""

    total_hours = sum(
        attendance_entry.hours_worked or 0.0
        for attendance_entry in attendance_entries
    )
    return f"{total_hours:.2f}" if total_hours > 0 else ""


def _set_docx_cell_text(cell: Any, value: str) -> None:
    """Replace one docx table cell with plain text while clearing placeholders."""

    cell.text = value


def _insert_signature_into_cell(cell: Any, signature_path: str) -> None:
    """Insert a saved signature image into one attendance register cell."""

    from docx.shared import Mm

    resolved_signature_path = Path(str(signature_path)).expanduser()
    if not signature_path or not resolved_signature_path.exists():
        _set_docx_cell_text(cell, "")
        return

    _set_docx_cell_text(cell, "")
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run().add_picture(str(resolved_signature_path), width=Mm(16))


def _replace_docx_token_everywhere(document: Document, token: str, value: str) -> None:
    """Replace one literal placeholder token across paragraphs and table cells."""

    for paragraph in document.paragraphs:
        if token in paragraph.text:
            paragraph.text = paragraph.text.replace(token, value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if token in cell.text:
                    cell.text = cell.text.replace(token, value)


def _populate_toolbox_talk_rows(
    document: Document,
    completions: List[ToolboxTalkCompletionDocument],
) -> None:
    """Fill the UHSF16.2 signature rows from the collected remote completions."""

    placeholder_row_specs: List[tuple[Any, int]] = []
    placeholder_tokens = ("{{ name }}", "{{ company }}", "{{ date }}", "{{ signature }}")
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)
            if any(token in row_text for token in placeholder_tokens):
                placeholder_row_specs.append((table, row_index))

    if not placeholder_row_specs:
        raise TemplateValidationError(
            "Toolbox talk template is missing the tagged signature rows."
        )

    target_table, template_row_index = placeholder_row_specs[0]
    placeholder_row_indices = [
        row_index
        for table, row_index in placeholder_row_specs
        if table is target_table
    ]
    template_row_xml = target_table.rows[template_row_index]._tr
    last_row_xml = target_table.rows[placeholder_row_indices[-1]]._tr
    missing_row_count = max(0, len(completions) - len(placeholder_row_indices))
    for _ in range(missing_row_count):
        cloned_row_xml = deepcopy(template_row_xml)
        last_row_xml.addnext(cloned_row_xml)
        last_row_xml = cloned_row_xml

    placeholder_rows = [
        row
        for row in target_table.rows
        if any(token in " ".join(cell.text for cell in row.cells) for token in placeholder_tokens)
    ]

    for row_index, row in enumerate(placeholder_rows):
        if row_index < len(completions):
            _apply_toolbox_talk_row_values(row, completions[row_index])
        else:
            _blank_toolbox_talk_row(row)


def _apply_toolbox_talk_row_values(
    row: Any,
    completion: ToolboxTalkCompletionDocument,
) -> None:
    """Fill one UHSF16.2 row with a single operative completion."""

    replacement_map = {
        "{{ name }}": completion.individual_name,
        "{{ company }}": completion.contractor_name,
        "{{ topic }}": completion.topic,
        "{{ date }}": completion.completed_at.strftime("%d/%m/%Y"),
    }
    for cell in row.cells:
        if "{{ signature }}" in cell.text:
            _insert_signature_into_cell(cell, completion.signature_image_path)
            continue
        updated_text = cell.text
        for token, value in replacement_map.items():
            updated_text = updated_text.replace(token, value)
        _set_docx_cell_text(cell, updated_text if "{{" not in updated_text else "")


def _blank_toolbox_talk_row(row: Any) -> None:
    """Clear one unused toolbox talk row in the export template."""

    for cell in row.cells:
        _set_docx_cell_text(cell, "")


def _list_site_rams_documents(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> List[RAMSDocument]:
    """Return live File 3 RAMS documents for one site."""

    return sorted(
        [
            document
            for document in repository.list_documents(
                document_type=RAMSDocument.document_type,
                site_name=site_name,
            )
            if isinstance(document, RAMSDocument)
            and document.status != DocumentStatus.ARCHIVED
        ],
        key=lambda document: (
            document.review_date,
            document.reference.casefold(),
            document.activity_description.casefold(),
        ),
        reverse=True,
    )


def _list_site_coshh_documents(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> List[COSHHDocument]:
    """Return live File 3 COSHH documents for one site."""

    return sorted(
        [
            document
            for document in repository.list_documents(
                document_type=COSHHDocument.document_type,
                site_name=site_name,
            )
            if isinstance(document, COSHHDocument)
            and document.status != DocumentStatus.ARCHIVED
        ],
        key=lambda document: (
            document.review_date,
            document.reference.casefold(),
            document.substance_name.casefold(),
        ),
        reverse=True,
    )


def generate_rams_register_document(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> GeneratedSafetyRegisterDocument:
    """Render the approved File 3 RAMS register template for one site."""

    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the printable File 3 RAMS register."
        ) from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError(
            "jinja2 is required to generate the printable File 3 RAMS register."
        ) from exc

    repository.create_schema()
    config.FILE_3_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    template_path = TemplateRegistry.resolve_template_path("rams_register")
    rams_documents = _list_site_rams_documents(repository, site_name=site_name)
    context = {
        "rams_list": [
            {
                "ref": document.reference,
                "title": document.activity_description,
                "company": document.contractor_name,
                "date": document.review_date.strftime("%d/%m/%y"),
                "checked_by": document.assessor_name,
                "manager_name": document.manager_name,
            }
            for document in rams_documents
        ]
    }

    output_name = Path(
        "16.4 RAMs Register - "
        f"{_sanitize_filename_fragment(site_name)} - "
        f"{date.today():%Y-%m-%d}.docx"
    )
    output_path = _build_available_destination(output_name, config.FILE_3_OUTPUT_DIR)

    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "rams-register-template.docx"
        _build_patched_docxtpl_template(template_path, repaired_template_path)

        discovered_placeholders = _discover_docx_template_tags(repaired_template_path)
        missing_placeholders = sorted(
            {"r.checked_by", "r.company", "r.date", "r.manager_name", "r.ref", "r.title"}
            - discovered_placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "RAMS register template is missing required placeholders: "
                + ", ".join(missing_placeholders)
            )

        document_template = DocxTemplate(str(repaired_template_path))
        clean_jinja_environment = Environment(autoescape=False)
        document_template.render(
            context,
            jinja_env=clean_jinja_environment,
            autoescape=False,
        )
        document_template.save(output_path)
        _normalise_permit_register_table_rows(output_path)

    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="rams_register_docx",
        file_group=FileGroup.FILE_3,
        site_name=site_name,
    )
    return GeneratedSafetyRegisterDocument(
        output_path=output_path,
        row_count=len(rams_documents),
        register_type="RAMS",
    )


def generate_coshh_register_document(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> GeneratedSafetyRegisterDocument:
    """Render the approved File 3 COSHH register template for one site."""

    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the printable File 3 COSHH register."
        ) from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError(
            "jinja2 is required to generate the printable File 3 COSHH register."
        ) from exc

    repository.create_schema()
    config.FILE_3_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    template_path = TemplateRegistry.resolve_template_path("coshh_register")
    coshh_documents = _list_site_coshh_documents(repository, site_name=site_name)
    context = {
        "coshh_list": [
            {
                "date": document.review_date.strftime("%d/%m/%y"),
                "manager_name": document.manager_name,
                "name": document.substance_name,
                "risk": ", ".join(document.hazard_pictograms) or "COSHH assessment",
                "status": document.status.label,
                "supplier": document.supplier_name or document.manufacturer or "Unknown Supplier",
                "use": document.intended_use or document.contractor_name,
            }
            for document in coshh_documents
        ]
    }

    output_name = Path(
        "COSHH Register - "
        f"{_sanitize_filename_fragment(site_name)} - "
        f"{date.today():%Y-%m-%d}.docx"
    )
    output_path = _build_available_destination(output_name, config.FILE_3_OUTPUT_DIR)

    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "coshh-register-template.docx"
        _build_patched_docxtpl_template(template_path, repaired_template_path)

        discovered_placeholders = _discover_docx_template_tags(repaired_template_path)
        missing_placeholders = sorted(
            {"c.date", "c.manager_name", "c.name", "c.risk", "c.status", "c.supplier", "c.use"}
            - discovered_placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "COSHH register template is missing required placeholders: "
                + ", ".join(missing_placeholders)
            )

        document_template = DocxTemplate(str(repaired_template_path))
        clean_jinja_environment = Environment(autoescape=False)
        document_template.render(
            context,
            jinja_env=clean_jinja_environment,
            autoescape=False,
        )
        document_template.save(output_path)
        _normalise_permit_register_table_rows(output_path)

    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="coshh_register_docx",
        file_group=FileGroup.FILE_3,
        site_name=site_name,
    )
    return GeneratedSafetyRegisterDocument(
        output_path=output_path,
        row_count=len(coshh_documents),
        register_type="COSHH",
    )


def create_ladder_permit_draft(
    repository: DocumentRepository,
    *,
    attendance_record: SiteAttendanceRecord,
    site_worker: Optional[SiteWorker] = None,
    worker_company_override: str = "",
    description_of_work: str,
    location_of_work: str,
    supervisor_name: str,
    safety_checklist: Optional[Mapping[int, bool]] = None,
    inspection_checked_by: str,
    inspection_rungs_ok: bool,
    inspection_stiles_ok: bool,
    inspection_feet_ok: bool,
    inspection_ok_to_use: bool,
    inspection_comments: str = "",
    site_name: Optional[str] = None,
    job_number: Optional[str] = None,
) -> GeneratedLadderPermit:
    """Generate a draft ladder permit from a live roster-backed attendance record."""

    cleaned_description = description_of_work.strip()
    cleaned_location = location_of_work.strip()
    cleaned_supervisor_name = supervisor_name.strip()
    cleaned_inspection_checked_by = inspection_checked_by.strip()
    cleaned_inspection_comments = inspection_comments.strip()
    resolved_checklist = _resolve_ladder_safety_checklist(safety_checklist)
    resolved_manager_name = DEFAULT_LADDER_PERMIT_MANAGER_NAME
    resolved_manager_position = DEFAULT_LADDER_PERMIT_MANAGER_POSITION

    if not cleaned_description:
        raise ValidationError("Description of work is required.")
    if not cleaned_location:
        raise ValidationError("Location on site is required.")
    if not cleaned_supervisor_name:
        raise ValidationError("Supervisor name is required.")
    if not cleaned_inspection_checked_by:
        raise ValidationError("Inspected by is required.")

    resolved_site_name = site_name or DEFAULT_SITE_NAME
    resolved_job_number = (job_number or "").strip()
    if not resolved_job_number:
        raise ValidationError(
            "Job number is required. Update Project Setup before issuing the permit."
        )
    resolved_worker_name = (
        site_worker.worker_name.strip()
        if site_worker is not None and site_worker.worker_name.strip()
        else attendance_record.workerName
    )
    resolved_worker_company = worker_company_override.strip() or (
        site_worker.company.strip()
        if site_worker is not None and site_worker.company.strip()
        else attendance_record.company
    )
    created_at = _current_permit_issue_datetime()
    permit_issue_date = created_at.date()
    permit_expiry_datetime = created_at + timedelta(hours=8)
    permit = LadderPermit(
        doc_id=_build_ladder_permit_doc_id(created_at, resolved_worker_name),
        site_name=resolved_site_name,
        created_at=created_at,
        status=DocumentStatus.DRAFT,
        permit_number=_build_ladder_permit_number(
            repository,
            site_name=resolved_site_name,
        ),
        project_name=resolved_site_name,
        project_number=resolved_job_number,
        location_of_work=cleaned_location,
        description_of_work=cleaned_description,
        valid_from_date=permit_issue_date,
        valid_from_time=created_at.time(),
        valid_to_date=permit_expiry_datetime.date(),
        valid_to_time=permit_expiry_datetime.time(),
        safer_alternative_eliminated=resolved_checklist[1],
        task_specific_rams_prepared_and_approved=resolved_checklist[2],
        personnel_briefed_and_understand_task=resolved_checklist[3],
        competent_supervisor_appointed=resolved_checklist[4],
        competent_supervisor_name=cleaned_supervisor_name,
        operatives_suitably_trained=resolved_checklist[5],
        ladder_length_suitable=resolved_checklist[6],
        conforms_to_bs_class_a=resolved_checklist[7],
        three_points_of_contact_maintained=resolved_checklist[8],
        harness_worn_and_secured_above_head_height=resolved_checklist[9],
        ladder_stabilisation_method=LadderStabilisationMethod.FOOTED,
        equipment_inspected_for_defects=resolved_checklist[11],
        ladder_stabilisation_confirmed=resolved_checklist[10],
        worker_name=resolved_worker_name,
        worker_company=resolved_worker_company,
        briefing_name=resolved_manager_name,
        manager_name=resolved_manager_name,
        manager_position=resolved_manager_position,
        issued_date=permit_issue_date,
    )
    permit.add_inspection_record(
        inspection_date=created_at.date(),
        inspected_by=cleaned_inspection_checked_by,
        rungs_ok=bool(inspection_rungs_ok),
        stiles_ok=bool(inspection_stiles_ok),
        feet_ok=bool(inspection_feet_ok),
        comments_or_action_taken=cleaned_inspection_comments,
        ok_to_use=bool(inspection_ok_to_use),
    )

    permit_output_dir = config.PERMITS_DESTINATION
    permit_output_dir.mkdir(parents=True, exist_ok=True)
    output_path = _build_available_destination(
        Path(f"{permit.permit_number}-{_slugify_identifier(resolved_worker_name)}.docx"),
        permit_output_dir,
    )

    rendered_path = TemplateManager(permit).render(output_path)
    _prefill_rendered_ladder_permit_sections(rendered_path, permit)
    repository.save(permit)
    repository.index_file(
        file_name=rendered_path.name,
        file_path=rendered_path,
        file_category="ladder_permit_docx",
        file_group=FileGroup.FILE_4,
        site_name=permit.site_name,
        related_doc_id=permit.doc_id,
    )
    return GeneratedLadderPermit(
        permit=permit,
        output_path=rendered_path,
    )


def _prefill_rendered_ladder_permit_sections(
    output_path: Path,
    permit: LadderPermit,
) -> None:
    """Fill blank management cells in the rendered ladder permit DOCX."""

    document = Document(output_path)
    if len(document.tables) < 5:
        document.save(output_path)
        return

    acceptance_table = document.tables[4]
    issued_date_text = (permit.issued_date or permit.valid_from_date).strftime("%d/%m/%Y")
    valid_from_time_text = permit.valid_from_time.strftime("%H:%M")
    valid_to_date_text = permit.valid_to_date.strftime("%d/%m/%Y")
    valid_to_time_text = permit.valid_to_time.strftime("%H:%M")
    manager_name = permit.manager_name or DEFAULT_LADDER_PERMIT_MANAGER_NAME
    manager_position = permit.manager_position or DEFAULT_LADDER_PERMIT_MANAGER_POSITION
    operative_name = permit.to_template_context().get("contractor_name", permit.worker_name)
    operative_company = permit.worker_company

    if len(document.tables) > 0:
        header_table = document.tables[0]
        if len(header_table.rows) > 3 and len(header_table.rows[3].cells) > 4:
            header_table.rows[3].cells[4].text = valid_from_time_text
        if len(header_table.rows) > 4 and len(header_table.rows[4].cells) > 4:
            header_table.rows[4].cells[1].text = valid_to_date_text
            header_table.rows[4].cells[4].text = valid_to_time_text

    # Section 1 site-manager authorisation row: keep tag-rendered name/date and add position.
    if len(acceptance_table.rows) > 1 and len(acceptance_table.rows[1].cells) > 3:
        acceptance_table.rows[1].cells[3].text = f"Position: {manager_position}"

    # Section 3 contractor / operative completion row is blank in the official template.
    if len(acceptance_table.rows) > 5 and len(acceptance_table.rows[5].cells) > 3:
        acceptance_table.rows[5].cells[0].text = f"Name: {operative_name}"
        acceptance_table.rows[5].cells[2].text = f"Date: (dd/mm/yyyy) {issued_date_text}"
        acceptance_table.rows[5].cells[3].text = f"Company: {operative_company}"

    # Section 4 completion acceptance row is blank in the official template and must be post-filled.
    if len(acceptance_table.rows) > 7 and len(acceptance_table.rows[7].cells) > 3:
        acceptance_table.rows[7].cells[0].text = f"Name: {manager_name}"
        acceptance_table.rows[7].cells[2].text = f"Date: (dd/mm/yyyy) {issued_date_text}"
        acceptance_table.rows[7].cells[3].text = f"Position: {manager_position}"

    for table in document.tables:
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                lowered_text = cell.text.casefold()
                if "name of person giving briefing" not in lowered_text:
                    continue
                if index + 1 >= len(row.cells):
                    continue
                if row.cells[index + 1].text.strip():
                    continue
                row.cells[index + 1].text = permit.briefing_name or manager_name

    document.save(output_path)


def _current_permit_issue_datetime() -> datetime:
    """Return the current issue timestamp rounded to the nearest minute."""

    return datetime.now().replace(second=0, microsecond=0)


def generate_permit_register_document(
    repository: DocumentRepository,
    *,
    site_name: str,
    job_number: str,
) -> GeneratedPermitRegisterDocument:
    """Render the approved File 4 permit register template for one site."""

    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:
        raise RuntimeError(
            "docxtpl is required to generate the printable File 4 permit register."
        ) from exc
    try:
        from jinja2 import Environment
    except ImportError as exc:
        raise RuntimeError(
            "jinja2 is required to generate the printable File 4 permit register."
        ) from exc

    repository.create_schema()
    output_directory = config.PERMITS_DESTINATION
    output_directory.mkdir(parents=True, exist_ok=True)

    template_path = TemplateRegistry.resolve_template_path("permit_register")
    permits = [
        document
        for document in repository.list_documents(
            document_type=LadderPermit.document_type,
            site_name=site_name,
        )
        if isinstance(document, LadderPermit)
    ]
    permits.sort(key=_ladder_permit_sort_key)

    context = {
        "site_name": site_name,
        "job_number": job_number.strip(),
        "permits": [
            {
                "ref": permit.permit_number,
                "date": _format_permit_register_date(permit),
                "type": "Ladder",
                "name_company": _format_name_company(permit),
                "location": permit.location_of_work,
                "contact": "",
                "time_issued": permit.valid_from_time.strftime("%H:%M"),
                "time_cancelled": "",
            }
            for permit in permits
        ],
    }

    output_name = Path(
        "UHSF21.00 Permit Register - "
        f"{_sanitize_filename_fragment(site_name)} - "
        f"{date.today():%Y-%m-%d}.docx"
    )
    output_path = _build_available_destination(output_name, output_directory)

    with tempfile.TemporaryDirectory() as temp_dir:
        repaired_template_path = Path(temp_dir) / "permit-register-template.docx"
        _build_patched_docxtpl_template(template_path, repaired_template_path)

        discovered_placeholders = _discover_docx_template_tags(repaired_template_path)
        missing_placeholders = sorted(
            {"site_name", "job_number"} - discovered_placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "Permit register template is missing required placeholders: "
                + ", ".join(missing_placeholders)
            )

        document_template = DocxTemplate(str(repaired_template_path))
        clean_jinja_environment = Environment(autoescape=False)
        document_template.render(
            context,
            jinja_env=clean_jinja_environment,
            autoescape=False,
        )
        document_template.save(output_path)
        _normalise_permit_register_table_rows(output_path)

    repository.index_file(
        file_name=output_path.name,
        file_path=output_path,
        file_category="permit_register_docx",
        file_group=FileGroup.FILE_4,
        site_name=site_name,
    )

    return GeneratedPermitRegisterDocument(
        output_path=output_path,
        permit_count=len(permits),
    )


def sync_file_4_permit_records(
    repository: DocumentRepository,
    *,
    site_name: Optional[str] = None,
) -> File4PermitSyncResult:
    """Remove File 4 ladder permit records whose physical DOCX files are missing."""

    repository.create_schema()
    permit_directory = config.PERMITS_DESTINATION
    permit_directory.mkdir(parents=True, exist_ok=True)

    removed_doc_ids: List[str] = []
    removed_indexed_files: List[Path] = []
    removed_doc_id_set = set()

    ladder_permits = [
        document
        for document in repository.list_documents(
            document_type=LadderPermit.document_type,
            site_name=site_name,
        )
        if isinstance(document, LadderPermit)
    ]
    indexed_ladder_files = repository.list_indexed_files(
        file_group=FileGroup.FILE_4,
        file_category="ladder_permit_docx",
    )

    indexed_by_doc_id: Dict[str, List[IndexedFileRecord]] = {}
    for indexed_file in indexed_ladder_files:
        if site_name is not None and indexed_file.site_name != site_name:
            continue
        if indexed_file.related_doc_id is None:
            continue
        indexed_by_doc_id.setdefault(indexed_file.related_doc_id, []).append(indexed_file)

    for permit in ladder_permits:
        related_files = indexed_by_doc_id.get(permit.doc_id, [])
        if related_files and any(indexed_file.file_path.exists() for indexed_file in related_files):
            continue

        repository.delete_document(permit.doc_id)
        removed_doc_id_set.add(permit.doc_id)
        removed_doc_ids.append(permit.doc_id)
        removed_indexed_files.extend(
            indexed_file.file_path for indexed_file in related_files if not indexed_file.file_path.exists()
        )

    for indexed_file in indexed_ladder_files:
        if site_name is not None and indexed_file.site_name != site_name:
            continue
        if indexed_file.file_path.exists():
            continue
        if indexed_file.related_doc_id in removed_doc_id_set:
            continue
        repository.delete_indexed_file(indexed_file.file_path)
        removed_indexed_files.append(indexed_file.file_path)

    return File4PermitSyncResult(
        removed_doc_ids=removed_doc_ids,
        removed_indexed_files=removed_indexed_files,
    )


def check_site_inductions(
    repository: DocumentRepository,
    *,
    on_date: Optional[date] = None,
    site_name: Optional[str] = None,
) -> SiteInductionAuditResult:
    """Cross-reference today's attendance against induction PDFs in File 3."""

    audit_date = on_date or date.today()
    induction_directory = config.INDUCTION_DIR
    induction_directory.mkdir(parents=True, exist_ok=True)

    attendance_registers = [
        document
        for document in repository.list_documents(
            document_type=SiteAttendanceRegister.document_type
        )
        if isinstance(document, SiteAttendanceRegister)
    ]
    if site_name is not None:
        attendance_registers = [
            register
            for register in attendance_registers
            if register.site_name.casefold() == site_name.casefold()
        ]

    worker_names_by_identifier: Dict[str, str] = {}
    resolved_site_name = site_name
    for register in attendance_registers:
        if resolved_site_name is None:
            resolved_site_name = register.site_name
        for record in register.attendance_records:
            if record.date != audit_date:
                continue
            worker_identifier = _normalize_worker_identifier(record.workerName)
            worker_names_by_identifier.setdefault(worker_identifier, record.workerName)

    worker_names = sorted(worker_names_by_identifier.values(), key=str.casefold)
    induction_files = sorted(
        path
        for path in induction_directory.iterdir()
        if path.is_file() and path.suffix.lower() == ".pdf"
    )

    matched_files: Dict[str, Path] = {}
    inducted_workers: List[str] = []
    missing_workers: List[str] = []

    for worker_name in worker_names:
        matched_file = _find_matching_induction_file(worker_name, induction_files)
        if matched_file is None:
            missing_workers.append(worker_name)
            continue
        inducted_workers.append(worker_name)
        matched_files[worker_name] = matched_file

    return SiteInductionAuditResult(
        audit_date=audit_date,
        site_name=resolved_site_name,
        workers_on_site=worker_names,
        inducted_workers=sorted(inducted_workers, key=str.casefold),
        missing_workers=sorted(missing_workers, key=str.casefold),
        matched_files=matched_files,
    )


def extract_expiry_date_from_pdf(pdf_path: Path) -> Optional[date]:
    """Read the first two PDF pages and infer the most likely expiry date."""

    page_text = _extract_pdf_text_pages(pdf_path)

    candidates: List[tuple[int, date]] = []
    for text in page_text:
        candidates.extend(_extract_expiry_candidates_from_text(text))

    candidates.extend(_extract_expiry_candidates_from_text(pdf_path.stem))

    if not candidates:
        return None

    candidates.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return candidates[0][1]


def extract_tonnage_from_ticket(pdf_path: Path) -> Optional[float]:
    """Read a waste ticket PDF and return its extracted tonnage in tonnes."""

    normalized_text = _normalize_text(" ".join(_extract_pdf_text_pages(pdf_path)))
    return _extract_tonnage_from_text(normalized_text)


def _extract_tonnage_from_text(normalized_text: str) -> Optional[float]:
    """Return the highest-confidence tonnage value from normalized ticket text."""

    candidates: List[Tuple[int, int, float]] = []

    for priority, pattern in TONNAGE_PATTERNS:
        for match in pattern.finditer(normalized_text):
            quantity_text = match.group("quantity").replace(",", "")
            try:
                quantity_value = float(quantity_text)
            except ValueError:
                continue

            unit = match.group("unit").casefold()
            if unit in {"kg", "kgs", "kilogram", "kilograms"}:
                quantity_tonnes = quantity_value / 1000.0
            else:
                quantity_tonnes = quantity_value

            candidates.append((priority, match.start(), round(quantity_tonnes, 3)))

    if not candidates:
        return None

    candidates.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return candidates[0][2]


def _discover_waste_kpi_workbooks() -> List[Path]:
    """Return the Excel workbooks filed under File 1 waste reports."""

    workbooks: List[Path] = []
    for pattern in ("*.xls", "*.xlsx"):
        workbooks.extend(config.WASTE_REPORTS_DESTINATION.glob(pattern))
    return sorted(
        path.resolve()
        for path in workbooks
        if path.is_file()
    )


def _read_waste_kpi_sheet_metadata(workbook_path: Path) -> WasteKpiSheetMetadata:
    """Parse workbook header values used by the File 1 waste register."""

    client_name = ""
    site_address = ""
    project_number = ""
    manager_name = ""

    for row in _iter_excel_sheet_rows(workbook_path):
        line = " | ".join(value for value in row if value)
        if not line:
            continue

        if not client_name and any(value.casefold() == "customer" for value in row):
            client_name = _value_after_matching_cell(row, "customer")
        if not site_address and "project name & address" in line.casefold():
            site_address = _value_after_matching_cell(row, "project name & address")
        if (
            not manager_name
            and "person responsible for waste management on site" in line.casefold()
        ):
            manager_name = _value_after_matching_cell(
                row,
                "person responsible for waste management on site",
            )
        if not project_number:
            project_number = _extract_project_number_from_row(row)

    return WasteKpiSheetMetadata(
        workbook_path=workbook_path.resolve(),
        client_name=client_name,
        site_address=site_address,
        project_number=project_number,
        manager_name=manager_name,
    )


def _iter_excel_sheet_rows(workbook_path: Path) -> Iterable[List[str]]:
    """Yield stripped text rows from the first worksheet of an Excel file."""

    suffix = workbook_path.suffix.lower()
    if suffix == ".xls":
        try:
            import xlrd
        except ImportError as exc:
            raise RuntimeError("xlrd is required to read .xls waste KPI workbooks.") from exc

        workbook = xlrd.open_workbook(workbook_path)
        sheet = workbook.sheet_by_index(0)
        for row_index in range(sheet.nrows):
            yield [
                str(value).strip()
                for value in sheet.row_values(row_index)
            ]
        return

    if suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("openpyxl is required to read .xlsx waste KPI workbooks.") from exc

        workbook = load_workbook(workbook_path, read_only=True, data_only=True)
        sheet = workbook.worksheets[0]
        for row in sheet.iter_rows(values_only=True):
            yield [str(value).strip() if value is not None else "" for value in row]
        return

    raise RuntimeError(f"Unsupported waste KPI workbook type: {workbook_path.suffix}")


def _next_non_empty_value(row: List[str], start_index: int) -> str:
    """Return the first non-empty cell from one row after the given index."""

    for value in row[start_index:]:
        cleaned_value = value.strip()
        if cleaned_value:
            return cleaned_value
    return ""


def _value_after_matching_cell(row: List[str], search_text: str) -> str:
    """Return the first non-empty cell after the matched label text."""

    lowered_search_text = search_text.casefold()
    for index, value in enumerate(row):
        if lowered_search_text in value.casefold():
            return _next_non_empty_value(row, index + 1)
    return ""


def _extract_project_number_from_row(row: List[str]) -> str:
    """Return a workbook project or job number when present in one row."""

    for candidate_label in ("project number", "project no", "job number", "job no"):
        project_number = _value_after_matching_cell(row, candidate_label)
        if project_number:
            return project_number

    line = " | ".join(value for value in row if value)
    for pattern in PROJECT_NUMBER_PATTERNS:
        match = pattern.search(line)
        if match is not None:
            return match.group("value").strip()
    return ""


def _score_waste_kpi_sheet_metadata(
    metadata: WasteKpiSheetMetadata,
    *,
    site_name: Optional[str],
    site_address: str,
) -> int:
    """Return a simple match score between workbook metadata and the active project."""

    score = 0
    candidate_text = " ".join(
        value
        for value in (
            metadata.client_name,
            metadata.site_address,
            metadata.project_number,
            metadata.workbook_path.name if metadata.workbook_path is not None else "",
        )
        if value
    ).casefold()
    if site_name:
        for token in re.findall(r"[a-z0-9]+", site_name.casefold()):
            if token and token in candidate_text:
                score += 2
    if site_address:
        for token in re.findall(r"[a-z0-9]+", site_address.casefold()):
            if token and token in candidate_text:
                score += 1
    return score


def _extract_waste_transfer_note_text(source_path: Path) -> str:
    """Return best-effort text from a WTN PDF or image."""

    suffix = source_path.suffix.lower()
    if suffix == ".pdf":
        try:
            embedded_text = "\n".join(_extract_pdf_text_pages(source_path)).strip()
        except RuntimeError:
            embedded_text = ""
        ocr_text = ""
        rendered_image_path = _render_pdf_first_page_to_image(source_path)
        if rendered_image_path is not None:
            try:
                ocr_text = _extract_image_text(rendered_image_path)
            finally:
                try:
                    rendered_image_path.unlink()
                    rendered_image_path.parent.rmdir()
                except OSError:
                    pass
        return "\n".join(
            text for text in (embedded_text, ocr_text) if text.strip()
        ).strip()

    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return _extract_image_text(source_path)

    return ""


def _render_pdf_first_page_to_image(pdf_path: Path) -> Optional[Path]:
    """Render the first PDF page to a temporary PNG for OCR fallback."""

    try:
        import fitz
    except ModuleNotFoundError:
        return None

    try:
        with fitz.open(pdf_path) as document:
            if document.page_count == 0:
                return None
            page = document.load_page(0)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            temp_dir = Path(tempfile.mkdtemp(prefix="uplands-wtn-ocr-"))
            output_path = temp_dir / f"{pdf_path.stem}.png"
            pixmap.save(str(output_path))
            return output_path
    except Exception:
        return None


def _extract_image_text(image_path: Path) -> str:
    """OCR an image using macOS Vision first, then optional pytesseract."""

    vision_text = _extract_image_text_with_macos_vision(image_path)
    if vision_text:
        return vision_text
    return _extract_image_text_with_pytesseract(image_path)


def _extract_image_text_with_macos_vision(image_path: Path) -> str:
    """Use the macOS Vision framework to OCR one image when available."""

    swift_source = """
import Foundation
import AppKit
import Vision

let imagePath = CommandLine.arguments[1]
let url = URL(fileURLWithPath: imagePath)
guard let image = NSImage(contentsOf: url) else { exit(2) }
var proposedRect = NSRect(origin: .zero, size: image.size)
guard let cgImage = image.cgImage(forProposedRect: &proposedRect, context: nil, hints: nil) else { exit(3) }
let request = VNRecognizeTextRequest()
request.recognitionLevel = .accurate
request.usesLanguageCorrection = true
let handler = VNImageRequestHandler(cgImage: cgImage, options: [:])
try handler.perform([request])
let observations = request.results as? [VNRecognizedTextObservation] ?? []
let lines = observations.compactMap { $0.topCandidates(1).first?.string }
print(lines.joined(separator: "\\n"))
"""
    try:
        completed_process = subprocess.run(
            ["swift", "-", str(image_path)],
            input=swift_source,
            text=True,
            capture_output=True,
            check=False,
        )
    except OSError:
        return ""

    if completed_process.returncode != 0:
        return ""
    return completed_process.stdout.strip()


def _extract_image_text_with_pytesseract(image_path: Path) -> str:
    """Fallback OCR path when pytesseract and the Tesseract binary are available."""

    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return ""

    try:
        return pytesseract.image_to_string(Image.open(image_path)).strip()
    except Exception:
        return ""


def _extract_vehicle_registration(normalized_text: str) -> str:
    """Return the most likely vehicle registration from normalized WTN text."""

    for pattern in VEHICLE_REG_PRIORITY_PATTERNS:
        match = pattern.search(normalized_text)
        if match is not None:
            return re.sub(r"\s+", " ", match.group("reg").upper()).strip()
    return ""


def _extract_vehicle_registration_from_pdf(
    pdf_path: Path,
    normalized_text: str,
    normalized_embedded_pdf_text: str,
) -> str:
    """Return the vehicle field value from a PDF ticket when available."""

    explicit_text_candidate = _extract_vehicle_registration(
        normalized_embedded_pdf_text
    )
    if explicit_text_candidate:
        return explicit_text_candidate

    field_value = _extract_labeled_pdf_field(pdf_path, "Vehicle")
    if _is_vehicle_field_label(field_value):
        field_value = ""
    normalized_field_value = _extract_uk_vehicle_registration_candidate(field_value)
    if normalized_field_value:
        return normalized_field_value

    region_value = _extract_vehicle_field_text_from_pdf_region(pdf_path, "Vehicle")
    normalized_region_value = _extract_uk_vehicle_registration_candidate(region_value)
    if normalized_region_value:
        return normalized_region_value

    normalized_region_value = _extract_targeted_vehicle_registration_from_ocr(
        pdf_path,
        "Vehicle",
    )
    if normalized_region_value:
        return normalized_region_value

    return ""


def _extract_ticket_number_from_pdf(
    pdf_path: Path,
    normalized_text: str,
) -> str:
    """Return the ticket number from a PDF field before falling back to filename text."""

    field_value = _extract_labeled_pdf_field(pdf_path, "Ticket No.")
    if field_value:
        cleaned_value = re.sub(r"[^A-Z0-9/_-]+", "", field_value.upper())
        if cleaned_value:
            return cleaned_value
    return _derive_waste_transfer_note_number(pdf_path, normalized_text)


def _extract_labeled_pdf_field(pdf_path: Path, label_text: str) -> str:
    """Return the block value immediately to the right of a known PDF label."""

    try:
        import fitz
    except ModuleNotFoundError:
        return ""

    lowered_label_text = label_text.casefold()
    try:
        with fitz.open(pdf_path) as document:
            for page_number in range(min(2, document.page_count)):
                page = document.load_page(page_number)
                blocks = page.get_text("blocks")
                for label_block in blocks:
                    if lowered_label_text not in str(label_block[4]).casefold():
                        continue
                    label_y = (float(label_block[1]) + float(label_block[3])) / 2.0
                    candidates: List[Tuple[float, str]] = []
                    for candidate_block in blocks:
                        if candidate_block == label_block:
                            continue
                        candidate_text = " ".join(str(candidate_block[4]).split()).strip()
                        if not candidate_text:
                            continue
                        if lowered_label_text in candidate_text.casefold():
                            continue
                        candidate_y = (
                            float(candidate_block[1]) + float(candidate_block[3])
                        ) / 2.0
                        if abs(candidate_y - label_y) > 16:
                            continue
                        if float(candidate_block[0]) < float(label_block[2]) - 4:
                            continue
                        candidates.append(
                            (
                                float(candidate_block[0]) - float(label_block[2]),
                                candidate_text,
                            )
                        )
                    if candidates:
                        candidates.sort(key=lambda item: item[0])
                        return candidates[0][1]
    except Exception:
        return ""
    return ""


def _extract_vehicle_field_text_from_pdf_region(pdf_path: Path, label_text: str) -> str:
    """Return text found inside the value area of the Abacus vehicle field."""

    label_block = _find_pdf_label_block(pdf_path, label_text)
    if label_block is None:
        return ""

    try:
        import fitz
    except ModuleNotFoundError:
        return ""

    try:
        with fitz.open(pdf_path) as document:
            page = document.load_page(label_block["page_number"])
            value_rect = _build_vehicle_value_rect(page, label_block)
            words = page.get_text("words", clip=value_rect)
    except Exception:
        return ""

    ordered_tokens = [
        str(word[4]).strip()
        for word in sorted(words, key=lambda item: (item[1], item[0]))
        if len(word) >= 5 and str(word[4]).strip()
    ]
    candidate_text = " ".join(ordered_tokens).strip()
    if _is_vehicle_field_label(candidate_text):
        return ""
    return candidate_text


def _ocr_labeled_pdf_region(pdf_path: Path, label_text: str) -> str:
    """OCR a high-contrast crop around one known PDF field label."""

    label_block = _find_pdf_label_block(pdf_path, label_text)
    if label_block is None:
        return ""

    try:
        import fitz
    except ModuleNotFoundError:
        return ""

    try:
        with fitz.open(pdf_path) as document:
            page = document.load_page(label_block["page_number"])
            clip_rect = fitz.Rect(
                label_block["x0"] - 8,
                label_block["y0"] - 12,
                min(page.rect.width * 0.62, label_block["x1"] + 250),
                label_block["y1"] + 24,
            )
            temp_dir = Path(tempfile.mkdtemp(prefix="uplands-vehicle-ocr-"))
            source_image_path = temp_dir / f"{pdf_path.stem}-vehicle.png"
            page.get_pixmap(
                matrix=fitz.Matrix(12, 12),
                clip=clip_rect,
                alpha=False,
            ).save(str(source_image_path))
    except Exception:
        return ""

    try:
        return _extract_high_contrast_ocr_text(source_image_path)
    finally:
        try:
            source_image_path.unlink()
            source_image_path.parent.rmdir()
        except OSError:
            pass


def _extract_targeted_vehicle_registration_from_ocr(
    pdf_path: Path,
    label_text: str,
) -> str:
    """Return a UK reg candidate only when OCR repeats it across variants."""

    label_block = _find_pdf_label_block(pdf_path, label_text)
    if label_block is None:
        return ""
    if not _vehicle_field_has_visible_content(pdf_path, label_block):
        return ""

    try:
        import fitz
    except ModuleNotFoundError:
        return ""

    try:
        with fitz.open(pdf_path) as document:
            page = document.load_page(label_block["page_number"])
            clip_rect = _build_vehicle_value_rect(page, label_block)
            temp_dir = Path(tempfile.mkdtemp(prefix="uplands-vehicle-ocr-"))
            source_image_path = temp_dir / f"{pdf_path.stem}-vehicle.png"
            page.get_pixmap(
                matrix=fitz.Matrix(12, 12),
                clip=clip_rect,
                alpha=False,
            ).save(str(source_image_path))
    except Exception:
        return ""

    try:
        candidate_counts: Dict[str, int] = {}
        for text_segment in _extract_high_contrast_ocr_segments(source_image_path):
            candidate = _extract_uk_vehicle_registration_candidate(text_segment)
            if not candidate:
                continue
            candidate_counts[candidate] = candidate_counts.get(candidate, 0) + 1
        if not candidate_counts:
            return ""
        best_candidate, best_count = max(
            candidate_counts.items(),
            key=lambda item: (item[1], item[0]),
        )
        return best_candidate if best_count >= 2 else ""
    finally:
        try:
            source_image_path.unlink()
            source_image_path.parent.rmdir()
        except OSError:
            pass


def _vehicle_field_has_visible_content(
    pdf_path: Path,
    label_block: Dict[str, Any],
) -> bool:
    """Return True when the interior of the vehicle field contains non-border ink."""

    try:
        import fitz
        from PIL import Image, ImageOps
    except ImportError:
        return True

    try:
        with fitz.open(pdf_path) as document:
            page = document.load_page(label_block["page_number"])
            content_rect = _build_vehicle_value_rect(page, label_block)
            temp_dir = Path(tempfile.mkdtemp(prefix="uplands-vehicle-content-"))
            content_image_path = temp_dir / f"{pdf_path.stem}-vehicle-content.png"
            page.get_pixmap(
                matrix=fitz.Matrix(20, 20),
                clip=content_rect,
                alpha=False,
            ).save(str(content_image_path))
    except Exception:
        return True

    try:
        grayscale_image = ImageOps.grayscale(Image.open(content_image_path))
        pixel_values = list(grayscale_image.getdata())
        if not pixel_values:
            return False
        dark_ratio = sum(1 for value in pixel_values if value < 220) / len(pixel_values)
        return dark_ratio >= 0.03
    finally:
        try:
            content_image_path.unlink()
            content_image_path.parent.rmdir()
        except OSError:
            pass


def _build_vehicle_value_rect(page: Any, label_block: Dict[str, Any]) -> Any:
    """Return the clipped area where the Abacus vehicle registration is expected."""

    try:
        import fitz
    except ModuleNotFoundError:
        raise RuntimeError("PyMuPDF is required to compute PDF clip regions.")

    return fitz.Rect(
        label_block["x1"] + 6,
        label_block["y0"] - 6,
        min(page.rect.width * 0.50, label_block["x1"] + 140),
        label_block["y1"] + 8,
    )


def _find_pdf_label_block(
    pdf_path: Path,
    label_text: str,
) -> Optional[Dict[str, Any]]:
    """Return the PDF text block for a known label."""

    try:
        import fitz
    except ModuleNotFoundError:
        return None

    lowered_label_text = label_text.casefold()
    try:
        with fitz.open(pdf_path) as document:
            for page_number in range(min(2, document.page_count)):
                page = document.load_page(page_number)
                for block in page.get_text("blocks"):
                    block_text = " ".join(str(block[4]).split())
                    if lowered_label_text not in block_text.casefold():
                        continue
                    return {
                        "page_number": page_number,
                        "x0": float(block[0]),
                        "y0": float(block[1]),
                        "x1": float(block[2]),
                        "y1": float(block[3]),
                    }
    except Exception:
        return None
    return None


def _extract_high_contrast_ocr_text(image_path: Path) -> str:
    """OCR an image and a few high-contrast variants, returning combined text."""

    return "\n".join(_extract_high_contrast_ocr_segments(image_path))


def _extract_high_contrast_ocr_segments(image_path: Path) -> List[str]:
    """OCR an image and high-contrast variants, returning non-empty segments."""

    try:
        from PIL import Image, ImageEnhance, ImageOps
    except ImportError:
        fallback_text = _extract_image_text(image_path)
        return [fallback_text] if fallback_text.strip() else []

    source_image = Image.open(image_path)
    grayscale_image = ImageOps.grayscale(source_image)
    image_variants = [
        source_image,
        grayscale_image,
        ImageOps.autocontrast(grayscale_image),
        ImageEnhance.Contrast(grayscale_image).enhance(3.0),
        grayscale_image.point(lambda pixel: 255 if pixel > 170 else 0),
        grayscale_image.point(lambda pixel: 255 if pixel > 150 else 0),
    ]

    extracted_segments: List[str] = []
    temp_dir = Path(tempfile.mkdtemp(prefix="uplands-ocr-variants-"))
    try:
        for index, image_variant in enumerate(image_variants):
            variant_path = temp_dir / f"variant-{index}.png"
            image_variant.save(variant_path)
            variant_text = _extract_image_text(variant_path)
            if variant_text.strip():
                extracted_segments.append(variant_text.strip())
    finally:
        for temp_file in temp_dir.glob("*"):
            try:
                temp_file.unlink()
            except OSError:
                pass
        try:
            temp_dir.rmdir()
        except OSError:
            pass

    return extracted_segments


def _is_vehicle_field_label(field_value: str) -> bool:
    """Return True when the extracted vehicle field is actually another label."""

    lowered_value = field_value.strip().casefold()
    if not lowered_value:
        return True
    disallowed_tokens = (
        "notes",
        "skip type",
        "movement type",
        "payment type",
        "driver",
        "weight",
        "date",
        "order number",
        "ticket no",
    )
    if any(token == lowered_value or token in lowered_value for token in disallowed_tokens):
        return True
    return False


def _is_plausible_vehicle_registration(field_value: str) -> bool:
    """Return True when one field looks like a real UK registration mark."""

    return bool(
        re.fullmatch(
            r"[A-Z]{2}\d{2}\s?[A-Z]{3}",
            field_value.strip().upper(),
        )
    )


def _extract_uk_vehicle_registration_candidate(raw_text: str) -> str:
    """Return one normalized UK-style vehicle registration candidate from OCR text."""

    if not raw_text:
        return ""

    explicit_match = re.search(r"\b([A-Z]{2}\d{2}\s?[A-Z]{3})\b", raw_text.upper())
    if explicit_match is not None:
        return explicit_match.group(1).replace("  ", " ").strip()

    compact_text = re.sub(r"[^A-Z0-9]+", " ", raw_text.upper())
    for token in compact_text.split():
        if len(token) < 6 or len(token) > 8:
            continue
        if not any(character.isdigit() for character in token):
            continue
        candidate = _normalize_vehicle_registration_token(token)
        if _is_plausible_vehicle_registration(candidate):
            return candidate

    compact_joined_text = re.sub(r"[^A-Z0-9]+", "", raw_text.upper())
    for start_index in range(0, max(0, len(compact_joined_text) - 6)):
        raw_token = compact_joined_text[start_index : start_index + 7]
        if not any(character.isdigit() for character in raw_token):
            continue
        candidate = _normalize_vehicle_registration_token(
            raw_token
        )
        if _is_plausible_vehicle_registration(candidate):
            return candidate
    return ""


def _normalize_vehicle_registration_token(token: str) -> str:
    """Normalize OCR confusions into a standard UK registration candidate."""

    cleaned_token = re.sub(r"[^A-Z0-9]+", "", token.upper())[:7]
    if len(cleaned_token) < 7:
        return ""

    letter_substitutions = {
        "0": "O",
        "1": "I",
        "2": "Z",
        "5": "S",
        "8": "B",
    }
    digit_substitutions = {
        "O": "0",
        "Q": "0",
        "D": "0",
        "I": "1",
        "L": "1",
        "Z": "2",
        "S": "5",
        "B": "8",
    }

    normalized_characters: List[str] = []
    for index, character in enumerate(cleaned_token):
        if index in (0, 1, 4, 5, 6):
            normalized_characters.append(letter_substitutions.get(character, character))
        else:
            normalized_characters.append(digit_substitutions.get(character, character))

    normalized_token = "".join(normalized_characters)
    if not _is_plausible_vehicle_registration(normalized_token):
        return ""
    return f"{normalized_token[:4]} {normalized_token[4:]}"


def _derive_waste_transfer_note_number(source_path: Path, normalized_text: str) -> str:
    """Return a stable waste note reference from the file name or scanned text."""

    if source_path.stem.strip():
        return source_path.stem.strip()

    match = re.search(
        r"\bticket\s*no\.?\s*(?P<value>[A-Z0-9/_-]+)\b",
        normalized_text,
        re.IGNORECASE,
    )
    if match is not None:
        return match.group("value").strip()

    match = re.search(
        r"\b(?:wtn|waste\s+transfer\s+note|ticket)\s*(?:number|no\.?)?\s*(?::|-)?\s*(?P<value>[A-Z0-9/_-]+)\b",
        normalized_text,
        re.IGNORECASE,
    )
    if match is not None:
        return match.group("value").strip()
    return f"WTN-{datetime.now():%Y%m%d%H%M%S}"


def _is_carrier_compliance_pdf(source_path: Path) -> bool:
    """Return True when the filename suggests a carrier licence or insurance PDF."""

    if source_path.suffix.lower() != ".pdf":
        return False

    lowered_name = source_path.name.lower()
    return "insurance" in lowered_name or "carrier" in lowered_name


def _extract_pdf_text_pages(pdf_path: Path) -> List[str]:
    """Return plain text from the first two pages of a PDF."""

    try:
        import fitz
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "PyMuPDF is required for carrier PDF expiry extraction."
        ) from exc

    try:
        with fitz.open(pdf_path) as document:
            return [
                document.load_page(page_number).get_text("text")
                for page_number in range(min(2, document.page_count))
            ]
    except Exception as exc:  # pragma: no cover - fitz error hierarchy is implementation-specific
        raise RuntimeError(f"Unable to read PDF text from {pdf_path}.") from exc


def _extract_expiry_candidates_from_text(text: str) -> List[tuple[int, date]]:
    """Return scored expiry candidates extracted from one text source."""

    normalized_text = _normalize_text(text)
    candidates: List[tuple[int, date]] = []

    for priority, pattern in PRIORITY_DATE_PATTERNS:
        for match in pattern.finditer(normalized_text):
            date_text = match.group("date")
            parsed_date = _parse_date_string(date_text)
            if parsed_date is None:
                continue
            if _has_ignored_date_context(
                normalized_text,
                match.start("date"),
                match.end("date"),
            ):
                continue
            candidates.append((priority, parsed_date))

    generic_dates = []
    if any(keyword in normalized_text.casefold() for keyword in LOW_PRIORITY_KEYWORDS):
        generic_dates = [
            parsed_date
            for parsed_date, start_index, end_index in _extract_dates_with_positions(
                normalized_text
            )
            if not _has_ignored_date_context(normalized_text, start_index, end_index)
        ]
        if generic_dates:
            candidates.append((20, max(generic_dates)))

    if not candidates:
        fallback_dates = [
            parsed_date
            for parsed_date, start_index, end_index in _extract_dates_with_positions(
                normalized_text
            )
            if not _has_ignored_date_context(normalized_text, start_index, end_index)
        ]
        if fallback_dates:
            candidates.append((10, max(fallback_dates)))

    return candidates


def _extract_dates_with_positions(text: str) -> List[tuple[date, int, int]]:
    """Return parsed dates together with their positions in the normalized text."""

    candidates: List[tuple[date, int, int]] = []

    for match in NUMERIC_DATE_PATTERN.finditer(text):
        parsed_date = _parse_date_string(match.group(0))
        if parsed_date is None:
            continue
        candidates.append((parsed_date, match.start(), match.end()))

    for match in TEXTUAL_DATE_PATTERN.finditer(text):
        parsed_date = _parse_date_string(match.group(0))
        if parsed_date is None:
            continue
        candidates.append((parsed_date, match.start(), match.end()))

    return candidates


def _parse_date_string(date_text: str) -> Optional[date]:
    """Parse one numeric or textual date string into a ``date``."""

    numeric_match = NUMERIC_DATE_PATTERN.fullmatch(date_text.strip())
    if numeric_match:
        day_text, month_text, year_text = numeric_match.groups()
        try:
            year_value = int(year_text)
            if len(year_text) == 2:
                year_value += 2000
            return date(year_value, int(month_text), int(day_text))
        except ValueError:
            return None

    textual_match = TEXTUAL_DATE_PATTERN.fullmatch(date_text.strip())
    if textual_match:
        day_text, month_name, year_text = textual_match.groups()
        try:
            return date(
                int(year_text),
                MONTH_NAME_MAP[month_name.casefold()],
                int(day_text),
            )
        except ValueError:
            return None

    return None


def _normalize_text(text: str) -> str:
    """Collapse whitespace to simplify phrase and date matching."""

    return re.sub(r"\s+", " ", text).strip()


def _has_ignored_date_context(text: str, start_index: int, end_index: int) -> bool:
    """Return True when a candidate date is tied to document generation metadata."""

    context_start = max(0, start_index - 40)
    context_end = min(len(text), end_index + 25)
    context = text[context_start:context_end].casefold()
    return any(keyword in context for keyword in IGNORE_DATE_CONTEXT_KEYWORDS)


def _upsert_carrier_compliance_document(
    repository: DocumentRepository,
    *,
    carrier_name: str,
    carrier_document_type: CarrierComplianceDocumentType,
    expiry_date: date,
    source_path: Path,
) -> CarrierComplianceDocument:
    """Create or update the carrier compliance record from an ingested PDF."""

    existing_document = _get_carrier_compliance_document(
        repository,
        carrier_name,
        carrier_document_type,
    )
    if existing_document is None:
        existing_document = _get_carrier_compliance_document_by_reference(
            repository,
            reference_number=source_path.stem,
            carrier_document_type=carrier_document_type,
        )
    document = CarrierComplianceDocument(
        doc_id=(
            existing_document.doc_id
            if existing_document is not None
            else _build_carrier_compliance_doc_id(carrier_name, carrier_document_type)
        ),
        site_name=(
            existing_document.site_name
            if existing_document is not None
            else _infer_default_site_name(repository)
        ),
        created_at=(
            existing_document.created_at
            if existing_document is not None
            else datetime.now()
        ),
        status=DocumentStatus.ACTIVE,
        carrier_name=carrier_name,
        carrier_document_type=carrier_document_type,
        reference_number=source_path.stem,
        expiry_date=expiry_date,
    )
    repository.save(document)
    _archive_duplicate_carrier_reference_documents(
        repository,
        canonical_document=document,
    )
    return document


def _upsert_waste_transfer_note_document(
    repository: DocumentRepository,
    pdf_path: Path,
) -> Optional[WasteTransferNoteDocument]:
    """Create or update a waste transfer note from a synced ticket PDF."""

    try:
        source_candidate = _build_waste_transfer_note_source_candidate(
            repository,
            pdf_path,
        )
    except RuntimeError:
        return None

    return _upsert_waste_transfer_note_document_from_candidate(
        repository,
        source_candidate,
    )


def _build_waste_transfer_note_source_candidate(
    repository: DocumentRepository,
    pdf_path: Path,
) -> WasteTransferNoteSourceCandidate:
    """Scan one physical waste-ticket source file into a comparable candidate."""

    scanned_waste_transfer_note = smart_scan_waste_transfer_note(
        repository,
        source_path=pdf_path,
    )
    normalized_text = _normalize_text(scanned_waste_transfer_note.extracted_text)
    existing_document = _get_waste_transfer_note_document(
        repository,
        scanned_waste_transfer_note.wtn_number,
        ticket_date=scanned_waste_transfer_note.ticket_date,
        collection_type=scanned_waste_transfer_note.collection_type,
    )
    site_name = (
        existing_document.site_name
        if existing_document is not None
        else _infer_waste_ticket_site_name(normalized_text, repository)
    )
    return WasteTransferNoteSourceCandidate(
        source_path=pdf_path,
        scanned_note=scanned_waste_transfer_note,
        site_name=site_name,
        source_mtime=pdf_path.stat().st_mtime if pdf_path.exists() else 0.0,
    )


def _build_waste_transfer_note_source_snapshot(
    source_candidate: WasteTransferNoteSourceCandidate,
) -> WasteTransferNoteSourceSnapshot:
    """Persist the useful parts of one physical WTN source file."""

    scanned_note = source_candidate.scanned_note
    return WasteTransferNoteSourceSnapshot(
        source_path=str(source_candidate.source_path.resolve()),
        source_file_name=source_candidate.source_path.name,
        ticket_date=scanned_note.ticket_date,
        collection_type=scanned_note.collection_type,
        quantity_tonnes=scanned_note.quantity_tonnes,
        carrier_name=scanned_note.carrier_name,
        vehicle_registration=scanned_note.vehicle_registration,
        waste_description=scanned_note.waste_description,
        ewc_code=scanned_note.ewc_code,
        destination_facility=scanned_note.destination_facility,
    )


def _build_waste_transfer_note_source_candidate_from_snapshot(
    *,
    waste_note: WasteTransferNoteDocument,
    source_snapshot: WasteTransferNoteSourceSnapshot,
) -> Optional[WasteTransferNoteSourceCandidate]:
    """Rehydrate one WTN source candidate from persisted sync metadata."""

    source_path = Path(source_snapshot.source_path)
    if not source_path.exists():
        return None

    scanned_note = SmartScannedWasteTransferNote(
        source_name=source_snapshot.source_file_name,
        wtn_number=waste_note.wtn_number,
        carrier_name=source_snapshot.carrier_name or waste_note.carrier_name,
        vehicle_registration=(
            source_snapshot.vehicle_registration or waste_note.vehicle_registration
        ),
        collection_type=source_snapshot.collection_type,
        waste_description=(
            source_snapshot.waste_description or waste_note.waste_description
        ),
        ticket_date=source_snapshot.ticket_date,
        quantity_tonnes=source_snapshot.quantity_tonnes,
        ewc_code=source_snapshot.ewc_code or waste_note.ewc_code,
        destination_facility=(
            source_snapshot.destination_facility or waste_note.destination_facility
        ),
        extracted_text="",
    )
    return WasteTransferNoteSourceCandidate(
        source_path=source_path.resolve(),
        scanned_note=scanned_note,
        site_name=waste_note.site_name,
        source_mtime=source_path.stat().st_mtime if source_path.exists() else 0.0,
    )


def _is_tanker_collection_type(collection_type: str) -> bool:
    """Return True when the scanned collection type should be treated as a tanker run."""

    normalized_collection_type = collection_type.casefold()
    return "tanker" in normalized_collection_type


def _build_waste_transfer_note_identity_key(
    *,
    wtn_number: str,
    ticket_date: date,
    collection_type: str = "",
) -> str:
    """Return the logical identity key for one waste ticket."""

    if _is_tanker_collection_type(collection_type):
        return f"{wtn_number}::{ticket_date.isoformat()}"
    return wtn_number


def _build_waste_transfer_note_doc_id(
    *,
    wtn_number: str,
    ticket_date: date,
    collection_type: str = "",
) -> str:
    """Return the persisted doc id for one waste ticket."""

    if _is_tanker_collection_type(collection_type):
        return f"WTN-{_slugify_identifier(wtn_number)}-{ticket_date.isoformat()}"
    return f"WTN-{wtn_number}"


def _waste_transfer_note_source_candidate_sort_key(
    source_candidate: WasteTransferNoteSourceCandidate,
) -> tuple[int, int, int, float, int, str]:
    """Return a descending sort key for selecting the canonical WTN source."""

    scanned_note = source_candidate.scanned_note
    extracted_text = _normalize_text(scanned_note.extracted_text)
    return (
        1 if (scanned_note.quantity_tonnes or 0.0) > 0 else 0,
        1 if _is_plausible_vehicle_registration(scanned_note.vehicle_registration) else 0,
        len(extracted_text),
        source_candidate.source_mtime,
        scanned_note.ticket_date.toordinal(),
        source_candidate.source_path.name.casefold(),
    )


def _select_canonical_waste_transfer_note_source(
    source_candidates: Iterable[WasteTransferNoteSourceCandidate],
) -> WasteTransferNoteSourceCandidate:
    """Choose the single source file that should drive the live WTN record."""

    resolved_candidates = list(source_candidates)
    if not resolved_candidates:
        raise ValueError("At least one waste-ticket source candidate is required.")
    return max(resolved_candidates, key=_waste_transfer_note_source_candidate_sort_key)


def _select_effective_waste_transfer_note_source(
    repository: DocumentRepository,
    source_candidates: Iterable[WasteTransferNoteSourceCandidate],
) -> WasteTransferNoteSourceCandidate:
    """Return the chosen source file, honoring any saved manual override first."""

    resolved_candidates = list(source_candidates)
    if not resolved_candidates:
        raise ValueError("At least one waste-ticket source candidate is required.")

    existing_document = _get_waste_transfer_note_document(
        repository,
        resolved_candidates[0].scanned_note.wtn_number,
        ticket_date=resolved_candidates[0].scanned_note.ticket_date,
        collection_type=resolved_candidates[0].scanned_note.collection_type,
    )
    if existing_document is not None and existing_document.source_file_override_path:
        override_path = Path(existing_document.source_file_override_path).resolve()
        for source_candidate in resolved_candidates:
            if source_candidate.source_path.resolve() == override_path:
                return source_candidate

    return _select_canonical_waste_transfer_note_source(resolved_candidates)


def _upsert_waste_transfer_note_document_from_candidate(
    repository: DocumentRepository,
    source_candidate: WasteTransferNoteSourceCandidate,
    *,
    source_file_override_path: Optional[Path] = None,
    source_candidate_snapshots: Optional[List[WasteTransferNoteSourceSnapshot]] = None,
) -> WasteTransferNoteDocument:
    """Create or update a waste transfer note from one chosen source candidate."""

    scanned_waste_transfer_note = source_candidate.scanned_note
    wtn_number = scanned_waste_transfer_note.wtn_number
    existing_document = _get_waste_transfer_note_document(
        repository,
        wtn_number,
        ticket_date=scanned_waste_transfer_note.ticket_date,
        collection_type=scanned_waste_transfer_note.collection_type,
    )
    quantity_tonnes = (
        scanned_waste_transfer_note.quantity_tonnes
        if scanned_waste_transfer_note.quantity_tonnes is not None
        else (
            existing_document.quantity_tonnes
            if existing_document is not None
            else 0.0
        )
    )
    tonnage_review_status = (
        existing_document.tonnage_review_status
        if existing_document is not None and quantity_tonnes <= 0
        else ""
    )
    document = WasteTransferNoteDocument(
        doc_id=(
            existing_document.doc_id
            if existing_document is not None
            else _build_waste_transfer_note_doc_id(
                wtn_number=wtn_number,
                ticket_date=scanned_waste_transfer_note.ticket_date,
                collection_type=scanned_waste_transfer_note.collection_type,
            )
        ),
        site_name=(
            existing_document.site_name
            if existing_document is not None
            else source_candidate.site_name
        ),
        created_at=(
            existing_document.created_at
            if existing_document is not None
            else datetime.now()
        ),
        status=DocumentStatus.ACTIVE,
        wtn_number=wtn_number,
        date=scanned_waste_transfer_note.ticket_date,
        waste_description=(
            existing_document.waste_description
            if existing_document is not None and existing_document.waste_description
            else scanned_waste_transfer_note.waste_description
        ),
        ewc_code=(
            existing_document.ewc_code
            if existing_document is not None and existing_document.ewc_code
            else scanned_waste_transfer_note.ewc_code
        ),
        quantity_tonnes=quantity_tonnes,
        carrier_name=(
            existing_document.carrier_name
            if existing_document is not None
            else scanned_waste_transfer_note.carrier_name
        ),
        destination_facility=(
            existing_document.destination_facility
            if existing_document is not None and existing_document.destination_facility
            else scanned_waste_transfer_note.destination_facility
        ),
        vehicle_registration=(
            scanned_waste_transfer_note.vehicle_registration
            if _is_plausible_vehicle_registration(
                scanned_waste_transfer_note.vehicle_registration
            )
            else ""
        ),
        source_file_override_path=(
            str(source_file_override_path.resolve())
            if source_file_override_path is not None
            else (
                existing_document.source_file_override_path
                if existing_document is not None
                else ""
            )
        ),
        canonical_source_path=str(source_candidate.source_path.resolve()),
        source_conflict_candidates=(
            list(source_candidate_snapshots)
            if source_candidate_snapshots is not None
            else (
                existing_document.source_conflict_candidates
                if existing_document is not None
                else []
            )
        ),
        tonnage_review_status=tonnage_review_status,
    )
    repository.save(document)
    return document


def set_waste_transfer_note_source_override(
    repository: DocumentRepository,
    *,
    source_document: WasteTransferNoteDocument,
    source_path: Path,
) -> LoggedWasteTransferNote:
    """Persist a manual source-file override for one WTN and refresh its live row."""

    repository.create_schema()
    source_candidate = _build_waste_transfer_note_source_candidate(repository, source_path)
    if source_candidate.scanned_note.wtn_number != source_document.wtn_number:
        raise ValidationError(
            "The selected source file does not match the chosen waste transfer note."
        )

    refreshed_waste_transfer_note = _upsert_waste_transfer_note_document_from_candidate(
        repository,
        source_candidate,
        source_file_override_path=source_candidate.source_path,
        source_candidate_snapshots=(
            source_document.source_conflict_candidates
            if source_document.source_conflict_candidates
            else [_build_waste_transfer_note_source_snapshot(source_candidate)]
        ),
    )
    register_document = _upsert_site_waste_register(
        repository,
        site_name=refreshed_waste_transfer_note.site_name,
    )
    return LoggedWasteTransferNote(
        waste_transfer_note=refreshed_waste_transfer_note,
        stored_file_path=source_candidate.source_path,
        register_document=register_document,
    )


def _list_waste_transfer_note_source_candidates(
    repository: DocumentRepository,
    *,
    waste_destination: Path,
) -> List[WasteTransferNoteSourceCandidate]:
    """Return all scannable waste-ticket source files currently filed in File 1."""

    source_candidates: List[WasteTransferNoteSourceCandidate] = []
    for pdf_path in sorted(waste_destination.iterdir(), key=lambda path: path.name.lower()):
        if not pdf_path.is_file() or pdf_path.suffix.lower() != ".pdf":
            continue
        try:
            source_candidate = _build_waste_transfer_note_source_candidate(repository, pdf_path)
        except RuntimeError:
            continue
        if not _looks_like_waste_ticket_source(
            pdf_path,
            source_candidate.scanned_note.extracted_text,
        ):
            continue
        source_candidates.append(source_candidate)
    return source_candidates


def list_waste_transfer_note_source_conflicts(
    repository: DocumentRepository,
    *,
    site_name: str,
    waste_destination: Optional[Path] = None,
) -> List[WasteTransferNoteSourceConflict]:
    """Return duplicate-source WTNs that need operator awareness in File 1."""

    resolved_destination = waste_destination or config.WASTE_DESTINATION
    if not resolved_destination.exists():
        return []

    grouped_candidates: Dict[tuple[str, str], List[WasteTransferNoteSourceCandidate]] = {}
    waste_note_lookup = {
        waste_note.doc_id: waste_note
        for waste_note in repository.list_documents(
            document_type=WasteTransferNoteDocument.document_type
        )
        if isinstance(waste_note, WasteTransferNoteDocument)
        and waste_note.site_name.casefold() == site_name.casefold()
    }
    persisted_conflicts: List[WasteTransferNoteSourceConflict] = []
    for waste_note in waste_note_lookup.values():
        if len(waste_note.source_conflict_candidates) < 2:
            continue
        source_candidates = [
            source_candidate
            for source_candidate in (
                _build_waste_transfer_note_source_candidate_from_snapshot(
                    waste_note=waste_note,
                    source_snapshot=source_snapshot,
                )
                for source_snapshot in waste_note.source_conflict_candidates
            )
            if source_candidate is not None
        ]
        if len(source_candidates) < 2:
            continue
        canonical_source_path = (
            Path(waste_note.canonical_source_path).resolve()
            if waste_note.canonical_source_path
            else None
        )
        canonical_source = next(
            (
                source_candidate
                for source_candidate in source_candidates
                if canonical_source_path is not None
                and source_candidate.source_path.resolve() == canonical_source_path
            ),
            None,
        )
        if canonical_source is None:
            canonical_source = _select_effective_waste_transfer_note_source(
                repository,
                source_candidates,
            )
        persisted_conflicts.append(
            WasteTransferNoteSourceConflict(
                wtn_number=waste_note.wtn_number,
                site_name=waste_note.site_name,
                canonical_source=canonical_source,
                source_candidates=sorted(
                    source_candidates,
                    key=_waste_transfer_note_source_candidate_sort_key,
                    reverse=True,
                ),
            )
        )

    if persisted_conflicts:
        return sorted(
            persisted_conflicts,
            key=lambda conflict: (
                conflict.canonical_source.scanned_note.ticket_date,
                conflict.wtn_number,
            ),
            reverse=True,
        )

    indexed_paths_by_group: Dict[tuple[str, str], Dict[str, Path]] = {}
    for indexed_file in repository.list_indexed_files(file_group=FileGroup.FILE_1):
        if indexed_file.file_category not in {"abucs_pdf", "waste_ticket_pdf"}:
            continue
        if indexed_file.related_doc_id not in waste_note_lookup:
            continue
        if indexed_file.site_name.casefold() != site_name.casefold():
            continue
        source_path = indexed_file.file_path
        if not source_path.exists():
            continue

        linked_waste_note = waste_note_lookup[indexed_file.related_doc_id]
        group_key = (
            linked_waste_note.site_name.casefold(),
            _build_waste_transfer_note_identity_key(
                wtn_number=linked_waste_note.wtn_number,
                ticket_date=linked_waste_note.date,
                collection_type=(
                    linked_waste_note.source_conflict_candidates[0].collection_type
                    if linked_waste_note.source_conflict_candidates
                    else ""
                ),
            ),
        )
        indexed_paths_by_group.setdefault(group_key, {})[
            str(source_path.resolve())
        ] = source_path.resolve()

    if indexed_paths_by_group:
        for group_key, source_paths_by_key in indexed_paths_by_group.items():
            if len(source_paths_by_key) < 2:
                continue
            source_candidates: List[WasteTransferNoteSourceCandidate] = []
            for source_path in sorted(
                source_paths_by_key.values(),
                key=lambda path: path.name.casefold(),
            ):
                try:
                    source_candidate = _build_waste_transfer_note_source_candidate(
                        repository,
                        source_path,
                    )
                except RuntimeError:
                    continue
                source_candidates.append(source_candidate)
            if len(source_candidates) >= 2:
                grouped_candidates[group_key] = source_candidates
        if not grouped_candidates:
            return []

    if not grouped_candidates:
        for source_candidate in _list_waste_transfer_note_source_candidates(
            repository,
            waste_destination=resolved_destination,
        ):
            group_key = (
                source_candidate.site_name.casefold(),
                _build_waste_transfer_note_identity_key(
                    wtn_number=source_candidate.scanned_note.wtn_number,
                    ticket_date=source_candidate.scanned_note.ticket_date,
                    collection_type=source_candidate.scanned_note.collection_type,
                ),
            )
            grouped_candidates.setdefault(group_key, []).append(source_candidate)

    conflicts: List[WasteTransferNoteSourceConflict] = []
    for (_, _), source_candidates in grouped_candidates.items():
        candidate_site_name = source_candidates[0].site_name.casefold()
        wtn_number = source_candidates[0].scanned_note.wtn_number
        if candidate_site_name != site_name.casefold() or len(source_candidates) < 2:
            continue
        canonical_source = _select_effective_waste_transfer_note_source(
            repository,
            source_candidates,
        )
        conflicts.append(
            WasteTransferNoteSourceConflict(
                wtn_number=wtn_number,
                site_name=canonical_source.site_name,
                canonical_source=canonical_source,
                source_candidates=sorted(
                    source_candidates,
                    key=_waste_transfer_note_source_candidate_sort_key,
                    reverse=True,
                ),
            )
        )

    return sorted(
        conflicts,
        key=lambda conflict: (
            conflict.canonical_source.scanned_note.ticket_date,
            conflict.wtn_number,
        ),
        reverse=True,
    )


def _sync_existing_waste_transfer_notes(
    repository: DocumentRepository,
    waste_destination: Path,
) -> None:
    """Backfill WTN documents from already-filed waste tickets."""

    synced_site_names = set()
    valid_doc_ids_by_site: Dict[str, set[str]] = {}
    source_candidates = _list_waste_transfer_note_source_candidates(
        repository,
        waste_destination=waste_destination,
    )
    source_candidates_by_wtn: Dict[tuple[str, str], List[WasteTransferNoteSourceCandidate]] = {}
    for source_candidate in source_candidates:
        group_key = (
            source_candidate.site_name.casefold(),
            _build_waste_transfer_note_identity_key(
                wtn_number=source_candidate.scanned_note.wtn_number,
                ticket_date=source_candidate.scanned_note.ticket_date,
                collection_type=source_candidate.scanned_note.collection_type,
            ),
        )
        source_candidates_by_wtn.setdefault(group_key, []).append(source_candidate)

    for source_candidates in source_candidates_by_wtn.values():
        canonical_source = _select_effective_waste_transfer_note_source(
            repository,
            source_candidates,
        )
        waste_transfer_note = _upsert_waste_transfer_note_document_from_candidate(
            repository,
            canonical_source,
            source_candidate_snapshots=[
                _build_waste_transfer_note_source_snapshot(source_candidate)
                for source_candidate in source_candidates
            ],
        )
        synced_site_names.add(waste_transfer_note.site_name)
        valid_doc_ids_by_site.setdefault(waste_transfer_note.site_name, set()).add(
            waste_transfer_note.doc_id
        )

        for source_candidate in source_candidates:
            repository.index_file(
                file_name=source_candidate.source_path.name,
                file_path=source_candidate.source_path,
                file_category=_classify_waste_ticket_file(source_candidate.source_path),
                file_group=FileGroup.FILE_1,
                site_name=waste_transfer_note.site_name,
                related_doc_id=waste_transfer_note.doc_id,
            )

    site_names_to_reconcile = {
        document.site_name
        for document in repository.list_documents(
            document_type=WasteTransferNoteDocument.document_type
        )
        if isinstance(document, WasteTransferNoteDocument)
        and document.status != DocumentStatus.ARCHIVED
    } | set(valid_doc_ids_by_site)

    for site_name in site_names_to_reconcile:
        _archive_stale_waste_transfer_note_documents(
            repository,
            site_name=site_name,
            valid_doc_ids=valid_doc_ids_by_site.get(site_name, set()),
        )

    for site_name in sorted(synced_site_names):
        _upsert_site_waste_register(repository, site_name=site_name)


def _archive_stale_waste_transfer_note_documents(
    repository: DocumentRepository,
    *,
    site_name: str,
    valid_doc_ids: set[str],
) -> None:
    """Archive active WTNs that no longer have a valid ticket source in File 1."""

    active_waste_transfer_notes = [
        document
        for document in repository.list_documents(
            document_type=WasteTransferNoteDocument.document_type,
            site_name=site_name,
        )
        if isinstance(document, WasteTransferNoteDocument)
        and document.status != DocumentStatus.ARCHIVED
    ]
    for waste_transfer_note in active_waste_transfer_notes:
        if waste_transfer_note.doc_id in valid_doc_ids:
            continue
        if valid_doc_ids:
            repository.save(
                replace(
                    waste_transfer_note,
                    status=DocumentStatus.ARCHIVED,
                )
            )
            continue
        indexed_files = repository.list_indexed_files(related_doc_id=waste_transfer_note.doc_id)
        still_has_valid_source = False
        for indexed_file in indexed_files:
            if indexed_file.file_group != FileGroup.FILE_1 or not indexed_file.file_path.exists():
                continue
            try:
                scanned_note = smart_scan_waste_transfer_note(
                    repository,
                    source_path=indexed_file.file_path,
                )
            except RuntimeError:
                continue
            if _looks_like_waste_ticket_source(
                indexed_file.file_path,
                scanned_note.extracted_text,
            ):
                still_has_valid_source = True
                break
        if still_has_valid_source:
            continue
        repository.save(
            replace(
                waste_transfer_note,
                status=DocumentStatus.ARCHIVED,
            )
        )


def _list_site_plant_assets(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> List[PlantAssetDocument]:
    """Return the current non-archived File 2 plant assets for one site."""

    plant_assets = [
        document
        for document in repository.list_documents(
            document_type=PlantAssetDocument.document_type,
            site_name=site_name,
        )
        if isinstance(document, PlantAssetDocument)
        and document.status != DocumentStatus.ARCHIVED
    ]
    return sorted(
        plant_assets,
        key=lambda asset: (
            _extract_plant_hire_sequence(asset.hire_num),
            asset.on_hire,
            asset.description.casefold(),
        ),
    )


def _list_site_waste_transfer_notes(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> List[WasteTransferNoteDocument]:
    """Return one WTN document per reference number for the active site."""

    waste_transfer_notes = [
        document
        for document in repository.list_documents(
            document_type=WasteTransferNoteDocument.document_type,
            site_name=site_name,
        )
        if isinstance(document, WasteTransferNoteDocument)
        and document.status != DocumentStatus.ARCHIVED
    ]
    return sorted(
        waste_transfer_notes,
        key=lambda note: (note.date, note.created_at, note.wtn_number),
    )


def _upsert_site_waste_register(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> WasteRegister:
    """Persist a site-level waste register snapshot for the current WTNs."""

    existing_registers = [
        document
        for document in repository.list_documents(
            document_type=WasteRegister.document_type,
            site_name=site_name,
        )
        if isinstance(document, WasteRegister)
    ]
    latest_register = (
        max(existing_registers, key=lambda register: register.created_at)
        if existing_registers
        else None
    )
    waste_register = WasteRegister(
        doc_id=(
            latest_register.doc_id
            if latest_register is not None
            else f"WASTE-REGISTER-{_slugify_identifier(site_name)}"
        ),
        site_name=site_name,
        created_at=(
            latest_register.created_at
            if latest_register is not None
            else datetime.now().replace(second=0, microsecond=0)
        ),
        status=DocumentStatus.ACTIVE,
        waste_transfer_notes=_list_site_waste_transfer_notes(
            repository,
            site_name=site_name,
        ),
    )
    repository.save(waste_register)
    return waste_register


def _get_waste_transfer_note_document(
    repository: DocumentRepository,
    wtn_number: str,
    *,
    ticket_date: Optional[date] = None,
    collection_type: str = "",
) -> Optional[WasteTransferNoteDocument]:
    """Return an existing WTN document by ticket number."""

    matching_documents = [
        document
        for document in repository.list_documents(
            document_type=WasteTransferNoteDocument.document_type,
            reference_number=wtn_number,
        )
        if isinstance(document, WasteTransferNoteDocument)
    ]
    if not matching_documents:
        return None

    if ticket_date is not None and _is_tanker_collection_type(collection_type):
        tanker_date_matches = [
            document for document in matching_documents if document.date == ticket_date
        ]
        if not tanker_date_matches:
            return None
        matching_documents = tanker_date_matches

    matching_documents.sort(
        key=lambda document: (
            document.status == DocumentStatus.ACTIVE,
            document.date,
            document.created_at,
        ),
        reverse=True,
    )
    return matching_documents[0]


def _get_waste_transfer_note_source_path(
    repository: DocumentRepository,
    waste_transfer_note: WasteTransferNoteDocument,
) -> Optional[Path]:
    """Return the filed PDF path linked to one waste transfer note when available."""

    indexed_files = [
        indexed_file
        for indexed_file in repository.list_indexed_files(related_doc_id=waste_transfer_note.doc_id)
        if indexed_file.file_group == FileGroup.FILE_1 and indexed_file.file_path.exists()
    ]
    if indexed_files:
        source_candidates: List[WasteTransferNoteSourceCandidate] = []
        for indexed_file in indexed_files:
            try:
                source_candidates.append(
                    _build_waste_transfer_note_source_candidate(
                        repository,
                        indexed_file.file_path,
                    )
                )
            except RuntimeError:
                continue
        if source_candidates:
            return _select_effective_waste_transfer_note_source(
                repository,
                source_candidates,
            ).source_path
        return indexed_files[0].file_path
    candidate_path = config.WASTE_DESTINATION / f"{waste_transfer_note.wtn_number}.pdf"
    if candidate_path.exists():
        return candidate_path
    upper_candidate_path = config.WASTE_DESTINATION / f"{waste_transfer_note.wtn_number}.PDF"
    if upper_candidate_path.exists():
        return upper_candidate_path
    return None


def _get_carrier_compliance_document(
    repository: DocumentRepository,
    carrier_name: str,
    carrier_document_type: CarrierComplianceDocumentType,
) -> Optional[CarrierComplianceDocument]:
    """Return the current compliance record for one carrier/type pair."""

    matching_documents = [
        document
        for document in repository.list_documents(
            document_type=CarrierComplianceDocument.document_type
        )
        if isinstance(document, CarrierComplianceDocument)
        and document.carrier_name.casefold() == carrier_name.casefold()
        and document.carrier_document_type == carrier_document_type
    ]
    if not matching_documents:
        return None

    matching_documents.sort(
        key=lambda document: (
            document.status == DocumentStatus.ACTIVE,
            document.expiry_date,
            document.created_at,
        ),
        reverse=True,
    )
    return matching_documents[0]


def _get_carrier_compliance_document_by_reference(
    repository: DocumentRepository,
    *,
    reference_number: str,
    carrier_document_type: CarrierComplianceDocumentType,
) -> Optional[CarrierComplianceDocument]:
    """Return an existing carrier document by its indexed source reference."""

    matching_documents = [
        document
        for document in repository.list_documents(
            document_type=CarrierComplianceDocument.document_type,
            reference_number=reference_number,
        )
        if isinstance(document, CarrierComplianceDocument)
        and document.carrier_document_type == carrier_document_type
    ]
    if not matching_documents:
        return None

    matching_documents.sort(
        key=lambda document: (
            document.status == DocumentStatus.ACTIVE,
            document.expiry_date,
            document.created_at,
        ),
        reverse=True,
    )
    return matching_documents[0]


def _archive_duplicate_carrier_reference_documents(
    repository: DocumentRepository,
    *,
    canonical_document: CarrierComplianceDocument,
) -> None:
    """Archive older active records that point to the same source PDF reference."""

    duplicate_documents = [
        document
        for document in repository.list_documents(
            document_type=CarrierComplianceDocument.document_type,
            reference_number=canonical_document.reference_number,
        )
        if isinstance(document, CarrierComplianceDocument)
        and document.carrier_document_type == canonical_document.carrier_document_type
        and document.doc_id != canonical_document.doc_id
        and document.status != DocumentStatus.ARCHIVED
    ]
    for duplicate_document in duplicate_documents:
        duplicate_document.status = DocumentStatus.ARCHIVED
        repository.save(duplicate_document)


def _infer_carrier_name(
    repository: DocumentRepository,
    source_path: Path,
    *,
    pdf_text: str = "",
) -> str:
    """Resolve a carrier name from the filename and known repository records."""

    searchable_text = f"{source_path.stem} {pdf_text}".casefold()
    for alias, canonical_name in sorted(
        CARRIER_NAME_ALIASES.items(),
        key=lambda item: len(item[0]),
        reverse=True,
    ):
        if alias in searchable_text:
            return canonical_name

    lowered_stem = source_path.stem.casefold()
    for carrier_name in _get_known_carrier_names(repository):
        variations = {
            carrier_name.casefold(),
            carrier_name.casefold().replace(" ", "_"),
            carrier_name.casefold().replace(" ", "-"),
            _slugify_identifier(carrier_name),
        }
        if any(variation and variation in searchable_text for variation in variations):
            return carrier_name

    cleaned_stem = re.sub(
        r"\b(insurance|carrier|licence|license|liability|certificate|policy|waste)\b",
        " ",
        source_path.stem,
        flags=re.IGNORECASE,
    )
    cleaned_stem = re.sub(r"\b\d{1,4}\b", " ", cleaned_stem)
    cleaned_stem = re.sub(r"[_-]+", " ", cleaned_stem)
    inferred_name = " ".join(token for token in cleaned_stem.split() if token)
    return inferred_name.title() if inferred_name else "Unknown Carrier"


def _infer_carrier_document_type(source_path: Path) -> CarrierComplianceDocumentType:
    """Resolve the document type from the carrier-doc filename."""

    lowered_name = source_path.name.casefold()
    if "insurance" in lowered_name:
        return CarrierComplianceDocumentType.INSURANCE
    return CarrierComplianceDocumentType.LICENCE


def _build_carrier_compliance_doc_id(
    carrier_name: str,
    carrier_document_type: CarrierComplianceDocumentType,
) -> str:
    """Build a stable doc id for a carrier compliance record."""

    return f"CCD-{_slugify_identifier(carrier_name)}-{carrier_document_type.value}"


def _get_known_carrier_names(repository: DocumentRepository) -> List[str]:
    """Return carriers already referenced in waste or compliance records."""

    carrier_names = {"Abucs"}
    carrier_names.update(
        document.carrier_name
        for document in repository.list_documents(
            document_type=CarrierComplianceDocument.document_type
        )
        if isinstance(document, CarrierComplianceDocument)
    )
    carrier_names.update(
        document.carrier_name
        for document in repository.list_documents(
            document_type=WasteTransferNoteDocument.document_type
        )
        if isinstance(document, WasteTransferNoteDocument)
    )
    for waste_register in repository.list_documents(document_type=WasteRegister.document_type):
        if not isinstance(waste_register, WasteRegister):
            continue
        carrier_names.update(
            waste_note.carrier_name for waste_note in waste_register.waste_transfer_notes
        )
    return sorted(carrier_name for carrier_name in carrier_names if carrier_name)


def _infer_default_site_name(repository: DocumentRepository) -> str:
    """Return a sensible site name for newly auto-created compliance records."""

    attendance_registers = [
        document
        for document in repository.list_documents(
            document_type=SiteAttendanceRegister.document_type
        )
        if isinstance(document, SiteAttendanceRegister)
    ]
    if attendance_registers:
        return max(attendance_registers, key=lambda register: register.created_at).site_name

    waste_notes = [
        document
        for document in repository.list_documents(
            document_type=WasteTransferNoteDocument.document_type
        )
        if isinstance(document, WasteTransferNoteDocument)
    ]
    if waste_notes:
        return max(waste_notes, key=lambda note: note.created_at).site_name

    carrier_documents = [
        document
        for document in repository.list_documents(
            document_type=CarrierComplianceDocument.document_type
        )
        if isinstance(document, CarrierComplianceDocument)
    ]
    if carrier_documents:
        return max(carrier_documents, key=lambda document: document.created_at).site_name

    return DEFAULT_SITE_NAME


def _extract_waste_ticket_date(normalized_text: str) -> Optional[date]:
    """Return the first date found on the waste ticket."""

    dated_candidates = sorted(_extract_dates_with_positions(normalized_text), key=lambda item: item[1])
    if not dated_candidates:
        return None
    return dated_candidates[0][0]


def _extract_waste_description(normalized_text: str) -> str:
    """Return the waste description shown on the ticket."""

    match = WASTE_TYPE_PATTERN.search(normalized_text)
    if not match:
        if _is_foul_waste_ticket_text(normalized_text):
            return DEFAULT_FOUL_WASTE_DESCRIPTION
        return DEFAULT_WASTE_DESCRIPTION

    description = " ".join(match.group("description").split())
    return description or DEFAULT_WASTE_DESCRIPTION


def _extract_waste_collection_type(normalized_text: str) -> str:
    """Return the skip / collection type shown on the ticket when present."""

    match = WASTE_COLLECTION_TYPE_PATTERN.search(normalized_text)
    if not match:
        return ""
    collection_type = " ".join(match.group("value").split())
    return collection_type.strip(" -:")


def _extract_ewc_code(normalized_text: str) -> str:
    """Return the first EWC code found on the ticket, falling back to a safe default."""

    for match in EWC_CODE_PATTERN.finditer(normalized_text):
        digits_only = "".join(character for character in match.group(0) if character.isdigit())
        hazard_suffix = "*" if "*" in match.group(0) else ""
        if len(digits_only) != 6:
            continue
        normalized_code = (
            f"{digits_only[:2]} {digits_only[2:4]} {digits_only[4:6]}"
            f"{hazard_suffix}"
        )
        if normalized_code in COMMON_CONSTRUCTION_EWC_CODES:
            return normalized_code
    if _is_foul_waste_ticket_text(normalized_text):
        return DEFAULT_FOUL_WASTE_EWC_CODE
    return DEFAULT_EWC_CODE


def _is_foul_waste_ticket_text(normalized_text: str) -> bool:
    """Return True when the ticket text reads like cess/septic/foul waste."""

    lowered_text = normalized_text.casefold()
    return any(keyword in lowered_text for keyword in FOUL_WASTE_KEYWORDS)


def _looks_like_waste_ticket_source(source_path: Path, source_text: str) -> bool:
    """Return True when a non-standard PDF still looks like a waste ticket."""

    if source_path.suffix.lower() != ".pdf":
        return False
    if ABUCS_PDF_PATTERN.match(source_path.name):
        return True

    normalized_text = _normalize_text(source_text)
    lowered_text = normalized_text.casefold()
    if not normalized_text:
        return False

    foul_keyword_hit = _is_foul_waste_ticket_text(normalized_text)
    ticket_anchor_keywords = (
        "ticket no",
        "ticket number",
        "transfer note",
        "skip type",
        "order number",
        "waste type",
    )
    has_ticket_anchor = any(keyword in lowered_text for keyword in ticket_anchor_keywords)
    has_ticket_date = _extract_waste_ticket_date(normalized_text) is not None
    return foul_keyword_hit or (has_ticket_anchor and has_ticket_date)


def _looks_like_waste_support_document(source_text: str) -> bool:
    """Return True when a Word document looks like a waste workbook/support file."""

    lowered_text = source_text.casefold()
    waste_keywords = (
        "waste register",
        "waste removal",
        "waste transfer",
        "wtn",
        "ewc",
        "carrier",
        "waste report",
    )
    return sum(1 for keyword in waste_keywords if keyword in lowered_text) >= 2


def _classify_waste_ticket_file(source_path: Path) -> str:
    """Return the indexed file category for one filed waste-ticket source."""

    if ABUCS_PDF_PATTERN.match(source_path.name):
        return "abucs_pdf"
    return "waste_ticket_pdf"


def _infer_waste_ticket_site_name(
    normalized_text: str,
    repository: DocumentRepository,
) -> str:
    """Infer the site name from the ticket text before falling back to repository defaults."""

    lowered_text = normalized_text.casefold()
    if "lovedean" in lowered_text:
        return DEFAULT_SITE_NAME
    return _infer_default_site_name(repository)


def _slugify_identifier(value: str) -> str:
    """Create a filesystem-safe lowercase identifier from free text."""

    return re.sub(r"[^a-z0-9]+", "-", value.strip().lower()).strip("-")


def _build_induction_doc_id(created_at: datetime, full_name: str) -> str:
    """Return a deterministic-ish induction record identifier."""

    return (
        f"IND-{created_at:%Y%m%d%H%M%S}-"
        f"{_slugify_identifier(full_name)}"
    )


def _build_daily_attendance_doc_id(created_at: datetime, full_name: str) -> str:
    """Return a deterministic-ish UHSF16.09 attendance identifier."""

    return (
        f"ATT-{created_at:%Y%m%d%H%M%S}-"
        f"{_slugify_identifier(full_name)}"
    )


def _normalise_uk_mobile_number(raw_phone_number: str) -> str:
    """Return one normalised UK mobile number or an empty string when invalid."""

    if not raw_phone_number:
        return ""

    compact_value = re.sub(r"[^\d+]+", "", raw_phone_number.strip())
    if compact_value.startswith("00"):
        compact_value = "+" + compact_value[2:]

    if compact_value.startswith("+447") and len(compact_value) == 13:
        return compact_value

    digits_only = re.sub(r"\D+", "", compact_value)
    if digits_only.startswith("447") and len(digits_only) == 12:
        return f"+{digits_only}"
    if digits_only.startswith("07") and len(digits_only) == 11:
        return f"+44{digits_only[1:]}"
    return ""


def _sanitize_filename_fragment(value: str) -> str:
    """Return a human-readable filename fragment with unsafe characters removed."""

    cleaned_value = re.sub(r'[\\/:*?"<>|]+', "-", value).strip()
    return cleaned_value or "Site"


def _normalize_worker_identifier(value: str) -> str:
    """Normalize worker names and filenames into a comparable token form."""

    return re.sub(r"[^a-z0-9]+", "_", value.strip().lower()).strip("_")


def _save_induction_signature_image(
    *,
    signature_image_data: Any,
    full_name: str,
    company: str,
    created_at: datetime,
) -> Path:
    """Persist one drawn kiosk signature as a PNG inside File 3."""

    return _save_signature_image(
        signature_image_data=signature_image_data,
        full_name=f"{full_name} - {company}",
        created_at=created_at,
        destination_directory=config.FILE_3_SIGNATURES_DIR,
        filename_prefix="signature",
        validation_label="induction",
    )


def _save_induction_competency_cards(
    *,
    competency_files: List[Mapping[str, Any]],
    full_name: str,
    company: str,
    created_at: datetime,
) -> List[Path]:
    """Persist uploaded competency cards alongside the File 3 induction record."""

    saved_paths: List[Path] = []
    for uploaded_file in competency_files:
        file_name = Path(str(uploaded_file.get("name", ""))).name
        file_stem = Path(file_name).stem
        file_suffix = Path(file_name).suffix or ".bin"
        competency_label = str(uploaded_file.get("label", "") or "").strip()
        file_bytes = uploaded_file.get("bytes", b"")
        if not file_name or not file_bytes:
            continue
        label_fragment = _sanitize_filename_fragment(competency_label)
        stem_fragment = _sanitize_filename_fragment(file_stem or file_name)
        destination_path = _build_available_destination(
            Path(
                f"{created_at:%Y-%m-%d}_"
                f"{_sanitize_filename_fragment(full_name)}_"
                f"{_sanitize_filename_fragment(company)}_"
                f"{label_fragment + '_' if label_fragment else ''}"
                f"{stem_fragment}{file_suffix}"
            ),
            config.FILE_3_COMPETENCY_CARDS_DIR,
        )
        destination_path.write_bytes(bytes(file_bytes))
        saved_paths.append(destination_path)
    return saved_paths


def _save_attendance_signature_image(
    *,
    signature_image_data: Any,
    full_name: str,
    created_at: datetime,
    action: str,
) -> Path:
    """Persist one drawn UHSF16.09 signature inside File 2."""

    return _save_signature_image(
        signature_image_data=signature_image_data,
        full_name=full_name,
        created_at=created_at,
        destination_directory=config.FILE_2_ATTENDANCE_SIGNATURES_DIR,
        filename_prefix=f"attendance-{action}",
        validation_label="attendance",
    )


def _save_toolbox_talk_signature_image(
    *,
    signature_image_data: Any,
    full_name: str,
    created_at: datetime,
    topic: str,
) -> Path:
    """Persist one remote UHSF16.2 signature inside File 2."""

    return _save_signature_image(
        signature_image_data=signature_image_data,
        full_name=f"{topic}-{full_name}",
        created_at=created_at,
        destination_directory=config.FILE_2_TBT_SIGNATURES_DIR,
        filename_prefix="tbt-signature",
        validation_label="toolbox talk",
    )


def read_toolbox_talk_document_bytes(
    toolbox_talk_document: ToolboxTalkDocument,
) -> Tuple[bytes, str]:
    """Return the stored toolbox talk file bytes and a best-effort mime type."""

    stored_path = Path(toolbox_talk_document.stored_file_path).expanduser()
    if not stored_path.exists():
        raise FileNotFoundError(f"Toolbox talk document is missing: {stored_path}")
    mime_type = mimetypes.guess_type(toolbox_talk_document.original_file_name)[0]
    return stored_path.read_bytes(), mime_type or "application/octet-stream"


def _save_signature_image(
    *,
    signature_image_data: Any,
    full_name: str,
    created_at: datetime,
    destination_directory: Path,
    filename_prefix: str,
    validation_label: str,
) -> Path:
    """Persist one drawn signature PNG in the requested workspace directory."""

    try:
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("Pillow is required to save signature images.") from exc

    if signature_image_data is None:
        raise ValidationError(
            f"Draw a signature before submitting the {validation_label} form."
        )

    try:
        signature_image = Image.fromarray(signature_image_data.astype("uint8")).convert(
            "RGBA"
        )
    except Exception as exc:  # pragma: no cover - third-party ndarray conversion path
        raise ValidationError("Signature capture could not be processed.") from exc

    if not _signature_image_has_content(signature_image):
        raise ValidationError(
            f"Draw a signature before submitting the {validation_label} form."
        )

    destination_directory.mkdir(parents=True, exist_ok=True)
    output_path = _build_available_destination(
        Path(
            f"{filename_prefix}-"
            f"{_sanitize_filename_fragment(full_name)}-"
            f"{created_at:%Y%m%d-%H%M}.png"
        ),
        destination_directory,
    )
    signature_image.convert("RGBA").save(output_path, format="PNG")
    return output_path


def _signature_image_has_content(signature_image: Any) -> bool:
    """Return True when the canvas image contains non-background strokes."""

    try:
        grayscale_image = signature_image.convert("L")
    except Exception:  # pragma: no cover - defensive path for unexpected image types
        return False

    pixel_values = list(grayscale_image.getdata())
    if not pixel_values:
        return False

    non_background_pixels = sum(1 for pixel_value in pixel_values if pixel_value < 245)
    return non_background_pixels > 64


def _build_ladder_permit_doc_id(created_at: datetime, worker_name: str) -> str:
    """Return a deterministic-enough document id for generated ladder permits."""

    return (
        f"LP-{created_at:%Y%m%d%H%M%S}-"
        f"{_slugify_identifier(worker_name)}"
    )


def _build_ladder_permit_number(
    repository: DocumentRepository,
    *,
    site_name: str,
) -> str:
    """Return the next sequential printed ladder permit reference for one site."""

    existing_permits = [
        document
        for document in repository.list_documents(
            document_type=LadderPermit.document_type,
            site_name=site_name,
        )
        if isinstance(document, LadderPermit)
    ]
    return f"LADD-{len(existing_permits) + 1:03d}"


def _ladder_permit_sort_key(permit: LadderPermit) -> Tuple[int, str]:
    """Return a stable sort key for printed ladder permit references."""

    match = re.search(r"(\d+)$", permit.permit_number)
    if match is None:
        return (0, permit.permit_number)
    return (int(match.group(1)), permit.permit_number)


def _format_permit_register_date(permit: LadderPermit) -> str:
    """Return the issued date in the register's dd/mm/yy format."""

    issued_date = permit.issued_date or permit.valid_from_date
    return issued_date.strftime("%d/%m/%y")


def _format_name_company(permit: LadderPermit) -> str:
    """Return the worker/company label for the printed register."""

    parts = [part for part in (permit.worker_name, permit.worker_company) if part]
    return " | ".join(parts)


def _format_waste_register_reference(
    waste_transfer_note: WasteTransferNoteDocument,
) -> str:
    """Return the combined vehicle / ticket reference for UHSF50.0."""

    parts = [
        part.strip()
        for part in (
            waste_transfer_note.vehicle_registration,
            waste_transfer_note.wtn_number,
        )
        if part and part.strip()
    ]
    return " / ".join(parts)


def _get_waste_register_collection_type(
    waste_transfer_note: WasteTransferNoteDocument,
) -> str:
    """Return the best available collection type for one printed File 1 row."""

    canonical_source_path = waste_transfer_note.canonical_source_path.strip()
    if canonical_source_path:
        for source_candidate in waste_transfer_note.source_conflict_candidates:
            if source_candidate.source_path.strip() == canonical_source_path:
                return source_candidate.collection_type.strip()

    for source_candidate in waste_transfer_note.source_conflict_candidates:
        if source_candidate.collection_type.strip():
            return source_candidate.collection_type.strip()
    return ""


def _format_waste_register_description(
    waste_transfer_note: WasteTransferNoteDocument,
) -> str:
    """Return the printed File 1 description with the collection type when available."""

    description = waste_transfer_note.waste_description.strip()
    collection_type = _get_waste_register_collection_type(waste_transfer_note)
    if not collection_type:
        return description
    if not description:
        return collection_type
    if collection_type.casefold() in description.casefold():
        return description
    return f"{collection_type} - {description}"


def _build_project_number(site_name: str) -> str:
    """Derive a compact project reference from the site name."""

    tokens = re.findall(r"[A-Za-z0-9]+", site_name.upper())
    if not tokens:
        return "UPLANDS-SITE"
    return "-".join(tokens[:4])


def _resolve_ladder_safety_checklist(
    safety_checklist: Optional[Mapping[int, bool]],
) -> Dict[int, bool]:
    """Normalize the 11 ladder permit checklist answers into a complete map."""

    resolved_checklist = {question_number: True for question_number in range(1, 12)}
    if safety_checklist is None:
        return resolved_checklist

    for question_number in range(1, 12):
        if question_number in safety_checklist:
            resolved_checklist[question_number] = bool(
                safety_checklist[question_number]
            )
    return resolved_checklist


def _find_matching_induction_file(
    worker_name: str,
    induction_files: List[Path],
) -> Optional[Path]:
    """Return the induction PDF that matches the supplied worker name."""

    worker_identifier = _normalize_worker_identifier(worker_name)
    for induction_file in induction_files:
        file_identifier = _normalize_worker_identifier(induction_file.stem)
        if file_identifier == worker_identifier:
            return induction_file
        if file_identifier.startswith(f"{worker_identifier}_"):
            return induction_file
        if f"{worker_identifier}_induction" in file_identifier:
            return induction_file
        if worker_identifier in file_identifier and "induction" in file_identifier:
            return induction_file
    return None


def _move_file(source_path: Path, destination_directory: Path) -> Path:
    """Move a file into its destination directory without overwriting an existing file."""

    destination_path = _build_available_destination(source_path, destination_directory)
    moved_path = Path(shutil.move(str(source_path), str(destination_path)))
    return moved_path.resolve()


def _build_available_destination(source_path: Path, destination_directory: Path) -> Path:
    """Return a destination path that avoids clobbering an existing file."""

    candidate = destination_directory / source_path.name
    counter = 1
    while candidate.exists():
        candidate = destination_directory / (
            f"{source_path.stem}-{counter}{source_path.suffix}"
        )
        counter += 1
    return candidate
