"""DOCX template filling utilities for Uplands documents."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, Iterable, List, Set, Union

from docx import Document
from docx.document import Document as DocxDocument
from docx.section import Section
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from uplands_site_command_centre.permits.models import BaseDocument, TemplateRegistry


PLACEHOLDER_PATTERN = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")


class TemplateValidationError(ValueError):
    """Raised when an approved template is missing required placeholders."""


class TemplateManager:
    """Populate DOCX templates using snake_case placeholders."""

    def __init__(self, document: BaseDocument) -> None:
        self.document = document
        try:
            self.template_path = TemplateRegistry.resolve_template_path(
                self.document.document_type
            )
        except (KeyError, ValueError) as exc:
            raise TemplateValidationError(str(exc)) from exc

    def discover_placeholders(self) -> List[str]:
        """Return placeholder names present in the template."""

        self._validate_template_path()
        document = Document(self.template_path)
        placeholders: Set[str] = set()
        for paragraph in self._iter_all_paragraphs(document):
            placeholders.update(PLACEHOLDER_PATTERN.findall(paragraph.text))
        return sorted(placeholders)

    def validate_template(self) -> List[str]:
        """Ensure the approved template contains every required placeholder."""

        placeholders = set(self.discover_placeholders())
        missing_placeholders = sorted(
            self.document.required_template_placeholders - placeholders
        )
        if missing_placeholders:
            raise TemplateValidationError(
                "Template %s is missing required placeholders for %s: %s"
                % (
                    self.template_path,
                    self.document.document_type,
                    ", ".join(missing_placeholders),
                )
            )
        return sorted(placeholders)

    def render(self, output_path: Union[str, Path]) -> Path:
        """Fill the approved template and save the populated draft DOCX."""

        self.validate_template()
        context = self.document.to_template_context()
        document = Document(self.template_path)
        for paragraph in self._iter_all_paragraphs(document):
            self._replace_placeholders_in_paragraph(paragraph, context)

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_file)
        return output_file

    def _validate_template_path(self) -> None:
        """Refuse to operate unless the approved official template exists."""

        if self.template_path.suffix.lower() != ".docx":
            raise TemplateValidationError(
                f"Approved template path must point to a .docx file: {self.template_path}"
            )
        if not self.template_path.exists():
            raise TemplateValidationError(
                f"Approved template file does not exist: {self.template_path}"
            )

    def _iter_all_paragraphs(self, document: DocxDocument) -> Iterable[Paragraph]:
        """Yield paragraphs from the document body, tables, headers, and footers."""

        for paragraph in self._iter_paragraphs_from_parent(document):
            yield paragraph

        for section in document.sections:
            for paragraph in self._iter_paragraphs_from_section(section):
                yield paragraph

    def _iter_paragraphs_from_section(self, section: Section) -> Iterable[Paragraph]:
        """Yield paragraphs from the section header and footer."""

        for paragraph in self._iter_paragraphs_from_parent(section.header):
            yield paragraph
        for paragraph in self._iter_paragraphs_from_parent(section.footer):
            yield paragraph

    def _iter_paragraphs_from_parent(self, parent: Union[DocxDocument, _Cell]) -> Iterable[Paragraph]:
        """Yield paragraphs from a document or table cell recursively."""

        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            for paragraph in self._iter_paragraphs_from_table(table):
                yield paragraph

    def _iter_paragraphs_from_table(self, table: Table) -> Iterable[Paragraph]:
        """Yield paragraphs from every cell in a table."""

        for row in table.rows:
            for cell in row.cells:
                for paragraph in self._iter_paragraphs_from_parent(cell):
                    yield paragraph

    def _replace_placeholders_in_paragraph(
        self,
        paragraph: Paragraph,
        context: Dict[str, str],
    ) -> None:
        """Replace placeholders even when Word has split text into runs."""

        if paragraph.runs:
            original_text = "".join(run.text for run in paragraph.runs)
        else:
            original_text = paragraph.text

        if not original_text:
            return

        replaced_text = PLACEHOLDER_PATTERN.sub(
            lambda match: context.get(match.group(1), match.group(0)),
            original_text,
        )

        if replaced_text == original_text:
            return

        if not paragraph.runs:
            paragraph.add_run(replaced_text)
            return

        paragraph.runs[0].text = replaced_text
        for run in paragraph.runs[1:]:
            run.text = ""
