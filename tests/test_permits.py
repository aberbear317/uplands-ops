"""Unit tests for the Uplands document hierarchy and repository flow."""

import json
from datetime import date, datetime, time, timedelta
from pathlib import Path
import tempfile
from typing import List, Optional
import unittest
from unittest.mock import patch

import app as app_module
from docx import Document
import fitz
import numpy as np

import uplands_site_command_centre.config as app_config
import uplands_site_command_centre.workspace as workspace_module
from uplands_site_command_centre import file_and_index_all
from uplands_site_command_centre.permits import (
    COSHHDocument,
    CarrierComplianceDocument,
    CarrierComplianceDocumentType,
    COMMON_CONSTRUCTION_EWC_CODES,
    BroadcastDispatchDocument,
    ComplianceAlertStatus,
    DailyAttendanceEntryDocument,
    DocumentNotFoundError,
    DocumentRepository,
    DocumentStatus,
    FileGroup,
    IncidentLogDocument,
    IncidentType,
    InductionDocument,
    IngestionEngine,
    LadderPermit,
    LadderStabilisationMethod,
    PlantAssetDocument,
    PlantInspectionType,
    RAMSDocument,
    SITE_CHECK_WEEKDAY_KEYS,
    SiteDiaryDocument,
    SiteCheckItem,
    SiteCheckRegister,
    SiteAttendanceRegister,
    SiteAttendanceRecord,
    SiteWorker,
    TemplateRegistry,
    TemplateManager,
    TemplateValidationError,
    ToolboxTalkCompletionDocument,
    ValidationError,
    VerificationStatus,
    WeeklySiteCheck,
    WeeklySiteCheckFrequency,
    WeeklySiteCheckRowState,
    WasteRegister,
    WasteTransferNoteDocument,
    check_carrier_compliance,
)
from uplands_site_command_centre.workspace import (
    _rewrite_inline_table_row_loops,
    build_site_worker_roster,
    build_live_site_broadcast_contacts,
    build_site_alert_sms_link,
    build_site_alert_sms_links,
    build_site_gate_access_code,
    build_pending_toolbox_talk_contacts,
    build_toolbox_talk_sms_message,
    build_toolbox_talk_document_view_url,
    build_toolbox_talk_url,
    calculate_haversine_distance_meters,
    check_site_inductions,
    complete_daily_attendance_sign_out,
    create_daily_attendance_sign_in,
    create_ladder_permit_draft,
    create_site_induction_document,
    add_site_induction_evidence_files,
    create_site_check_checklist_draft,
    create_weekly_site_check_checklist_draft,
    detect_public_tunnel_url_from_log,
    generate_attendance_register_document,
    generate_site_diary_document,
    generate_toolbox_talk_register_document,
    extract_expiry_date_from_pdf,
    extract_tonnage_from_ticket,
    format_plant_inspection_reference,
    get_latest_toolbox_talk_document,
    generate_site_induction_poster,
    generate_plant_register_document,
    generate_waste_register_document,
    generate_permit_register_document,
    get_daily_contractor_headcount,
    get_site_induction_url,
    get_waste_kpi_sheet_metadata,
    get_valid_template_tags,
    get_weekly_site_check_row_definitions,
    infer_plant_inspection_type,
    load_app_settings,
    log_uploaded_waste_transfer_note,
    list_daily_attendance_entries,
    list_broadcast_dispatches,
    list_toolbox_talk_documents,
    lookup_uk_postcode_details,
    lookup_uk_postcode_coordinates,
    log_toolbox_talk_completion,
    list_toolbox_talk_completions,
    read_toolbox_talk_document_bytes,
    run_workspace_diagnostic,
    save_app_settings,
    save_toolbox_talk_document,
    smart_scan_waste_transfer_note,
    sync_file_4_permit_records,
    launch_messages_sms_broadcast,
    log_broadcast_dispatch,
    update_daily_attendance_entry,
    update_site_induction_document,
    validate_site_gate_access_code,
)


class LadderPermitTests(unittest.TestCase):
    def build_permit(self) -> LadderPermit:
        return LadderPermit(
            doc_id="LP-001",
            site_name="Uplands - Newport",
            created_at=datetime(2026, 3, 10, 8, 0),
            status=DocumentStatus.DRAFT,
            permit_number="UHSF21.09-001",
            project_name="Newport Retail Fit-Out",
            project_number="UP-24020",
            location_of_work="Plant room access",
            description_of_work="Lamp replacement above service riser.",
            valid_from_date=date(2026, 3, 10),
            valid_from_time=time(8, 30),
            valid_to_date=date(2026, 3, 10),
            valid_to_time=time(17, 0),
            safer_alternative_eliminated=True,
            task_specific_rams_prepared_and_approved=True,
            personnel_briefed_and_understand_task=True,
            competent_supervisor_appointed=True,
            competent_supervisor_name="J. Evans",
            operatives_suitably_trained=True,
            ladder_length_suitable=True,
            conforms_to_bs_class_a=True,
            three_points_of_contact_maintained=True,
            harness_worn_and_secured_above_head_height=False,
            ladder_stabilisation_method=LadderStabilisationMethod.TIED_AT_BOTTOM,
            equipment_inspected_for_defects=True,
        )

    def test_file_path_uses_four_file_structure(self) -> None:
        permit = self.build_permit()

        self.assertEqual(
            permit.get_file_path(),
            Path("File 4") / "uplands-newport" / "ladder_permit" / "LP-001",
        )
        self.assertEqual(permit.file_group, FileGroup.FILE_4)

    def test_add_inspection_record_appends_typed_entry(self) -> None:
        permit = self.build_permit()

        record = permit.add_inspection_record(
            inspection_date=date(2026, 3, 10),
            inspected_by="J. Evans",
            rungs_ok=True,
            stiles_ok=True,
            feet_ok=False,
            comments_or_action_taken="Replace worn foot before use.",
            ok_to_use=False,
        )

        self.assertEqual(len(permit.inspection_records), 1)
        self.assertEqual(record.inspected_by, "J. Evans")
        self.assertFalse(record.ok_to_use)

    def test_supervisor_name_is_required_when_supervisor_is_appointed(self) -> None:
        with self.assertRaises(ValueError):
            LadderPermit(
                doc_id="LP-002",
                site_name="Uplands - Swansea",
                created_at=datetime(2026, 3, 10, 8, 0),
                status=DocumentStatus.DRAFT,
                permit_number="UHSF21.09-002",
                project_name="Swansea Maintenance",
                project_number="UP-24021",
                location_of_work="Rear stair core",
                description_of_work="Visual snagging check.",
                valid_from_date=date(2026, 3, 10),
                valid_from_time=time(8, 30),
                valid_to_date=date(2026, 3, 10),
                valid_to_time=time(17, 0),
                safer_alternative_eliminated=True,
                task_specific_rams_prepared_and_approved=True,
                personnel_briefed_and_understand_task=True,
                competent_supervisor_appointed=True,
                competent_supervisor_name="",
                operatives_suitably_trained=True,
                ladder_length_suitable=True,
                conforms_to_bs_class_a=True,
                three_points_of_contact_maintained=True,
                harness_worn_and_secured_above_head_height=False,
                ladder_stabilisation_method=LadderStabilisationMethod.FOOTED,
                equipment_inspected_for_defects=True,
            )

    def test_ladder_permit_context_maps_tagged_template_fields(self) -> None:
        permit = self.build_permit()
        permit.add_inspection_record(
            inspection_date=date(2026, 3, 10),
            inspected_by="J. Evans",
            rungs_ok=True,
            stiles_ok=False,
            feet_ok=True,
            comments_or_action_taken="Replace foot before reuse.",
            ok_to_use=False,
        )

        context = permit.to_template_context()

        self.assertEqual(context["job_number"], "UP-24020")
        self.assertEqual(context["date_issued"], "2026-03-10")
        self.assertEqual(context["task_description"], "Lamp replacement above service riser.")
        self.assertEqual(context["supervisor_name"], "J. Evans")
        self.assertEqual(context["manager_name"], "J. Evans")
        self.assertEqual(context["ladder_id"], "Plant room access")
        self.assertEqual(context["q1_yes"], "✔")
        self.assertEqual(context["q1_no"], "")
        self.assertEqual(context["q9_yes"], "")
        self.assertEqual(context["q9_no"], "✔")
        self.assertEqual(context["q10_yes"], "✔")
        self.assertEqual(context["q11_yes"], "✔")
        self.assertEqual(context["insp_name"], "J. Evans")
        self.assertEqual(context["insp_rungs"], "✔")
        self.assertEqual(context["insp_stiles"], "✘")
        self.assertEqual(context["insp_ok"], "✘")


class IncidentLogDocumentTests(unittest.TestCase):
    def build_incident(self, *, status: DocumentStatus = DocumentStatus.ACTIVE) -> IncidentLogDocument:
        return IncidentLogDocument(
            doc_id="INC-001",
            site_name="Uplands - Newport",
            created_at=datetime(2026, 3, 10, 11, 0),
            status=status,
            incident_type=IncidentType.NEAR_MISS,
            location="North compound access route",
            description="Operative slipped on wet boarding but avoided injury.",
            witness_list=["A. Hughes", "J. Evans"],
        )

    def test_incident_log_template_context_includes_witnesses(self) -> None:
        incident = self.build_incident()

        context = incident.to_template_context()

        self.assertEqual(context["incident_type_label"], "Near Miss")
        self.assertEqual(context["witness_count"], "2")
        self.assertEqual(context["witness_1"], "A. Hughes")
        self.assertIn("A. Hughes\nJ. Evans", context["witness_list"])


class WasteDocumentTests(unittest.TestCase):
    def build_wtn(
        self,
        *,
        doc_id: str = "WTN-001",
        wtn_number: str = "UWTN-001",
        note_date: date = date(2026, 3, 10),
        ewc_code: str = "17 09 04",
        quantity_tonnes: float = 3.25,
        carrier_name: str = "Green Haul Ltd",
    ) -> WasteTransferNoteDocument:
        return WasteTransferNoteDocument(
            doc_id=doc_id,
            site_name="Uplands - Newport",
            created_at=datetime(2026, 3, 10, 14, 0),
            status=DocumentStatus.ACTIVE,
            wtn_number=wtn_number,
            date=note_date,
            waste_description="Mixed construction and demolition waste.",
            ewc_code=ewc_code,
            quantity_tonnes=quantity_tonnes,
            carrier_name=carrier_name,
            destination_facility="Newport Recovery Centre",
        )

    def test_wtn_uses_file_1_path_and_metadata(self) -> None:
        wtn = self.build_wtn()

        self.assertEqual(wtn.file_group, FileGroup.FILE_1)
        self.assertEqual(
            wtn.get_file_path(),
            Path("File 1") / "uplands-newport" / "waste_transfer_note" / "WTN-001",
        )
        self.assertEqual(
            wtn.get_repository_metadata(),
            {
                "wtn_number": "UWTN-001",
                "reference_number": "UWTN-001",
                "carrier_name": "Green Haul Ltd",
            },
        )

    def test_wtn_validation_rejects_unknown_ewc_code(self) -> None:
        self.assertIn("17 09 04", COMMON_CONSTRUCTION_EWC_CODES)

        with self.assertRaises(ValidationError):
            self.build_wtn(ewc_code="99 99 99")

    def test_waste_register_summarises_monthly_tonnage(self) -> None:
        register = WasteRegister(
            doc_id="WR-001",
            site_name="Uplands - Newport",
            created_at=datetime(2026, 3, 10, 15, 0),
            status=DocumentStatus.ACTIVE,
            waste_transfer_notes=[
                self.build_wtn(
                    doc_id="WTN-001",
                    wtn_number="UWTN-001",
                    note_date=date(2026, 3, 2),
                    ewc_code="17 09 04",
                    quantity_tonnes=3.25,
                ),
                self.build_wtn(
                    doc_id="WTN-002",
                    wtn_number="UWTN-002",
                    note_date=date(2026, 3, 18),
                    ewc_code="17 02 01",
                    quantity_tonnes=1.75,
                ),
                self.build_wtn(
                    doc_id="WTN-003",
                    wtn_number="UWTN-003",
                    note_date=date(2026, 2, 28),
                    ewc_code="17 09 04",
                    quantity_tonnes=2.0,
                ),
            ],
        )

        summary = register.get_monthly_tonnage_summary(month=3, year=2026)

        self.assertEqual(summary["note_count"], 2)
        self.assertEqual(summary["total_tonnage"], 5.0)
        self.assertEqual(summary["by_ewc_code"], {"17 02 01": 1.75, "17 09 04": 3.25})


class CarrierComplianceDocumentTests(unittest.TestCase):
    def build_compliance_document(
        self,
        *,
        doc_id: str = "CCD-001",
        carrier_name: str = "Abucs",
        carrier_document_type: CarrierComplianceDocumentType = CarrierComplianceDocumentType.LICENCE,
        reference_number: str = "WCL-001",
        expiry_date: date = date(2026, 5, 1),
    ) -> CarrierComplianceDocument:
        return CarrierComplianceDocument(
            doc_id=doc_id,
            site_name="Uplands - Newport",
            created_at=datetime(2026, 3, 10, 10, 0),
            status=DocumentStatus.ACTIVE,
            carrier_name=carrier_name,
            carrier_document_type=carrier_document_type,
            reference_number=reference_number,
            expiry_date=expiry_date,
        )

    def test_carrier_compliance_document_uses_file_1_and_reference_index(self) -> None:
        compliance_document = self.build_compliance_document()

        self.assertEqual(compliance_document.file_group, FileGroup.FILE_1)
        self.assertEqual(
            compliance_document.get_repository_metadata(),
            {"carrier_name": "Abucs", "reference_number": "WCL-001"},
        )
        self.assertEqual(
            compliance_document.to_template_context()["document_type"],
            "Licence",
        )

    def test_check_carrier_compliance_flags_expiring_and_missing_documents(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            repository.save(
                self.build_compliance_document(
                    doc_id="CCD-010",
                    carrier_name="Abucs",
                    carrier_document_type=CarrierComplianceDocumentType.LICENCE,
                    reference_number="WCL-010",
                    expiry_date=date(2026, 4, 1),
                )
            )
            repository.save(
                WasteTransferNoteDocument(
                    doc_id="WTN-CARRIER-1",
                    site_name="Uplands - Newport",
                    created_at=datetime(2026, 3, 10, 11, 0),
                    status=DocumentStatus.ACTIVE,
                    wtn_number="UWTN-CARRIER-1",
                    date=date(2026, 3, 10),
                    waste_description="Mixed construction waste.",
                    ewc_code="17 09 04",
                    quantity_tonnes=2.5,
                    carrier_name="Abucs",
                    destination_facility="Newport Recovery Centre",
                )
            )

            findings = check_carrier_compliance(repository, on_date=date(2026, 3, 10))
            findings_by_type = {
                finding.carrier_document_type: finding for finding in findings
            }

            self.assertEqual(len(findings), 2)
            self.assertEqual(
                findings_by_type[CarrierComplianceDocumentType.LICENCE].status,
                ComplianceAlertStatus.CRITICAL,
            )
            self.assertFalse(
                findings_by_type[CarrierComplianceDocumentType.LICENCE].blocking
            )
            self.assertEqual(
                findings_by_type[CarrierComplianceDocumentType.INSURANCE].status,
                ComplianceAlertStatus.CRITICAL,
            )
            self.assertTrue(
                findings_by_type[CarrierComplianceDocumentType.INSURANCE].blocking
            )


class File3DocumentTests(unittest.TestCase):
    def build_rams(self) -> RAMSDocument:
        return RAMSDocument(
            doc_id="RAMS-001",
            site_name="Uplands - Bristol",
            created_at=datetime(2026, 3, 10, 9, 0),
            status=DocumentStatus.ACTIVE,
            contractor_name="Acme Interiors",
            activity_description="Partition wall installation and ceiling fixings.",
            approval_date=date(2025, 3, 1),
        )

    def build_coshh(self) -> COSHHDocument:
        return COSHHDocument(
            doc_id="COSHH-001",
            site_name="Uplands - Bristol",
            created_at=datetime(2026, 3, 10, 9, 30),
            status=DocumentStatus.ACTIVE,
            contractor_name="Acme Interiors",
            substance_name="Solvent Adhesive",
            hazard_pictograms=["Flammable", "Irritant"],
            ppe_required=["Gloves", "Eye Protection"],
            emergency_first_aid="Move to fresh air and flush eyes with clean water.",
        )

    def build_induction(self) -> InductionDocument:
        return InductionDocument(
            doc_id="IND-001",
            site_name="Uplands - Bristol",
            created_at=datetime(2026, 3, 10, 10, 0),
            status=DocumentStatus.ACTIVE,
            contractor_name="Acme Interiors",
            individual_name="S. Carter",
            linked_rams_doc_id="RAMS-001",
        )

    def test_rams_expiry_defaults_to_twelve_month_window(self) -> None:
        rams = self.build_rams()

        self.assertTrue(rams.has_expired(on_date=date(2026, 3, 10)))
        self.assertFalse(
            rams.has_expired(
                on_date=date(2026, 3, 10),
                max_age=timedelta(days=400),
            )
        )

    def test_coshh_template_context_flattens_lists(self) -> None:
        coshh = self.build_coshh()

        context = coshh.to_template_context()

        self.assertEqual(context["substance_name"], "Solvent Adhesive")
        self.assertEqual(context["hazard_pictogram_1"], "Flammable")
        self.assertEqual(context["ppe_required_2"], "Eye Protection")
        self.assertIn("Gloves\nEye Protection", context["ppe_required"])

    def test_induction_document_keeps_rams_link(self) -> None:
        induction = self.build_induction()

        self.assertEqual(induction.linked_rams_doc_id, "RAMS-001")
        self.assertEqual(induction.file_group, FileGroup.FILE_3)


class SiteAttendanceRecordTests(unittest.TestCase):
    def test_site_attendance_record_maps_json_fields(self) -> None:
        record = SiteAttendanceRecord.from_json_row(
            {
                "date": "2026-03-10",
                "company": "Acme Interiors",
                "workerName": "S. Carter",
                "timeIn": "07:30",
                "timeOut": "16:00",
                "totalHours": 8.5,
            }
        )

        self.assertEqual(record.workerName, "S. Carter")
        self.assertEqual(record.totalHours, 8.5)

    def test_site_attendance_record_accepts_kpi_export_date_format(self) -> None:
        record = SiteAttendanceRecord.from_json_row(
            {
                "date": "10/03/2026",
                "company": "Acme Interiors",
                "workerName": "S. Carter",
                "timeIn": "07:30",
                "timeOut": "16:00",
                "totalHours": "8.5",
            }
        )

        self.assertEqual(record.date, date(2026, 3, 10))
        self.assertEqual(record.totalHours, 8.5)


class SiteCheckRegisterTests(unittest.TestCase):
    def test_site_check_register_uses_file_2_structure(self) -> None:
        register = SiteCheckRegister(
            doc_id="SCR-001",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 11, 6, 45),
            status=DocumentStatus.ACTIVE,
            week_commencing=date(2026, 3, 9),
            checked_at=datetime(2026, 3, 11, 6, 45),
            checked_by="Ceri Edwards",
            check_items=[
                SiteCheckItem(
                    check_name="Site access and egress routes are clear.",
                    frequency="Daily",
                    passed=True,
                    day_results={"tue": True},
                )
            ],
            overall_safe_to_start=True,
        )

        self.assertEqual(register.file_group, FileGroup.FILE_2)
        self.assertEqual(
            register.get_file_path(),
            Path("File 2") / "ng-lovedean-substation" / "site_check_register" / "SCR-001",
        )

    def test_site_check_register_round_trips_through_repository(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            register = SiteCheckRegister(
                doc_id="SCR-002",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 11, 6, 45),
                status=DocumentStatus.ACTIVE,
                week_commencing=date(2026, 3, 9),
                checked_at=datetime(2026, 3, 11, 6, 45),
                checked_by="Ceri Edwards",
                check_items=[
                    SiteCheckItem(
                        check_name="Fire points, extinguishers, and emergency routes are in place.",
                        frequency="Daily",
                        passed=True,
                        day_results={"mon": True, "tue": True},
                    ),
                    SiteCheckItem(
                        check_name="Plant condition and lifting/inspection records have been reviewed.",
                        frequency="Weekly",
                        passed=False,
                        day_results={"tue": False},
                    ),
                ],
                overall_safe_to_start=False,
            )

            repository.save(register)
            loaded = repository.get("SCR-002")

            self.assertIsInstance(loaded, SiteCheckRegister)
            self.assertEqual(loaded.checked_by, "Ceri Edwards")
            self.assertEqual(len(loaded.check_items), 2)
            self.assertFalse(loaded.overall_safe_to_start)
            self.assertEqual(loaded.week_commencing, date(2026, 3, 9))
            self.assertTrue(loaded.check_items[0].day_results["mon"])
            self.assertFalse(loaded.check_items[1].day_results["tue"])

    def test_site_check_register_exposes_weekly_template_context(self) -> None:
        register = SiteCheckRegister(
            doc_id="SCR-003",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 13, 6, 45),
            status=DocumentStatus.ACTIVE,
            week_commencing=date(2026, 3, 9),
            checked_at=datetime(2026, 3, 13, 6, 45),
            checked_by="Ceri Edwards",
            check_items=[
                SiteCheckItem(
                    check_name="Site access and egress routes are clear.",
                    frequency="Daily",
                    passed=True,
                    day_results={"mon": True, "tue": False, "wed": True},
                )
            ],
            overall_safe_to_start=True,
        )

        context = register.to_template_context()

        self.assertEqual(context["week_commencing"], "09/03/2026")
        self.assertEqual(context["site_check_1_mon"], "✓")
        self.assertEqual(context["site_check_1_tue"], "✗")
        self.assertEqual(context["site_check_1_wed"], "✓")
        self.assertEqual(context["site_check_1_sun"], "")


class SiteCheckChecklistGenerationTests(unittest.TestCase):
    def test_create_site_check_checklist_draft_renders_docx_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "workspace"
            output_directory = workspace_root / "FILE_2_Output"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "UHSF19.1 Daily-Weekly Checklist.docx"

            template_document = Document()
            template_document.add_paragraph("Week commencing {{week_commencing}}")
            template_document.add_paragraph("Site {{site_name}}")
            template_document.add_paragraph("Checked by {{checked_by}}")
            template_document.add_paragraph("Checked at {{checked_at}}")
            table = template_document.add_table(
                rows=1 + 7,
                cols=1 + len(SITE_CHECK_WEEKDAY_KEYS),
            )
            table.cell(0, 0).text = "Check"
            for day_offset, day_key in enumerate(SITE_CHECK_WEEKDAY_KEYS, start=1):
                table.cell(0, day_offset).text = day_key.title()
            for row_index in range(1, 8):
                table.cell(row_index, 0).text = f"{{{{site_check_{row_index}_name}}}}"
                for day_offset, day_key in enumerate(SITE_CHECK_WEEKDAY_KEYS, start=1):
                    table.cell(row_index, day_offset).text = (
                        f"{{{{site_check_{row_index}_{day_key}}}}}"
                    )
            template_document.save(template_path)

            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_output_dir = app_config.FILE_2_OUTPUT_DIR
            try:
                TemplateRegistry.TEMPLATE_PATHS["site_check_register"] = template_path
                app_config.FILE_2_OUTPUT_DIR = output_directory

                repository = DocumentRepository(database_path)
                repository.create_schema()
                register = SiteCheckRegister(
                    doc_id="SCR-PRINT-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 11, 6, 45),
                    status=DocumentStatus.ACTIVE,
                    week_commencing=date(2026, 3, 9),
                    checked_at=datetime(2026, 3, 11, 6, 45),
                    checked_by="Ceri Edwards",
                    check_items=[
                        SiteCheckItem(
                            check_name="Site access and egress routes are clear.",
                            frequency="Daily",
                            passed=True,
                            day_results={"mon": True, "tue": True},
                        ),
                        SiteCheckItem(
                            check_name="Housekeeping standards are acceptable across the workface.",
                            frequency="Daily",
                            passed=False,
                            day_results={"mon": True, "tue": False},
                        ),
                    ],
                    overall_safe_to_start=False,
                )
                repository.save(register)

                result = create_site_check_checklist_draft(
                    repository,
                    register=register,
                )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                app_config.FILE_2_OUTPUT_DIR = original_output_dir

            self.assertTrue(result.output_path.exists())
            self.assertEqual(result.output_path.parent, output_directory)

            rendered_document = Document(result.output_path)
            self.assertIn("Week commencing 09/03/2026", rendered_document.paragraphs[0].text)
            self.assertIn("Site NG Lovedean Substation", rendered_document.paragraphs[1].text)
            self.assertEqual(rendered_document.tables[0].cell(1, 1).text, "✓")
            self.assertEqual(rendered_document.tables[0].cell(2, 2).text, "✗")

            indexed_files = repository.list_indexed_files(related_doc_id=register.doc_id)
            self.assertEqual(len(indexed_files), 1)
            self.assertEqual(indexed_files[0].file_group, FileGroup.FILE_2)


class WeeklySiteCheckTests(unittest.TestCase):
    def test_weekly_site_check_context_maps_matrix_and_signoff_fields(self) -> None:
        weekly_site_check = WeeklySiteCheck(
            doc_id="WSC-20260309",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 11, 6, 45),
            status=DocumentStatus.ACTIVE,
            week_commencing=date(2026, 3, 9),
            checked_at=datetime(2026, 3, 11, 6, 45),
            checked_by="Ceri Edwards",
            active_day_key="tue",
            row_states=[
                WeeklySiteCheckRowState(
                    row_number=1,
                    values={"mon": True, "tue": False, "weekly": True},
                ),
                WeeklySiteCheckRowState(
                    row_number=31,
                    values={"tue": True},
                ),
            ],
            daily_initials={"tue": "CE"},
            daily_time_markers={"tue": "AM"},
            overall_safe_to_start=False,
        )

        context = weekly_site_check.to_template_context()

        self.assertEqual(context["week_commencing"], "09/03/2026")
        self.assertEqual(context["mon_1"], "✔")
        self.assertEqual(context["tue_1"], "✘")
        self.assertEqual(context["weekly_1"], "✔")
        self.assertEqual(context["tue_31"], "✔")
        self.assertEqual(context["initials_tue"], "CE")
        self.assertEqual(context["time_tue"], "AM")

    def test_weekly_site_check_context_blanks_incompatible_frequency_slots(self) -> None:
        weekly_site_check = WeeklySiteCheck(
            doc_id="WSC-20260309",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 11, 6, 45),
            status=DocumentStatus.ACTIVE,
            week_commencing=date(2026, 3, 9),
            checked_at=datetime(2026, 3, 11, 6, 45),
            checked_by="Ceri Edwards",
            active_day_key="tue",
            row_states=[
                WeeklySiteCheckRowState(
                    row_number=3,
                    values={"mon": True, "weekly": True},
                ),
                WeeklySiteCheckRowState(
                    row_number=6,
                    values={"tue": False, "weekly": True},
                ),
            ],
            overall_safe_to_start=False,
        )

        context = weekly_site_check.to_template_context()

        self.assertEqual(context["mon_3"], "")
        self.assertEqual(context["weekly_3"], "✔")
        self.assertEqual(context["tue_6"], "✘")
        self.assertEqual(context["weekly_6"], "")

    def test_weekly_site_check_template_rows_include_frequency_rules(self) -> None:
        get_weekly_site_check_row_definitions.cache_clear()
        row_lookup = {
            row_definition.row_number: row_definition
            for row_definition in get_weekly_site_check_row_definitions()
        }

        self.assertEqual(
            row_lookup[3].frequency,
            WeeklySiteCheckFrequency.WEEKLY_ONLY,
        )
        self.assertEqual(
            row_lookup[6].frequency,
            WeeklySiteCheckFrequency.DAILY_ONLY,
        )
        self.assertEqual(
            row_lookup[14].frequency,
            WeeklySiteCheckFrequency.WEEKLY_ONLY,
        )
        self.assertEqual(
            row_lookup[26].frequency,
            WeeklySiteCheckFrequency.DAILY_ONLY,
        )

    def test_weekly_site_check_template_rows_match_official_template(self) -> None:
        get_weekly_site_check_row_definitions.cache_clear()
        row_definitions = get_weekly_site_check_row_definitions()

        self.assertEqual(len(row_definitions), 31)
        self.assertEqual(row_definitions[0].section, "Information Displayed")
        self.assertIn("H&S law poster", row_definitions[0].prompt)
        self.assertEqual(row_definitions[-1].section, "Environment")

    def test_weekly_site_check_valid_tags_follow_pruned_template(self) -> None:
        get_valid_template_tags.cache_clear()
        valid_tags = get_valid_template_tags()

        self.assertIn("week_commencing", valid_tags)
        self.assertIn("checked_by", valid_tags)
        self.assertIn("mon_1", valid_tags)
        self.assertIn("tue_1", valid_tags)
        self.assertIn("weekly_1", valid_tags)
        self.assertIn("mon_2", valid_tags)
        self.assertIn("weekly_2", valid_tags)
        self.assertIn("initials_wed", valid_tags)

    def test_weekly_site_check_from_storage_backfills_blank_checked_by(self) -> None:
        loaded = WeeklySiteCheck.from_storage_dict(
            {
                "doc_id": "WSC-LEGACY-001",
                "site_name": "NG Lovedean Substation",
                "created_at": "2026-03-13T04:25:09",
                "status": "active",
                "week_commencing": "2026-03-09",
                "checked_at": "2026-03-13T04:25:09",
                "checked_by": "",
                "active_day_key": "thu",
                "row_states": [],
                "daily_initials": {"thu": "CE"},
                "daily_time_markers": {"thu": "AM"},
                "overall_safe_to_start": False,
            }
        )

        self.assertEqual(loaded.checked_by, "Ceri Edwards")
        self.assertEqual(loaded.active_day_key, "thu")

    def test_create_weekly_site_check_checklist_uses_official_tagged_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_directory = Path(temp_dir) / "outputs" / "FILE_2_Checklists"
            database_path = Path(temp_dir) / "documents.sqlite3"
            original_output_dir = app_config.FILE_2_CHECKLIST_OUTPUT_DIR
            try:
                app_config.FILE_2_CHECKLIST_OUTPUT_DIR = output_directory
                repository = DocumentRepository(database_path)
                repository.create_schema()

                weekly_site_check = WeeklySiteCheck(
                    doc_id="WSC-20260309",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 11, 6, 45),
                    status=DocumentStatus.ACTIVE,
                    week_commencing=date(2026, 3, 9),
                    checked_at=datetime(2026, 3, 11, 6, 45),
                    checked_by="Ceri Edwards",
                    active_day_key="tue",
                    row_states=[
                        WeeklySiteCheckRowState(
                            row_number=1,
                            values={"mon": True, "tue": True, "weekly": True},
                        ),
                        WeeklySiteCheckRowState(
                            row_number=2,
                            values={"tue": False},
                        ),
                    ],
                    daily_initials={"mon": "CE", "tue": "CE", "wed": "CE"},
                    daily_time_markers={"tue": "AM", "wed": "PM"},
                    overall_safe_to_start=False,
                )
                repository.save(weekly_site_check)

                generated = create_weekly_site_check_checklist_draft(
                    repository,
                    weekly_site_check=weekly_site_check,
                )
            finally:
                app_config.FILE_2_CHECKLIST_OUTPUT_DIR = original_output_dir

            self.assertTrue(generated.output_path.exists())
            self.assertEqual(generated.output_path.parent, output_directory)

            rendered_document = Document(generated.output_path)
            rendered_table = rendered_document.tables[0]
            self.assertIn("NG Lovedean Substation", rendered_table.cell(0, 1).text)
            self.assertIn("09/03/2026", rendered_table.cell(0, 1).text)
            self.assertEqual(rendered_table.cell(1, 10).text, "✔")
            self.assertEqual(rendered_table.cell(2, 4).text, "✘")
            self.assertEqual(rendered_table.cell(33, 5).text.strip(), "CE")
            self.assertEqual(rendered_table.cell(34, 5).text.strip(), "PM")

            indexed_files = repository.list_indexed_files(
                related_doc_id=weekly_site_check.doc_id
            )
            self.assertEqual(len(indexed_files), 1)
            self.assertEqual(indexed_files[0].file_group, FileGroup.FILE_2)


class DocumentRepositoryTests(unittest.TestCase):
    def build_permit(self, *, status: DocumentStatus = DocumentStatus.ACTIVE) -> LadderPermit:
        permit = LadderPermit(
            doc_id="LP-003",
            site_name="Uplands - Caerphilly",
            created_at=datetime(2026, 3, 10, 9, 15),
            status=status,
            permit_number="UHSF21.09-003",
            project_name="Caerphilly Upgrade",
            project_number="UP-24022",
            location_of_work="Warehouse loading area",
            description_of_work="Camera alignment at mezzanine edge.",
            valid_from_date=date(2026, 3, 10),
            valid_from_time=time(9, 30),
            valid_to_date=date(2026, 3, 10),
            valid_to_time=time(15, 0),
            safer_alternative_eliminated=True,
            task_specific_rams_prepared_and_approved=True,
            personnel_briefed_and_understand_task=True,
            competent_supervisor_appointed=True,
            competent_supervisor_name="L. Morgan",
            operatives_suitably_trained=True,
            ladder_length_suitable=True,
            conforms_to_bs_class_a=True,
            three_points_of_contact_maintained=True,
            harness_worn_and_secured_above_head_height=False,
            ladder_stabilisation_method=LadderStabilisationMethod.FOOTED,
            equipment_inspected_for_defects=True,
        )
        permit.add_inspection_record(
            inspection_date=date(2026, 3, 10),
            inspected_by="L. Morgan",
            rungs_ok=True,
            stiles_ok=True,
            feet_ok=True,
            comments_or_action_taken="Serviceable.",
        )
        return permit

    def build_incident(
        self,
        *,
        doc_id: str = "INC-002",
        site_name: str = "Uplands - Caerphilly",
        status: DocumentStatus = DocumentStatus.ACTIVE,
    ) -> IncidentLogDocument:
        return IncidentLogDocument(
            doc_id=doc_id,
            site_name=site_name,
            created_at=datetime(2026, 3, 10, 12, 0),
            status=status,
            incident_type=IncidentType.PROPERTY_DAMAGE,
            location="Delivery bay shutter",
            description="Forklift clipped shutter guide rail.",
            witness_list=["L. Morgan"],
        )

    def build_rams(
        self,
        *,
        doc_id: str = "RAMS-100",
        contractor_name: str = "Acme Interiors",
    ) -> RAMSDocument:
        return RAMSDocument(
            doc_id=doc_id,
            site_name="Uplands - Caerphilly",
            created_at=datetime(2026, 3, 10, 7, 30),
            status=DocumentStatus.ACTIVE,
            contractor_name=contractor_name,
            activity_description="Dry-lining and suspended ceiling installation.",
            approval_date=date(2026, 2, 28),
        )

    def build_coshh(
        self,
        *,
        doc_id: str = "COSHH-100",
        contractor_name: str = "Acme Interiors",
    ) -> COSHHDocument:
        return COSHHDocument(
            doc_id=doc_id,
            site_name="Uplands - Caerphilly",
            created_at=datetime(2026, 3, 10, 8, 15),
            status=DocumentStatus.ACTIVE,
            contractor_name=contractor_name,
            substance_name="Foam Cleaner",
            hazard_pictograms=["Compressed Gas"],
            ppe_required=["Gloves"],
            emergency_first_aid="Rinse affected skin and seek medical advice if irritation persists.",
        )

    def build_induction(
        self,
        *,
        doc_id: str = "IND-100",
        contractor_name: str = "Acme Interiors",
        linked_rams_doc_id: str = "RAMS-100",
    ) -> InductionDocument:
        return InductionDocument(
            doc_id=doc_id,
            site_name="Uplands - Caerphilly",
            created_at=datetime(2026, 3, 10, 8, 45),
            status=DocumentStatus.ACTIVE,
            contractor_name=contractor_name,
            individual_name="P. Lewis",
            linked_rams_doc_id=linked_rams_doc_id,
        )

    def build_wtn(
        self,
        *,
        doc_id: str = "WTN-100",
        wtn_number: str = "UWTN-100",
        carrier_name: str = "Green Haul Ltd",
    ) -> WasteTransferNoteDocument:
        return WasteTransferNoteDocument(
            doc_id=doc_id,
            site_name="Uplands - Caerphilly",
            created_at=datetime(2026, 3, 10, 11, 15),
            status=DocumentStatus.ACTIVE,
            wtn_number=wtn_number,
            date=date(2026, 3, 10),
            waste_description="Mixed construction and demolition waste.",
            ewc_code="17 09 04",
            quantity_tonnes=4.5,
            carrier_name=carrier_name,
            destination_facility="Caerphilly Recycling Hub",
        )

    def build_carrier_compliance_document(
        self,
        *,
        doc_id: str = "CCD-100",
        carrier_name: str = "Abucs",
        carrier_document_type: CarrierComplianceDocumentType = CarrierComplianceDocumentType.LICENCE,
        reference_number: str = "REF-100",
        expiry_date: date = date(2026, 6, 30),
    ) -> CarrierComplianceDocument:
        return CarrierComplianceDocument(
            doc_id=doc_id,
            site_name="Uplands - Caerphilly",
            created_at=datetime(2026, 3, 10, 9, 0),
            status=DocumentStatus.ACTIVE,
            carrier_name=carrier_name,
            carrier_document_type=carrier_document_type,
            reference_number=reference_number,
            expiry_date=expiry_date,
        )

    def test_round_trip_with_sqlite_persists_inspection_records(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            permit = self.build_permit()
            repository.save(permit)
            loaded = repository.get("LP-003")

            self.assertIsInstance(loaded, LadderPermit)
            self.assertEqual(len(loaded.inspection_records), 1)
            self.assertEqual(
                loaded.inspection_records[0].comments_or_action_taken,
                "Serviceable.",
            )

    def test_repository_auto_discovers_incident_log_document(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            incident = self.build_incident()
            repository.save(incident)
            loaded = repository.get(incident.doc_id)

            self.assertIsInstance(loaded, IncidentLogDocument)
            self.assertEqual(loaded.file_group, FileGroup.FILE_1)
            self.assertEqual(loaded.document_type, "incident_log")

    def test_site_compliance_summary_groups_by_file_group(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            active_incident = self.build_incident(doc_id="INC-100", status=DocumentStatus.ACTIVE)
            draft_incident = self.build_incident(doc_id="INC-101", status=DocumentStatus.DRAFT)
            active_permit = self.build_permit(status=DocumentStatus.ACTIVE)

            repository.save(active_incident)
            repository.save(draft_incident)
            repository.save(active_permit)

            summary = repository.get_site_compliance_summary("Uplands - Caerphilly")

            self.assertEqual(summary["files"]["File 1"]["count"], 1)
            self.assertEqual(summary["files"]["File 1"]["total_count"], 2)
            self.assertEqual(
                summary["files"]["File 1"]["status_counts"],
                {"draft": 1, "active": 1, "archived": 0},
            )
            self.assertEqual(summary["files"]["File 4"]["count"], 1)
            self.assertEqual(
                summary["files"]["File 4"]["active_documents"][0]["document_type"],
                "ladder_permit",
            )

    def test_get_raises_for_missing_document(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            with self.assertRaises(DocumentNotFoundError):
                repository.get("missing")

    def test_search_by_contractor_name_returns_all_linked_file3_documents(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            rams = self.build_rams()
            coshh = self.build_coshh()
            induction = self.build_induction()
            other_rams = self.build_rams(
                doc_id="RAMS-101",
                contractor_name="Different Contractor Ltd",
            )

            repository.save(rams)
            repository.save(coshh)
            repository.save(induction)
            repository.save(other_rams)

            results = repository.search_by_contractor_name("acme interiors")

            self.assertEqual(len(results), 3)
            self.assertEqual(
                {document.document_type for document in results},
                {"rams", "coshh", "induction"},
            )
            self.assertEqual(
                {getattr(document, "contractor_name", "") for document in results},
                {"Acme Interiors"},
            )

    def test_delete_document_and_files_removes_induction_and_linked_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            repository = DocumentRepository(temp_path / "documents.sqlite3")
            repository.create_schema()
            signature_path = temp_path / "signature.png"
            completed_doc_path = temp_path / "induction.docx"
            competency_card_path = temp_path / "cscs-card.jpg"
            signature_path.write_bytes(b"sig")
            completed_doc_path.write_bytes(b"doc")
            competency_card_path.write_bytes(b"card")

            induction = InductionDocument(
                doc_id="IND-DEL-001",
                site_name="Uplands - Caerphilly",
                created_at=datetime(2026, 3, 10, 8, 45),
                status=DocumentStatus.ACTIVE,
                contractor_name="Acme Interiors",
                individual_name="P. Lewis",
                signature_image_path=str(signature_path),
                completed_document_path=str(completed_doc_path),
                competency_card_paths=str(competency_card_path),
            )
            repository.save(induction)
            repository.index_file(
                file_name=signature_path.name,
                file_path=signature_path,
                file_category="induction_signature_png",
                file_group=FileGroup.FILE_3,
                site_name=induction.site_name,
                related_doc_id=induction.doc_id,
            )
            repository.index_file(
                file_name=completed_doc_path.name,
                file_path=completed_doc_path,
                file_category="completed_induction_docx",
                file_group=FileGroup.FILE_3,
                site_name=induction.site_name,
                related_doc_id=induction.doc_id,
            )

            deleted_paths = repository.delete_document_and_files(induction.doc_id)

            self.assertEqual(
                {path.resolve() for path in deleted_paths},
                {
                    signature_path.resolve(),
                    completed_doc_path.resolve(),
                    competency_card_path.resolve(),
                },
            )
            self.assertFalse(signature_path.exists())
            self.assertFalse(completed_doc_path.exists())
            self.assertFalse(competency_card_path.exists())
            self.assertEqual(
                repository.list_documents(document_type=InductionDocument.document_type),
                [],
            )

    def test_delete_documents_and_files_removes_multiple_inductions(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            repository = DocumentRepository(temp_path / "documents.sqlite3")
            repository.create_schema()

            first_signature = temp_path / "first-signature.png"
            second_signature = temp_path / "second-signature.png"
            first_signature.write_bytes(b"sig-1")
            second_signature.write_bytes(b"sig-2")

            first_induction = InductionDocument(
                doc_id="IND-BULK-001",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 7, 0),
                status=DocumentStatus.ACTIVE,
                contractor_name="Acme Interiors",
                individual_name="First Operative",
                signature_image_path=str(first_signature),
            )
            second_induction = InductionDocument(
                doc_id="IND-BULK-002",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 7, 5),
                status=DocumentStatus.ACTIVE,
                contractor_name="Acme Interiors",
                individual_name="Second Operative",
                signature_image_path=str(second_signature),
            )
            repository.save(first_induction)
            repository.save(second_induction)

            deleted_paths = repository.delete_documents_and_files(
                [first_induction.doc_id, second_induction.doc_id]
            )

            self.assertEqual(
                {path.resolve() for path in deleted_paths},
                {first_signature.resolve(), second_signature.resolve()},
            )
            self.assertEqual(
                repository.list_documents(document_type=InductionDocument.document_type),
                [],
            )

    def test_delete_document_and_files_removes_site_diary_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            repository = DocumentRepository(temp_path / "documents.sqlite3")
            repository.create_schema()
            generated_doc_path = temp_path / "daily-site-diary.docx"
            generated_doc_path.write_bytes(b"doc")

            diary = SiteDiaryDocument(
                doc_id="SITE-DIARY-DEL-001",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 17, 0),
                status=DocumentStatus.ACTIVE,
                date=date(2026, 3, 15),
                incidents_details="None",
                area_handovers="None",
                todays_comments="None",
                generated_document_path=str(generated_doc_path),
            )
            repository.save(diary)

            deleted_paths = repository.delete_document_and_files(diary.doc_id)

            self.assertEqual(
                {path.resolve() for path in deleted_paths},
                {generated_doc_path.resolve()},
            )
            self.assertFalse(generated_doc_path.exists())
            self.assertEqual(
                repository.list_documents(document_type=SiteDiaryDocument.document_type),
                [],
            )

    def test_search_by_wtn_number_and_carrier_name_returns_waste_transfer_notes(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            waste_note = self.build_wtn()
            other_waste_note = self.build_wtn(
                doc_id="WTN-101",
                wtn_number="UWTN-101",
                carrier_name="Other Carrier Ltd",
            )

            repository.save(waste_note)
            repository.save(other_waste_note)

            by_number = repository.search_by_wtn_number("uwtn-100")
            by_carrier = repository.search_by_carrier_name("green haul ltd")

            self.assertEqual(len(by_number), 1)
            self.assertIsInstance(by_number[0], WasteTransferNoteDocument)
            self.assertEqual(by_number[0].wtn_number, "UWTN-100")

            self.assertEqual(len(by_carrier), 1)
            self.assertIsInstance(by_carrier[0], WasteTransferNoteDocument)
            self.assertEqual(by_carrier[0].carrier_name, "Green Haul Ltd")

    def test_waste_transfer_note_is_marked_unverified_when_carrier_docs_are_missing(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            waste_note = self.build_wtn(carrier_name="Abucs")
            repository.save(waste_note)

            loaded = repository.get(waste_note.doc_id)

            self.assertIsInstance(loaded, WasteTransferNoteDocument)
            self.assertEqual(loaded.verification_status, VerificationStatus.UNVERIFIED)
            self.assertIn("missing or expired", loaded.verification_notes.lower())

    def test_waste_transfer_note_is_verified_when_carrier_docs_are_current(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            repository.save(
                self.build_carrier_compliance_document(
                    doc_id="CCD-200",
                    carrier_name="Abucs",
                    carrier_document_type=CarrierComplianceDocumentType.LICENCE,
                    reference_number="WCL-200",
                    expiry_date=date(2026, 12, 31),
                )
            )
            repository.save(
                self.build_carrier_compliance_document(
                    doc_id="CCD-201",
                    carrier_name="Abucs",
                    carrier_document_type=CarrierComplianceDocumentType.INSURANCE,
                    reference_number="LI-201",
                    expiry_date=date(2026, 12, 31),
                )
            )

            waste_note = self.build_wtn(carrier_name="Abucs")
            repository.save(waste_note)

            loaded = repository.get(waste_note.doc_id)

            self.assertIsInstance(loaded, WasteTransferNoteDocument)
            self.assertEqual(loaded.verification_status, VerificationStatus.VERIFIED)
            self.assertIn("valid", loaded.verification_notes.lower())

    def test_saving_carrier_docs_rechecks_existing_waste_transfer_notes(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            waste_note = self.build_wtn(carrier_name="Abucs")
            repository.save(waste_note)
            self.assertEqual(
                repository.get(waste_note.doc_id).verification_status,
                VerificationStatus.UNVERIFIED,
            )

            repository.save(
                self.build_carrier_compliance_document(
                    doc_id="CCD-300",
                    carrier_name="Abucs",
                    carrier_document_type=CarrierComplianceDocumentType.LICENCE,
                    reference_number="WCL-300",
                    expiry_date=date(2026, 12, 31),
                )
            )
            self.assertEqual(
                repository.get(waste_note.doc_id).verification_status,
                VerificationStatus.UNVERIFIED,
            )

            repository.save(
                self.build_carrier_compliance_document(
                    doc_id="CCD-301",
                    carrier_name="Abucs",
                    carrier_document_type=CarrierComplianceDocumentType.INSURANCE,
                    reference_number="LI-301",
                    expiry_date=date(2026, 12, 31),
                )
            )
            self.assertEqual(
                repository.get(waste_note.doc_id).verification_status,
                VerificationStatus.VERIFIED,
            )


class IngestionEngineTests(unittest.TestCase):
    def test_ingestion_engine_creates_file2_register_and_aggregates_split_shifts(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            engine = IngestionEngine(repository)
            json_path = Path(temp_dir) / "attendance.json"

            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    [
                        {
                            "date": "2026-03-10",
                            "company": "Acme Interiors",
                            "workerName": "S. Carter",
                            "timeIn": "07:30",
                            "timeOut": "10:30",
                            "totalHours": 3.0,
                        },
                        {
                            "date": "2026-03-10",
                            "company": "Acme Interiors",
                            "workerName": "S. Carter",
                            "timeIn": "12:00",
                            "timeOut": "15:18",
                            "totalHours": 3.3,
                        },
                        {
                            "date": "2026-03-10",
                            "company": "Acme Interiors",
                            "workerName": "J. Evans",
                            "timeIn": "08:00",
                            "timeOut": "16:30",
                            "totalHours": 8.0,
                        },
                    ],
                    file_handle,
                )

            register = engine.ingest_site_attendance_json(
                json_path,
                site_name="Uplands - Cardiff",
            )

            self.assertIsInstance(register, SiteAttendanceRegister)
            self.assertEqual(register.file_group, FileGroup.FILE_2)
            self.assertEqual(len(register.attendance_records), 2)
            s_carter = next(
                record
                for record in register.attendance_records
                if record.workerName == "S. Carter"
            )
            self.assertEqual(s_carter.timeIn.strftime("%H:%M"), "07:30")
            self.assertEqual(s_carter.timeOut.strftime("%H:%M"), "15:18")
            self.assertAlmostEqual(s_carter.totalHours, 6.3)

            loaded = repository.get(register.doc_id)
            self.assertIsInstance(loaded, SiteAttendanceRegister)
            self.assertEqual(len(loaded.attendance_records), 2)
            loaded_s_carter = next(
                record
                for record in loaded.attendance_records
                if record.workerName == "S. Carter"
            )
            self.assertAlmostEqual(loaded_s_carter.totalHours, 6.3)

            second_pass = engine.ingest_site_attendance_json(
                json_path,
                site_name="Uplands - Cardiff",
            )
            self.assertEqual(len(second_pass.attendance_records), 2)
            second_pass_s_carter = next(
                record
                for record in second_pass.attendance_records
                if record.workerName == "S. Carter"
            )
            self.assertAlmostEqual(second_pass_s_carter.totalHours, 6.3)

    def test_ingestion_engine_deduplicates_identical_rows_before_aggregation(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            engine = IngestionEngine(repository)
            json_path = Path(temp_dir) / "attendance.json"

            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    [
                        {
                            "date": "2026-03-10",
                            "company": "Acme Interiors",
                            "workerName": "S. Carter",
                            "timeIn": "07:30",
                            "timeOut": "16:00",
                            "totalHours": 8.5,
                        },
                        {
                            "date": "2026-03-10",
                            "company": "Acme Interiors",
                            "workerName": "S. Carter",
                            "timeIn": "07:30",
                            "timeOut": "16:00",
                            "totalHours": 8.5,
                        },
                    ],
                    file_handle,
                )

            register = engine.ingest_site_attendance_json(
                json_path,
                site_name="Uplands - Cardiff",
            )

            self.assertEqual(len(register.attendance_records), 1)
            self.assertAlmostEqual(register.attendance_records[0].totalHours, 8.5)

    def test_ingestion_engine_reimport_updates_existing_day_total_for_split_shift_correction(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            engine = IngestionEngine(repository)
            json_path = Path(temp_dir) / "attendance.json"

            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    [
                        {
                            "date": "2026-03-04",
                            "company": "Acme Interiors",
                            "workerName": "C Ward",
                            "timeIn": "07:30",
                            "timeOut": "10:30",
                            "totalHours": 3.1,
                        }
                    ],
                    file_handle,
                )

            first_register = engine.ingest_site_attendance_json(
                json_path,
                site_name="Uplands - Cardiff",
            )
            self.assertAlmostEqual(first_register.attendance_records[0].totalHours, 3.1)

            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    [
                        {
                            "date": "2026-03-04",
                            "company": "Acme Interiors",
                            "workerName": "C Ward",
                            "timeIn": "07:30",
                            "timeOut": "10:30",
                            "totalHours": 3.1,
                        },
                        {
                            "date": "2026-03-04",
                            "company": "Acme Interiors",
                            "workerName": "C Ward",
                            "timeIn": "11:45",
                            "timeOut": "14:57",
                            "totalHours": 3.2,
                        },
                    ],
                    file_handle,
                )

            corrected_register = engine.ingest_site_attendance_json(
                json_path,
                site_name="Uplands - Cardiff",
            )

            self.assertEqual(len(corrected_register.attendance_records), 1)
            self.assertAlmostEqual(corrected_register.attendance_records[0].totalHours, 6.3)
            self.assertEqual(
                corrected_register.attendance_records[0].timeOut.strftime("%H:%M"),
                "14:57",
            )

    def test_ingestion_engine_reads_nested_kpi_export_shape(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            engine = IngestionEngine(repository)
            json_path = Path(temp_dir) / "attendance_export.json"

            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    {
                        "version": 1,
                        "settings": {"siteName": "Uplands - Cardiff"},
                        "extractedRows": {
                            "weekly": [
                                {
                                    "id": "weekly-1",
                                    "date": "10/03/2026",
                                    "company": "Acme Interiors",
                                    "workerName": "S. Carter",
                                    "timeIn": "07:30",
                                    "timeOut": "16:00",
                                    "totalHours": 8.5,
                                    "isSeniorManager": False,
                                }
                            ],
                            "eom": [
                                {
                                    "id": "eom-1",
                                    "date": "10/03/2026",
                                    "company": "Acme Interiors",
                                    "workerName": "S. Carter",
                                    "timeIn": "07:30",
                                    "timeOut": "16:00",
                                    "totalHours": 8.5,
                                    "isSeniorManager": False,
                                },
                                {
                                    "id": "eom-2",
                                    "date": "10/03/2026",
                                    "company": "Acme Interiors",
                                    "workerName": "J. Evans",
                                    "timeIn": "08:00",
                                    "timeOut": "16:30",
                                    "totalHours": 8.0,
                                    "isSeniorManager": False,
                                },
                            ],
                        },
                    },
                    file_handle,
                )

            register = engine.ingest_site_attendance_json(
                json_path,
                site_name="Uplands - Cardiff",
            )

            self.assertEqual(len(register.attendance_records), 2)
            self.assertEqual(
                {record.workerName for record in register.attendance_records},
                {"S. Carter", "J. Evans"},
            )
            s_carter = next(
                record
                for record in register.attendance_records
                if record.workerName == "S. Carter"
            )
            self.assertAlmostEqual(s_carter.totalHours, 8.5)
            self.assertEqual(register.site_name, "Uplands - Cardiff")

    def test_ingestion_engine_recovers_short_kpi_dates_from_row_id(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            engine = IngestionEngine(repository)
            json_path = Path(temp_dir) / "attendance_export.json"

            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    {
                        "version": 1,
                        "settings": {"siteName": "Uplands - Cardiff"},
                        "extractedRows": {
                            "weekly": [
                                {
                                    "id": "extracted:weekly:2026-03-12|url|c edwards|10:40|12:00",
                                    "date": "12/03",
                                    "company": "URL",
                                    "workerName": "C Edwards",
                                    "timeIn": "10:40",
                                    "timeOut": "12:00",
                                    "totalHours": 1.33,
                                    "isSeniorManager": False,
                                }
                            ],
                            "eom": [],
                        },
                    },
                    file_handle,
                )

            register = engine.ingest_site_attendance_json(
                json_path,
                site_name="Uplands - Cardiff",
            )

            self.assertEqual(len(register.attendance_records), 1)
            self.assertEqual(register.attendance_records[0].date, date(2026, 3, 12))

    def test_ingestion_engine_accepts_year_month_day_slash_dates(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            engine = IngestionEngine(repository)
            json_path = Path(temp_dir) / "attendance_export.json"

            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    {
                        "version": 1,
                        "settings": {"siteName": "Uplands - Cardiff"},
                        "extractedRows": {
                            "weekly": [
                                {
                                    "id": "extracted:weekly:2026/03/13|150|r minard|07:30|13:00",
                                    "date": "2026/03/13",
                                    "company": "150",
                                    "workerName": "R Minard",
                                    "timeIn": "07:30",
                                    "timeOut": "13:00",
                                    "totalHours": 5.5,
                                    "isSeniorManager": False,
                                }
                            ],
                            "eom": [],
                        },
                    },
                    file_handle,
                )

            register = engine.ingest_site_attendance_json(
                json_path,
                site_name="Uplands - Cardiff",
            )

            self.assertEqual(len(register.attendance_records), 1)
            self.assertEqual(register.attendance_records[0].date, date(2026, 3, 13))


class SiteWorkerRosterTests(unittest.TestCase):
    def test_build_site_worker_roster_reads_all_active_kpi_arrays(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            json_path = Path(temp_dir) / "site-kpi-backup-1.json"
            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    {
                        "settings": {"siteName": "NG Lovedean Substation"},
                        "extractedRows": {
                            "weekly": [
                                {
                                    "date": "03/03/2026",
                                    "company": "Acme Interiors",
                                    "workerName": "S. Carter",
                                }
                            ],
                            "liveNow": [
                                {
                                    "date": "05/03/2026",
                                    "company": "Acme Interiors",
                                    "workerName": "S. Carter",
                                },
                                {
                                    "date": "04/03/2026",
                                    "company": "Beacon Civils",
                                    "workerName": "J. Evans",
                                },
                            ],
                        },
                    },
                    file_handle,
                )

            roster = build_site_worker_roster(
                site_name="NG Lovedean Substation",
                source_paths=[json_path],
            )

            self.assertEqual(len(roster), 2)
            s_carter = next(worker for worker in roster if worker.worker_name == "S. Carter")
            self.assertEqual(s_carter.company, "Acme Interiors")
            self.assertEqual(s_carter.last_on_site_date, date(2026, 3, 5))
            self.assertEqual(s_carter.induction_status, "Verified (Paper Record)")

    def test_build_site_worker_roster_filters_other_sites(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            lovedean_json_path = Path(temp_dir) / "site-kpi-backup-lovedean.json"
            with lovedean_json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    {
                        "settings": {"siteName": "NG Lovedean Substation"},
                        "extractedRows": {
                            "weekly": [
                                {
                                    "date": "03/03/2026",
                                    "company": "Acme Interiors",
                                    "workerName": "S. Carter",
                                }
                            ]
                        },
                    },
                    file_handle,
                )

            other_json_path = Path(temp_dir) / "site-kpi-backup-other.json"
            with other_json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    {
                        "settings": {"siteName": "Another Project"},
                        "extractedRows": {
                            "weekly": [
                                {
                                    "date": "03/03/2026",
                                    "company": "Other Co",
                                    "workerName": "A. Other",
                                }
                            ]
                        },
                    },
                    file_handle,
                )

            roster = build_site_worker_roster(
                site_name="NG Lovedean Substation",
                source_paths=[lovedean_json_path, other_json_path],
            )

            self.assertEqual(len(roster), 1)
            self.assertEqual(roster[0].worker_name, "S. Carter")

    def test_build_site_worker_roster_recovers_short_kpi_dates_from_row_id(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            json_path = Path(temp_dir) / "site-kpi-backup-short-date.json"
            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    {
                        "settings": {"siteName": "NG Lovedean Substation"},
                        "extractedRows": {
                            "weekly": [
                                {
                                    "id": "extracted:weekly:2026-03-12|url|c edwards|10:40|12:00",
                                    "date": "12/03",
                                    "company": "URL",
                                    "workerName": "C Edwards",
                                }
                            ]
                        },
                    },
                    file_handle,
                )

            roster = build_site_worker_roster(
                site_name="NG Lovedean Substation",
                source_paths=[json_path],
            )

            self.assertEqual(len(roster), 1)
            self.assertEqual(roster[0].worker_name, "C Edwards")
            self.assertEqual(roster[0].last_on_site_date, date(2026, 3, 12))


class WorkspaceFileIndexingTests(unittest.TestCase):
    def _build_ladder_tagged_template(self, template_path: Path) -> None:
        template_document = Document()
        template_document.add_paragraph("Permit {{permit_number}}")
        template_document.add_paragraph("{{company_name}}")
        template_document.add_paragraph("{{contractor_name}}")
        template_document.add_paragraph("{{job_number}}")
        template_document.add_paragraph("{{date_issued}}")
        template_document.add_paragraph("{{site_name}}")
        template_document.add_paragraph("{{task_description}}")
        template_document.add_paragraph("{{supervisor_name}}")
        for question_number in range(1, 12):
            template_document.add_paragraph(
                f"Q{question_number} {{{{q{question_number}_yes}}}} {{{{q{question_number}_no}}}}"
            )
        template_document.add_paragraph("{{manager_name}}")
        template_document.add_paragraph("{{manager_signature}}")
        template_document.add_paragraph("{{ladder_id}}")
        table = template_document.add_table(rows=1, cols=7)
        table.cell(0, 0).text = "{{insp_date}}"
        table.cell(0, 1).text = "{{insp_name}}"
        table.cell(0, 2).text = "{{insp_rungs}}"
        table.cell(0, 3).text = "{{insp_stiles}}"
        table.cell(0, 4).text = "{{insp_feet}}"
        table.cell(0, 5).text = "{{insp_comments}}"
        table.cell(0, 6).text = "{{insp_ok}}"
        template_document.save(template_path)

    def _build_ladder_management_sections_template(self, template_path: Path) -> None:
        template_document = Document()
        header_table = template_document.add_table(rows=5, cols=5)
        header_table.cell(0, 1).text = "{{permit_number}}"
        header_table.cell(0, 4).text = "{{site_name}}"
        header_table.cell(1, 1).text = "{{job_number}}"
        header_table.cell(1, 4).text = "{{ladder_id}}"
        header_table.cell(2, 1).text = "{{task_description}}"
        header_table.cell(3, 1).text = "{{date_issued}}"

        template_document.add_paragraph("{{company_name}}")
        template_document.add_paragraph("{{contractor_name}}")
        template_document.add_paragraph("{{supervisor_name}}")
        for question_number in range(1, 12):
            template_document.add_paragraph(
                f"Q{question_number} {{{{q{question_number}_yes}}}} {{{{q{question_number}_no}}}}"
            )
        template_document.add_paragraph("{{manager_name}}")
        template_document.add_paragraph("{{manager_signature}}")
        template_document.add_paragraph("{{ladder_id}}")

        inspection_table = template_document.add_table(rows=1, cols=7)
        inspection_table.cell(0, 0).text = "{{insp_date}}"
        inspection_table.cell(0, 1).text = "{{insp_name}}"
        inspection_table.cell(0, 2).text = "{{insp_rungs}}"
        inspection_table.cell(0, 3).text = "{{insp_stiles}}"
        inspection_table.cell(0, 4).text = "{{insp_feet}}"
        inspection_table.cell(0, 5).text = "{{insp_comments}}"
        inspection_table.cell(0, 6).text = "{{insp_ok}}"

        for _ in range(2):
            template_document.add_table(rows=1, cols=1)

        acceptance_table = template_document.add_table(rows=8, cols=4)
        acceptance_table.cell(1, 0).text = "Name: {{manager_name}}"
        acceptance_table.cell(1, 1).text = "Sign: {{manager_signature}}"
        acceptance_table.cell(1, 2).text = "Date: (dd/mm/yyyy) {{date_issued}}"
        acceptance_table.cell(1, 3).text = "Position:"
        acceptance_table.cell(3, 0).text = "Name: {{contractor_name}}"
        acceptance_table.cell(3, 2).text = "Date: (dd/mm/yyyy) {{date_issued}}"
        acceptance_table.cell(3, 3).text = "Company: {{company_name}}"
        acceptance_table.cell(7, 0).text = "Name:"
        acceptance_table.cell(7, 1).text = "Sign:"
        acceptance_table.cell(7, 2).text = "Date: (dd/mm/yyyy)"
        acceptance_table.cell(7, 3).text = "Position:"

        template_document.save(template_path)

    def _build_permit_register_template(self, template_path: Path) -> None:
        template_document = Document()
        template_document.add_paragraph("{{site_name}}")
        template_document.add_paragraph("{{job_number}}")
        register_table = template_document.add_table(rows=4, cols=7)
        header_titles = [
            "Ref",
            "Date",
            "Type",
            "Name / Company",
            "Location",
            "Contact",
            "Issued",
        ]
        for index, title in enumerate(header_titles):
            register_table.cell(0, index).text = title

        start_paragraph = register_table.cell(1, 0).paragraphs[0]
        start_paragraph.add_run("{% tr ")
        start_paragraph.add_run("for p in permits ")
        start_paragraph.add_run("%}")

        register_table.cell(2, 0).text = "{{p.ref}}"
        register_table.cell(2, 1).text = "{{p.date}}"
        register_table.cell(2, 2).text = "{{p.type}}"
        register_table.cell(2, 3).text = "{{p.name_company}}"
        register_table.cell(2, 4).text = "{{p.location}}"
        register_table.cell(2, 5).text = "{{p.contact}}"
        register_table.cell(2, 6).text = "{{p.time_issued}} {{p.time_cancelled}}"

        end_paragraph = register_table.cell(3, 0).paragraphs[0]
        end_paragraph.add_run("{% tr ")
        end_paragraph.add_run("endfor ")
        end_paragraph.add_run("%}")
        template_document.save(template_path)

    def _build_site_induction_template(self, template_path: Path) -> None:
        template_document = Document()
        template_document.add_paragraph("{{site_name}}")
        template_document.add_paragraph("{{date}}")
        template_document.add_paragraph("{{full_name}}")
        template_document.add_paragraph("{{company}}")
        template_document.add_paragraph("{{home_address}}")
        template_document.add_paragraph("{{contact_number}}")
        template_document.add_paragraph("{{occupation}}")
        template_document.add_paragraph("{{emergency_contact}}")
        template_document.add_paragraph("{{emergency_tel}}")
        template_document.add_paragraph("{{medical}}")
        template_document.add_paragraph("{{cscs_no}}")
        template_document.add_paragraph("{{cscs_expiry}}")
        template_document.add_paragraph("{{asbestos_cert}}")
        template_document.add_paragraph("{{erect_scaffold}}")
        template_document.add_paragraph("{{cisrs_no}}")
        template_document.add_paragraph("{{cisrs_expiry}}")
        template_document.add_paragraph("{{operate_plant}}")
        template_document.add_paragraph("{{cpcs_no}}")
        template_document.add_paragraph("{{cpcs_expiry}}")
        template_document.add_paragraph("{{client_training_desc}}")
        template_document.add_paragraph("{{client_training_date}}")
        template_document.add_paragraph("{{client_training_expiry}}")
        template_document.add_paragraph("{{first_aider}}")
        template_document.add_paragraph("{{fire_warden}}")
        template_document.add_paragraph("{{supervisor}}")
        template_document.add_paragraph("{{smsts}}")
        template_document.add_paragraph("{{inductor_name_date}}")
        template_document.add_paragraph("{{inductor_title}}")
        template_document.add_paragraph("{{signature_image}}")
        template_document.save(template_path)

    def test_extract_tonnage_from_ticket_converts_kg_to_tonnes(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / "waste-ticket.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text(
                (72, 72),
                "Net Weight 3520 KG",
            )
            document.save(pdf_path)
            document.close()

            self.assertEqual(
                extract_tonnage_from_ticket(pdf_path),
                3.52,
            )

    def test_check_site_inductions_flags_missing_workers_against_induction_pdfs(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            induction_directory = workspace_root / "FILE_3_Inductions"
            database_path = workspace_root / "documents.sqlite3"
            induction_directory.mkdir(parents=True, exist_ok=True)
            (induction_directory / "S_Carter_Induction.pdf").write_bytes(b"%PDF-1.4\n")

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.INDUCTION_DIR,
                app_config.DATABASE_PATH,
            )

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INDUCTION_DIR = induction_directory
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                repository.save(
                    SiteAttendanceRegister(
                        doc_id="SAR-001",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 10, 7, 0),
                        status=DocumentStatus.ACTIVE,
                        attendance_records=[
                            SiteAttendanceRecord(
                                date=date(2026, 3, 10),
                                company="Acme Interiors",
                                workerName="S. Carter",
                                timeIn="07:30",
                                timeOut="16:00",
                                totalHours=8.5,
                            ),
                            SiteAttendanceRecord(
                                date=date(2026, 3, 10),
                                company="Acme Interiors",
                                workerName="J. Evans",
                                timeIn="07:30",
                                timeOut="16:00",
                                totalHours=8.5,
                            ),
                        ],
                    )
                )

                result = check_site_inductions(
                    repository,
                    on_date=date(2026, 3, 10),
                    site_name="NG Lovedean Substation",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.INDUCTION_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            self.assertFalse(result.is_compliant)
            self.assertEqual(result.workers_on_site, ["J. Evans", "S. Carter"])
            self.assertEqual(result.inducted_workers, ["S. Carter"])
            self.assertEqual(result.missing_workers, ["J. Evans"])
            self.assertEqual(
                result.matched_files["S. Carter"].name,
                "S_Carter_Induction.pdf",
            )

    def test_create_ladder_permit_draft_renders_and_indexes_without_induction_pdf(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            today = date.today()
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_template.docx"
            self._build_ladder_tagged_template(template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                attendance_register = SiteAttendanceRegister(
                    doc_id="SAR-PERMIT-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime.combine(today, time(7, 0)),
                    status=DocumentStatus.ACTIVE,
                    attendance_records=[
                        SiteAttendanceRecord(
                            date=today,
                            company="Abucs",
                            workerName="S. Carter",
                            timeIn="07:30",
                            timeOut="16:00",
                            totalHours=8.5,
                        )
                    ],
                )
                repository.save(attendance_register)

                result = create_ladder_permit_draft(
                    repository,
                    attendance_record=attendance_register.attendance_records[0],
                    description_of_work="Installing CCTV cameras",
                    location_of_work="Transformer bay",
                    supervisor_name="Ceri Edwards",
                    safety_checklist={
                        question_number: question_number != 10
                        for question_number in range(1, 12)
                    },
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=False,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=False,
                    inspection_comments="Replace worn stile before use",
                    site_name=attendance_register.site_name,
                    job_number="JOB-001",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertTrue(result.output_path.exists())
            self.assertEqual(result.output_path.parent, permits_directory)
            self.assertEqual(result.permit.worker_name, "S. Carter")
            self.assertEqual(result.permit.worker_company, "Abucs")
            self.assertEqual(result.permit.issued_date, today)
            self.assertEqual(result.permit.project_number, "JOB-001")
            self.assertEqual(result.permit.permit_number, "LADD-001")
            self.assertEqual(result.permit.description_of_work, "Installing CCTV cameras")
            self.assertEqual(result.permit.competent_supervisor_name, "Ceri Edwards")
            self.assertEqual(result.permit.to_template_context()["q10_yes"], "")
            self.assertEqual(result.permit.to_template_context()["q10_no"], "✔")
            self.assertEqual(result.permit.to_template_context()["manager_name"], "Ceri Edwards")
            self.assertEqual(result.permit.to_template_context()["auth_name"], "Ceri Edwards")
            self.assertEqual(result.permit.to_template_context()["issue_name"], "Ceri Edwards")
            self.assertEqual(
                result.permit.to_template_context()["manager_position"],
                "Project Manager",
            )
            self.assertEqual(result.permit.to_template_context()["insp_name"], "Ceri Edwards")
            self.assertEqual(result.permit.to_template_context()["insp_rungs"], "✔")
            self.assertEqual(result.permit.to_template_context()["insp_stiles"], "✘")
            self.assertEqual(result.permit.to_template_context()["insp_feet"], "✔")
            self.assertEqual(result.permit.to_template_context()["insp_ok"], "✘")
            self.assertEqual(
                result.permit.to_template_context()["insp_comments"],
                "Replace worn stile before use",
            )
            self.assertIsNone(result.induction_file)

            saved_permit = repository.get(result.permit.doc_id)
            self.assertIsInstance(saved_permit, LadderPermit)
            indexed_files = repository.list_indexed_files(related_doc_id=result.permit.doc_id)
            self.assertEqual(len(indexed_files), 1)
            self.assertEqual(indexed_files[0].file_group, FileGroup.FILE_4)

    def test_create_ladder_permit_draft_prefills_management_sections_in_output_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            today = date.today()
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_management_template.docx"
            self._build_ladder_management_sections_template(template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                attendance_record = SiteAttendanceRecord(
                    date=today,
                    company="Abucs",
                    workerName="S. Carter",
                    timeIn="07:30",
                    timeOut="16:00",
                    totalHours=8.5,
                )

                result = create_ladder_permit_draft(
                    repository,
                    attendance_record=attendance_record,
                    description_of_work="Installing CCTV cameras",
                    location_of_work="Transformer bay",
                    supervisor_name="Shift Supervisor",
                    safety_checklist={
                        question_number: True for question_number in range(1, 12)
                    },
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=True,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=True,
                    inspection_comments="No defects found",
                    site_name="NG Lovedean Substation",
                    job_number="JOB-001",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            rendered = Document(result.output_path)
            header_table = rendered.tables[0]
            acceptance_table = rendered.tables[4]
            self.assertEqual(header_table.rows[3].cells[4].text, result.permit.valid_from_time.strftime("%H:%M"))
            self.assertEqual(
                header_table.rows[4].cells[1].text,
                result.permit.valid_to_date.strftime("%d/%m/%Y"),
            )
            self.assertEqual(
                header_table.rows[4].cells[4].text,
                result.permit.valid_to_time.strftime("%H:%M"),
            )
            self.assertEqual(acceptance_table.rows[1].cells[3].text, "Position: Project Manager")
            self.assertEqual(acceptance_table.rows[5].cells[0].text, "Name: S. Carter")
            self.assertEqual(
                acceptance_table.rows[5].cells[2].text,
                f"Date: (dd/mm/yyyy) {today.strftime('%d/%m/%Y')}",
            )
            self.assertEqual(acceptance_table.rows[5].cells[3].text, "Company: Abucs")
            self.assertEqual(acceptance_table.rows[7].cells[0].text, "Name: Ceri Edwards")
            self.assertEqual(
                acceptance_table.rows[7].cells[2].text,
                f"Date: (dd/mm/yyyy) {today.strftime('%d/%m/%Y')}",
            )
            self.assertEqual(acceptance_table.rows[7].cells[3].text, "Position: Project Manager")

    def test_create_ladder_permit_draft_accepts_latest_non_today_attendance_row(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            attendance_date = date(2026, 3, 4)
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_template.docx"
            self._build_ladder_tagged_template(template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                attendance_record = SiteAttendanceRecord(
                    date=attendance_date,
                    company="Abucs",
                    workerName="J. Evans",
                    timeIn="07:30",
                    timeOut="16:00",
                    totalHours=8.5,
                )

                result = create_ladder_permit_draft(
                    repository,
                    attendance_record=attendance_record,
                    description_of_work="Fire alarm sensor replacement",
                    location_of_work="Switch room",
                    supervisor_name="Ceri Edwards",
                    safety_checklist={question_number: True for question_number in range(1, 12)},
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=True,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=True,
                    inspection_comments="No defects found",
                    site_name="NG Lovedean Substation",
                    job_number="JOB-002",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertEqual(result.permit.valid_from_date, result.permit.issued_date)
            self.assertEqual(
                result.permit.valid_to_datetime - result.permit.valid_from_datetime,
                timedelta(hours=8),
            )
            self.assertTrue(result.output_path.exists())
            self.assertEqual(result.permit.project_number, "JOB-002")

    def test_create_ladder_permit_draft_uses_roster_name_for_operative_context(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            attendance_date = date(2026, 3, 4)
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_template.docx"
            self._build_ladder_tagged_template(template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                attendance_record = SiteAttendanceRecord(
                    date=attendance_date,
                    company="Abucs",
                    workerName="Abucs Ltd",
                    timeIn="07:30",
                    timeOut="16:00",
                    totalHours=8.5,
                )
                selected_worker = SiteWorker(
                    company="Abucs",
                    worker_name="Sean Michael Carter",
                    last_on_site_date=attendance_date,
                )

                result = create_ladder_permit_draft(
                    repository,
                    attendance_record=attendance_record,
                    site_worker=selected_worker,
                    description_of_work="Fire alarm sensor replacement",
                    location_of_work="Switch room",
                    supervisor_name="Ceri Edwards",
                    safety_checklist={
                        question_number: True for question_number in range(1, 12)
                    },
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=True,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=True,
                    inspection_comments="No defects found",
                    site_name="NG Lovedean Substation",
                    job_number="JOB-002",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            context = result.permit.to_template_context()
            self.assertEqual(result.permit.worker_name, "Sean Michael Carter")
            self.assertEqual(result.permit.worker_company, "Abucs")
            self.assertEqual(context["company_name"], "Abucs")
            self.assertEqual(context["contractor_name"], "Sean Carter")
            self.assertEqual(context["op_name"], "Sean Carter")

    def test_create_ladder_permit_draft_uses_current_system_time_for_valid_from(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            attendance_date = date(2026, 3, 4)
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_template.docx"
            self._build_ladder_tagged_template(template_path)
            fixed_issue_time = datetime(2026, 3, 11, 8, 30)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                attendance_record = SiteAttendanceRecord(
                    date=attendance_date,
                    company="Abucs",
                    workerName="J. Evans",
                    timeIn="07:30",
                    timeOut="16:00",
                    totalHours=8.5,
                )

                with patch(
                    "uplands_site_command_centre.workspace._current_permit_issue_datetime",
                    return_value=fixed_issue_time,
                ):
                    result = create_ladder_permit_draft(
                        repository,
                        attendance_record=attendance_record,
                        description_of_work="Fire alarm sensor replacement",
                        location_of_work="Switch room",
                        supervisor_name="Ceri Edwards",
                        safety_checklist={
                            question_number: True for question_number in range(1, 12)
                        },
                        inspection_checked_by="Ceri Edwards",
                        inspection_rungs_ok=True,
                        inspection_stiles_ok=True,
                        inspection_feet_ok=True,
                        inspection_ok_to_use=True,
                        inspection_comments="No defects found",
                        site_name="NG Lovedean Substation",
                        job_number="JOB-002",
                    )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertEqual(result.permit.valid_from_time, time(8, 30))
            self.assertEqual(result.permit.valid_to_time, time(16, 30))
            self.assertEqual(result.permit.issued_date, fixed_issue_time.date())
            self.assertEqual(result.permit.valid_from_date, fixed_issue_time.date())
            self.assertEqual(result.permit.valid_to_date, fixed_issue_time.date())

    def test_create_ladder_permit_draft_rolls_expiry_date_when_eight_hours_crosses_midnight(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            attendance_date = date(2026, 3, 4)
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_template.docx"
            self._build_ladder_tagged_template(template_path)
            fixed_issue_time = datetime(2026, 3, 11, 20, 30)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                attendance_record = SiteAttendanceRecord(
                    date=attendance_date,
                    company="Abucs",
                    workerName="J. Evans",
                    timeIn="07:30",
                    timeOut="16:00",
                    totalHours=8.5,
                )

                with patch(
                    "uplands_site_command_centre.workspace._current_permit_issue_datetime",
                    return_value=fixed_issue_time,
                ):
                    result = create_ladder_permit_draft(
                        repository,
                        attendance_record=attendance_record,
                        description_of_work="Fire alarm sensor replacement",
                        location_of_work="Switch room",
                        supervisor_name="Ceri Edwards",
                        safety_checklist={
                            question_number: True for question_number in range(1, 12)
                        },
                        inspection_checked_by="Ceri Edwards",
                        inspection_rungs_ok=True,
                        inspection_stiles_ok=True,
                        inspection_feet_ok=True,
                        inspection_ok_to_use=True,
                        inspection_comments="No defects found",
                        site_name="NG Lovedean Substation",
                        job_number="JOB-002",
                    )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertEqual(result.permit.valid_from_time, time(20, 30))
            self.assertEqual(result.permit.valid_to_time, time(4, 30))
            self.assertEqual(result.permit.valid_from_date, date(2026, 3, 11))
            self.assertEqual(result.permit.valid_to_date, date(2026, 3, 12))

    def test_create_ladder_permit_draft_requires_job_number_from_project_setup(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_template.docx"
            self._build_ladder_tagged_template(template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                attendance_record = SiteAttendanceRecord(
                    date=date.today(),
                    company="Abucs",
                    workerName="S. Carter",
                    timeIn="07:30",
                    timeOut="16:00",
                    totalHours=8.5,
                )

                with self.assertRaises(ValidationError):
                    create_ladder_permit_draft(
                        repository,
                        attendance_record=attendance_record,
                        description_of_work="Installing CCTV cameras",
                        location_of_work="Switch room",
                        supervisor_name="Ceri Edwards",
                        safety_checklist={
                            question_number: True for question_number in range(1, 12)
                        },
                        inspection_checked_by="Ceri Edwards",
                        inspection_rungs_ok=True,
                        inspection_stiles_ok=True,
                        inspection_feet_ok=True,
                        inspection_ok_to_use=True,
                        inspection_comments="No defects found",
                        site_name="NG Lovedean Substation",
                        job_number="",
                    )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

    def test_create_ladder_permit_draft_auto_numbers_sequentially_for_site(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            today = date.today()
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_template.docx"
            self._build_ladder_tagged_template(template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                first_record = SiteAttendanceRecord(
                    date=today,
                    company="Abucs",
                    workerName="S. Carter",
                    timeIn="07:30",
                    timeOut="16:00",
                    totalHours=8.5,
                )
                second_record = SiteAttendanceRecord(
                    date=today,
                    company="Abucs",
                    workerName="J. Evans",
                    timeIn="08:00",
                    timeOut="16:30",
                    totalHours=8.0,
                )

                first_permit = create_ladder_permit_draft(
                    repository,
                    attendance_record=first_record,
                    description_of_work="Install tray supports",
                    location_of_work="Bay 1",
                    supervisor_name="Ceri Edwards",
                    safety_checklist={question_number: True for question_number in range(1, 12)},
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=True,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=True,
                    inspection_comments="No defects found",
                    site_name="NG Lovedean Substation",
                    job_number="JOB-010",
                )
                second_permit = create_ladder_permit_draft(
                    repository,
                    attendance_record=second_record,
                    description_of_work="Inspect cable ladder",
                    location_of_work="Bay 2",
                    supervisor_name="Ceri Edwards",
                    safety_checklist={question_number: True for question_number in range(1, 12)},
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=True,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=True,
                    inspection_comments="No defects found",
                    site_name="NG Lovedean Substation",
                    job_number="JOB-010",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertEqual(first_permit.permit.permit_number, "LADD-001")
            self.assertEqual(second_permit.permit.permit_number, "LADD-002")

    def test_generate_permit_register_document_renders_and_indexes_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            today = date.today()
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            ladder_template_path = Path(temp_dir) / "ladder_template.docx"
            register_template_path = Path(temp_dir) / "permit_register_template.docx"
            self._build_ladder_tagged_template(ladder_template_path)
            self._build_permit_register_template(register_template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = ladder_template_path
                TemplateRegistry.TEMPLATE_PATHS["permit_register"] = register_template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                first_record = SiteAttendanceRecord(
                    date=today,
                    company="Abucs",
                    workerName="S. Carter",
                    timeIn="07:30",
                    timeOut="16:00",
                    totalHours=8.5,
                )
                second_record = SiteAttendanceRecord(
                    date=today,
                    company="Abucs",
                    workerName="J. Evans",
                    timeIn="08:00",
                    timeOut="16:30",
                    totalHours=8.0,
                )

                create_ladder_permit_draft(
                    repository,
                    attendance_record=first_record,
                    description_of_work="Install tray supports",
                    location_of_work="Bay 1",
                    supervisor_name="Ceri Edwards",
                    safety_checklist={question_number: True for question_number in range(1, 12)},
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=True,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=True,
                    inspection_comments="No defects found",
                    site_name="NG Lovedean Substation",
                    job_number="JOB-777",
                )
                create_ladder_permit_draft(
                    repository,
                    attendance_record=second_record,
                    description_of_work="Inspect cable ladder",
                    location_of_work="Bay 2",
                    supervisor_name="Ceri Edwards",
                    safety_checklist={question_number: True for question_number in range(1, 12)},
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=True,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=True,
                    inspection_comments="No defects found",
                    site_name="NG Lovedean Substation",
                    job_number="JOB-777",
                )

                result = generate_permit_register_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    job_number="JOB-777",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertTrue(result.output_path.exists())
            self.assertEqual(result.output_path.parent, permits_directory)
            self.assertEqual(result.permit_count, 2)

            rendered_document = Document(result.output_path)
            rendered_text = "\n".join(
                paragraph.text for paragraph in rendered_document.paragraphs
            )
            table_text = "\n".join(
                cell.text
                for table in rendered_document.tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertIn("NG Lovedean Substation", rendered_text)
            self.assertIn("JOB-777", rendered_text)
            self.assertIn("LADD-001", table_text)
            self.assertIn("LADD-002", table_text)
            self.assertIn("S. Carter | Abucs", table_text)
            self.assertIn("J. Evans | Abucs", table_text)
            register_table = next(
                table
                for table in rendered_document.tables
                if table.rows
                and table.cell(0, 0).text.strip()
                in {"Permit Reference Number", "Ref"}
            )
            first_data_row = register_table.rows[1]
            permit_ref_run = first_data_row.cells[0].paragraphs[0].runs[0]
            issued_to_run = first_data_row.cells[3].paragraphs[0].runs[0]
            self.assertEqual(permit_ref_run.font.name, "Arial")
            self.assertAlmostEqual(permit_ref_run.font.size.pt, 9.0)
            self.assertEqual(issued_to_run.font.name, "Arial")
            self.assertAlmostEqual(issued_to_run.font.size.pt, 9.0)

            indexed_files = repository.list_indexed_files(
                file_category="permit_register_docx"
            )
            self.assertEqual(len(indexed_files), 1)
            self.assertEqual(indexed_files[0].file_group, FileGroup.FILE_4)

    def test_sync_file_4_permit_records_removes_ghost_permits(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            today = date.today()
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            permits_directory = workspace_root / "FILE_4_Permits"
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "ladder_template.docx"
            self._build_ladder_tagged_template(template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.PERMITS_DESTINATION,
                app_config.DATABASE_PATH,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.PERMITS_DESTINATION = permits_directory
                app_config.DATABASE_PATH = database_path
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                generated_permit = create_ladder_permit_draft(
                    repository,
                    attendance_record=SiteAttendanceRecord(
                        date=today,
                        company="Abucs",
                        workerName="S. Carter",
                        timeIn="07:30",
                        timeOut="16:00",
                        totalHours=8.5,
                    ),
                    description_of_work="Installing CCTV cameras",
                    location_of_work="Switch room",
                    supervisor_name="Ceri Edwards",
                    safety_checklist={
                        question_number: True for question_number in range(1, 12)
                    },
                    inspection_checked_by="Ceri Edwards",
                    inspection_rungs_ok=True,
                    inspection_stiles_ok=True,
                    inspection_feet_ok=True,
                    inspection_ok_to_use=True,
                    inspection_comments="No defects found",
                    site_name="NG Lovedean Substation",
                    job_number="JOB-900",
                )
                generated_permit.output_path.unlink()

                sync_result = sync_file_4_permit_records(
                    repository,
                    site_name="NG Lovedean Substation",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.PERMITS_DESTINATION,
                    app_config.DATABASE_PATH,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertEqual(sync_result.removed_count, 1)
            self.assertEqual(
                repository.list_documents(
                    document_type=LadderPermit.document_type,
                    site_name="NG Lovedean Substation",
                ),
                [],
            )
            self.assertEqual(
                repository.list_indexed_files(file_category="ladder_permit_docx"),
                [],
            )

    def test_extract_expiry_date_from_pdf_prefers_expiry_phrase_and_ignores_generation_dates(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / "Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text(
                (72, 72),
                (
                    "Date of registration 13 September 2023\n"
                    "Expiry date of registration (unless revoked) 25 October 2026\n"
                    "This certificate was created on 13 September 2023. "
                    "These details are correct at the time of certificate generation."
                ),
            )
            document.save(pdf_path)
            document.close()

            self.assertEqual(
                extract_expiry_date_from_pdf(pdf_path),
                date(2026, 10, 25),
            )

    def test_extract_expiry_date_from_pdf_uses_filename_when_expiry_is_not_in_visible_text(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = (
                Path(temp_dir)
                / "Liability Insurance - Biffa Ltd Subsidiaries - Expiry 31.03.2026.pdf"
            )
            document = fitz.open()
            page = document.new_page()
            page.insert_text(
                (72, 72),
                (
                    "Attachment to letter dated 17-March-2025\n"
                    "This certificate was created on 17 March 2025."
                ),
            )
            document.save(pdf_path)
            document.close()

            self.assertEqual(
                extract_expiry_date_from_pdf(pdf_path),
                date(2026, 3, 31),
            )

    def test_file_and_index_all_overwrites_stale_carrier_record_using_reference_number(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            inbox = workspace_root / "ingest"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            carrier_docs_destination = workspace_root / "FILE_1_Environment" / "Carrier_Docs"
            waste_reports_destination = workspace_root / "FILE_1_Environment" / "Waste_Reports"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            plant_hire_directory = workspace_root / "FILE_2_Registers" / "Plant_Hire_Register"
            database_path = workspace_root / "documents.sqlite3"
            inbox.mkdir(parents=True, exist_ok=True)

            ls_pdf_path = inbox / "Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877.pdf"
            ls_pdf_document = fitz.open()
            ls_pdf_page = ls_pdf_document.new_page()
            ls_pdf_page.insert_text(
                (72, 72),
                (
                    "Name of registered carrier L&S WASTE MANAGEMENT LIMITED\n"
                    "Date of registration 13 September 2023\n"
                    "Expiry date of registration (unless revoked) 25 October 2026\n"
                    "This certificate was created on 13 September 2023. "
                    "These details are correct at the time of certificate generation."
                ),
            )
            ls_pdf_document.save(ls_pdf_path)
            ls_pdf_document.close()

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.INBOX,
                app_config.WASTE_DESTINATION,
                app_config.CARRIER_DOCS_DESTINATION,
                app_config.WASTE_REPORTS_DESTINATION,
                app_config.ATTENDANCE_DESTINATION,
                app_config.PLANT_HIRE_REGISTER_DIR,
                app_config.INDUCTION_DIR,
                app_config.RAMS_DESTINATION,
                app_config.COSHH_DESTINATION,
                app_config.FILE_3_REVIEW_DIR,
                app_config.FILE_3_OUTPUT_DIR,
                app_config.DATABASE_PATH,
            )

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INBOX = inbox
                app_config.WASTE_DESTINATION = waste_destination
                app_config.CARRIER_DOCS_DESTINATION = carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = plant_hire_directory
                app_config.INDUCTION_DIR = workspace_root / "FILE_3_Inductions"
                app_config.FILE_3_REVIEW_DIR = workspace_root / "FILE_3_Inductions" / "Needs_Review"
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                repository.save(
                    CarrierComplianceDocument(
                        doc_id="CCD-carriers-exp-cbdu198877-licence",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 10, 8, 0),
                        status=DocumentStatus.ACTIVE,
                        carrier_name="Carriers Exp . . Cbdu198877",
                        carrier_document_type=CarrierComplianceDocumentType.LICENCE,
                        reference_number="Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877",
                        expiry_date=date(2023, 9, 13),
                    )
                )

                filed_assets = file_and_index_all(repository)
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.INBOX,
                    app_config.WASTE_DESTINATION,
                    app_config.CARRIER_DOCS_DESTINATION,
                    app_config.WASTE_REPORTS_DESTINATION,
                    app_config.ATTENDANCE_DESTINATION,
                    app_config.PLANT_HIRE_REGISTER_DIR,
                    app_config.INDUCTION_DIR,
                    app_config.RAMS_DESTINATION,
                    app_config.COSHH_DESTINATION,
                    app_config.FILE_3_REVIEW_DIR,
                    app_config.FILE_3_OUTPUT_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            carrier_documents = repository.list_documents(
                document_type=CarrierComplianceDocument.document_type
            )
            self.assertEqual(len(carrier_documents), 1)
            self.assertEqual(carrier_documents[0].doc_id, "CCD-carriers-exp-cbdu198877-licence")
            self.assertEqual(carrier_documents[0].carrier_name, "Abucs")
            self.assertEqual(
                carrier_documents[0].carrier_document_type,
                CarrierComplianceDocumentType.LICENCE,
            )
            self.assertEqual(carrier_documents[0].expiry_date, date(2026, 10, 25))
            self.assertEqual(len(filed_assets), 1)

    def test_file_and_index_all_archives_duplicate_active_record_for_same_reference(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            inbox = workspace_root / "ingest"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            carrier_docs_destination = workspace_root / "FILE_1_Environment" / "Carrier_Docs"
            waste_reports_destination = workspace_root / "FILE_1_Environment" / "Waste_Reports"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            plant_hire_directory = workspace_root / "FILE_2_Registers" / "Plant_Hire_Register"
            database_path = workspace_root / "documents.sqlite3"
            inbox.mkdir(parents=True, exist_ok=True)

            ls_pdf_path = inbox / "Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877.pdf"
            ls_pdf_document = fitz.open()
            ls_pdf_page = ls_pdf_document.new_page()
            ls_pdf_page.insert_text(
                (72, 72),
                (
                    "Name of registered carrier L&S WASTE MANAGEMENT LIMITED\n"
                    "Date of registration 13 September 2023\n"
                    "Expiry date of registration (unless revoked) 25 October 2026\n"
                    "This certificate was created on 13 September 2023. "
                    "These details are correct at the time of certificate generation."
                ),
            )
            ls_pdf_document.save(ls_pdf_path)
            ls_pdf_document.close()

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.INBOX,
                app_config.WASTE_DESTINATION,
                app_config.CARRIER_DOCS_DESTINATION,
                app_config.WASTE_REPORTS_DESTINATION,
                app_config.ATTENDANCE_DESTINATION,
                app_config.PLANT_HIRE_REGISTER_DIR,
                app_config.INDUCTION_DIR,
                app_config.RAMS_DESTINATION,
                app_config.COSHH_DESTINATION,
                app_config.FILE_3_REVIEW_DIR,
                app_config.FILE_3_OUTPUT_DIR,
                app_config.DATABASE_PATH,
            )

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INBOX = inbox
                app_config.WASTE_DESTINATION = waste_destination
                app_config.CARRIER_DOCS_DESTINATION = carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = plant_hire_directory
                app_config.INDUCTION_DIR = workspace_root / "FILE_3_Inductions"
                app_config.RAMS_DESTINATION = workspace_root / "FILE_3_Safety" / "RAMS"
                app_config.COSHH_DESTINATION = workspace_root / "FILE_3_Safety" / "COSHH"
                app_config.FILE_3_REVIEW_DIR = workspace_root / "FILE_3_Safety" / "Needs_Review"
                app_config.FILE_3_OUTPUT_DIR = workspace_root / "FILE_3_Safety" / "Registers"
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                repository.save(
                    CarrierComplianceDocument(
                        doc_id="CCD-abucs-licence",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 10, 9, 0),
                        status=DocumentStatus.ACTIVE,
                        carrier_name="Abucs",
                        carrier_document_type=CarrierComplianceDocumentType.LICENCE,
                        reference_number="Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877",
                        expiry_date=date(2026, 10, 25),
                    )
                )
                repository.save(
                    CarrierComplianceDocument(
                        doc_id="CCD-carriers-exp-cbdu198877-licence",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 10, 8, 0),
                        status=DocumentStatus.ACTIVE,
                        carrier_name="Carriers Exp . . Cbdu198877",
                        carrier_document_type=CarrierComplianceDocumentType.LICENCE,
                        reference_number="Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877",
                        expiry_date=date(2023, 9, 13),
                    )
                )

                file_and_index_all(repository)
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.INBOX,
                    app_config.WASTE_DESTINATION,
                    app_config.CARRIER_DOCS_DESTINATION,
                    app_config.WASTE_REPORTS_DESTINATION,
                    app_config.ATTENDANCE_DESTINATION,
                    app_config.PLANT_HIRE_REGISTER_DIR,
                    app_config.INDUCTION_DIR,
                    app_config.RAMS_DESTINATION,
                    app_config.COSHH_DESTINATION,
                    app_config.FILE_3_REVIEW_DIR,
                    app_config.FILE_3_OUTPUT_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            carrier_documents = sorted(
                repository.list_documents(
                    document_type=CarrierComplianceDocument.document_type
                ),
                key=lambda document: document.doc_id,
            )
            self.assertEqual(len(carrier_documents), 2)
            self.assertEqual(carrier_documents[0].doc_id, "CCD-abucs-licence")
            self.assertEqual(carrier_documents[0].status, DocumentStatus.ACTIVE)
            self.assertEqual(carrier_documents[1].doc_id, "CCD-carriers-exp-cbdu198877-licence")
            self.assertEqual(carrier_documents[1].status, DocumentStatus.ARCHIVED)

    def test_file_and_index_all_moves_supported_files_and_indexes_paths(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            inbox = workspace_root / "ingest"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            carrier_docs_destination = workspace_root / "FILE_1_Environment" / "Carrier_Docs"
            waste_reports_destination = workspace_root / "FILE_1_Environment" / "Waste_Reports"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            database_path = workspace_root / "documents.sqlite3"
            inbox.mkdir(parents=True, exist_ok=True)

            pdf_path = inbox / "31194.PDF"
            waste_pdf_document = fitz.open()
            waste_pdf_page = waste_pdf_document.new_page()
            waste_pdf_page.insert_text(
                (72, 72),
                (
                    "Ticket No.\n"
                    "10/03/2026\n"
                    "31194\n"
                    "National Grid Broadway Lane Lovedean Waterlooville\n"
                    "Waste Type:Mixed Construction\n"
                    "17 09 04\n"
                    "Weight\n"
                    "3.52 Tonnes\n"
                    "DUTY OF CARE WASTE TRANSFER NOTE / INVOICE\n"
                ),
            )
            waste_pdf_document.save(pdf_path)
            waste_pdf_document.close()

            biffa_pdf_path = (
                inbox
                / "Liability Insurance   - EL 20m PL  PI  10m  -  Biffa Ltd  Subsidiaries - Expiry 31.03.2026.pdf"
            )
            biffa_pdf_document = fitz.open()
            biffa_pdf_page = biffa_pdf_document.new_page()
            biffa_pdf_page.insert_text(
                (72, 72),
                (
                    "Attachment to letter dated 17-March-2025\n"
                    "This certificate was created on 17 March 2025."
                ),
            )
            biffa_pdf_document.save(biffa_pdf_path)
            biffa_pdf_document.close()

            ls_pdf_path = inbox / "Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877.pdf"
            ls_pdf_document = fitz.open()
            ls_pdf_page = ls_pdf_document.new_page()
            ls_pdf_page.insert_text(
                (72, 72),
                (
                    "Name of registered carrier L&S WASTE MANAGEMENT LIMITED\n"
                    "Date of registration 13 September 2023\n"
                    "Expiry date of registration (unless revoked) 25 October 2026\n"
                    "This certificate was created on 13 September 2023. "
                    "These details are correct at the time of certificate generation."
                ),
            )
            ls_pdf_document.save(ls_pdf_path)
            ls_pdf_document.close()

            waste_report_path = inbox / "March_Waste_Report.xlsx"
            waste_report_path.write_bytes(b"PK\x03\x04")

            json_path = inbox / "site-kpi-backup.json"
            with json_path.open("w", encoding="utf-8") as file_handle:
                json.dump(
                    {
                        "settings": {"siteName": "Uplands - Cardiff"},
                        "extractedRows": {
                            "weekly": [
                                {
                                    "date": "10/03/2026",
                                    "company": "Acme Interiors",
                                    "workerName": "S. Carter",
                                    "timeIn": "07:30",
                                    "timeOut": "16:00",
                                    "totalHours": 8.5,
                                }
                            ],
                            "eom": [],
                        },
                    },
                    file_handle,
                )

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.INBOX,
                app_config.WASTE_DESTINATION,
                app_config.CARRIER_DOCS_DESTINATION,
                app_config.WASTE_REPORTS_DESTINATION,
                app_config.ATTENDANCE_DESTINATION,
                app_config.PLANT_HIRE_REGISTER_DIR,
                app_config.INDUCTION_DIR,
                app_config.RAMS_DESTINATION,
                app_config.COSHH_DESTINATION,
                app_config.FILE_3_REVIEW_DIR,
                app_config.FILE_3_OUTPUT_DIR,
                app_config.DATABASE_PATH,
            )

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INBOX = inbox
                app_config.WASTE_DESTINATION = waste_destination
                app_config.CARRIER_DOCS_DESTINATION = carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = workspace_root / "FILE_2_Registers" / "Plant_Hire_Register"
                app_config.INDUCTION_DIR = workspace_root / "FILE_3_Inductions"
                app_config.RAMS_DESTINATION = workspace_root / "FILE_3_Safety" / "RAMS"
                app_config.COSHH_DESTINATION = workspace_root / "FILE_3_Safety" / "COSHH"
                app_config.FILE_3_REVIEW_DIR = workspace_root / "FILE_3_Safety" / "Needs_Review"
                app_config.FILE_3_OUTPUT_DIR = workspace_root / "FILE_3_Safety" / "Registers"
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                filed_assets = file_and_index_all(repository)
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.INBOX,
                    app_config.WASTE_DESTINATION,
                    app_config.CARRIER_DOCS_DESTINATION,
                    app_config.WASTE_REPORTS_DESTINATION,
                    app_config.ATTENDANCE_DESTINATION,
                    app_config.PLANT_HIRE_REGISTER_DIR,
                    app_config.INDUCTION_DIR,
                    app_config.RAMS_DESTINATION,
                    app_config.COSHH_DESTINATION,
                    app_config.FILE_3_REVIEW_DIR,
                    app_config.FILE_3_OUTPUT_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            moved_pdf_path = waste_destination / "31194.PDF"
            moved_biffa_pdf_path = (
                carrier_docs_destination
                / "Liability Insurance   - EL 20m PL  PI  10m  -  Biffa Ltd  Subsidiaries - Expiry 31.03.2026.pdf"
            )
            moved_ls_pdf_path = (
                carrier_docs_destination
                / "Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877.pdf"
            )
            moved_waste_report_path = waste_reports_destination / "March_Waste_Report.xlsx"
            moved_json_path = attendance_destination / "site-kpi-backup.json"

            self.assertFalse(pdf_path.exists())
            self.assertFalse(biffa_pdf_path.exists())
            self.assertFalse(ls_pdf_path.exists())
            self.assertFalse(waste_report_path.exists())
            self.assertFalse(json_path.exists())
            self.assertTrue(moved_pdf_path.exists())
            self.assertTrue(moved_biffa_pdf_path.exists())
            self.assertTrue(moved_ls_pdf_path.exists())
            self.assertTrue(moved_waste_report_path.exists())
            self.assertTrue(moved_json_path.exists())
            self.assertEqual(
                {asset.file_category for asset in filed_assets},
                {
                    "abucs_pdf",
                    "carrier_doc_pdf",
                    "waste_report_excel",
                    "kpi_json",
                },
            )

            indexed_files = repository.list_indexed_files()
            self.assertEqual(len(indexed_files), 5)
            self.assertEqual(
                {record.file_path for record in indexed_files},
                {
                    moved_pdf_path.resolve(),
                    moved_biffa_pdf_path.resolve(),
                    moved_ls_pdf_path.resolve(),
                    moved_waste_report_path.resolve(),
                    moved_json_path.resolve(),
                },
            )
            carrier_documents = repository.list_documents(
                document_type=CarrierComplianceDocument.document_type
            )
            self.assertEqual(len(carrier_documents), 2)
            documents_by_type = {
                document.carrier_document_type: document
                for document in carrier_documents
                if isinstance(document, CarrierComplianceDocument)
            }
            self.assertEqual(
                set(documents_by_type.keys()),
                {
                    CarrierComplianceDocumentType.INSURANCE,
                    CarrierComplianceDocumentType.LICENCE,
                },
            )
            self.assertEqual(
                documents_by_type[CarrierComplianceDocumentType.INSURANCE].carrier_name,
                "Abucs",
            )
            self.assertEqual(
                documents_by_type[CarrierComplianceDocumentType.INSURANCE].expiry_date,
                date(2026, 3, 31),
            )
            self.assertEqual(
                documents_by_type[CarrierComplianceDocumentType.LICENCE].carrier_name,
                "Abucs",
            )
            self.assertEqual(
                documents_by_type[CarrierComplianceDocumentType.LICENCE].expiry_date,
                date(2026, 10, 25),
            )
            self.assertEqual(
                documents_by_type[
                    CarrierComplianceDocumentType.INSURANCE
                ].reference_number,
                "Liability Insurance   - EL 20m PL  PI  10m  -  Biffa Ltd  Subsidiaries - Expiry 31.03.2026",
            )
            self.assertEqual(
                documents_by_type[
                    CarrierComplianceDocumentType.LICENCE
                ].reference_number,
                "Waste-Carriers-Certificate-Exp-25.10.-2026-CBDU198877",
            )
            carrier_indexed_files = repository.list_indexed_files(
                file_category="carrier_doc_pdf"
            )
            self.assertEqual(len(carrier_indexed_files), 2)
            self.assertEqual(
                {record.related_doc_id for record in carrier_indexed_files},
                {
                    documents_by_type[CarrierComplianceDocumentType.INSURANCE].doc_id,
                    documents_by_type[CarrierComplianceDocumentType.LICENCE].doc_id,
                },
            )

            carrier_assets = [
                asset for asset in filed_assets if asset.file_category == "carrier_doc_pdf"
            ]
            self.assertEqual(len(carrier_assets), 2)
            captured_by_type = {
                asset.auto_captured_document_type: asset
                for asset in carrier_assets
            }
            self.assertEqual(
                captured_by_type[
                    CarrierComplianceDocumentType.INSURANCE
                ].auto_captured_expiry_date,
                date(2026, 3, 31),
            )
            self.assertEqual(
                captured_by_type[
                    CarrierComplianceDocumentType.INSURANCE
                ].auto_captured_carrier_name,
                "Abucs",
            )
            self.assertEqual(
                captured_by_type[
                    CarrierComplianceDocumentType.LICENCE
                ].auto_captured_expiry_date,
                date(2026, 10, 25),
            )
            self.assertEqual(
                captured_by_type[
                    CarrierComplianceDocumentType.LICENCE
                ].auto_captured_carrier_name,
                "Abucs",
            )

            waste_transfer_notes = repository.list_documents(
                document_type=WasteTransferNoteDocument.document_type
            )
            self.assertEqual(len(waste_transfer_notes), 1)
            self.assertEqual(waste_transfer_notes[0].wtn_number, "31194")
            self.assertEqual(waste_transfer_notes[0].quantity_tonnes, 3.52)
            self.assertEqual(waste_transfer_notes[0].date, date(2026, 3, 10))
            self.assertEqual(waste_transfer_notes[0].ewc_code, "17 09 04")
            self.assertEqual(waste_transfer_notes[0].carrier_name, "Abucs")
            self.assertEqual(waste_transfer_notes[0].status, DocumentStatus.ACTIVE)
            self.assertEqual(
                repository.list_indexed_files(file_category="abucs_pdf")[0].related_doc_id,
                waste_transfer_notes[0].doc_id,
            )

            attendance_registers = repository.list_documents(
                document_type=SiteAttendanceRegister.document_type
            )
            self.assertEqual(len(attendance_registers), 1)
            self.assertEqual(attendance_registers[0].site_name, "Uplands - Cardiff")
            self.assertEqual(
                repository.list_indexed_files(file_category="kpi_json")[0].related_doc_id,
                attendance_registers[0].doc_id,
            )



class TemplateManagerTests(unittest.TestCase):
    def build_permit(self) -> LadderPermit:
        permit = LadderPermit(
            doc_id="LP-004",
            site_name="Uplands - Bridgend",
            created_at=datetime(2026, 3, 10, 10, 0),
            status=DocumentStatus.DRAFT,
            permit_number="UHSF21.09-004",
            project_name="Bridgend Remedials",
            project_number="UP-24023",
            location_of_work="External stair tower",
            description_of_work="Bracket inspection at first-floor landing.",
            valid_from_date=date(2026, 3, 10),
            valid_from_time=time(10, 30),
            valid_to_date=date(2026, 3, 10),
            valid_to_time=time(16, 0),
            safer_alternative_eliminated=True,
            task_specific_rams_prepared_and_approved=True,
            personnel_briefed_and_understand_task=True,
            competent_supervisor_appointed=True,
            competent_supervisor_name="R. Davies",
            operatives_suitably_trained=True,
            ladder_length_suitable=True,
            conforms_to_bs_class_a=True,
            three_points_of_contact_maintained=False,
            harness_worn_and_secured_above_head_height=True,
            ladder_stabilisation_method=LadderStabilisationMethod.TIED_AT_TOP,
            equipment_inspected_for_defects=True,
            worker_name="R. Davies",
            worker_company="Uplands",
            issued_date=date(2026, 3, 10),
        )
        permit.add_inspection_record(
            inspection_date=date(2026, 3, 10),
            inspected_by="R. Davies",
            rungs_ok=True,
            stiles_ok=True,
            feet_ok=True,
            comments_or_action_taken="Ready for use.",
        )
        return permit

    def build_incident(self) -> IncidentLogDocument:
        return IncidentLogDocument(
            doc_id="INC-300",
            site_name="Uplands - Bridgend",
            created_at=datetime(2026, 3, 10, 13, 0),
            status=DocumentStatus.ACTIVE,
            incident_type=IncidentType.ACCIDENT,
            location="Front reception",
            description="Minor hand injury during unloading.",
            witness_list=["R. Davies", "C. Thomas"],
        )

    def _build_ladder_tagged_template(self, template_path: Path) -> None:
        document = Document()
        document.add_paragraph("Permit {{permit_number}} at {{site_name}}")
        document.add_paragraph("{{company_name}} / {{contractor_name}}")
        document.add_paragraph("{{op_name}}")
        document.add_paragraph("{{job_number}}")
        document.add_paragraph("{{date_issued}}")
        document.add_paragraph("{{task_description}}")
        document.add_paragraph("{{supervisor_name}}")
        for question_number in range(1, 12):
            document.add_paragraph(
                f"Q{question_number} {{{{q{question_number}_yes}}}} / {{{{q{question_number}_no}}}}"
            )
        document.add_paragraph("{{manager_name}}")
        document.add_paragraph("{{manager_signature}}")
        document.add_paragraph("{{ladder_id}}")
        table = document.add_table(rows=1, cols=7)
        table.cell(0, 0).text = "{{insp_date}}"
        table.cell(0, 1).text = "{{insp_name}}"
        table.cell(0, 2).text = "{{insp_rungs}}"
        table.cell(0, 3).text = "{{insp_stiles}}"
        table.cell(0, 4).text = "{{insp_feet}}"
        table.cell(0, 5).text = "{{insp_comments}}"
        table.cell(0, 6).text = "{{insp_ok}}"
        document.save(template_path)

    def test_template_manager_replaces_placeholders_for_ladder_permit(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "permit_template.docx"
            output_path = Path(temp_dir) / "permit_filled.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            self._build_ladder_tagged_template(template_path)

            try:
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path
                manager = TemplateManager(self.build_permit())
                manager.render(output_path)
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            rendered = Document(output_path)
            self.assertIn("UHSF21.09-004", rendered.paragraphs[0].text)
            self.assertEqual(rendered.paragraphs[1].text, "Uplands / R. Davies")
            self.assertEqual(rendered.paragraphs[2].text, "R. Davies")
            self.assertIn("UP-24023", rendered.paragraphs[3].text)
            self.assertEqual(rendered.tables[0].cell(0, 1).text, "R. Davies")
            self.assertEqual(rendered.tables[0].cell(0, 6).text, "✔")

    def test_template_manager_replaces_placeholders_for_incident_log(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "incident_template.docx"
            output_path = Path(temp_dir) / "incident_filled.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            document = Document()
            document.add_paragraph(
                "{{site_name}}"
            )
            document.add_paragraph(
                "{{incident_type_label}} at {{location}} witnessed by {{witness_1}}"
            )
            document.add_paragraph("{{description}}")
            document.add_paragraph("{{witness_list}}")
            document.save(template_path)

            try:
                TemplateRegistry.TEMPLATE_PATHS["incident_log"] = template_path
                manager = TemplateManager(self.build_incident())
                manager.render(output_path)
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            rendered = Document(output_path)
            self.assertIn("Uplands - Bridgend", rendered.paragraphs[0].text)
            self.assertIn("Accident at Front reception", rendered.paragraphs[1].text)
            self.assertIn("R. Davies", rendered.paragraphs[3].text)
            self.assertIn("C. Thomas", rendered.paragraphs[3].text)

    def test_template_manager_rejects_template_missing_required_placeholders(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "invalid_template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            document = Document()
            document.add_paragraph("Permit {{permit_number}} only")
            document.save(template_path)

            try:
                TemplateRegistry.TEMPLATE_PATHS["ladder_permit"] = template_path
                manager = TemplateManager(self.build_permit())
                with self.assertRaises(TemplateValidationError):
                    manager.render(Path(temp_dir) / "should_not_exist.docx")
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry


class PlantRegisterAutomationTests(unittest.TestCase):
    def _build_hss_order_confirmation_pdf(
        self,
        pdf_path: Path,
        *,
        order_ref: str = "H-3YXFCBWH",
        purchase_order: str = "81888",
        description: str = "DUST EXTRACTOR M CLASS 110V",
        stock_code: str = "52538",
        start_date: str = "26/01/2026",
        end_date: str = "01/02/2026",
        quantity: int = 1,
    ) -> None:
        document = fitz.open()
        page = document.new_page()
        page.insert_text(
            (72, 72),
            "\n".join(
                [
                    "Order Confirmation",
                    "Product",
                    "1",
                    "UPLANDS RETAIL LTD",
                    "UP0037",
                    purchase_order,
                    order_ref,
                    "07970335636",
                    "ceri.edwards@uplandsretail.co.uk",
                    "Ceri Edwards",
                    "07970335636",
                    "PO8 0SJ",
                    "Waterlooville",
                    "National Grid",
                    "Sub Station",
                    description,
                    stock_code,
                    f"{start_date} {end_date}*4dw",
                    "£25.00",
                    str(quantity),
                    "delivery",
                    "0161 749 4090",
                ]
            ),
        )
        document.save(pdf_path)
        document.close()

    def _build_plant_collection_pdf(
        self,
        pdf_path: Path,
        *,
        contract_ref: str = "5538-02105",
        collection_code: str = "5538-06878",
        description: str = "FIRE POINT WIRELESS SUBSIDIARY",
        stock_code: str = "50767",
        on_hire: str = "12/01/2026",
        off_hire: str = "27/03/2026",
        collection_date: str = "27/03/2026",
        actual_quantity: int = 4,
        serials: Optional[List[str]] = None,
    ) -> None:
        serial_values = serials or ["HOWY3155", "HOWY3331", "HOW2168", "645X9643"]
        lines = [
            "Proof of collection",
            "Contract:",
            contract_ref,
            "Collection Code:",
            collection_code,
            "Customer Name:UPLANDS RETAIL LTD",
            "Site:",
            "NATIONAL GRID",
            "BROADWAY LA",
            "LOVEDEAN, WATERLOOVILLE",
            "PO8 0SJ",
            "On Hire:",
            on_hire,
            "Off Hire:",
            off_hire,
            "Collection:",
            collection_date,
            "Comm Code",
            "Description",
            "E/Code",
            "O/S Qty",
            "Advised Qty",
            "Actual Qty",
            "Damage Qty",
            "Dirty Qty",
            stock_code,
            description,
            "0",
            str(actual_quantity),
            str(actual_quantity),
            "0",
            "0",
        ]
        for serial_value in serial_values:
            lines.extend(
                [
                    stock_code,
                    description,
                    serial_value,
                    "0",
                    "0",
                    "1",
                    "0",
                    "0",
                ]
            )
        lines.extend(
            [
                "Collection Slot:",
                "08:00 to 17:30",
                "Driver Name:",
                "PAUL BRYANT",
                "Vehicle Registration:",
                "BF24YNS",
                f"Date: {collection_date}",
            ]
        )

        document = fitz.open()
        page = document.new_page(width=595, height=1200)
        page.insert_text((72, 72), "\n".join(lines), fontsize=10)
        document.save(pdf_path)
        document.close()

    def _build_plant_register_template(self, template_path: Path) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=9)
        headers = [
            "Hire Number",
            "Plant description",
            "Hire Company",
            "Telephone number",
            "On Hire Date",
            "Who Hired",
            "Serial Number",
            "Last or Next LOLER Inspection",
            "Plant Certificate in H&S File 2",
        ]
        for column_index, header in enumerate(headers):
            table.cell(0, column_index).text = header
        table.cell(1, 0).text = "{% tr for p in plant_assets %}{{p.hire_num}}"
        table.cell(1, 1).text = "{{p.description}}"
        table.cell(1, 2).text = "{{p.company}}"
        table.cell(1, 3).text = "{{p.phone}}"
        table.cell(1, 4).text = "{{p.on_hire}}"
        table.cell(1, 5).text = "{{p.hired_by}}"
        table.cell(1, 6).text = "{{p.serial}}"
        table.cell(1, 7).text = "{{p.inspection}}"
        table.cell(1, 8).text = "{{p.in_file}}{% tr endfor %}"
        document.save(template_path)

    def _build_site_induction_template(self, template_path: Path) -> None:
        document = Document()
        document.add_paragraph("{{site_name}}")
        document.add_paragraph("{{date}}")
        document.add_paragraph("{{full_name}}")
        document.add_paragraph("{{company}}")
        document.add_paragraph("{{home_address}}")
        document.add_paragraph("{{contact_number}}")
        document.add_paragraph("{{occupation}}")
        document.add_paragraph("{{emergency_contact}}")
        document.add_paragraph("{{emergency_tel}}")
        document.add_paragraph("{{medical}}")
        document.add_paragraph("{{cscs_no}}")
        document.add_paragraph("{{cscs_expiry}}")
        document.add_paragraph("{{asbestos_cert}}")
        document.add_paragraph("{{erect_scaffold}}")
        document.add_paragraph("{{cisrs_no}}")
        document.add_paragraph("{{cisrs_expiry}}")
        document.add_paragraph("{{operate_plant}}")
        document.add_paragraph("{{cpcs_no}}")
        document.add_paragraph("{{cpcs_expiry}}")
        document.add_paragraph("{{client_training_desc}}")
        document.add_paragraph("{{client_training_date}}")
        document.add_paragraph("{{client_training_expiry}}")
        document.add_paragraph("{{first_aider}}")
        document.add_paragraph("{{fire_warden}}")
        document.add_paragraph("{{supervisor}}")
        document.add_paragraph("{{smsts}}")
        document.add_paragraph("{{inductor_name_date}}")
        document.add_paragraph("{{inductor_title}}")
        document.add_paragraph("{{signature_image}}")
        document.save(template_path)

    def test_file_and_index_all_creates_pending_plant_asset_from_hss_contract(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            carrier_docs_destination = workspace_root / "FILE_1_Environment" / "Carrier_Docs"
            waste_reports_destination = workspace_root / "FILE_1_Environment" / "Waste_Reports"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            plant_hire_directory = workspace_root / "FILE_2_Registers" / "Plant_Hire_Register"
            induction_directory = workspace_root / "FILE_3_Inductions"
            inbox = workspace_root / "ingest"
            database_path = workspace_root / "documents.sqlite3"

            for directory in (
                waste_destination,
                carrier_docs_destination,
                waste_reports_destination,
                attendance_destination,
                plant_hire_directory,
                induction_directory,
                inbox,
            ):
                directory.mkdir(parents=True, exist_ok=True)

            (workspace_root / "project_setup.json").write_text(
                json.dumps(
                    {
                        "current_site_name": "NG Lovedean Substation",
                        "job_number": "JOB-4471",
                        "site_address": "Broadway Lane",
                        "client_name": "National Grid",
                    }
                ),
                encoding="utf-8",
            )

            contract_path = inbox / "Contract-H-3YXFCBWH.pdf"
            self._build_hss_order_confirmation_pdf(contract_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.INBOX,
                app_config.WASTE_DESTINATION,
                app_config.CARRIER_DOCS_DESTINATION,
                app_config.WASTE_REPORTS_DESTINATION,
                app_config.ATTENDANCE_DESTINATION,
                app_config.PLANT_HIRE_REGISTER_DIR,
                app_config.INDUCTION_DIR,
                app_config.FILE_3_REVIEW_DIR,
                app_config.DATABASE_PATH,
            )

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INBOX = inbox
                app_config.WASTE_DESTINATION = waste_destination
                app_config.CARRIER_DOCS_DESTINATION = carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = plant_hire_directory
                app_config.INDUCTION_DIR = induction_directory
                app_config.FILE_3_REVIEW_DIR = induction_directory / "Needs_Review"
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                file_and_index_all(repository)
                plant_assets = repository.list_documents(
                    document_type=PlantAssetDocument.document_type,
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.INBOX,
                    app_config.WASTE_DESTINATION,
                    app_config.CARRIER_DOCS_DESTINATION,
                    app_config.WASTE_REPORTS_DESTINATION,
                    app_config.ATTENDANCE_DESTINATION,
                    app_config.PLANT_HIRE_REGISTER_DIR,
                    app_config.INDUCTION_DIR,
                    app_config.FILE_3_REVIEW_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            self.assertEqual(len(plant_assets), 1)
            plant_asset = plant_assets[0]
            self.assertEqual(plant_asset.hire_num, "JOB-4471-01")
            self.assertEqual(plant_asset.description, "DUST EXTRACTOR M CLASS 110V")
            self.assertEqual(plant_asset.company, "HSS")
            self.assertEqual(plant_asset.phone, "0161 749 4090")
            self.assertEqual(plant_asset.on_hire, date(2026, 1, 26))
            self.assertEqual(plant_asset.source_reference, "H-3YXFCBWH")
            self.assertEqual(plant_asset.purchase_order, "81888")
            self.assertEqual(plant_asset.hired_by, "URL (Uplands)")
            self.assertEqual(plant_asset.status, DocumentStatus.ACTIVE)
            self.assertEqual(plant_asset.stock_code, "52538")
            self.assertEqual(plant_asset.serial, "52538")
            self.assertEqual(plant_asset.inspection_type, PlantInspectionType.SERVICE)
            self.assertEqual(plant_asset.inspection, "Inspection / cert ref not logged")

    def test_generate_plant_register_document_renders_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_directory = Path(temp_dir) / "Plant_Hire_Register"
            database_path = Path(temp_dir) / "documents.sqlite3"
            template_path = Path(temp_dir) / "plant-register-template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_plant_hire_directory = app_config.PLANT_HIRE_REGISTER_DIR
            repository = DocumentRepository(database_path)
            repository.create_schema()
            self._build_plant_register_template(template_path)

            repository.save(
                PlantAssetDocument(
                    doc_id="PLANT-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 12, 8, 0),
                    status=DocumentStatus.ACTIVE,
                    hire_num="JOB-4471-01",
                    description="DUST EXTRACTOR M CLASS 110V",
                    company="HSS",
                    phone="0161 749 4090",
                    on_hire=date(2026, 1, 26),
                    hired_by="TDE",
                    serial="SER-001",
                    inspection=(
                        "Exam 01/03/2026 | Next due 14/03/2026 | "
                        "Asset ID 001 | Report 123"
                    ),
                    source_reference="H-3YXFCBWH",
                    purchase_order="81888",
                )
            )

            try:
                TemplateRegistry.TEMPLATE_PATHS["plant_register"] = template_path
                app_config.PLANT_HIRE_REGISTER_DIR = output_directory
                generated = generate_plant_register_document(
                    repository,
                    site_name="NG Lovedean Substation",
                )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                app_config.PLANT_HIRE_REGISTER_DIR = original_plant_hire_directory

            self.assertTrue(generated.output_path.exists())
            self.assertEqual(generated.asset_count, 1)

            rendered = Document(generated.output_path)
            self.assertEqual(rendered.tables[0].cell(1, 0).text, "JOB-4471-01")
            self.assertEqual(
                rendered.tables[0].cell(1, 1).text,
                "DUST EXTRACTOR M CLASS 110V",
            )
            self.assertEqual(rendered.tables[0].cell(1, 8).text, "Yes")
            data_run = rendered.tables[0].cell(1, 0).paragraphs[0].runs[0]
            self.assertEqual(data_run.font.name, "Arial")
            self.assertEqual(data_run.font.size.pt, 10.0)

    def test_file_and_index_all_archives_matching_plant_asset_from_collection_note(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            carrier_docs_destination = workspace_root / "FILE_1_Environment" / "Carrier_Docs"
            waste_reports_destination = workspace_root / "FILE_1_Environment" / "Waste_Reports"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            plant_hire_directory = workspace_root / "FILE_2_Registers" / "Plant_Hire_Register"
            induction_directory = workspace_root / "FILE_3_Inductions"
            inbox = workspace_root / "ingest"
            database_path = workspace_root / "documents.sqlite3"

            for directory in (
                waste_destination,
                carrier_docs_destination,
                waste_reports_destination,
                attendance_destination,
                plant_hire_directory,
                induction_directory,
                inbox,
            ):
                directory.mkdir(parents=True, exist_ok=True)

            (workspace_root / "project_setup.json").write_text(
                json.dumps(
                    {
                        "current_site_name": "NG Lovedean Substation",
                        "job_number": "81888",
                        "site_address": "Broadway Lane",
                        "client_name": "National Grid",
                    }
                ),
                encoding="utf-8",
            )

            contract_path = inbox / "Contract-H-ZSR77G3B.pdf"
            self._build_hss_order_confirmation_pdf(
                contract_path,
                order_ref="H-ZSR77G3B",
                description="FIRE POINT WIRELESS SUBSIDIARY",
                stock_code="50767",
                start_date="12/01/2026",
                end_date="19/01/2026",
                quantity=4,
            )
            collection_path = inbox / "5538-02105.pdf"
            self._build_plant_collection_pdf(collection_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.INBOX,
                app_config.WASTE_DESTINATION,
                app_config.CARRIER_DOCS_DESTINATION,
                app_config.WASTE_REPORTS_DESTINATION,
                app_config.ATTENDANCE_DESTINATION,
                app_config.PLANT_HIRE_REGISTER_DIR,
                app_config.INDUCTION_DIR,
                app_config.FILE_3_REVIEW_DIR,
                app_config.DATABASE_PATH,
            )

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INBOX = inbox
                app_config.WASTE_DESTINATION = waste_destination
                app_config.CARRIER_DOCS_DESTINATION = carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = plant_hire_directory
                app_config.INDUCTION_DIR = induction_directory
                app_config.FILE_3_REVIEW_DIR = induction_directory / "Needs_Review"
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                file_and_index_all(repository)
                plant_assets = repository.list_documents(
                    document_type=PlantAssetDocument.document_type,
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.INBOX,
                    app_config.WASTE_DESTINATION,
                    app_config.CARRIER_DOCS_DESTINATION,
                    app_config.WASTE_REPORTS_DESTINATION,
                    app_config.ATTENDANCE_DESTINATION,
                    app_config.PLANT_HIRE_REGISTER_DIR,
                    app_config.INDUCTION_DIR,
                    app_config.FILE_3_REVIEW_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            self.assertEqual(len(plant_assets), 1)
            plant_asset = plant_assets[0]
            self.assertEqual(plant_asset.description, "FIRE POINT WIRELESS SUBSIDIARY (x4)")
            self.assertEqual(plant_asset.stock_code, "50767")
            self.assertEqual(plant_asset.status, DocumentStatus.ARCHIVED)
            self.assertEqual(plant_asset.off_hire, date(2026, 3, 27))
            self.assertEqual(
                plant_asset.serial,
                "HOWY3155, HOWY3331, HOW2168, 645X9643",
            )

    def test_create_site_induction_document_renders_signature_and_logs_record(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            signatures_directory = workspace_root / "FILE_3_Safety" / "Signatures"
            completed_directory = workspace_root / "FILE_3_Safety" / "Completed_Inductions"
            competency_cards_directory = (
                workspace_root / "FILE_3_Safety" / "Competency_Cards"
            )
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "site_induction_template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.FILE_3_COMPETENCY_CARDS_DIR,
                app_config.FILE_3_SIGNATURES_DIR,
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                app_config.DATABASE_PATH,
            )

            self._build_site_induction_template(template_path)

            try:
                TemplateRegistry.TEMPLATE_PATHS["site_induction"] = template_path
                app_config.BASE_DATA_DIR = workspace_root
                app_config.FILE_3_COMPETENCY_CARDS_DIR = competency_cards_directory
                app_config.FILE_3_SIGNATURES_DIR = signatures_directory
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR = completed_directory
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                signature_image_data = np.full((200, 420, 4), 255, dtype=np.uint8)
                signature_image_data[92:98, 40:220, :3] = 0

                generated_document = create_site_induction_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    full_name="Sean Carter",
                    home_address="1 Test Street",
                    contact_number="07123 456789",
                    company="A. Archer Electrical",
                    occupation="Electrician",
                    emergency_contact="Jane Carter",
                    emergency_tel="07999 888777",
                    medical="None declared",
                    cscs_number="CSCS-1234",
                    cscs_expiry=date(2027, 3, 1),
                    asbestos_cert=True,
                    erect_scaffold=True,
                    cisrs_no="CISRS-9981",
                    cisrs_expiry=date(2026, 12, 31),
                    operate_plant=True,
                    cpcs_no="CPCS-4455",
                    cpcs_expiry=date(2027, 1, 15),
                    client_training_desc="National Grid substation access briefing",
                    client_training_date=date(2026, 3, 13),
                    client_training_expiry=date(2027, 3, 13),
                    first_aider=True,
                    fire_warden=False,
                    supervisor=True,
                    smsts=False,
                    competency_expiry_date=date(2027, 4, 1),
                    competency_files=[
                        {
                            "label": "CSCS Card",
                            "name": "cscs-front.jpg",
                            "bytes": b"front-card-binary",
                        },
                        {
                            "label": "Manual Handling Certificate",
                            "name": "manual-handling.pdf",
                            "bytes": b"manual-handling-binary",
                        }
                    ],
                    signature_image_data=signature_image_data,
                )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                (
                    app_config.BASE_DATA_DIR,
                    app_config.FILE_3_COMPETENCY_CARDS_DIR,
                    app_config.FILE_3_SIGNATURES_DIR,
                    app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            self.assertTrue(generated_document.signature_path.exists())
            self.assertTrue(generated_document.output_path.exists())
            self.assertEqual(generated_document.signature_path.parent, signatures_directory)
            self.assertEqual(generated_document.output_path.parent, completed_directory)

            persisted_documents = repository.list_documents(
                document_type=InductionDocument.document_type,
                site_name="NG Lovedean Substation",
            )
            self.assertEqual(len(persisted_documents), 1)
            induction_document = persisted_documents[0]
            self.assertIsInstance(induction_document, InductionDocument)
            self.assertEqual(induction_document.individual_name, "Sean Carter")
            self.assertEqual(induction_document.contractor_name, "A. Archer Electrical")
            self.assertTrue(induction_document.first_aider)
            self.assertTrue(induction_document.supervisor)
            self.assertEqual(induction_document.competency_expiry_date, date(2027, 4, 1))
            saved_competency_paths = [
                Path(path_text)
                for path_text in induction_document.competency_card_paths.split(",")
                if path_text
            ]
            self.assertEqual(len(saved_competency_paths), 2)
            for saved_path in saved_competency_paths:
                self.assertTrue(saved_path.exists())
                self.assertEqual(saved_path.parent, competency_cards_directory)
            self.assertTrue(
                any("cscs card" in saved_path.name.casefold() for saved_path in saved_competency_paths)
            )
            self.assertTrue(
                any(
                    "manual handling certificate" in saved_path.name.casefold()
                    for saved_path in saved_competency_paths
                )
            )
            self.assertEqual(
                Path(induction_document.completed_document_path),
                generated_document.output_path,
            )

            rendered_document = Document(generated_document.output_path)
            rendered_text = "\n".join(
                paragraph.text for paragraph in rendered_document.paragraphs
            )
            self.assertIn("Sean Carter", rendered_text)
            self.assertIn("A. Archer Electrical", rendered_text)
            self.assertIn("1 Test Street", rendered_text)
            self.assertIn("01/03/2027", rendered_text)
            self.assertIn("CISRS-9981", rendered_text)
            self.assertIn("CPCS-4455", rendered_text)
            self.assertIn("National Grid substation access briefing", rendered_text)
            self.assertIn("Ceri Edwards", rendered_text)
            self.assertIn("Site Manager", rendered_text)
            self.assertIn(generated_document.induction_document.created_at.strftime("%d/%m/%Y"), rendered_text)

    def test_update_site_induction_document_regenerates_saved_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            signatures_directory = workspace_root / "FILE_3_Safety" / "Signatures"
            completed_directory = workspace_root / "FILE_3_Safety" / "Completed_Inductions"
            competency_cards_directory = (
                workspace_root / "FILE_3_Safety" / "Competency_Cards"
            )
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "site_induction_template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.FILE_3_COMPETENCY_CARDS_DIR,
                app_config.FILE_3_SIGNATURES_DIR,
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                app_config.DATABASE_PATH,
            )

            self._build_site_induction_template(template_path)

            try:
                TemplateRegistry.TEMPLATE_PATHS["site_induction"] = template_path
                app_config.BASE_DATA_DIR = workspace_root
                app_config.FILE_3_COMPETENCY_CARDS_DIR = competency_cards_directory
                app_config.FILE_3_SIGNATURES_DIR = signatures_directory
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR = completed_directory
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                signature_image_data = np.full((200, 420, 4), 255, dtype=np.uint8)
                signature_image_data[92:98, 40:220, :3] = 0

                generated_document = create_site_induction_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    full_name="Sean Carter",
                    home_address="1 Test Street",
                    contact_number="07123 456789",
                    company="A. Archer Electrical",
                    occupation="Electrician",
                    emergency_contact="Jane Carter",
                    emergency_tel="07999 888777",
                    medical="None declared",
                    cscs_number="CSCS-1234",
                    cscs_expiry=date(2027, 3, 1),
                    asbestos_cert=True,
                    erect_scaffold=False,
                    operate_plant=False,
                    client_training_desc="Original training note",
                    first_aider=False,
                    fire_warden=False,
                    supervisor=False,
                    smsts=False,
                    competency_files=[
                        {
                            "label": "Manual Handling Certificate",
                            "name": "manual-handling.pdf",
                            "bytes": b"manual-handling-binary",
                        }
                    ],
                    signature_image_data=signature_image_data,
                )

                updated_document = update_site_induction_document(
                    repository,
                    induction_doc_id=generated_document.induction_document.doc_id,
                    full_name="Sean Carter",
                    home_address="22 Updated Street",
                    contact_number="07000 111222",
                    company="A. Archer Electrical",
                    occupation="Chargehand",
                    emergency_contact="Janet Carter",
                    emergency_tel="07000 333444",
                    medical="No changes",
                    cscs_number="CSCS-9876",
                    cscs_expiry=date(2028, 1, 15),
                    asbestos_cert=True,
                    erect_scaffold=False,
                    operate_plant=False,
                    client_training_desc="Updated training note",
                    client_training_date=date(2026, 3, 14),
                    client_training_expiry=date(2027, 3, 14),
                    first_aider=True,
                    fire_warden=True,
                    supervisor=False,
                    smsts=False,
                    competency_expiry_date=date(2028, 2, 1),
                )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                (
                    app_config.BASE_DATA_DIR,
                    app_config.FILE_3_COMPETENCY_CARDS_DIR,
                    app_config.FILE_3_SIGNATURES_DIR,
                    app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            self.assertEqual(
                updated_document.output_path,
                generated_document.output_path,
            )
            self.assertTrue(updated_document.output_path.exists())
            persisted_document = repository.get(generated_document.induction_document.doc_id)
            self.assertIsInstance(persisted_document, InductionDocument)
            self.assertEqual(persisted_document.contact_number, "07000 111222")
            self.assertEqual(persisted_document.home_address, "22 Updated Street")
            self.assertTrue(persisted_document.first_aider)
            self.assertTrue(persisted_document.fire_warden)

            rendered_document = Document(updated_document.output_path)
            rendered_text = "\n".join(
                paragraph.text for paragraph in rendered_document.paragraphs
            )
            self.assertIn("22 Updated Street", rendered_text)
            self.assertIn("07000 111222", rendered_text)
            self.assertIn("Updated training note", rendered_text)

    def test_add_site_induction_evidence_files_appends_extra_saved_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            signatures_directory = workspace_root / "FILE_3_Safety" / "Signatures"
            completed_directory = workspace_root / "FILE_3_Safety" / "Completed_Inductions"
            competency_cards_directory = (
                workspace_root / "FILE_3_Safety" / "Competency_Cards"
            )
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "site_induction_template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.FILE_3_COMPETENCY_CARDS_DIR,
                app_config.FILE_3_SIGNATURES_DIR,
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                app_config.DATABASE_PATH,
            )

            self._build_site_induction_template(template_path)

            try:
                TemplateRegistry.TEMPLATE_PATHS["site_induction"] = template_path
                app_config.BASE_DATA_DIR = workspace_root
                app_config.FILE_3_COMPETENCY_CARDS_DIR = competency_cards_directory
                app_config.FILE_3_SIGNATURES_DIR = signatures_directory
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR = completed_directory
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                signature_image_data = np.full((200, 420, 4), 255, dtype=np.uint8)
                signature_image_data[92:98, 40:220, :3] = 0

                generated_document = create_site_induction_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    full_name="Sean Carter",
                    home_address="1 Test Street",
                    contact_number="07123 456789",
                    company="A. Archer Electrical",
                    occupation="Electrician",
                    emergency_contact="Jane Carter",
                    emergency_tel="07999 888777",
                    medical="None declared",
                    cscs_number="CSCS-1234",
                    first_aider=False,
                    fire_warden=False,
                    supervisor=False,
                    smsts=False,
                    competency_files=[
                        {
                            "label": "Manual Handling Certificate",
                            "name": "manual-handling.pdf",
                            "bytes": b"manual-handling-binary",
                        }
                    ],
                    signature_image_data=signature_image_data,
                )

                updated_document = add_site_induction_evidence_files(
                    repository,
                    induction_doc_id=generated_document.induction_document.doc_id,
                    competency_files=[
                        {
                            "label": "First Aid Certificate",
                            "name": "first-aid-cert.pdf",
                            "bytes": b"first-aid-binary",
                        },
                        {
                            "label": "Face Fit Certificate",
                            "name": "face-fit.pdf",
                            "bytes": b"face-fit-binary",
                        },
                    ],
                )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                (
                    app_config.BASE_DATA_DIR,
                    app_config.FILE_3_COMPETENCY_CARDS_DIR,
                    app_config.FILE_3_SIGNATURES_DIR,
                    app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            saved_competency_paths = [
                Path(path_text)
                for path_text in updated_document.competency_card_paths.split(",")
                if path_text
            ]
            self.assertEqual(len(saved_competency_paths), 3)
            self.assertTrue(
                any("manual handling certificate" in saved_path.name.casefold() for saved_path in saved_competency_paths)
            )
            self.assertTrue(
                any("first aid certificate" in saved_path.name.casefold() for saved_path in saved_competency_paths)
            )
            self.assertTrue(
                any("face fit certificate" in saved_path.name.casefold() for saved_path in saved_competency_paths)
            )

    def test_create_site_induction_document_requires_home_address_and_contact_number(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            signatures_directory = workspace_root / "FILE_3_Safety" / "Signatures"
            completed_directory = workspace_root / "FILE_3_Safety" / "Completed_Inductions"
            competency_cards_directory = (
                workspace_root / "FILE_3_Safety" / "Competency_Cards"
            )
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "site_induction_template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.FILE_3_COMPETENCY_CARDS_DIR,
                app_config.FILE_3_SIGNATURES_DIR,
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                app_config.DATABASE_PATH,
            )

            self._build_site_induction_template(template_path)

            try:
                TemplateRegistry.TEMPLATE_PATHS["site_induction"] = template_path
                app_config.BASE_DATA_DIR = workspace_root
                app_config.FILE_3_COMPETENCY_CARDS_DIR = competency_cards_directory
                app_config.FILE_3_SIGNATURES_DIR = signatures_directory
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR = completed_directory
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                signature_image_data = np.full((200, 420, 4), 255, dtype=np.uint8)
                signature_image_data[92:98, 40:220, :3] = 0

                with self.assertRaises(ValidationError) as error_context:
                    create_site_induction_document(
                        repository,
                        site_name="NG Lovedean Substation",
                        full_name="Sean Carter",
                        home_address="",
                        contact_number="",
                        company="A. Archer Electrical",
                        occupation="Electrician",
                        emergency_contact="Jane Carter",
                        emergency_tel="07999 888777",
                        medical="None declared",
                        cscs_number="CSCS-1234",
                        first_aider=False,
                        fire_warden=False,
                        supervisor=False,
                        smsts=False,
                        signature_image_data=signature_image_data,
                    )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                (
                    app_config.BASE_DATA_DIR,
                    app_config.FILE_3_COMPETENCY_CARDS_DIR,
                    app_config.FILE_3_SIGNATURES_DIR,
                    app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            self.assertIn("Home Address is required", str(error_context.exception))

    def test_create_site_induction_document_requires_manual_handling_certificate(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            signatures_directory = workspace_root / "FILE_3_Safety" / "Signatures"
            completed_directory = workspace_root / "FILE_3_Safety" / "Completed_Inductions"
            competency_cards_directory = (
                workspace_root / "FILE_3_Safety" / "Competency_Cards"
            )
            database_path = workspace_root / "documents.sqlite3"
            template_path = Path(temp_dir) / "site_induction_template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.FILE_3_COMPETENCY_CARDS_DIR,
                app_config.FILE_3_SIGNATURES_DIR,
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                app_config.DATABASE_PATH,
            )

            self._build_site_induction_template(template_path)

            try:
                TemplateRegistry.TEMPLATE_PATHS["site_induction"] = template_path
                app_config.BASE_DATA_DIR = workspace_root
                app_config.FILE_3_COMPETENCY_CARDS_DIR = competency_cards_directory
                app_config.FILE_3_SIGNATURES_DIR = signatures_directory
                app_config.FILE_3_COMPLETED_INDUCTIONS_DIR = completed_directory
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                repository.create_schema()
                signature_image_data = np.full((200, 420, 4), 255, dtype=np.uint8)
                signature_image_data[92:98, 40:220, :3] = 0

                with self.assertRaises(ValidationError) as error_context:
                    create_site_induction_document(
                        repository,
                        site_name="NG Lovedean Substation",
                        full_name="Sean Carter",
                        home_address="1 Test Street",
                        contact_number="07123 456789",
                        company="A. Archer Electrical",
                        occupation="Electrician",
                        emergency_contact="Jane Carter",
                        emergency_tel="07999 888777",
                        medical="None declared",
                        cscs_number="CSCS-1234",
                        first_aider=False,
                        fire_warden=False,
                        supervisor=False,
                        smsts=False,
                        signature_image_data=signature_image_data,
                    )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                (
                    app_config.BASE_DATA_DIR,
                    app_config.FILE_3_COMPETENCY_CARDS_DIR,
                    app_config.FILE_3_SIGNATURES_DIR,
                    app_config.FILE_3_COMPLETED_INDUCTIONS_DIR,
                    app_config.DATABASE_PATH,
                ) = original_config

            self.assertIn(
                "Manual Handling Certificate upload is required",
                str(error_context.exception),
            )

    def test_daily_attendance_entry_document_flags_live_status_and_uplands_company(
        self,
    ) -> None:
        attendance_entry = DailyAttendanceEntryDocument(
            doc_id="ATT-001",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 13, 7, 30),
            status=DocumentStatus.ACTIVE,
            linked_induction_doc_id="IND-001",
            individual_name="Sean Carter",
            contractor_name="Uplands Construction Group",
            vehicle_registration="AB12 CDE",
            distance_travelled="14 miles",
            time_in=datetime(2026, 3, 13, 7, 30),
            sign_in_signature_path="/tmp/sign-in.png",
        )

        self.assertTrue(attendance_entry.is_on_site)
        self.assertTrue(attendance_entry.is_uplands_employee)
        self.assertEqual(
            attendance_entry.document_name,
            "Site Attendance Register (UHSF16.09)",
        )

    def test_daily_attendance_entry_recognises_uplands_aliases(self) -> None:
        for alias_company_name in ("URL", "Uplands Retail", "Uplands Retail Limited"):
            attendance_entry = DailyAttendanceEntryDocument(
                doc_id=f"ATT-{alias_company_name}",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 13, 7, 30),
                status=DocumentStatus.ACTIVE,
                linked_induction_doc_id="IND-001",
                individual_name="Sean Carter",
                contractor_name=alias_company_name,
                time_in=datetime(2026, 3, 13, 7, 30),
            )
            self.assertTrue(attendance_entry.is_uplands_employee)

    def test_file_4_company_options_include_roster_and_trusted_company_values(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            repository.save(
                InductionDocument(
                    doc_id="IND-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 13, 7, 0),
                    status=DocumentStatus.ACTIVE,
                    contractor_name="A. Archer Electrical",
                    individual_name="A Hustler",
                )
            )
            repository.save(
                SiteAttendanceRegister(
                    doc_id="SAR-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 13, 17, 0),
                    status=DocumentStatus.ACTIVE,
                    attendance_records=[
                        SiteAttendanceRecord(
                            date=date(2026, 3, 13),
                            company="150",
                            workerName="A Hustler",
                            timeIn=time(7, 30),
                            timeOut=time(16, 0),
                            totalHours=8.5,
                        )
                    ],
                )
            )

            with patch.object(
                app_module,
                "build_site_worker_roster",
                return_value=[
                    SiteWorker(
                        company="150",
                        worker_name="A Hustler",
                        last_on_site_date=date(2026, 3, 13),
                    )
                ],
            ):
                company_options = app_module._build_file_4_company_options(
                    repository,
                    site_name="NG Lovedean Substation",
                    worker_name="A Hustler",
                    default_company="150",
                )

            self.assertIn("150", company_options)
            self.assertIn("A. Archer Electrical", company_options)
            self.assertEqual(company_options[-1], app_module.FILE_4_NEW_COMPANY_OPTION)

    def test_build_file_4_manual_worker_context_returns_synthetic_worker_and_record(
        self,
    ) -> None:
        worker, attendance_record = app_module._build_file_4_manual_worker_context(
            worker_name="Jamie Stone",
            company_name="One Fifty Enterprises",
        )

        self.assertEqual(worker.worker_name, "Jamie Stone")
        self.assertEqual(worker.company, "One Fifty Enterprises")
        self.assertEqual(attendance_record.workerName, "Jamie Stone")
        self.assertEqual(attendance_record.company, "One Fifty Enterprises")
        self.assertEqual(attendance_record.totalHours, 0.0)

    def test_build_site_induction_readiness_snapshot_tracks_required_fields(self) -> None:
        readiness_snapshot = app_module._build_site_induction_readiness_snapshot(
            full_name="Jamie Stone",
            company="One Fifty Enterprises",
            home_address="1 Test Street",
            contact_number="07700111222",
            occupation="Electrician",
            emergency_contact="Jane Stone",
            emergency_tel="07700999888",
            cscs_number="CSCS-123",
            manual_handling_uploaded=True,
            signature_ready=True,
            asbestos_cert=False,
            asbestos_evidence_uploaded=False,
            erect_scaffold=False,
            cisrs_evidence_uploaded=False,
            first_aider=False,
            first_aider_evidence_uploaded=False,
            fire_warden=False,
            fire_warden_evidence_uploaded=False,
            supervisor=False,
            supervisor_evidence_uploaded=False,
            smsts=False,
            smsts_evidence_uploaded=False,
            operate_plant=False,
            cpcs_evidence_uploaded=False,
        )

        self.assertTrue(readiness_snapshot["ready_to_submit"])
        self.assertEqual(readiness_snapshot["required_complete_count"], 6)
        self.assertEqual(readiness_snapshot["missing_required_labels"], [])

    def test_build_site_induction_readiness_snapshot_flags_selected_role_evidence(
        self,
    ) -> None:
        readiness_snapshot = app_module._build_site_induction_readiness_snapshot(
            full_name="Jamie Stone",
            company="One Fifty Enterprises",
            home_address="1 Test Street",
            contact_number="07700111222",
            occupation="",
            emergency_contact="",
            emergency_tel="",
            cscs_number="",
            manual_handling_uploaded=True,
            signature_ready=False,
            asbestos_cert=True,
            asbestos_evidence_uploaded=False,
            erect_scaffold=True,
            cisrs_evidence_uploaded=False,
            first_aider=True,
            first_aider_evidence_uploaded=False,
            fire_warden=False,
            fire_warden_evidence_uploaded=False,
            supervisor=False,
            supervisor_evidence_uploaded=False,
            smsts=False,
            smsts_evidence_uploaded=False,
            operate_plant=True,
            cpcs_evidence_uploaded=False,
        )

        self.assertFalse(readiness_snapshot["ready_to_submit"])
        self.assertIn("Operative Signature", readiness_snapshot["missing_required_labels"])
        self.assertIn("Asbestos Certificate", readiness_snapshot["missing_recommended_labels"])
        self.assertIn("CISRS Card", readiness_snapshot["missing_recommended_labels"])
        self.assertIn("First Aid Certificate", readiness_snapshot["missing_recommended_labels"])
        self.assertIn("CPCS Card", readiness_snapshot["missing_recommended_labels"])

    def test_file_4_worker_options_include_permit_history_entries(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            repository.save(
                LadderPermit(
                    doc_id="LP-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 13, 7, 30),
                    status=DocumentStatus.DRAFT,
                    permit_number="UHSF21.09-001",
                    project_name="NG Lovedean Substation",
                    project_number="81888",
                    location_of_work="Switch room",
                    description_of_work="Fire alarm sensor replacement",
                    valid_from_date=date(2026, 3, 13),
                    valid_from_time=time(7, 30),
                    valid_to_date=date(2026, 3, 13),
                    valid_to_time=time(15, 30),
                    safer_alternative_eliminated=True,
                    task_specific_rams_prepared_and_approved=True,
                    personnel_briefed_and_understand_task=True,
                    competent_supervisor_appointed=True,
                    competent_supervisor_name="Ceri Edwards",
                    operatives_suitably_trained=True,
                    ladder_length_suitable=True,
                    conforms_to_bs_class_a=True,
                    three_points_of_contact_maintained=True,
                    harness_worn_and_secured_above_head_height=False,
                    ladder_stabilisation_method=LadderStabilisationMethod.FOOTED,
                    equipment_inspected_for_defects=True,
                    ladder_stabilisation_confirmed=True,
                    worker_name="Jamie Stone",
                    worker_company="One Fifty Enterprises",
                    briefing_name="Ceri Edwards",
                    manager_name="Ceri Edwards",
                    manager_position="Project Manager",
                    issued_date=date(2026, 3, 13),
                )
            )

            with patch.object(app_module, "build_site_worker_roster", return_value=[]):
                worker_options = app_module._build_file_4_worker_options(
                    repository,
                    site_name="NG Lovedean Substation",
                )

            self.assertIn("Jamie Stone (One Fifty Enterprises)", worker_options)

    def test_daily_attendance_sign_in_and_sign_out_round_trip(self) -> None:
        class SignInDateTime(datetime):
            @classmethod
            def now(cls, tz=None):  # type: ignore[override]
                return cls(2026, 3, 13, 7, 30, tzinfo=tz)

        class SignOutDateTime(datetime):
            @classmethod
            def now(cls, tz=None):  # type: ignore[override]
                return cls(2026, 3, 13, 15, 0, tzinfo=tz)

        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            attendance_signatures_directory = (
                workspace_root / "FILE_2_Registers" / "Attendance" / "Signatures"
            )
            database_path = workspace_root / "documents.sqlite3"
            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.DATABASE_PATH,
                app_config.ATTENDANCE_DESTINATION,
                app_config.FILE_2_ATTENDANCE_SIGNATURES_DIR,
            )

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.DATABASE_PATH = database_path
                app_config.ATTENDANCE_DESTINATION = (
                    workspace_root / "FILE_2_Registers" / "Attendance"
                )
                app_config.FILE_2_ATTENDANCE_SIGNATURES_DIR = (
                    attendance_signatures_directory
                )

                repository = DocumentRepository(database_path)
                repository.create_schema()
                induction_document = InductionDocument(
                    doc_id="IND-202603130700-sean-carter",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 13, 7, 0),
                    status=DocumentStatus.ACTIVE,
                    contractor_name="A. Archer Electrical",
                    individual_name="Sean Carter",
                )
                repository.save(induction_document)

                signature_image_data = np.full((200, 420, 4), 255, dtype=np.uint8)
                signature_image_data[92:98, 40:220, :3] = 0

                with patch.object(workspace_module, "datetime", SignInDateTime):
                    sign_in_result = create_daily_attendance_sign_in(
                        repository,
                        site_name="NG Lovedean Substation",
                        induction_document=induction_document,
                        vehicle_registration="ab12 cde",
                        distance_travelled="14 miles",
                        signature_image_data=signature_image_data,
                    )

                self.assertTrue(sign_in_result.signature_path.exists())
                self.assertEqual(
                    sign_in_result.attendance_entry.vehicle_registration,
                    "AB12 CDE",
                )
                active_entries = list_daily_attendance_entries(
                    repository,
                    site_name="NG Lovedean Substation",
                    on_date=date(2026, 3, 13),
                    active_only=True,
                )
                self.assertEqual(len(active_entries), 1)
                self.assertEqual(active_entries[0].individual_name, "Sean Carter")

                with patch.object(workspace_module, "datetime", SignOutDateTime):
                    sign_out_result = complete_daily_attendance_sign_out(
                        repository,
                        attendance_doc_id=sign_in_result.attendance_entry.doc_id,
                        signature_image_data=signature_image_data,
                    )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.DATABASE_PATH,
                    app_config.ATTENDANCE_DESTINATION,
                    app_config.FILE_2_ATTENDANCE_SIGNATURES_DIR,
                ) = original_config

            self.assertTrue(sign_out_result.signature_path.exists())
            self.assertEqual(
                sign_out_result.attendance_entry.status,
                DocumentStatus.ARCHIVED,
            )
            self.assertEqual(sign_out_result.attendance_entry.hours_worked, 7.5)
            self.assertIsNotNone(sign_out_result.attendance_entry.time_out)
            self.assertFalse(
                list_daily_attendance_entries(
                    repository,
                    site_name="NG Lovedean Substation",
                    on_date=date(2026, 3, 13),
                    active_only=True,
                )
            )

    def test_update_daily_attendance_entry_reopens_signed_out_record(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            attendance_signatures_directory = (
                workspace_root / "FILE_2_Registers" / "Attendance" / "Signatures"
            )
            attendance_signatures_directory.mkdir(parents=True, exist_ok=True)
            database_path = workspace_root / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            sign_out_signature_path = attendance_signatures_directory / "sign-out.png"
            sign_out_signature_path.write_bytes(b"fake-signout")

            attendance_entry = DailyAttendanceEntryDocument(
                doc_id="ATT-20260313-1",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 13, 7, 30),
                status=DocumentStatus.ARCHIVED,
                linked_induction_doc_id="IND-001",
                individual_name="Sean Carter",
                contractor_name="A. Archer Electrical",
                vehicle_registration="AB12 CDE",
                distance_travelled="14 miles",
                gate_verification_method="gps",
                time_in=datetime(2026, 3, 13, 7, 30),
                time_out=datetime(2026, 3, 13, 15, 0),
                hours_worked=7.5,
                sign_out_signature_path=str(sign_out_signature_path),
            )
            repository.save(attendance_entry)
            repository.index_file(
                file_name=sign_out_signature_path.name,
                file_path=sign_out_signature_path,
                file_category="attendance_sign_out_signature_png",
                file_group=FileGroup.FILE_2,
                site_name=attendance_entry.site_name,
                related_doc_id=attendance_entry.doc_id,
            )

            updated_entry = update_daily_attendance_entry(
                repository,
                attendance_doc_id=attendance_entry.doc_id,
                contractor_name="One Fifty Enterprises",
                vehicle_registration="xy34 zzz",
                distance_travelled="22 miles",
                gate_verification_method="manager_correction",
                gate_verification_note="Reopened after accidental sign-out",
                time_in=datetime(2026, 3, 13, 7, 25),
                time_out=None,
            )

            self.assertEqual(updated_entry.status, DocumentStatus.ACTIVE)
            self.assertIsNone(updated_entry.time_out)
            self.assertIsNone(updated_entry.hours_worked)
            self.assertEqual(updated_entry.sign_out_signature_path, "")
            self.assertEqual(updated_entry.contractor_name, "One Fifty Enterprises")
            self.assertEqual(updated_entry.vehicle_registration, "XY34 ZZZ")
            self.assertEqual(updated_entry.distance_travelled, "22 miles")
            self.assertFalse(sign_out_signature_path.exists())

    def test_update_daily_attendance_entry_recalculates_hours(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            attendance_entry = DailyAttendanceEntryDocument(
                doc_id="ATT-20260313-2",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 13, 7, 30),
                status=DocumentStatus.ACTIVE,
                linked_induction_doc_id="IND-001",
                individual_name="Sean Carter",
                contractor_name="A. Archer Electrical",
                vehicle_registration="AB12 CDE",
                distance_travelled="14 miles",
                time_in=datetime(2026, 3, 13, 7, 30),
            )
            repository.save(attendance_entry)

            updated_entry = update_daily_attendance_entry(
                repository,
                attendance_doc_id=attendance_entry.doc_id,
                contractor_name="150",
                vehicle_registration="ab12 cde",
                distance_travelled="18 miles",
                gate_verification_method="gate_code",
                gate_verification_note="Manager gate recovery",
                time_in=datetime(2026, 3, 13, 7, 15),
                time_out=datetime(2026, 3, 13, 16, 0),
            )

            self.assertEqual(updated_entry.status, DocumentStatus.ARCHIVED)
            self.assertEqual(updated_entry.contractor_name, "150")
            self.assertEqual(updated_entry.vehicle_registration, "AB12 CDE")
            self.assertEqual(updated_entry.distance_travelled, "18 miles")
            self.assertEqual(updated_entry.gate_verification_method, "gate_code")
            self.assertAlmostEqual(updated_entry.hours_worked, 8.75)

    def test_attendance_helper_returns_latest_vehicle_registration_for_induction(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            induction_document = InductionDocument(
                doc_id="IND-202603130700-sean-carter",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 13, 7, 0),
                status=DocumentStatus.ACTIVE,
                contractor_name="A. Archer Electrical",
                individual_name="Sean Carter",
            )
            repository.save(induction_document)
            repository.save(
                DailyAttendanceEntryDocument(
                    doc_id="ATT-20260313-1",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 13, 7, 30),
                    status=DocumentStatus.ARCHIVED,
                    linked_induction_doc_id=induction_document.doc_id,
                    individual_name="Sean Carter",
                    contractor_name="A. Archer Electrical",
                    vehicle_registration="AB12 CDE",
                    distance_travelled="14 miles",
                    time_in=datetime(2026, 3, 13, 7, 30),
                    time_out=datetime(2026, 3, 13, 15, 30),
                    hours_worked=8.0,
                )
            )
            latest_entry = DailyAttendanceEntryDocument(
                doc_id="ATT-20260314-1",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 14, 7, 20),
                status=DocumentStatus.ACTIVE,
                linked_induction_doc_id=induction_document.doc_id,
                individual_name="Sean Carter",
                contractor_name="A. Archer Electrical",
                vehicle_registration="XY34 ZZZ",
                distance_travelled="14 miles",
                time_in=datetime(2026, 3, 14, 7, 20),
            )
            repository.save(latest_entry)

            resolved_entry = app_module._get_latest_daily_attendance_entry_for_induction(
                repository,
                induction_document,
                site_name="NG Lovedean Substation",
            )

            self.assertIsNotNone(resolved_entry)
            self.assertEqual(resolved_entry.doc_id, latest_entry.doc_id)
            self.assertEqual(resolved_entry.vehicle_registration, "XY34 ZZZ")
            self.assertEqual(resolved_entry.distance_travelled, "14 miles")

    def test_resolve_attendance_sign_in_selection_prefers_pending_doc_id(self) -> None:
        records = [
            InductionDocument(
                doc_id="IND-001",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 13, 7, 0),
                status=DocumentStatus.ACTIVE,
                contractor_name="A. Archer Electrical",
                individual_name="Sean Carter",
            ),
            InductionDocument(
                doc_id="IND-002",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 13, 7, 5),
                status=DocumentStatus.ACTIVE,
                contractor_name="Uplands Construction Group",
                individual_name="Luke Green",
            ),
        ]

        resolved_doc_id = app_module._resolve_attendance_sign_in_selection(
            filtered_records=records,
            current_doc_id="",
            pending_doc_id="IND-002",
        )

        self.assertEqual(resolved_doc_id, "IND-002")

    def test_resolve_attendance_sign_in_selection_auto_selects_single_match(self) -> None:
        records = [
            InductionDocument(
                doc_id="IND-001",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 13, 7, 0),
                status=DocumentStatus.ACTIVE,
                contractor_name="A. Archer Electrical",
                individual_name="Sean Carter",
            )
        ]

        resolved_doc_id = app_module._resolve_attendance_sign_in_selection(
            filtered_records=records,
            current_doc_id="",
            pending_doc_id="",
        )

        self.assertEqual(resolved_doc_id, "IND-001")

    def test_resolve_attendance_sign_out_selection_preserves_current_match(self) -> None:
        entries = [
            DailyAttendanceEntryDocument(
                doc_id="ATT-001",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 7, 20),
                status=DocumentStatus.ACTIVE,
                linked_induction_doc_id="IND-001",
                individual_name="Sean Carter",
                contractor_name="A. Archer Electrical",
                time_in=datetime(2026, 3, 15, 7, 20),
            ),
            DailyAttendanceEntryDocument(
                doc_id="ATT-002",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 7, 25),
                status=DocumentStatus.ACTIVE,
                linked_induction_doc_id="IND-002",
                individual_name="Luke Green",
                contractor_name="Uplands Construction Group",
                time_in=datetime(2026, 3, 15, 7, 25),
            ),
        ]

        resolved_doc_id = app_module._resolve_attendance_sign_out_selection(
            filtered_entries=entries,
            current_doc_id="ATT-002",
        )

        self.assertEqual(resolved_doc_id, "ATT-002")

    def test_resolve_attendance_sign_out_selection_auto_selects_single_match(self) -> None:
        entries = [
            DailyAttendanceEntryDocument(
                doc_id="ATT-001",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 7, 20),
                status=DocumentStatus.ACTIVE,
                linked_induction_doc_id="IND-001",
                individual_name="Sean Carter",
                contractor_name="A. Archer Electrical",
                time_in=datetime(2026, 3, 15, 7, 20),
            )
        ]

        resolved_doc_id = app_module._resolve_attendance_sign_out_selection(
            filtered_entries=entries,
            current_doc_id="",
        )

        self.assertEqual(resolved_doc_id, "ATT-001")

    def test_generate_attendance_register_document_writes_today_sheet_to_file_2(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            attendance_signatures_directory = attendance_destination / "Signatures"
            attendance_output_directory = attendance_destination / "Registers"
            database_path = workspace_root / "documents.sqlite3"

            for directory in (
                attendance_destination,
                attendance_signatures_directory,
                attendance_output_directory,
            ):
                directory.mkdir(parents=True, exist_ok=True)

            sign_in_signature_path = attendance_signatures_directory / "sign-in.png"
            sign_out_signature_path = attendance_signatures_directory / "sign-out.png"
            fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 120, 40), 0).save(
                str(sign_in_signature_path)
            )
            fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 120, 40), 0).save(
                str(sign_out_signature_path)
            )

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.DATABASE_PATH,
                app_config.ATTENDANCE_DESTINATION,
                app_config.FILE_2_ATTENDANCE_SIGNATURES_DIR,
                app_config.FILE_2_ATTENDANCE_OUTPUT_DIR,
            )

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.DATABASE_PATH = database_path
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.FILE_2_ATTENDANCE_SIGNATURES_DIR = (
                    attendance_signatures_directory
                )
                app_config.FILE_2_ATTENDANCE_OUTPUT_DIR = attendance_output_directory

                repository = DocumentRepository(database_path)
                repository.create_schema()
                repository.save(
                    InductionDocument(
                        doc_id="IND-001",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 13, 7, 0),
                        status=DocumentStatus.ACTIVE,
                        contractor_name="A. Archer Electrical",
                        individual_name="Sean Carter",
                        contact_number="07700111222",
                    )
                )
                repository.save(
                    InductionDocument(
                        doc_id="IND-002",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 13, 7, 5),
                        status=DocumentStatus.ACTIVE,
                        contractor_name="Uplands Construction Group",
                        individual_name="Luke Green",
                        contact_number="07700999444",
                    )
                )
                repository.save(
                    DailyAttendanceEntryDocument(
                        doc_id="ATT-001",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 13, 7, 30),
                        status=DocumentStatus.ARCHIVED,
                        linked_induction_doc_id="IND-001",
                        individual_name="Sean Carter",
                        contractor_name="A. Archer Electrical",
                        vehicle_registration="AB12 CDE",
                        distance_travelled="14 miles",
                        time_in=datetime(2026, 3, 13, 7, 30),
                        time_out=datetime(2026, 3, 13, 15, 0),
                        hours_worked=7.5,
                        sign_in_signature_path=str(sign_in_signature_path),
                        sign_out_signature_path=str(sign_out_signature_path),
                    )
                )
                repository.save(
                    DailyAttendanceEntryDocument(
                        doc_id="ATT-002",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 13, 8, 0),
                        status=DocumentStatus.ACTIVE,
                        linked_induction_doc_id="IND-002",
                        individual_name="Luke Green",
                        contractor_name="Uplands Construction Group",
                        vehicle_registration="XY34 ZZZ",
                        distance_travelled="8 miles",
                        time_in=datetime(2026, 3, 13, 8, 0),
                        sign_in_signature_path=str(sign_in_signature_path),
                    )
                )

                generated_register = generate_attendance_register_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    on_date=date(2026, 3, 13),
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.DATABASE_PATH,
                    app_config.ATTENDANCE_DESTINATION,
                    app_config.FILE_2_ATTENDANCE_SIGNATURES_DIR,
                    app_config.FILE_2_ATTENDANCE_OUTPUT_DIR,
                ) = original_config

            self.assertTrue(generated_register.output_path.exists())
            self.assertEqual(
                generated_register.output_path.parent,
                attendance_output_directory,
            )
            self.assertEqual(generated_register.row_count, 2)

            indexed_files = repository.list_indexed_files(
                file_group=FileGroup.FILE_2,
                file_category="attendance_register_docx",
            )
            self.assertEqual(len(indexed_files), 1)
            self.assertEqual(
                indexed_files[0].file_path.resolve(),
                generated_register.output_path.resolve(),
            )

            rendered_document = Document(generated_register.output_path)
            attendance_table = rendered_document.tables[1]
            first_row = [cell.text.strip() for cell in attendance_table.rows[1].cells]
            second_row = [cell.text.strip() for cell in attendance_table.rows[2].cells]
            blank_row = [cell.text.strip() for cell in attendance_table.rows[3].cells]

            self.assertEqual(first_row[0], "13/03/2026")
            self.assertEqual(first_row[1], "Sean Carter")
            self.assertEqual(first_row[2], "A. Archer Electrical")
            self.assertEqual(first_row[3], "07700111222")
            self.assertEqual(first_row[4], "14 miles")
            self.assertEqual(first_row[5], "AB12 CDE")
            self.assertEqual(first_row[6], "07:30")
            self.assertEqual(first_row[8], "15:00")
            self.assertEqual(first_row[9], "7.50")

            self.assertEqual(second_row[1], "Luke Green")
            self.assertEqual(second_row[2], "Uplands Construction Group")
            self.assertEqual(second_row[3], "07700999444")
            self.assertEqual(second_row[5], "XY34 ZZZ")
            self.assertEqual(second_row[8], "")
            self.assertEqual(second_row[9], "")
            self.assertTrue(all(value == "" for value in blank_row))

            summary_table = rendered_document.tables[3]
            self.assertEqual(summary_table.rows[0].cells[1].text.strip(), "1")
            self.assertEqual(summary_table.rows[0].cells[4].text.strip(), "1")

    def test_get_daily_contractor_headcount_groups_active_entries_by_company(self) -> None:
        repository = DocumentRepository(Path(tempfile.mkdtemp()) / "documents.sqlite3")
        repository.create_schema()
        repository.save(
            DailyAttendanceEntryDocument(
                doc_id="ATT-001",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 7, 30),
                status=DocumentStatus.ACTIVE,
                linked_induction_doc_id="IND-001",
                individual_name="Sean Carter",
                contractor_name="A. Archer Electrical",
                vehicle_registration="AB12 CDE",
                distance_travelled="14 miles",
                time_in=datetime(2026, 3, 15, 7, 30),
            )
        )
        repository.save(
            DailyAttendanceEntryDocument(
                doc_id="ATT-002",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 7, 40),
                status=DocumentStatus.ACTIVE,
                linked_induction_doc_id="IND-002",
                individual_name="Luke Green",
                contractor_name="A. Archer Electrical",
                vehicle_registration="XY34 ZZZ",
                distance_travelled="8 miles",
                time_in=datetime(2026, 3, 15, 7, 40),
            )
        )
        repository.save(
            DailyAttendanceEntryDocument(
                doc_id="ATT-003",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 8, 0),
                status=DocumentStatus.ARCHIVED,
                linked_induction_doc_id="IND-003",
                individual_name="Pat Visitor",
                contractor_name="Visitor Co",
                vehicle_registration="",
                distance_travelled="",
                time_in=datetime(2026, 3, 15, 8, 0),
                time_out=datetime(2026, 3, 15, 8, 30),
                hours_worked=0.5,
            )
        )

        contractor_counts = get_daily_contractor_headcount(
            repository,
            "NG Lovedean Substation",
            date(2026, 3, 15),
        )

        self.assertEqual(
            contractor_counts,
            [{"company": "A. Archer Electrical", "days": 2, "nights": 0}],
        )

    def test_generate_site_diary_document_renders_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            output_directory = project_root / "FILE_2_Daily_Site_Diary"
            database_path = project_root / "documents.sqlite3"
            template_path = project_root / "UHSF15.63_Template.docx"
            output_directory.mkdir(parents=True, exist_ok=True)

            template_document = Document()
            template_document.add_paragraph("Date: {{ date }}")
            template_document.add_paragraph("Uplands Days: {{ uplands_days }}")
            template_document.add_paragraph("Uplands Nights: {{ uplands_nights }}")
            template_document.add_paragraph("Skip Exchange: {{ skip_exchange }}")
            template_document.add_paragraph("Day On: {{ fire_day_on }}")
            template_document.add_paragraph("Day Off: {{ fire_day_off }}")
            template_document.add_paragraph("Night On: {{ fire_night_on }}")
            template_document.add_paragraph("Night Off: {{ fire_night_off }}")
            template_document.add_paragraph("Dry: {{ weather_dry }}")
            template_document.add_paragraph("Mixed: {{ weather_mixed }}")
            template_document.add_paragraph("Wet: {{ weather_wet }}")
            template_document.add_paragraph(
                "{% for c in contractors %}Contractor: {{ c.company }} | {{ c.days }} | {{ c.nights }}\n{% endfor %}"
            )
            template_document.add_paragraph("Incidents: {{ incidents_details }}")
            template_document.add_paragraph("H&S: {{ hs_reported_tick }}")
            template_document.add_paragraph(
                "{% for v in visitors %}Visitor: {{ v.name }} | {{ v.company }}\n{% endfor %}"
            )
            template_document.add_paragraph("Handovers: {{ area_handovers }}")
            template_document.add_paragraph("Comments: {{ todays_comments }}")
            template_document.save(template_path)

            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_output_dir = app_config.FILE_2_DIARY_OUTPUT_DIR
            try:
                TemplateRegistry.TEMPLATE_PATHS["site_diary"] = template_path
                app_config.FILE_2_DIARY_OUTPUT_DIR = output_directory

                repository = DocumentRepository(database_path)
                repository.create_schema()
                site_diary_document = SiteDiaryDocument(
                    doc_id="SITE-DIARY-ng-lovedean-substation-20260315",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 15, 9, 0),
                    status=DocumentStatus.ACTIVE,
                    date=date(2026, 3, 15),
                    uplands_days=3,
                    uplands_nights=0,
                    skip_exchange="8-yard open skip",
                    fire_day_on=True,
                    fire_day_off=False,
                    fire_night_on=False,
                    fire_night_off=True,
                    weather_dry=True,
                    weather_mixed=False,
                    weather_wet=False,
                    contractors=[
                        {"company": "A. Archer Electrical", "days": 4, "nights": 0},
                        {"company": "Uplands Retail", "days": 3, "nights": 0},
                    ],
                    visitors=[{"name": "Pat Visitor", "company": "National Grid"}],
                    incidents_details="No incidents reported.",
                    hs_reported_tick=True,
                    area_handovers="Mess room handed back.",
                    todays_comments="Progress on programme.",
                )

                generated_diary = generate_site_diary_document(
                    repository,
                    site_diary_document=site_diary_document,
                )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                app_config.FILE_2_DIARY_OUTPUT_DIR = original_output_dir

            self.assertTrue(generated_diary.output_path.exists())
            self.assertEqual(generated_diary.output_path.parent, output_directory)
            self.assertEqual(generated_diary.contractor_count, 2)
            self.assertEqual(generated_diary.visitor_count, 1)
            saved_diary = repository.get(site_diary_document.doc_id)
            self.assertIsInstance(saved_diary, SiteDiaryDocument)
            self.assertEqual(Path(saved_diary.generated_document_path), generated_diary.output_path)

            rendered_diary = Document(generated_diary.output_path)
            rendered_text = "\n".join(
                paragraph.text.strip()
                for paragraph in rendered_diary.paragraphs
                if paragraph.text.strip()
            )
            self.assertIn("15/03/2026", rendered_text)
            self.assertIn("A. Archer Electrical", rendered_text)
            self.assertIn("Uplands Retail", rendered_text)
            self.assertIn("Pat Visitor", rendered_text)
            self.assertIn("No incidents reported.", rendered_text)
            self.assertIn("Mess room handed back.", rendered_text)
            self.assertIn("Progress on programme.", rendered_text)

    def test_rewrite_inline_table_row_loops_supports_endtr(self) -> None:
        original_xml = (
            "<w:tr>"
            "{% tr for c in contractors %}"
            "<w:tc><w:p><w:r><w:t>{{ c.company }}</w:t></w:r></w:p></w:tc>"
            "{% endtr %}"
            "</w:tr>"
        )

        rewritten_xml = _rewrite_inline_table_row_loops(original_xml)

        self.assertIn("{% for c in contractors %}", rewritten_xml)
        self.assertIn("{{ c.company }}", rewritten_xml)
        self.assertIn("{% endfor %}", rewritten_xml)
        self.assertNotIn("{% endtr %}", rewritten_xml)

    def test_build_live_site_broadcast_contacts_uses_active_mobile_numbers(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            repository.save(
                InductionDocument(
                    doc_id="IND-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 7, 0),
                    status=DocumentStatus.ACTIVE,
                    contractor_name="A. Archer Electrical",
                    individual_name="Sean Carter",
                    contact_number="07700 111222",
                )
            )
            repository.save(
                InductionDocument(
                    doc_id="IND-002",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 7, 5),
                    status=DocumentStatus.ACTIVE,
                    contractor_name="Visitor",
                    individual_name="Site Visitor",
                    contact_number="02392 123456",
                )
            )
            repository.save(
                DailyAttendanceEntryDocument(
                    doc_id="ATT-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 7, 30),
                    status=DocumentStatus.ACTIVE,
                    linked_induction_doc_id="IND-001",
                    individual_name="Sean Carter",
                    contractor_name="A. Archer Electrical",
                    vehicle_registration="AB12 CDE",
                    distance_travelled="14 miles",
                    time_in=datetime(2026, 3, 14, 7, 30),
                )
            )
            repository.save(
                DailyAttendanceEntryDocument(
                    doc_id="ATT-002",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 8, 0),
                    status=DocumentStatus.ACTIVE,
                    linked_induction_doc_id="IND-002",
                    individual_name="Site Visitor",
                    contractor_name="Visitor",
                    vehicle_registration="",
                    distance_travelled="",
                    time_in=datetime(2026, 3, 14, 8, 0),
                )
            )

            contacts = build_live_site_broadcast_contacts(
                repository,
                site_name="NG Lovedean Substation",
                on_date=date(2026, 3, 14),
            )

            self.assertEqual(len(contacts), 1)
            self.assertEqual(contacts[0].individual_name, "Sean Carter")
            self.assertEqual(contacts[0].mobile_number, "+447700111222")
            self.assertEqual(contacts[0].vehicle_registration, "AB12 CDE")

    def test_build_site_alert_sms_link_prefills_numbers_and_message(self) -> None:
        self.assertEqual(
            build_site_alert_sms_link(
                ["+447700111222", "+447700999444"],
                message="TBT in the canteen in 5 mins",
            ),
            "sms:+447700111222,+447700999444&body=TBT%20in%20the%20canteen%20in%205%20mins",
        )

    def test_build_site_alert_sms_links_chunks_large_audiences(self) -> None:
        mobile_numbers = [
            f"+44770011{index:04d}"
            for index in range(30)
        ]

        sms_links = build_site_alert_sms_links(
            mobile_numbers,
            message="Stand down at welfare",
            max_recipients_per_chunk=10,
        )

        self.assertEqual(len(sms_links), 3)
        self.assertTrue(all(link.startswith("sms:+44770011") for link in sms_links))
        self.assertTrue(all("&body=Stand%20down%20at%20welfare" in link for link in sms_links))

    def test_build_toolbox_talk_sms_message_uses_topic_and_link(self) -> None:
        self.assertEqual(
            build_toolbox_talk_sms_message(
                "High winds",
                "https://uplands-site-induction.omegaleague.win?station=tbt&topic=High+winds",
            ),
            "Toolbox Talk: High winds. Please click this link to read the document and sign the register: https://uplands-site-induction.omegaleague.win?station=tbt&topic=High+winds",
        )

    def test_save_app_settings_persists_recent_site_histories(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            original_settings_path = app_config.SETTINGS_PATH
            try:
                app_config.SETTINGS_PATH = Path(temp_dir) / "settings.json"
                save_app_settings(
                    public_tunnel_url="https://uplands-site-induction.omegaleague.win",
                    broadcast_message_history_by_site={
                        "NG Lovedean Substation": [
                            "Stand down in the canteen",
                            "Stand down in the canteen",
                            "High winds - crane suspended",
                        ]
                    },
                    tbt_topic_history_by_site={
                        "NG Lovedean Substation": [
                            "High winds",
                            "Face fit refresh",
                        ]
                    },
                )

                loaded_settings = load_app_settings()
            finally:
                app_config.SETTINGS_PATH = original_settings_path

        self.assertEqual(
            loaded_settings["public_tunnel_url"],
            "https://uplands-site-induction.omegaleague.win",
        )
        self.assertEqual(
            loaded_settings["broadcast_message_history_by_site"][
                "NG Lovedean Substation"
            ],
            [
                "Stand down in the canteen",
                "High winds - crane suspended",
            ],
        )
        self.assertEqual(
            loaded_settings["tbt_topic_history_by_site"]["NG Lovedean Substation"],
            [
                "High winds",
                "Face fit refresh",
            ],
        )

    def test_build_toolbox_talk_url_prefills_topic(self) -> None:
        self.assertEqual(
            build_toolbox_talk_url(
                "Working in high winds",
                public_url="https://uplands-site-induction.omegaleague.win",
            ),
            "https://uplands-site-induction.omegaleague.win?station=tbt&topic=Working+in+high+winds",
        )

    def test_build_toolbox_talk_url_recovers_base_url_from_component_iframe_link(self) -> None:
        self.assertEqual(
            build_toolbox_talk_url(
                "Working in high winds",
                public_url=(
                    "https://uplands-site-induction.omegaleague.win/"
                    "component/streamlit_js_eval.streamlit_js_eval/index.html"
                    "?streamlitUrl=https%3A%2F%2Fuplands-site-induction.omegaleague.win%2F"
                ),
            ),
            "https://uplands-site-induction.omegaleague.win?station=tbt&topic=Working+in+high+winds",
        )

    def test_build_toolbox_talk_document_view_url_uses_public_gps_route(self) -> None:
        self.assertEqual(
            build_toolbox_talk_document_view_url(
                "TBTDOC-123",
                public_url="https://uplands-site-induction.omegaleague.win",
            ),
            "https://uplands-site-induction.omegaleague.win/gps/tbt-preview?doc_id=TBTDOC-123",
        )

    def test_save_app_settings_normalizes_component_iframe_public_url(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            original_settings_path = app_config.SETTINGS_PATH
            try:
                app_config.SETTINGS_PATH = Path(temp_dir) / "settings.json"
                save_app_settings(
                    public_tunnel_url=(
                        "https://uplands-site-induction.omegaleague.win/"
                        "component/streamlit_js_eval.streamlit_js_eval/index.html"
                        "?streamlitUrl=https%3A%2F%2Fuplands-site-induction.omegaleague.win%2F"
                    )
                )
                loaded_settings = load_app_settings()
            finally:
                app_config.SETTINGS_PATH = original_settings_path

        self.assertEqual(
            loaded_settings["public_tunnel_url"],
            "https://uplands-site-induction.omegaleague.win",
        )

    def test_calculate_haversine_distance_meters_matches_expected_site_scale(self) -> None:
        self.assertEqual(
            calculate_haversine_distance_meters(50.917, -1.036, 50.917, -1.036),
            0.0,
        )
        self.assertGreater(
            calculate_haversine_distance_meters(50.917, -1.036, 50.918, -1.036),
            100.0,
        )
        self.assertLess(
            calculate_haversine_distance_meters(50.917, -1.036, 50.918, -1.036),
            120.0,
        )

    def test_site_gate_access_code_validates_current_and_previous_slot(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            settings_path = Path(temp_dir) / "settings.json"
            with patch.object(app_config, "SETTINGS_PATH", settings_path):
                generated_code, minutes_remaining = build_site_gate_access_code(
                    "NG Lovedean Substation",
                    at_time=datetime(2026, 3, 15, 10, 10),
                    slot_minutes=30,
                )
                self.assertEqual(len(generated_code), 6)
                self.assertGreaterEqual(minutes_remaining, 1)
                self.assertTrue(
                    validate_site_gate_access_code(
                        "NG Lovedean Substation",
                        generated_code,
                        at_time=datetime(2026, 3, 15, 10, 10),
                        slot_minutes=30,
                    )
                )
                self.assertTrue(
                    validate_site_gate_access_code(
                        "NG Lovedean Substation",
                        generated_code,
                        at_time=datetime(2026, 3, 15, 10, 35),
                        slot_minutes=30,
                    )
                )
                self.assertFalse(
                    validate_site_gate_access_code(
                        "NG Lovedean Substation",
                        generated_code,
                        at_time=datetime(2026, 3, 15, 11, 15),
                        slot_minutes=30,
                    )
                )
                self.assertFalse(
                    validate_site_gate_access_code(
                        "11 station close",
                        generated_code,
                        at_time=datetime(2026, 3, 15, 10, 10),
                        slot_minutes=30,
                    )
                )

    def test_build_pending_toolbox_talk_contacts_returns_only_unsigned_live_people(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            repository.save(
                InductionDocument(
                    doc_id="IND-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 7, 0),
                    status=DocumentStatus.ACTIVE,
                    contractor_name="A. Archer Electrical",
                    individual_name="Sean Carter",
                    contact_number="07700 111222",
                )
            )
            repository.save(
                InductionDocument(
                    doc_id="IND-002",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 7, 5),
                    status=DocumentStatus.ACTIVE,
                    contractor_name="Uplands Construction Group",
                    individual_name="Luke Green",
                    contact_number="07700 999444",
                )
            )
            repository.save(
                DailyAttendanceEntryDocument(
                    doc_id="ATT-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 7, 30),
                    status=DocumentStatus.ACTIVE,
                    linked_induction_doc_id="IND-001",
                    individual_name="Sean Carter",
                    contractor_name="A. Archer Electrical",
                    vehicle_registration="AB12 CDE",
                    distance_travelled="14 miles",
                    time_in=datetime(2026, 3, 14, 7, 30),
                )
            )
            repository.save(
                DailyAttendanceEntryDocument(
                    doc_id="ATT-002",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 7, 45),
                    status=DocumentStatus.ACTIVE,
                    linked_induction_doc_id="IND-002",
                    individual_name="Luke Green",
                    contractor_name="Uplands Construction Group",
                    vehicle_registration="XY34 ZZZ",
                    distance_travelled="5 miles",
                    time_in=datetime(2026, 3, 14, 7, 45),
                )
            )
            repository.save(
                ToolboxTalkCompletionDocument(
                    doc_id="TBT-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 8, 30),
                    status=DocumentStatus.ACTIVE,
                    topic="High winds",
                    linked_induction_doc_id="IND-001",
                    individual_name="Sean Carter",
                    contractor_name="A. Archer Electrical",
                    completed_at=datetime(2026, 3, 14, 8, 30),
                    signature_image_path="/tmp/signature.png",
                    document_read_confirmed=True,
                )
            )

            pending_contacts = build_pending_toolbox_talk_contacts(
                repository,
                site_name="NG Lovedean Substation",
                topic="High winds",
                on_date=date(2026, 3, 14),
            )

            self.assertEqual(len(pending_contacts), 1)
            self.assertEqual(pending_contacts[0].individual_name, "Luke Green")

    def test_launch_messages_sms_broadcast_opens_messages_drafts(self) -> None:
        mobile_numbers = [f"+44770011{index:04d}" for index in range(25)]

        with patch.object(
            workspace_module,
            "_launch_messages_group_draft_via_gui_automation",
            return_value=(False, "Accessibility not enabled."),
        ):
            with patch.object(workspace_module.time_module, "sleep") as mocked_sleep:
                with patch.object(workspace_module.subprocess, "run") as mocked_run:
                    mocked_run.side_effect = [
                        workspace_module.subprocess.CompletedProcess(
                            ["open", "-Ra", "Messages"],
                            0,
                            "",
                            "",
                        )
                    ] + [
                        workspace_module.subprocess.CompletedProcess(
                            ["open", "-a", "Messages", f"sms:+44770011{index:04d}"],
                            0,
                            "",
                            "",
                        )
                        for index in range(25)
                    ]

                    launch_result = launch_messages_sms_broadcast(
                        mobile_numbers,
                        message="Toolbox talk in ten minutes",
                        max_recipients_per_chunk=20,
                    )

        self.assertTrue(launch_result.launched_successfully)
        self.assertEqual(launch_result.recipient_count, 25)
        self.assertEqual(launch_result.chunk_count, 25)
        self.assertEqual(mocked_run.call_count, 26)
        self.assertEqual(
            mocked_run.call_args_list[1].args[0][:3],
            ["open", "-a", "Messages"],
        )
        self.assertIn("one draft per recipient", launch_result.error_message)
        self.assertEqual(mocked_sleep.call_count, 24)

    def test_launch_messages_sms_broadcast_prefers_grouped_messages_draft(self) -> None:
        mobile_numbers = ["+447700110001", "+447700110002"]

        with patch.object(
            workspace_module,
            "_launch_messages_group_draft_via_gui_automation",
            return_value=(True, ""),
        ) as mocked_group_launch:
            with patch.object(workspace_module.subprocess, "run") as mocked_run:
                mocked_run.side_effect = [
                    workspace_module.subprocess.CompletedProcess(
                        ["open", "-Ra", "Messages"],
                        0,
                        "",
                        "",
                    )
                ]

                launch_result = launch_messages_sms_broadcast(
                    mobile_numbers,
                    message="Toolbox talk in ten minutes",
                )

        self.assertTrue(launch_result.launched_successfully)
        self.assertEqual(launch_result.recipient_count, 2)
        self.assertEqual(launch_result.chunk_count, 1)
        self.assertEqual(mocked_run.call_count, 1)
        mocked_group_launch.assert_called_once()
        self.assertEqual(
            launch_result.error_message,
            "Opened one grouped Messages draft for the live audience.",
        )

    def test_log_broadcast_dispatch_round_trip(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            launch_result = workspace_module.MessagesDraftLaunchResult(
                draft_links=["sms:+447700111222&body=Stand%20down"],
                recipient_count=1,
                chunk_count=1,
                launched_successfully=True,
                error_message="",
            )
            logged_dispatch = log_broadcast_dispatch(
                repository,
                site_name="NG Lovedean Substation",
                dispatch_kind="mass_broadcast",
                audience_label="Everyone On Site",
                subject="Stand down",
                message_body="Stand down in the canteen",
                recipient_numbers=["+447700111222"],
                recipient_names=["Sean Carter"],
                launch_result=launch_result,
            )

            stored_dispatches = list_broadcast_dispatches(
                repository,
                site_name="NG Lovedean Substation",
            )

            self.assertEqual(len(stored_dispatches), 1)
            self.assertIsInstance(stored_dispatches[0], BroadcastDispatchDocument)
            self.assertEqual(stored_dispatches[0].subject, "Stand down")
            self.assertEqual(stored_dispatches[0].recipient_numbers, ["+447700111222"])
            self.assertTrue(logged_dispatch.dispatch_document.launched_successfully)

    def test_toolbox_talk_completion_and_export_round_trip(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            workspace_root = project_root / "Uplands_Workspace"
            tbt_register_dir = workspace_root / "FILE_2_Registers" / "Toolbox_Talk_Register"
            tbt_active_docs_dir = tbt_register_dir / "Active_Docs"
            tbt_signature_dir = tbt_register_dir / "Signatures"
            tbt_output_dir = tbt_register_dir / "Registers"
            templates_dir = project_root / "templates"
            database_path = workspace_root / "documents.sqlite3"

            for directory in (
                tbt_active_docs_dir,
                tbt_signature_dir,
                tbt_output_dir,
                templates_dir,
            ):
                directory.mkdir(parents=True, exist_ok=True)

            template_path = templates_dir / "UHSF16.2_Template.docx"
            document = Document()
            document.add_paragraph("Topic {{ topic }}")
            table = document.add_table(rows=2, cols=4)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Company"
            table.cell(0, 2).text = "Date"
            table.cell(0, 3).text = "Signature"
            table.cell(1, 0).text = "{{ name }}"
            table.cell(1, 1).text = "{{ company }}"
            table.cell(1, 2).text = "{{ date }}"
            table.cell(1, 3).text = "{{ signature }}"
            document.save(template_path)

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.DATABASE_PATH,
                app_config.TOOLBOX_TALK_REGISTER_DIR,
                app_config.FILE_2_TBT_ACTIVE_DOCS_DIR,
                app_config.FILE_2_TBT_SIGNATURES_DIR,
                app_config.FILE_2_TBT_OUTPUT_DIR,
            )
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.DATABASE_PATH = database_path
                app_config.TOOLBOX_TALK_REGISTER_DIR = tbt_register_dir
                app_config.FILE_2_TBT_ACTIVE_DOCS_DIR = tbt_active_docs_dir
                app_config.FILE_2_TBT_SIGNATURES_DIR = tbt_signature_dir
                app_config.FILE_2_TBT_OUTPUT_DIR = tbt_output_dir
                TemplateRegistry.PROJECT_ROOT = project_root
                TemplateRegistry.TEMPLATE_PATHS["toolbox_talk_register"] = Path(
                    "templates/UHSF16.2_Template.docx"
                )

                repository = DocumentRepository(database_path)
                repository.create_schema()

                saved_toolbox_talk_document = save_toolbox_talk_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    topic="Working in high winds",
                    uploaded_file_name="Working-in-high-winds.pdf",
                    uploaded_file_bytes=b"%PDF-1.4 fake toolbox talk",
                )

                repository.save(
                    InductionDocument(
                        doc_id="IND-001",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 14, 7, 0),
                        status=DocumentStatus.ACTIVE,
                        contractor_name="A. Archer Electrical",
                        individual_name="Sean Carter",
                        contact_number="07700111222",
                    )
                )
                first_attendance = DailyAttendanceEntryDocument(
                    doc_id="ATT-001",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 7, 30),
                    status=DocumentStatus.ACTIVE,
                    linked_induction_doc_id="IND-001",
                    individual_name="Sean Carter",
                    contractor_name="A. Archer Electrical",
                    time_in=datetime(2026, 3, 14, 7, 30),
                )
                repository.save(first_attendance)

                signature_image_data = np.full((200, 420, 4), 255, dtype=np.uint8)
                signature_image_data[92:98, 40:220, :3] = 0

                first_completion = log_toolbox_talk_completion(
                    repository,
                    site_name="NG Lovedean Substation",
                    topic="Working in high winds",
                    attendance_entry=first_attendance,
                    signature_image_data=signature_image_data,
                    document_read_confirmed=True,
                )

                repository.save(
                    InductionDocument(
                        doc_id="IND-002",
                        site_name="NG Lovedean Substation",
                        created_at=datetime(2026, 3, 14, 7, 5),
                        status=DocumentStatus.ACTIVE,
                        contractor_name="Uplands Construction Group",
                        individual_name="Luke Green",
                        contact_number="07700999444",
                    )
                )
                second_attendance = DailyAttendanceEntryDocument(
                    doc_id="ATT-002",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 14, 8, 0),
                    status=DocumentStatus.ACTIVE,
                    linked_induction_doc_id="IND-002",
                    individual_name="Luke Green",
                    contractor_name="Uplands Construction Group",
                    time_in=datetime(2026, 3, 14, 8, 0),
                )
                repository.save(second_attendance)
                log_toolbox_talk_completion(
                    repository,
                    site_name="NG Lovedean Substation",
                    topic="Working in high winds",
                    attendance_entry=second_attendance,
                    signature_image_data=signature_image_data,
                    document_read_confirmed=True,
                )

                source_documents = list_toolbox_talk_documents(
                    repository,
                    site_name="NG Lovedean Substation",
                    topic="Working in high winds",
                )
                latest_toolbox_talk_document = get_latest_toolbox_talk_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    topic="Working in high winds",
                )
                source_document_bytes, source_document_mime_type = (
                    read_toolbox_talk_document_bytes(saved_toolbox_talk_document.toolbox_talk_document)
                )
                completions = list_toolbox_talk_completions(
                    repository,
                    site_name="NG Lovedean Substation",
                    topic="Working in high winds",
                )
                generated_register = generate_toolbox_talk_register_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    topic="Working in high winds",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.DATABASE_PATH,
                    app_config.TOOLBOX_TALK_REGISTER_DIR,
                    app_config.FILE_2_TBT_ACTIVE_DOCS_DIR,
                    app_config.FILE_2_TBT_SIGNATURES_DIR,
                    app_config.FILE_2_TBT_OUTPUT_DIR,
                ) = original_config
                TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertEqual(len(source_documents), 1)
            self.assertIsNotNone(latest_toolbox_talk_document)
            self.assertEqual(
                latest_toolbox_talk_document.original_file_name,
                "Working-in-high-winds.pdf",
            )
            self.assertEqual(source_document_bytes, b"%PDF-1.4 fake toolbox talk")
            self.assertEqual(source_document_mime_type, "application/pdf")
            self.assertEqual(len(completions), 2)
            self.assertTrue(first_completion.signature_path.exists())
            self.assertTrue(generated_register.output_path.exists())
            self.assertEqual(generated_register.row_count, 2)

            rendered_document = Document(generated_register.output_path)
            rendered_text = "\n".join(
                paragraph.text for paragraph in rendered_document.paragraphs
            )
            self.assertIn("Working in high winds", rendered_text)
            table = rendered_document.tables[0]
            first_row = [cell.text.strip() for cell in table.rows[1].cells]
            second_row = [cell.text.strip() for cell in table.rows[2].cells]
            self.assertEqual(first_row[0], "Sean Carter")
            self.assertEqual(first_row[1], "A. Archer Electrical")
            self.assertEqual(second_row[0], "Luke Green")
            self.assertEqual(second_row[1], "Uplands Construction Group")

    def test_plant_asset_document_parses_inspection_due_date(self) -> None:
        plant_asset = PlantAssetDocument(
            doc_id="PLANT-002",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 12, 8, 0),
            status=DocumentStatus.ACTIVE,
            hire_num="JOB-4471-02",
            description="Rechargeable Worklight LED",
            company="HSS",
            phone="0161 749 4090",
            on_hire=date(2026, 3, 12),
            hired_by="TDE",
            serial="SER-002",
            inspection="Exam 12/03/2026 | Next due 18/03/2026 | Asset 52 | Report 900",
        )

        self.assertEqual(plant_asset.inspection_due_date(), date(2026, 3, 18))
        self.assertTrue(
            plant_asset.inspection_requires_attention(on_date=date(2026, 3, 12))
        )

    def test_infer_plant_inspection_type_uses_description(self) -> None:
        self.assertEqual(
            infer_plant_inspection_type("LOOSE EXTENSION LEAD 110V 15M 32AM"),
            PlantInspectionType.PAT,
        )
        self.assertEqual(
            infer_plant_inspection_type("FIRE EXTINGUISHER 6L FOAM"),
            PlantInspectionType.EXTINGUISHER,
        )
        self.assertEqual(
            infer_plant_inspection_type("DUST EXTRACTOR M CLASS 110V"),
            PlantInspectionType.SERVICE,
        )

    def test_format_plant_inspection_reference_combines_type_and_reference(self) -> None:
        self.assertEqual(
            format_plant_inspection_reference(
                PlantInspectionType.PAT,
                "Label 51 | Next due 18/03/2026",
            ),
            "PAT - Label 51 | Next due 18/03/2026",
        )
        self.assertEqual(
            format_plant_inspection_reference(
                PlantInspectionType.EXTINGUISHER,
                "",
            ),
            "Extinguisher - cert ref not logged",
        )


class WasteRegisterAutomationTests(unittest.TestCase):
    def _build_waste_kpi_workbook(self, workbook_path: Path) -> None:
        import xlwt

        workbook = xlwt.Workbook()
        sheet = workbook.add_sheet("Sheet1")
        sheet.write(6, 0, "Customer")
        sheet.write(6, 1, "Uplands")
        sheet.write(10, 0, "Project Number")
        sheet.write(10, 1, "JOB-4471")
        sheet.write(12, 1, "Project Name & Address")
        sheet.write(12, 3, "National Grid, Broadway Lane, Waterlooville, Hampshire, PO8 0SJ")
        sheet.write(16, 0, "Person responsible for waste management on site (incl. Job title)")
        sheet.write(16, 4, "Ceri Edwards")
        workbook.save(str(workbook_path))

    def _build_waste_register_template(self, template_path: Path) -> None:
        document = Document()
        document.add_paragraph("Client {{client_name}}")
        document.add_paragraph("Site {{site_address}}")
        document.add_paragraph("Manager {{manager_name}}")
        table = document.add_table(rows=2, cols=4)
        table.cell(0, 0).text = "Carrier"
        table.cell(0, 1).text = "Date"
        table.cell(0, 2).text = "Description"
        table.cell(0, 3).text = "Reg"
        table.cell(1, 0).text = "{% tr for w in waste_entries %}{{w.carrier}}"
        table.cell(1, 1).text = "{{w.date}}"
        table.cell(1, 2).text = "{{w.description}}"
        table.cell(1, 3).text = "{{w.reg_no}}{% tr endfor %}"
        document.save(template_path)

    def test_get_waste_kpi_sheet_metadata_reads_header_values(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            waste_reports_directory = Path(temp_dir) / "Waste_Reports"
            waste_reports_directory.mkdir(parents=True, exist_ok=True)
            workbook_path = waste_reports_directory / "PO8 0SJ - MAR 2026.xls"
            original_waste_reports_destination = app_config.WASTE_REPORTS_DESTINATION
            self._build_waste_kpi_workbook(workbook_path)

            try:
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_directory
                metadata = get_waste_kpi_sheet_metadata(
                    site_name="NG Lovedean Substation",
                    site_address="Broadway Lane, Waterlooville, Hampshire, PO8 0SJ",
                    fallback_project_number="",
                )
            finally:
                app_config.WASTE_REPORTS_DESTINATION = original_waste_reports_destination

            self.assertEqual(metadata.client_name, "Uplands")
            self.assertEqual(metadata.project_number, "JOB-4471")
            self.assertIn("Broadway Lane", metadata.site_address)
            self.assertEqual(metadata.manager_name, "Ceri Edwards")

    def test_smart_scan_waste_transfer_note_extracts_pdf_fields(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            pdf_path = Path(temp_dir) / "31194.pdf"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            pdf_document = fitz.open()
            pdf_page = pdf_document.new_page()
            pdf_page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "Carrier: L&S Waste Management Limited",
                        "Vehicle Reg: AB12 CDE",
                        "Waste Type: Mixed Construction Payment Type",
                        "Net Weight: 2400 kg",
                        "Date: 11/03/2026",
                        "17 09 04",
                    ]
                ),
            )
            pdf_document.save(pdf_path)
            pdf_document.close()

            scanned = smart_scan_waste_transfer_note(
                repository,
                source_path=pdf_path,
            )

            self.assertEqual(scanned.carrier_name, "Abucs")
            self.assertEqual(scanned.wtn_number, "31194")
            self.assertEqual(scanned.vehicle_registration, "AB12 CDE")
            self.assertEqual(scanned.waste_description, "Mixed Construction")
            self.assertEqual(scanned.quantity_tonnes, 2.4)
            self.assertEqual(scanned.ewc_code, "17 09 04")

    def test_smart_scan_waste_transfer_note_recognises_foul_waste_ticket(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            pdf_path = Path(temp_dir) / "foul-ticket.pdf"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            pdf_document = fitz.open()
            pdf_page = pdf_document.new_page()
            pdf_page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "Carrier: Acme Tankers Ltd",
                        "Ticket No: FP-101",
                        "Cess Pit Emptying",
                        "Foul Water Collection",
                        "Date: 16/03/2026",
                    ]
                ),
            )
            pdf_document.save(pdf_path)
            pdf_document.close()

            scanned = smart_scan_waste_transfer_note(
                repository,
                source_path=pdf_path,
            )

            self.assertEqual(scanned.collection_type, "")
            self.assertEqual(scanned.waste_description, "Cess Pit / Foul Waste")
            self.assertEqual(scanned.ewc_code, "20 03 04")

    def test_smart_scan_waste_transfer_note_extracts_collection_type(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            pdf_path = Path(temp_dir) / "30879-1.PDF"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            pdf_document = fitz.open()
            pdf_page = pdf_document.new_page()
            pdf_page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "Ticket No.",
                        "16/03/2026",
                        "30879",
                        "Vehicle Skip Type Tanker-Municipal W WasteTruck Notes Movement Type",
                        "Waste Type:Mixed Construction",
                        "17 09 04",
                    ]
                ),
            )
            pdf_document.save(pdf_path)
            pdf_document.close()

            scanned = smart_scan_waste_transfer_note(
                repository,
                source_path=pdf_path,
            )

            self.assertEqual(scanned.wtn_number, "30879")
            self.assertEqual(scanned.collection_type, "Tanker-Municipal W")

    def test_file_and_index_all_backfills_weightless_abacus_ticket_into_database(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            carrier_docs_destination = workspace_root / "FILE_1_Environment" / "Carrier_Docs"
            waste_reports_destination = workspace_root / "FILE_1_Environment" / "Waste_Reports"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            induction_directory = workspace_root / "FILE_3_Inductions"
            inbox = workspace_root / "ingest"
            database_path = workspace_root / "documents.sqlite3"

            for directory in (
                waste_destination,
                carrier_docs_destination,
                waste_reports_destination,
                attendance_destination,
                induction_directory,
                inbox,
            ):
                directory.mkdir(parents=True, exist_ok=True)

            pdf_path = waste_destination / "30649.PDF"
            document = fitz.open()
            page = document.new_page()
            page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "ABACUS BRISTOL LTD",
                        "Ticket No.",
                        "19/01/2026",
                        "30649",
                        "Vehicle",
                        "Waste Type:Mixed Construction",
                        "17 09 04",
                        "Weight",
                    ]
                ),
            )
            document.save(pdf_path)
            document.close()

            original_base_data_dir = app_config.BASE_DATA_DIR
            original_inbox = app_config.INBOX
            original_waste_destination = app_config.WASTE_DESTINATION
            original_carrier_docs_destination = app_config.CARRIER_DOCS_DESTINATION
            original_waste_reports_destination = app_config.WASTE_REPORTS_DESTINATION
            original_attendance_destination = app_config.ATTENDANCE_DESTINATION
            original_plant_hire_directory = app_config.PLANT_HIRE_REGISTER_DIR
            original_induction_directory = app_config.INDUCTION_DIR
            original_review_directory = app_config.FILE_3_REVIEW_DIR
            original_database_path = app_config.DATABASE_PATH

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INBOX = inbox
                app_config.WASTE_DESTINATION = waste_destination
                app_config.CARRIER_DOCS_DESTINATION = carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = workspace_root / "FILE_2_Registers" / "Plant_Hire_Register"
                app_config.INDUCTION_DIR = induction_directory
                app_config.FILE_3_REVIEW_DIR = induction_directory / "Needs_Review"
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                file_and_index_all(repository)
                waste_transfer_notes = repository.list_documents(
                    document_type=WasteTransferNoteDocument.document_type,
                )
            finally:
                app_config.BASE_DATA_DIR = original_base_data_dir
                app_config.INBOX = original_inbox
                app_config.WASTE_DESTINATION = original_waste_destination
                app_config.CARRIER_DOCS_DESTINATION = original_carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = original_waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = original_attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = original_plant_hire_directory
                app_config.INDUCTION_DIR = original_induction_directory
                app_config.FILE_3_REVIEW_DIR = original_review_directory
                app_config.DATABASE_PATH = original_database_path

            self.assertEqual(len(waste_transfer_notes), 1)
            self.assertEqual(waste_transfer_notes[0].wtn_number, "30649")
            self.assertEqual(waste_transfer_notes[0].carrier_name, "Abucs")
            self.assertEqual(waste_transfer_notes[0].vehicle_registration, "")
            self.assertEqual(waste_transfer_notes[0].quantity_tonnes, 0.0)

    def test_file_and_index_all_routes_foul_waste_pdf_into_file_1(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            carrier_docs_destination = workspace_root / "FILE_1_Environment" / "Carrier_Docs"
            waste_reports_destination = workspace_root / "FILE_1_Environment" / "Waste_Reports"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            induction_directory = workspace_root / "FILE_3_Inductions"
            inbox = workspace_root / "ingest"
            database_path = workspace_root / "documents.sqlite3"

            for directory in (
                waste_destination,
                carrier_docs_destination,
                waste_reports_destination,
                attendance_destination,
                induction_directory,
                inbox,
            ):
                directory.mkdir(parents=True, exist_ok=True)

            pdf_path = inbox / "cess-pit-ticket.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "Carrier: Acme Tankers Ltd",
                        "Ticket No: FP-101",
                        "Cess Pit Emptying",
                        "Foul Water Collection",
                        "Date: 16/03/2026",
                    ]
                ),
            )
            document.save(pdf_path)
            document.close()

            original_base_data_dir = app_config.BASE_DATA_DIR
            original_inbox = app_config.INBOX
            original_waste_destination = app_config.WASTE_DESTINATION
            original_carrier_docs_destination = app_config.CARRIER_DOCS_DESTINATION
            original_waste_reports_destination = app_config.WASTE_REPORTS_DESTINATION
            original_attendance_destination = app_config.ATTENDANCE_DESTINATION
            original_plant_hire_directory = app_config.PLANT_HIRE_REGISTER_DIR
            original_induction_directory = app_config.INDUCTION_DIR
            original_review_directory = app_config.FILE_3_REVIEW_DIR
            original_database_path = app_config.DATABASE_PATH

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INBOX = inbox
                app_config.WASTE_DESTINATION = waste_destination
                app_config.CARRIER_DOCS_DESTINATION = carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = workspace_root / "FILE_2_Registers" / "Plant_Hire_Register"
                app_config.INDUCTION_DIR = induction_directory
                app_config.FILE_3_REVIEW_DIR = induction_directory / "Needs_Review"
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                file_and_index_all(repository)
                waste_transfer_notes = repository.list_documents(
                    document_type=WasteTransferNoteDocument.document_type,
                )
                indexed_files = repository.list_indexed_files(file_group=FileGroup.FILE_1)
            finally:
                app_config.BASE_DATA_DIR = original_base_data_dir
                app_config.INBOX = original_inbox
                app_config.WASTE_DESTINATION = original_waste_destination
                app_config.CARRIER_DOCS_DESTINATION = original_carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = original_waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = original_attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = original_plant_hire_directory
                app_config.INDUCTION_DIR = original_induction_directory
                app_config.FILE_3_REVIEW_DIR = original_review_directory
                app_config.DATABASE_PATH = original_database_path

            self.assertEqual(len(waste_transfer_notes), 1)
            self.assertEqual(waste_transfer_notes[0].waste_description, "Cess Pit / Foul Waste")
            self.assertEqual(waste_transfer_notes[0].ewc_code, "20 03 04")
            self.assertEqual(indexed_files[0].file_category, "waste_ticket_pdf")

    def test_sync_existing_waste_transfer_notes_ignores_support_pdf_in_waste_notes(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            waste_destination = Path(temp_dir) / "Waste_Notes"
            waste_destination.mkdir(parents=True, exist_ok=True)
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            pdf_path = waste_destination / "KE2428 datasheet.pdf"
            pdf_document = fitz.open()
            pdf_page = pdf_document.new_page()
            pdf_page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "Rapid PVC/CV, Lino and Carpet Adhesive",
                        "MAIN APPLICATION FIELD",
                        "Ready to use, water based dispersion adhesive.",
                    ]
                ),
            )
            pdf_document.save(pdf_path)
            pdf_document.close()

            workspace_module._sync_existing_waste_transfer_notes(
                repository,
                waste_destination,
            )

            waste_transfer_notes = repository.list_documents(
                document_type=WasteTransferNoteDocument.document_type,
            )
            self.assertEqual(len(waste_transfer_notes), 0)

    def test_sync_existing_waste_transfer_notes_archives_stale_support_doc_record(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            waste_destination = Path(temp_dir) / "Waste_Notes"
            waste_destination.mkdir(parents=True, exist_ok=True)
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            pdf_path = waste_destination / "KE2428 datasheet.pdf"
            pdf_document = fitz.open()
            pdf_page = pdf_document.new_page()
            pdf_page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "Rapid PVC/CV, Lino and Carpet Adhesive",
                        "MAIN APPLICATION FIELD",
                        "Ready to use, water based dispersion adhesive.",
                    ]
                ),
            )
            pdf_document.save(pdf_path)
            pdf_document.close()

            stale_document = WasteTransferNoteDocument(
                doc_id="WTN-KE2428",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 17, 8, 0),
                status=DocumentStatus.ACTIVE,
                wtn_number="KE2428 datasheet",
                date=date(2026, 3, 17),
                waste_description="Mixed Construction",
                ewc_code="17 09 04",
                quantity_tonnes=0.0,
                carrier_name="Ke2428 Datasheet",
                destination_facility="Not captured from ticket PDF",
                vehicle_registration="",
            )
            repository.save(stale_document)
            repository.index_file(
                file_name=pdf_path.name,
                file_path=pdf_path,
                file_category="waste_ticket_pdf",
                file_group=FileGroup.FILE_1,
                site_name=stale_document.site_name,
                related_doc_id=stale_document.doc_id,
            )

            workspace_module._sync_existing_waste_transfer_notes(
                repository,
                waste_destination,
            )

            refreshed_document = next(
                document
                for document in repository.list_documents(
                    document_type=WasteTransferNoteDocument.document_type,
                    site_name=stale_document.site_name,
                )
                if document.doc_id == stale_document.doc_id
            )
            self.assertEqual(refreshed_document.status, DocumentStatus.ARCHIVED)

    def test_log_uploaded_waste_transfer_note_moves_file_and_updates_register(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            database_path = workspace_root / "documents.sqlite3"
            source_path = Path(temp_dir) / "scan.pdf"
            source_path.write_bytes(b"%PDF-1.4")
            original_waste_destination = app_config.WASTE_DESTINATION
            repository = DocumentRepository(database_path)

            try:
                app_config.WASTE_DESTINATION = waste_destination
                logged = log_uploaded_waste_transfer_note(
                    repository,
                    upload_path=source_path,
                    original_filename="scan.pdf",
                    site_name="NG Lovedean Substation",
                    carrier_name="Abucs",
                    vehicle_registration="AB12 CDE",
                    waste_description="Mixed Construction",
                    ticket_date=date(2026, 3, 11),
                    quantity_tonnes=2.4,
                    ewc_code="17 09 04",
                    wtn_number="31194",
                )
            finally:
                app_config.WASTE_DESTINATION = original_waste_destination

            self.assertTrue(logged.stored_file_path.exists())
            self.assertFalse(source_path.exists())
            self.assertEqual(logged.waste_transfer_note.vehicle_registration, "AB12 CDE")
            self.assertIsNotNone(logged.register_document)
            self.assertEqual(len(logged.register_document.waste_transfer_notes), 1)

    def test_sync_existing_waste_transfer_notes_prefers_stronger_duplicate_source(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            waste_destination = Path(temp_dir) / "Waste_Notes"
            waste_destination.mkdir(parents=True, exist_ok=True)
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            first_pdf_path = waste_destination / "30879.PDF"
            first_document = fitz.open()
            first_page = first_document.new_page()
            first_page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "ABACUS BRISTOL LTD",
                        "Ticket No.",
                        "10/03/2026",
                        "30879",
                        "Vehicle",
                        "Waste Type:Mixed Construction",
                        "17 09 04",
                    ]
                ),
            )
            first_document.save(first_pdf_path)
            first_document.close()

            second_pdf_path = waste_destination / "30879-1.PDF"
            second_document = fitz.open()
            second_page = second_document.new_page()
            second_page.insert_text(
                (72, 72),
                "\n".join(
                    [
                        "ABACUS BRISTOL LTD",
                        "Ticket No.",
                        "16/03/2026",
                        "30879",
                        "Vehicle Reg: AB12 CDE",
                        "Waste Type:Mixed Construction",
                        "Net Weight: 3200 kg",
                        "17 09 04",
                    ]
                ),
            )
            second_document.save(second_pdf_path)
            second_document.close()

            workspace_module._sync_existing_waste_transfer_notes(
                repository,
                waste_destination,
            )

            waste_transfer_notes = repository.list_documents(
                document_type=WasteTransferNoteDocument.document_type,
            )
            self.assertEqual(len(waste_transfer_notes), 1)
            self.assertEqual(waste_transfer_notes[0].wtn_number, "30879")
            self.assertEqual(waste_transfer_notes[0].date, date(2026, 3, 16))
            self.assertEqual(waste_transfer_notes[0].quantity_tonnes, 3.2)
            self.assertEqual(waste_transfer_notes[0].vehicle_registration, "AB12 CDE")
            self.assertEqual(
                Path(waste_transfer_notes[0].canonical_source_path).name,
                "30879-1.PDF",
            )
            self.assertEqual(len(waste_transfer_notes[0].source_conflict_candidates), 2)

            indexed_files = repository.list_indexed_files(
                related_doc_id=waste_transfer_notes[0].doc_id
            )
            self.assertEqual(len(indexed_files), 2)

            conflicts = workspace_module.list_waste_transfer_note_source_conflicts(
                repository,
                site_name=waste_transfer_notes[0].site_name,
                waste_destination=waste_destination,
            )
            self.assertEqual(len(conflicts), 1)
            self.assertEqual(conflicts[0].wtn_number, "30879")
            self.assertEqual(
                conflicts[0].canonical_source.source_path.name,
                "30879-1.PDF",
            )

    def test_set_waste_transfer_note_source_override_persists_manual_choice(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            waste_destination = Path(temp_dir) / "Waste_Notes"
            waste_destination.mkdir(parents=True, exist_ok=True)
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            for file_name, ticket_date in (
                ("30879.PDF", "10/03/2026"),
                ("30879-1.PDF", "16/03/2026"),
            ):
                pdf_path = waste_destination / file_name
                pdf_document = fitz.open()
                pdf_page = pdf_document.new_page()
                pdf_page.insert_text(
                    (72, 72),
                    "\n".join(
                        [
                            "ABACUS BRISTOL LTD",
                            "Ticket No.",
                            ticket_date,
                            "30879",
                            "Waste Type:Mixed Construction",
                            "17 09 04",
                        ]
                    ),
                )
                pdf_document.save(pdf_path)
                pdf_document.close()

            workspace_module._sync_existing_waste_transfer_notes(
                repository,
                waste_destination,
            )
            waste_transfer_note = repository.list_documents(
                document_type=WasteTransferNoteDocument.document_type,
            )[0]
            self.assertEqual(waste_transfer_note.date, date(2026, 3, 16))

            workspace_module.set_waste_transfer_note_source_override(
                repository,
                source_document=waste_transfer_note,
                source_path=waste_destination / "30879.PDF",
            )
            workspace_module._sync_existing_waste_transfer_notes(
                repository,
                waste_destination,
            )

            refreshed_note = repository.list_documents(
                document_type=WasteTransferNoteDocument.document_type,
            )[0]
            self.assertEqual(refreshed_note.date, date(2026, 3, 10))
            self.assertEqual(
                Path(refreshed_note.source_file_override_path).name,
                "30879.PDF",
            )
            self.assertEqual(
                Path(refreshed_note.canonical_source_path).name,
                "30879.PDF",
            )
            self.assertEqual(len(refreshed_note.source_conflict_candidates), 2)

    def test_sync_existing_waste_transfer_notes_splits_tanker_runs_by_date(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            waste_destination = Path(temp_dir) / "Waste_Notes"
            waste_destination.mkdir(parents=True, exist_ok=True)
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()

            for file_name, ticket_date in (
                ("30879.PDF", "10/03/2026"),
                ("30879-2.PDF", "10/03/2026"),
                ("30879-1.PDF", "16/03/2026"),
            ):
                pdf_path = waste_destination / file_name
                pdf_document = fitz.open()
                pdf_page = pdf_document.new_page()
                pdf_page.insert_text(
                    (72, 72),
                    "\n".join(
                        [
                            "ABACUS BRISTOL LTD",
                            "Ticket No.",
                            ticket_date,
                            "30879",
                            "Skip Type Tanker-Municipal W",
                            "Waste Type:Mixed Construction",
                            "17 09 04",
                        ]
                    ),
                )
                pdf_document.save(pdf_path)
                pdf_document.close()

            workspace_module._sync_existing_waste_transfer_notes(
                repository,
                waste_destination,
            )

            waste_transfer_notes = sorted(
                repository.list_documents(
                    document_type=WasteTransferNoteDocument.document_type,
                ),
                key=lambda note: note.date,
            )
            self.assertEqual(len(waste_transfer_notes), 2)
            self.assertEqual(
                [waste_transfer_note.date for waste_transfer_note in waste_transfer_notes],
                [date(2026, 3, 10), date(2026, 3, 16)],
            )
            self.assertEqual(len(waste_transfer_notes[0].source_conflict_candidates), 2)
            self.assertEqual(len(waste_transfer_notes[1].source_conflict_candidates), 1)

            conflicts = workspace_module.list_waste_transfer_note_source_conflicts(
                repository,
                site_name=waste_transfer_notes[0].site_name,
                waste_destination=waste_destination,
            )
            self.assertEqual(len(conflicts), 1)
            self.assertEqual(conflicts[0].wtn_number, "30879")
            self.assertEqual(
                conflicts[0].canonical_source.scanned_note.ticket_date,
                date(2026, 3, 10),
            )

    def test_build_live_waste_register_rows_flags_only_real_waste_qa_issues(self) -> None:
        ready_note = WasteTransferNoteDocument(
            doc_id="WTN-READY",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 17, 8, 0),
            status=DocumentStatus.ACTIVE,
            wtn_number="31420",
            date=date(2026, 3, 12),
            waste_description="Mixed Construction",
            ewc_code="17 09 04",
            quantity_tonnes=1.42,
            carrier_name="Abucs",
            destination_facility="Not captured from ticket PDF",
            vehicle_registration="",
        )
        missing_tonnage_note = WasteTransferNoteDocument(
            doc_id="WTN-MISSING",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 17, 8, 5),
            status=DocumentStatus.ACTIVE,
            wtn_number="30649",
            date=date(2026, 1, 19),
            waste_description="Mixed Construction",
            ewc_code="17 09 04",
            quantity_tonnes=0.0,
            carrier_name="Abucs",
            destination_facility="Not captured from ticket PDF",
            vehicle_registration="",
        )
        register_rows = app_module._build_live_waste_register_rows(
            [ready_note, missing_tonnage_note],
            waste_source_conflict_lookup={"31420": object()},
        )

        rows_by_ticket = {row["Ticket No"]: row for row in register_rows}
        self.assertEqual(rows_by_ticket["31420"]["QA"], "Source Conflict")
        self.assertEqual(rows_by_ticket["31420"]["Tonnes"], "1.42")
        self.assertEqual(rows_by_ticket["30649"]["QA"], "Needs Review")
        self.assertEqual(rows_by_ticket["30649"]["Tonnes"], "Needs review")

    def test_build_live_waste_register_rows_shows_reviewed_missing_tonnage_status(self) -> None:
        reviewed_note = WasteTransferNoteDocument(
            doc_id="WTN-TANKER",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 17, 8, 5),
            status=DocumentStatus.ACTIVE,
            wtn_number="30879",
            date=date(2026, 3, 16),
            waste_description="Mixed Construction",
            ewc_code="17 09 04",
            quantity_tonnes=0.0,
            carrier_name="Abucs",
            destination_facility="Not captured from ticket PDF",
            vehicle_registration="",
            tonnage_review_status="Weight not shown on supplier ticket",
        )

        register_rows = app_module._build_live_waste_register_rows(
            [reviewed_note],
            waste_source_conflict_lookup={},
        )

        self.assertEqual(register_rows[0]["QA"], "Weight not shown on supplier ticket")
        self.assertEqual(register_rows[0]["Tonnes"], "Not shown on ticket")

    def test_build_file_1_waste_review_queue_rows_only_returns_unresolved_tickets(self) -> None:
        ready_note = WasteTransferNoteDocument(
            doc_id="WTN-READY",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 17, 8, 0),
            status=DocumentStatus.ACTIVE,
            wtn_number="31420",
            date=date(2026, 3, 12),
            waste_description="Mixed Construction",
            ewc_code="17 09 04",
            quantity_tonnes=1.42,
            carrier_name="Abucs",
            destination_facility="Not captured from ticket PDF",
            vehicle_registration="",
        )
        reviewed_tanker_note = WasteTransferNoteDocument(
            doc_id="WTN-TANKER",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 17, 8, 4),
            status=DocumentStatus.ACTIVE,
            wtn_number="30879",
            date=date(2026, 3, 16),
            waste_description="Mixed Construction",
            ewc_code="17 09 04",
            quantity_tonnes=0.0,
            carrier_name="Abucs",
            destination_facility="Not captured from ticket PDF",
            vehicle_registration="",
            tonnage_review_status="Awaiting monthly waste report",
        )
        unresolved_note = WasteTransferNoteDocument(
            doc_id="WTN-UNRESOLVED",
            site_name="NG Lovedean Substation",
            created_at=datetime(2026, 3, 17, 8, 5),
            status=DocumentStatus.ACTIVE,
            wtn_number="30649",
            date=date(2026, 1, 19),
            waste_description="Mixed Construction",
            ewc_code="17 09 04",
            quantity_tonnes=0.0,
            carrier_name="Abucs",
            destination_facility="Not captured from ticket PDF",
            vehicle_registration="",
        )
        queue_rows = app_module._build_file_1_waste_review_queue_rows(
            [ready_note, reviewed_tanker_note, unresolved_note],
            waste_source_conflict_lookup={},
        )

        self.assertEqual(len(queue_rows), 1)
        self.assertEqual(queue_rows[0]["WTN"], "30649")
        self.assertEqual(queue_rows[0]["Issue"], "Missing Tonnage")

    def test_update_logged_waste_transfer_note_persists_missing_tonnage_resolution(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            source_document = WasteTransferNoteDocument(
                doc_id="WTN-30879-2026-03-16",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 17, 8, 5),
                status=DocumentStatus.ACTIVE,
                wtn_number="30879",
                date=date(2026, 3, 16),
                waste_description="Mixed Construction",
                ewc_code="17 09 04",
                quantity_tonnes=0.0,
                carrier_name="Abucs",
                destination_facility="Not captured from ticket PDF",
                vehicle_registration="",
            )
            repository.save(source_document)

            logged_waste_note = workspace_module.update_logged_waste_transfer_note(
                repository,
                source_document=source_document,
                site_name="NG Lovedean Substation",
                carrier_name="Abucs",
                vehicle_registration="",
                waste_description="Mixed Construction",
                ticket_date=date(2026, 3, 16),
                quantity_tonnes=0.0,
                ewc_code="17 09 04",
                destination_facility="Not captured from ticket PDF",
                tonnage_review_status="Resolved by manager",
            )

            self.assertEqual(
                logged_waste_note.waste_transfer_note.tonnage_review_status,
                "Resolved by manager",
            )

    def test_generate_waste_register_document_renders_and_indexes_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_directory = Path(temp_dir) / "output"
            database_path = Path(temp_dir) / "documents.sqlite3"
            template_path = Path(temp_dir) / "waste-register-template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_file_1_output_dir = app_config.FILE_1_OUTPUT_DIR
            repository = DocumentRepository(database_path)
            repository.create_schema()
            self._build_waste_register_template(template_path)

            repository.save(
                WasteTransferNoteDocument(
                    doc_id="WTN-31194",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 11, 8, 0),
                    status=DocumentStatus.ACTIVE,
                    wtn_number="31194",
                    date=date(2026, 3, 11),
                    waste_description="Mixed Construction",
                    ewc_code="17 09 04",
                    quantity_tonnes=2.4,
                    carrier_name="Abucs",
                    destination_facility="Not captured from ticket PDF",
                    vehicle_registration="AB12 CDE",
                    canonical_source_path="/tmp/31194.PDF",
                    source_conflict_candidates=[
                        {
                            "source_path": "/tmp/31194.PDF",
                            "ticket_date": "2026-03-11",
                            "collection_type": "Tanker-Municipal W",
                            "waste_description": "Mixed Construction",
                            "quantity_tonnes": 2.4,
                        }
                    ],
                )
            )

            try:
                TemplateRegistry.TEMPLATE_PATHS["waste_register"] = template_path
                app_config.FILE_1_OUTPUT_DIR = output_directory
                generated = generate_waste_register_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    client_name="Uplands",
                    site_address="National Grid, Broadway Lane, Waterlooville, Hampshire, PO8 0SJ",
                    manager_name="Ceri Edwards",
                )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                app_config.FILE_1_OUTPUT_DIR = original_file_1_output_dir

            self.assertTrue(generated.output_path.exists())
            self.assertEqual(generated.row_count, 1)

            rendered = Document(generated.output_path)
            self.assertIn("Client Uplands", rendered.paragraphs[0].text)
            self.assertEqual(rendered.tables[0].cell(1, 0).text, "Abucs")
            self.assertEqual(
                rendered.tables[0].cell(1, 2).text,
                "Tanker-Municipal W - Mixed Construction",
            )
            self.assertEqual(rendered.tables[0].cell(1, 3).text, "AB12 CDE / 31194")

    def test_generate_waste_register_document_includes_skip_type_in_description(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_directory = Path(temp_dir) / "output"
            database_path = Path(temp_dir) / "documents.sqlite3"
            template_path = Path(temp_dir) / "waste-register-template.docx"
            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            original_file_1_output_dir = app_config.FILE_1_OUTPUT_DIR
            repository = DocumentRepository(database_path)
            repository.create_schema()
            self._build_waste_register_template(template_path)

            repository.save(
                WasteTransferNoteDocument(
                    doc_id="WTN-31420",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 12, 8, 0),
                    status=DocumentStatus.ACTIVE,
                    wtn_number="31420",
                    date=date(2026, 3, 12),
                    waste_description="Mixed Construction",
                    ewc_code="17 09 04",
                    quantity_tonnes=1.42,
                    carrier_name="Abucs",
                    destination_facility="Not captured from ticket PDF",
                    vehicle_registration="",
                    canonical_source_path="/tmp/31420.PDF",
                    source_conflict_candidates=[
                        {
                            "source_path": "/tmp/31420.PDF",
                            "ticket_date": "2026-03-12",
                            "collection_type": "12 Yard ENCLOSED Exchange",
                            "waste_description": "Mixed Construction",
                            "quantity_tonnes": 1.42,
                        }
                    ],
                )
            )

            try:
                TemplateRegistry.TEMPLATE_PATHS["waste_register"] = template_path
                app_config.FILE_1_OUTPUT_DIR = output_directory
                generated = generate_waste_register_document(
                    repository,
                    site_name="NG Lovedean Substation",
                    client_name="Uplands",
                    site_address="National Grid, Broadway Lane, Waterlooville, Hampshire, PO8 0SJ",
                    manager_name="Ceri Edwards",
                )
            finally:
                TemplateRegistry.TEMPLATE_PATHS = original_registry
                app_config.FILE_1_OUTPUT_DIR = original_file_1_output_dir

            rendered = Document(generated.output_path)
            self.assertEqual(
                rendered.tables[0].cell(1, 2).text,
                "12 Yard ENCLOSED Exchange - Mixed Construction",
            )


class SafetyScannerAutomationTests(unittest.TestCase):
    def _build_word_document(self, document_path: Path, lines: list[str]) -> None:
        document = Document()
        for line in lines:
            document.add_paragraph(line)
        document.save(document_path)

    def test_file_and_index_all_imports_file_3_word_documents(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            inbox = workspace_root / "ingest"
            waste_destination = workspace_root / "FILE_1_Environment" / "Waste_Notes"
            carrier_docs_destination = workspace_root / "FILE_1_Environment" / "Carrier_Docs"
            waste_reports_destination = workspace_root / "FILE_1_Environment" / "Waste_Reports"
            attendance_destination = workspace_root / "FILE_2_Registers" / "Attendance"
            plant_hire_destination = workspace_root / "FILE_2_Registers" / "Plant_Hire_Register"
            induction_directory = workspace_root / "FILE_3_Inductions"
            rams_destination = workspace_root / "FILE_3_Safety" / "RAMS"
            coshh_destination = workspace_root / "FILE_3_Safety" / "COSHH"
            file_3_output_directory = workspace_root / "FILE_3_Safety" / "Registers"
            database_path = workspace_root / "documents.sqlite3"

            for directory in (
                inbox,
                waste_destination,
                carrier_docs_destination,
                waste_reports_destination,
                attendance_destination,
                plant_hire_destination,
                induction_directory,
                rams_destination,
                coshh_destination,
                file_3_output_directory,
            ):
                directory.mkdir(parents=True, exist_ok=True)

            (workspace_root / "project_setup.json").write_text(
                json.dumps(
                    {
                        "current_site_name": "NG Lovedean Substation",
                        "job_number": "81888",
                        "site_address": "Broadway Lane, Waterlooville, Hampshire, PO8 0SJ",
                        "client_name": "National Grid",
                    }
                ),
                encoding="utf-8",
            )

            rams_docx_path = inbox / "TDE Method Statement.docx"
            coshh_docx_path = inbox / "CT1 COSHH Assessment.docx"
            self._build_word_document(
                rams_docx_path,
                [
                    "Risk Assessment and Method Statement",
                    "Reference: RAMS-47",
                    "Version: 3.1",
                    "Activity Description: Cable tray installation in switch room",
                    "Contractor: TDE",
                    "Review Date: 12/03/2026",
                ],
            )
            self._build_word_document(
                coshh_docx_path,
                [
                    "COSHH Assessment",
                    "Substance Name: CT1 Sealant",
                    "Supplier: C-Tec",
                    "Reference: COSHH-12",
                    "Version: 2.0",
                    "Intended Use: General sealing works",
                    "Review Date: 12/03/2026",
                ],
            )

            original_base_data_dir = app_config.BASE_DATA_DIR
            original_inbox = app_config.INBOX
            original_waste_destination = app_config.WASTE_DESTINATION
            original_carrier_docs_destination = app_config.CARRIER_DOCS_DESTINATION
            original_waste_reports_destination = app_config.WASTE_REPORTS_DESTINATION
            original_attendance_destination = app_config.ATTENDANCE_DESTINATION
            original_plant_hire_directory = app_config.PLANT_HIRE_REGISTER_DIR
            original_induction_directory = app_config.INDUCTION_DIR
            original_rams_destination = app_config.RAMS_DESTINATION
            original_coshh_destination = app_config.COSHH_DESTINATION
            original_review_directory = app_config.FILE_3_REVIEW_DIR
            original_file_3_output_dir = app_config.FILE_3_OUTPUT_DIR
            original_database_path = app_config.DATABASE_PATH

            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.INBOX = inbox
                app_config.WASTE_DESTINATION = waste_destination
                app_config.CARRIER_DOCS_DESTINATION = carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = plant_hire_destination
                app_config.INDUCTION_DIR = induction_directory
                app_config.RAMS_DESTINATION = rams_destination
                app_config.COSHH_DESTINATION = coshh_destination
                app_config.FILE_3_REVIEW_DIR = workspace_root / "FILE_3_Inductions" / "Needs_Review"
                app_config.FILE_3_OUTPUT_DIR = file_3_output_directory
                app_config.DATABASE_PATH = database_path

                repository = DocumentRepository(database_path)
                filed_assets = file_and_index_all(repository)
                rams_documents = repository.list_documents(
                    document_type=RAMSDocument.document_type,
                )
                coshh_documents = repository.list_documents(
                    document_type=COSHHDocument.document_type,
                )
            finally:
                app_config.BASE_DATA_DIR = original_base_data_dir
                app_config.INBOX = original_inbox
                app_config.WASTE_DESTINATION = original_waste_destination
                app_config.CARRIER_DOCS_DESTINATION = original_carrier_docs_destination
                app_config.WASTE_REPORTS_DESTINATION = original_waste_reports_destination
                app_config.ATTENDANCE_DESTINATION = original_attendance_destination
                app_config.PLANT_HIRE_REGISTER_DIR = original_plant_hire_directory
                app_config.INDUCTION_DIR = original_induction_directory
                app_config.RAMS_DESTINATION = original_rams_destination
                app_config.COSHH_DESTINATION = original_coshh_destination
                app_config.FILE_3_REVIEW_DIR = original_review_directory
                app_config.FILE_3_OUTPUT_DIR = original_file_3_output_dir
                app_config.DATABASE_PATH = original_database_path

            self.assertEqual(len(rams_documents), 1)
            self.assertEqual(rams_documents[0].reference, "RAMS-47")
            self.assertEqual(rams_documents[0].version, "3.1")
            self.assertIn("Cable tray installation", rams_documents[0].activity_description)
            self.assertEqual(len(coshh_documents), 1)
            self.assertEqual(coshh_documents[0].reference, "COSHH-12")
            self.assertEqual(coshh_documents[0].version, "2.0")
            self.assertEqual(coshh_documents[0].substance_name, "CT1 Sealant")
            self.assertEqual(coshh_documents[0].manufacturer, "C-Tec")
            self.assertTrue((rams_destination / rams_docx_path.name).exists())
            self.assertTrue((coshh_destination / coshh_docx_path.name).exists())
            filed_categories = {asset.file_category for asset in filed_assets}
            self.assertIn("rams_docx", filed_categories)
            self.assertIn("coshh_docx", filed_categories)

    def test_safe_extract_word_text_reads_legacy_doc_via_conversion(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            converted_docx_path = Path(temp_dir) / "converted.docx"
            legacy_doc_path = Path(temp_dir) / "legacy.doc"
            legacy_doc_path.write_bytes(b"legacy")
            self._build_word_document(
                converted_docx_path,
                [
                    "COSHH Assessment",
                    "Reference: COSHH-21",
                    "Version: 4.0",
                ],
            )

            def _fake_convert(source_path: Path, destination_path: Path) -> None:
                self.assertEqual(source_path, legacy_doc_path)
                destination_path.write_bytes(converted_docx_path.read_bytes())

            with patch.object(
                workspace_module,
                "_convert_legacy_word_document_to_docx",
                side_effect=_fake_convert,
            ):
                extracted_text = workspace_module._safe_extract_word_text(legacy_doc_path)

            self.assertIn("COSHH Assessment", extracted_text)
            self.assertIn("Reference: COSHH-21", extracted_text)
            self.assertIn("Version: 4.0", extracted_text)

    def test_file_3_filename_first_parser_ignores_blacklisted_site_terms(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()
            source_path = Path(
                "NG-National Grid-Lovedean-(Bluecord)-RAMS-CCTV Camera Install-Rev 3.docx"
            )
            noisy_text = (
                "and people at risk "
                "and people at risk "
                "and people at risk "
                "and people at risk"
            )

            contractor_name = workspace_module._guess_file_3_contractor_name(
                repository,
                site_name="NG Lovedean Substation",
                pdf_text=noisy_text,
                source_path=source_path,
            )
            activity_description = workspace_module._extract_rams_activity_description(
                noisy_text,
                source_path,
            )
            version = workspace_module._extract_safety_version(
                noisy_text,
                source_path=source_path,
            )

            self.assertEqual(contractor_name, "Bluecord")
            self.assertEqual(activity_description, "CCTV Camera Install")
            self.assertEqual(version, "3")
            self.assertNotIn(contractor_name.casefold(), {"ng", "national grid", "lovedean"})

    def test_file_3_version_fallback_rejects_reviewed_text_noise(self) -> None:
        version = workspace_module._extract_safety_version(
            "Reviewed by Ceri Edwards\nReview Date: 12/03/2026\nVersion: 2a",
        )
        noisy_version = workspace_module._extract_safety_version(
            "Reviewed by Ceri Edwards\nReview Date: 12/03/2026",
        )

        self.assertEqual(version, "2a")
        self.assertEqual(noisy_version, "1.0")

    def test_file_3_text_fallback_rejects_blacklisted_contractor_names(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()

            contractor_name = workspace_module._guess_file_3_contractor_name(
                repository,
                site_name="NG Lovedean Substation",
                pdf_text="Client: National Grid\nSite: Lovedean Substation\nReviewed by NG",
                source_path=Path("RAMS_Method_Statement.docx"),
            )

            self.assertEqual(contractor_name, "Site Contractor")

    def test_file_3_anchor_labels_override_filename_parser(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "documents.sqlite3"
            repository = DocumentRepository(database_path)
            repository.create_schema()
            source_path = Path("NG-Lovedean-UHSF20.1-Review-Form-Rev-iewed.docx")
            anchored_text = (
                "COMPANY NAME:\nUplands\n"
                "RAMS TITLE:\nSwitch Room Cable Tray Install\n"
                "RAMS VERSION:\n03"
            )

            contractor_name = workspace_module._guess_file_3_contractor_name(
                repository,
                site_name="NG Lovedean Substation",
                pdf_text=anchored_text,
                source_path=source_path,
            )
            activity_description = workspace_module._extract_rams_activity_description(
                anchored_text,
                source_path,
            )
            version = workspace_module._extract_safety_version(
                anchored_text,
                source_path=source_path,
            )

            self.assertEqual(contractor_name, "Uplands")
            self.assertEqual(activity_description, "Switch Room Cable Tray Install")
            self.assertEqual(version, "03")

    def test_is_rams_safety_source_excludes_review_forms_from_live_register(self) -> None:
        self.assertFalse(
            workspace_module._is_rams_safety_source(
                Path("UHSF20.1 Review Form - Bluecord RAMS.docx"),
                "Risk Assessment and Method Statement",
            )
        )
        self.assertFalse(
            workspace_module._is_rams_safety_source(
                Path("Bluecord Review Form.docx"),
                "Review Form only",
            )
        )
        self.assertTrue(
            workspace_module._is_rams_safety_source(
                Path("Bluecord RAMS Rev 2.docx"),
                "Risk Assessment and Method Statement",
            )
        )

    def test_is_coshh_safety_source_rejects_rams_documents_with_generic_labels(self) -> None:
        self.assertFalse(
            workspace_module._is_coshh_safety_source(
                Path("MS26-003-002 - Ventilation Install National Grid Lovedean.docx"),
                (
                    "COMPANY NAME: Bluecord\n"
                    "BRIEF DESCRIPTION OF WORK: Installation of ventilation services\n"
                    "Risk Assessment and Method Statement"
                ),
            )
        )

    def test_is_coshh_safety_source_rejects_construction_phase_plan(self) -> None:
        self.assertFalse(
            workspace_module._is_coshh_safety_source(
                Path("UHSF15.1 Construction Phase Plan - NG Lovedean rev 2.docx"),
                "Construction Phase Plan\nProject Details\nSite safety management",
            )
        )

    def test_is_rams_safety_source_does_not_capture_safety_data_sheet(self) -> None:
        self.assertFalse(
            workspace_module._is_rams_safety_source(
                Path("Mapei Latexplan NA safety sheet.pdf"),
                (
                    "Safety Data Sheet\n"
                    "Product Name(s): LATEXPLAN NA\n"
                    "Supplier Name & Address: Mapei UK Ltd\n"
                    "Recommended use: Cement based levelling mortar"
                ),
            )
        )

    def test_rebuild_file_3_safety_inventory_reclassifies_and_holds_sources(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            safety_root = workspace_root / "FILE_3_Inductions"
            rams_destination = safety_root / "RAMS"
            coshh_destination = safety_root / "COSHH"
            review_directory = safety_root / "Needs_Review"
            database_path = workspace_root / "documents.sqlite3"
            for directory in (rams_destination, coshh_destination, review_directory):
                directory.mkdir(parents=True, exist_ok=True)

            def _write_docx(path: Path, lines: list[str]) -> None:
                document = Document()
                for line in lines:
                    document.add_paragraph(line)
                document.save(path)

            misfiled_rams_path = coshh_destination / "Bluecord-RAMS-Cable Tray Install-Rev 2.docx"
            review_form_path = rams_destination / "UHSF20.1 Review Form - Bluecord.docx"
            _write_docx(
                misfiled_rams_path,
                [
                    "Risk Assessment and Method Statement",
                    "Activity Description: Cable Tray Install",
                    "Version: 2",
                    "Contractor: Bluecord",
                ],
            )
            _write_docx(
                review_form_path,
                [
                    "Review Form",
                    "RAMS TITLE: Cable Tray Install",
                    "Version: 1",
                ],
            )

            original_config = (
                app_config.BASE_DATA_DIR,
                app_config.DATABASE_PATH,
                app_config.INDUCTION_DIR,
                app_config.FILE_3_SAFETY_DIR,
                app_config.RAMS_DESTINATION,
                app_config.COSHH_DESTINATION,
                app_config.FILE_3_REVIEW_DIR,
            )
            try:
                app_config.BASE_DATA_DIR = workspace_root
                app_config.DATABASE_PATH = database_path
                app_config.INDUCTION_DIR = safety_root
                app_config.FILE_3_SAFETY_DIR = safety_root
                app_config.RAMS_DESTINATION = rams_destination
                app_config.COSHH_DESTINATION = coshh_destination
                app_config.FILE_3_REVIEW_DIR = review_directory

                repository = DocumentRepository(database_path)
                repository.create_schema()

                rebuild_result = workspace_module.rebuild_file_3_safety_inventory(
                    repository,
                    site_name="NG Lovedean Substation",
                )
            finally:
                (
                    app_config.BASE_DATA_DIR,
                    app_config.DATABASE_PATH,
                    app_config.INDUCTION_DIR,
                    app_config.FILE_3_SAFETY_DIR,
                    app_config.RAMS_DESTINATION,
                    app_config.COSHH_DESTINATION,
                    app_config.FILE_3_REVIEW_DIR,
                ) = original_config

            self.assertEqual(rebuild_result.total_sources, 2)
            self.assertEqual(rebuild_result.rams_records, 1)
            self.assertEqual(rebuild_result.coshh_records, 0)
            self.assertEqual(rebuild_result.ignored_sources, 1)
            self.assertTrue((rams_destination / misfiled_rams_path.name).exists())
            self.assertTrue((review_directory / review_form_path.name).exists())

            rams_documents = repository.list_documents(
                document_type=RAMSDocument.document_type,
                site_name="NG Lovedean Substation",
            )
            coshh_documents = repository.list_documents(
                document_type=COSHHDocument.document_type,
                site_name="NG Lovedean Substation",
            )
            self.assertEqual(len(rams_documents), 1)
            self.assertFalse(coshh_documents)

    def test_park_file_3_document_for_review_moves_linked_source(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            workspace_root = Path(temp_dir) / "Uplands_Workspace"
            rams_destination = workspace_root / "FILE_3_Inductions" / "RAMS"
            review_directory = workspace_root / "FILE_3_Inductions" / "Needs_Review"
            database_path = workspace_root / "documents.sqlite3"
            rams_destination.mkdir(parents=True, exist_ok=True)
            review_directory.mkdir(parents=True, exist_ok=True)

            source_path = rams_destination / "Bluecord-RAMS.docx"
            document = Document()
            document.add_paragraph("Risk Assessment and Method Statement")
            document.save(source_path)

            original_review_directory = app_config.FILE_3_REVIEW_DIR
            try:
                app_config.FILE_3_REVIEW_DIR = review_directory
                repository = DocumentRepository(database_path)
                repository.create_schema()
                rams_document = RAMSDocument(
                    doc_id="RAMS-bluecord-cable-tray",
                    site_name="NG Lovedean Substation",
                    created_at=datetime(2026, 3, 15, 8, 0),
                    status=DocumentStatus.ACTIVE,
                    contractor_name="Bluecord",
                    activity_description="Cable Tray Install",
                    approval_date=date(2026, 3, 15),
                    reference="RAMS-01",
                    version="1",
                )
                repository.save(rams_document)
                repository.index_file(
                    file_name=source_path.name,
                    file_path=source_path,
                    file_category="rams_docx",
                    file_group=FileGroup.FILE_3,
                    site_name=rams_document.site_name,
                    related_doc_id=rams_document.doc_id,
                )

                parked_paths = workspace_module.park_file_3_document_for_review(
                    repository,
                    rams_document.doc_id,
                )
            finally:
                app_config.FILE_3_REVIEW_DIR = original_review_directory

            self.assertEqual(len(parked_paths), 1)
            self.assertTrue(parked_paths[0].exists())
            self.assertEqual(parked_paths[0].parent.resolve(), review_directory.resolve())
            self.assertFalse(repository.list_documents(document_type=RAMSDocument.document_type))

    def test_build_file_3_review_candidates_flags_noisy_live_rows(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = DocumentRepository(Path(temp_dir) / "documents.sqlite3")
            repository.create_schema()
            noisy_rams = RAMSDocument(
                doc_id="RAMS-flagged",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 8, 0),
                status=DocumentStatus.ACTIVE,
                contractor_name="Electrical Work",
                activity_description="Severity (S)",
                approval_date=date(2026, 3, 15),
                reference="EC26 - Electrical Work",
                version="1.0",
            )
            clean_rams = RAMSDocument(
                doc_id="RAMS-clean",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 8, 0),
                status=DocumentStatus.ACTIVE,
                contractor_name="Uplands",
                activity_description="Barrier Matting",
                approval_date=date(2026, 3, 15),
                reference="RAMS-02",
                version="1.0",
            )

            candidates = app_module._build_file_3_review_candidates(
                repository,
                rams_documents=[noisy_rams, clean_rams],
                coshh_documents=[],
            )

            self.assertEqual(len(candidates), 1)
            self.assertEqual(candidates[0].doc_id, noisy_rams.doc_id)
            self.assertIn("Company", candidates[0].findings)
            self.assertIn("Activity", candidates[0].findings)

    def test_filter_file_3_review_candidates_supports_type_finding_and_search(self) -> None:
        candidates = [
            app_module.File3ReviewCandidate(
                document_type="RAMS",
                doc_id="RAMS-01",
                company="Electrical Work",
                title="Severity (S)",
                reference="EC26 - Electrical Work",
                version="1.0",
                findings=("Company", "Activity"),
                source_path=Path("/tmp/Electrical-Work.docx"),
            ),
            app_module.File3ReviewCandidate(
                document_type="COSHH",
                doc_id="COSHH-01",
                company="Limits",
                title="Mapei Latexplan NA",
                reference="COSHH ASSESSMENT - Mapei Latexplan NA",
                version="1.0",
                findings=("Supplier",),
                source_path=Path("/tmp/Mapei.pdf"),
            ),
        ]

        filtered_by_type = app_module._filter_file_3_review_candidates(
            candidates,
            document_type_filter="RAMS",
        )
        filtered_by_finding = app_module._filter_file_3_review_candidates(
            candidates,
            finding_filter="Title / Substance",
        )
        filtered_by_search = app_module._filter_file_3_review_candidates(
            candidates,
            search_query="mapei pdf",
        )

        self.assertEqual([candidate.doc_id for candidate in filtered_by_type], ["RAMS-01"])
        self.assertEqual([candidate.doc_id for candidate in filtered_by_finding], ["RAMS-01"])
        self.assertEqual([candidate.doc_id for candidate in filtered_by_search], ["COSHH-01"])

    def test_get_file_3_review_adjacent_doc_ids_returns_previous_and_next(self) -> None:
        candidates = [
            app_module.File3ReviewCandidate(
                document_type="RAMS",
                doc_id="RAMS-01",
                company="Uplands",
                title="Barrier Matting",
                reference="RAMS-01",
                version="1",
                findings=("Company",),
            ),
            app_module.File3ReviewCandidate(
                document_type="RAMS",
                doc_id="RAMS-02",
                company="Bluecord",
                title="Roof Works",
                reference="RAMS-02",
                version="1",
                findings=("Reference",),
            ),
            app_module.File3ReviewCandidate(
                document_type="COSHH",
                doc_id="COSHH-01",
                company="Limits",
                title="Latexplan",
                reference="COSHH-01",
                version="1",
                findings=("Supplier",),
            ),
        ]

        previous_doc_id, next_doc_id = app_module._get_file_3_review_adjacent_doc_ids(
            candidates,
            "RAMS-02",
        )

        self.assertEqual(previous_doc_id, "RAMS-01")
        self.assertEqual(next_doc_id, "COSHH-01")


class WorkspaceDoctorTests(unittest.TestCase):
    def test_run_workspace_diagnostic_reports_healthy_workspace(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            workspace_root = project_root / "Uplands_Workspace"
            templates_dir = project_root / "templates"
            file_2_output_dir = workspace_root / "FILE_2_Output"
            review_dir = workspace_root / "FILE_3_Inductions" / "Needs_Review"
            signatures_dir = workspace_root / "FILE_3_Inductions" / "Signatures"
            completed_dir = workspace_root / "FILE_3_Inductions" / "Completed_Inductions"
            database_path = workspace_root / "documents.sqlite3"
            template_path = templates_dir / "UHSF16.01_Template.docx"

            for directory in (
                workspace_root,
                templates_dir,
                file_2_output_dir,
                review_dir,
                signatures_dir,
                completed_dir,
            ):
                directory.mkdir(parents=True, exist_ok=True)
            database_path.touch()
            template_path.touch()

            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            with patch.object(app_config, "PROJECT_ROOT", project_root), patch.object(
                app_config, "BASE_DATA_DIR", workspace_root
            ), patch.object(app_config, "FILE_2_OUTPUT_DIR", file_2_output_dir), patch.object(
                app_config, "FILE_3_REVIEW_DIR", review_dir
            ), patch.object(
                app_config, "FILE_3_SIGNATURES_DIR", signatures_dir
            ), patch.object(
                app_config,
                "FILE_3_COMPLETED_INDUCTIONS_DIR",
                completed_dir,
            ), patch.object(
                app_config, "DATABASE_PATH", database_path
            ), patch.object(
                TemplateRegistry, "PROJECT_ROOT", project_root
            ):
                TemplateRegistry.TEMPLATE_PATHS = {
                    "site_induction": Path("templates/UHSF16.01_Template.docx")
                }
                try:
                    diagnostic_checks = run_workspace_diagnostic()
                finally:
                    TemplateRegistry.TEMPLATE_PATHS = original_registry

            self.assertTrue(diagnostic_checks)
            self.assertTrue(all(check.exists for check in diagnostic_checks))

    def test_run_workspace_diagnostic_reports_missing_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            workspace_root = project_root / "Uplands_Workspace"
            templates_dir = project_root / "templates"
            file_2_output_dir = workspace_root / "FILE_2_Output"
            review_dir = workspace_root / "FILE_3_Inductions" / "Needs_Review"
            signatures_dir = workspace_root / "FILE_3_Inductions" / "Signatures"
            completed_dir = workspace_root / "FILE_3_Inductions" / "Completed_Inductions"
            database_path = workspace_root / "documents.sqlite3"

            for directory in (
                workspace_root,
                templates_dir,
                file_2_output_dir,
                review_dir,
                signatures_dir,
                completed_dir,
            ):
                directory.mkdir(parents=True, exist_ok=True)
            database_path.touch()

            original_registry = dict(TemplateRegistry.TEMPLATE_PATHS)
            with patch.object(app_config, "PROJECT_ROOT", project_root), patch.object(
                app_config, "BASE_DATA_DIR", workspace_root
            ), patch.object(app_config, "FILE_2_OUTPUT_DIR", file_2_output_dir), patch.object(
                app_config, "FILE_3_REVIEW_DIR", review_dir
            ), patch.object(
                app_config, "FILE_3_SIGNATURES_DIR", signatures_dir
            ), patch.object(
                app_config,
                "FILE_3_COMPLETED_INDUCTIONS_DIR",
                completed_dir,
            ), patch.object(
                app_config, "DATABASE_PATH", database_path
            ), patch.object(
                TemplateRegistry, "PROJECT_ROOT", project_root
            ):
                TemplateRegistry.TEMPLATE_PATHS = {
                    "site_induction": Path("templates/UHSF16.01_Template.docx")
                }
                try:
                    diagnostic_checks = run_workspace_diagnostic()
                finally:
                    TemplateRegistry.TEMPLATE_PATHS = original_registry

            missing_checks = [check for check in diagnostic_checks if not check.exists]
            self.assertEqual(len(missing_checks), 1)
            self.assertEqual(
                missing_checks[0].display_path,
                "templates/UHSF16.01_Template.docx",
            )


class SiteDiaryDictationTests(unittest.TestCase):
    def tearDown(self) -> None:
        app_module.st.session_state.clear()

    def test_apply_site_diary_dictation_result_appends_transcript_once(self) -> None:
        target_key = "file2_site_diary_comments_2026-03-15"
        app_module.st.session_state.clear()
        app_module.st.session_state[target_key] = "Existing note"

        app_module._apply_site_diary_dictation_result(
            {
                "target": target_key,
                "transcript": "Additional dictated update",
                "nonce": "12345",
            },
            friendly_label="Today's Comments",
        )
        app_module._apply_site_diary_dictation_result(
            {
                "target": target_key,
                "transcript": "Additional dictated update",
                "nonce": "12345",
            },
            friendly_label="Today's Comments",
        )

        self.assertEqual(
            app_module.st.session_state[target_key],
            "Existing note\nAdditional dictated update",
        )
        self.assertEqual(
            app_module.st.session_state["file2_site_diary_dictation_flash"],
            "Dictation added to Today's Comments.",
        )

    def test_apply_site_diary_dictation_result_surfaces_browser_error(self) -> None:
        app_module.st.session_state.clear()

        app_module._apply_site_diary_dictation_result(
            {
                "target": "file2_site_diary_incidents_2026-03-15",
                "error": "Microphone permission was blocked by the browser.",
                "nonce": "error-1",
            },
            friendly_label="Incidents Details",
        )

        self.assertEqual(
            app_module.st.session_state["file2_site_diary_dictation_warning"],
            "Microphone permission was blocked by the browser.",
        )


class SiteInductionPrintPackTests(unittest.TestCase):
    def test_build_induction_evidence_rows_labels_manual_handling_certificate(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            manual_handling_pdf = temp_root / "2026-03-16_sean_arter_manual_handling_certificate_manual-handling.pdf"
            cscs_pdf = temp_root / "2026-03-16_sean_arter_cscs_card_cscs-front.pdf"
            manual_handling_pdf.write_bytes(b"pdf")
            cscs_pdf.write_bytes(b"pdf")

            induction = InductionDocument(
                doc_id="induction-1",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 16, 8, 0),
                status=DocumentStatus.ACTIVE,
                contractor_name="150",
                individual_name="Sean Carter",
                competency_card_paths=",".join([str(manual_handling_pdf), str(cscs_pdf)]),
            )

            evidence_rows = app_module._build_induction_evidence_rows(induction)

            self.assertEqual(evidence_rows[0]["Evidence"], "CSCS Card")
            self.assertEqual(
                evidence_rows[1]["Evidence"],
                "Manual Handling Certificate",
            )

    def test_build_induction_print_pack_paths_groups_docx_and_preview_files(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            induction_docx = temp_root / "induction.docx"
            cscs_pdf = temp_root / "cscs.pdf"
            first_aid_jpg = temp_root / "first-aid.jpg"
            extra_docx = temp_root / "certificate.docx"
            induction_docx.write_bytes(b"docx")
            cscs_pdf.write_bytes(b"pdf")
            first_aid_jpg.write_bytes(b"jpg")
            extra_docx.write_bytes(b"other-docx")

            induction = InductionDocument(
                doc_id="induction-1",
                site_name="NG Lovedean Substation",
                created_at=datetime(2026, 3, 15, 8, 0),
                status=DocumentStatus.ACTIVE,
                contractor_name="A. Archer Electrical",
                individual_name="Sean Carter",
                competency_card_paths=",".join(
                    [
                        str(cscs_pdf),
                        str(first_aid_jpg),
                        str(extra_docx),
                        str(temp_root / "missing.pdf"),
                    ]
                ),
                completed_document_path=str(induction_docx),
            )

            pack_paths = app_module._build_induction_print_pack_paths(induction)

            self.assertEqual(pack_paths["default"], [induction_docx, extra_docx])
            self.assertEqual(pack_paths["preview"], [cscs_pdf, first_aid_jpg])


class SiteInductionPosterTests(unittest.TestCase):
    def test_lookup_uk_postcode_details_returns_formatted_area(self) -> None:
        class _FakeResponse:
            def __enter__(self) -> "_FakeResponse":
                return self

            def __exit__(self, exc_type, exc, tb) -> bool:
                return False

            def read(self) -> bytes:
                return json.dumps(
                    {
                        "status": 200,
                        "result": {
                            "postcode": "CF44 9TZ",
                            "latitude": 51.742077,
                            "longitude": -3.505021,
                            "parish": "Hirwaun",
                            "admin_district": "Rhondda Cynon Taf",
                            "country": "Wales",
                        },
                    }
                ).encode("utf-8")

        with patch.object(workspace_module, "urlopen", return_value=_FakeResponse()):
            self.assertEqual(
                lookup_uk_postcode_details("cf449tz"),
                {
                    "postcode": "CF44 9TZ",
                    "latitude": 51.742077,
                    "longitude": -3.505021,
                    "locality": "Hirwaun",
                    "district": "Rhondda Cynon Taf",
                    "country": "Wales",
                    "formatted_address": "Hirwaun, Rhondda Cynon Taf, CF44 9TZ",
                },
            )

    def test_lookup_uk_postcode_coordinates_returns_lat_lng(self) -> None:
        class _FakeResponse:
            def __enter__(self) -> "_FakeResponse":
                return self

            def __exit__(self, exc_type, exc, tb) -> bool:
                return False

            def read(self) -> bytes:
                return json.dumps(
                    {
                        "status": 200,
                        "result": {
                            "postcode": "PO8 0SJ",
                            "latitude": 50.917,
                            "longitude": -1.036,
                        },
                    }
                ).encode("utf-8")

        with patch.object(workspace_module, "urlopen", return_value=_FakeResponse()):
            self.assertEqual(
                lookup_uk_postcode_coordinates("po8 0sj"),
                (50.917, -1.036, "PO8 0SJ"),
            )

    def test_lookup_uk_postcode_coordinates_returns_none_for_bad_response(self) -> None:
        class _FakeResponse:
            def __enter__(self) -> "_FakeResponse":
                return self

            def __exit__(self, exc_type, exc, tb) -> bool:
                return False

            def read(self) -> bytes:
                return json.dumps({"status": 404, "result": None}).encode("utf-8")

        with patch.object(workspace_module, "urlopen", return_value=_FakeResponse()):
            self.assertIsNone(lookup_uk_postcode_coordinates("BAD"))

    def test_load_project_setup_normalizes_lovedean_profile(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_setup_path = Path(temp_dir) / "project_setup.json"
            project_setup_path.write_text(
                json.dumps(
                    {
                        "current_site_name": "NG Lovedean Substation",
                        "job_number": "81888",
                        "site_address": "National Grid Sub Station Horndean Broadway Lane Waterlooville PO8 0SJ, CF44 9TZ",
                        "client_name": "National Grid",
                        "site_latitude": 51.742077,
                        "site_longitude": -3.505021,
                        "known_sites": [
                            {
                                "site_name": "NG Lovedean Substation",
                                "site_address": "National Grid Sub Station Horndean Broadway Lane Waterlooville PO8 0SJ, CF44 9TZ",
                                "client_name": "National Grid",
                                "job_number": "81888",
                                "site_latitude": 51.742077,
                                "site_longitude": -3.505021,
                                "geofence_radius_meters": 500,
                                "last_used_at": "2026-03-15T01:35:52",
                            },
                            {
                                "site_name": "11 station close",
                                "site_address": "11 station close, Hirwaun, Rhondda Cynon Taf, CF44 9TZ",
                                "client_name": "National Grid",
                                "job_number": "81888",
                                "site_latitude": 51.742077,
                                "site_longitude": -3.505021,
                                "geofence_radius_meters": 500,
                                "last_used_at": "2026-03-15T01:35:52",
                            },
                        ],
                    }
                ),
                encoding="utf-8",
            )

            with patch.object(app_module, "PROJECT_SETUP_PATH", project_setup_path):
                loaded_setup = app_module._load_project_setup()

            self.assertEqual(
                loaded_setup.site_address,
                app_module.LOVEDEAN_SITE_ADDRESS,
            )
            self.assertEqual(loaded_setup.site_latitude, app_module.LOVEDEAN_SITE_LATITUDE)
            self.assertEqual(loaded_setup.site_longitude, app_module.LOVEDEAN_SITE_LONGITUDE)
            lovedean_profile = next(
                profile
                for profile in loaded_setup.known_sites
                if profile.site_name == "NG Lovedean Substation"
            )
            self.assertEqual(lovedean_profile.site_address, app_module.LOVEDEAN_SITE_ADDRESS)
            self.assertEqual(lovedean_profile.site_latitude, app_module.LOVEDEAN_SITE_LATITUDE)
            self.assertEqual(lovedean_profile.site_longitude, app_module.LOVEDEAN_SITE_LONGITUDE)

    def test_get_site_induction_url_uses_local_ip_and_induction_station(self) -> None:
        with patch.object(
            workspace_module,
            "get_local_ip_address",
            return_value="192.168.1.88",
        ):
            self.assertEqual(
                get_site_induction_url(),
                "http://192.168.1.88:8501/?station=induction",
            )

    def test_generate_site_induction_poster_uses_kiosk_mode_with_local_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            logo_path = Path(temp_dir) / "uplands-logo.png"
            fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 50, 20), 0).save(str(logo_path))

            with patch.object(
                workspace_module,
                "get_local_ip_address",
                return_value="192.168.1.88",
            ):
                poster = generate_site_induction_poster(
                    site_name="NG Lovedean Substation",
                    logo_path=logo_path,
                )

        self.assertEqual(
            poster.induction_url,
            "https://uplands-site-induction.omegaleague.win/?station=induction&mode=kiosk",
        )
        self.assertTrue(poster.qr_code_png.startswith(b"\x89PNG\r\n\x1a\n"))
        self.assertTrue(poster.poster_png.startswith(b"\x89PNG\r\n\x1a\n"))

    def test_generate_site_induction_poster_prefers_public_tunnel_url(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            logo_path = Path(temp_dir) / "uplands-logo.png"
            fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 50, 20), 0).save(str(logo_path))

            poster = generate_site_induction_poster(
                site_name="NG Lovedean Substation",
                logo_path=logo_path,
                public_url="https://uplands.example.workers.dev",
            )

        self.assertEqual(
            poster.induction_url,
            "https://uplands-site-induction.omegaleague.win/?station=induction&mode=kiosk",
        )
        self.assertTrue(poster.qr_code_png.startswith(b"\x89PNG\r\n\x1a\n"))
        self.assertTrue(poster.poster_png.startswith(b"\x89PNG\r\n\x1a\n"))

    def test_detect_public_tunnel_url_from_log_accepts_custom_domain(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            original_tunnel_log_path = app_config.TUNNEL_LOG_PATH
            try:
                tunnel_log_path = Path(temp_dir) / "tunnel.log"
                tunnel_log_path.write_text(
                    "INFO Connected\nhttps://uplands-site-induction.omegaleague.win\n",
                    encoding="utf-8",
                )
                app_config.TUNNEL_LOG_PATH = tunnel_log_path
                self.assertEqual(
                    detect_public_tunnel_url_from_log(),
                    "https://uplands-site-induction.omegaleague.win",
                )
            finally:
                app_config.TUNNEL_LOG_PATH = original_tunnel_log_path

    def test_detect_public_tunnel_url_from_log_prioritizes_permanent_domain(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            original_tunnel_log_path = app_config.TUNNEL_LOG_PATH
            try:
                tunnel_log_path = Path(temp_dir) / "tunnel.log"
                tunnel_log_path.write_text(
                    (
                        "INFO Connected\n"
                        "https://random-subdomain.trycloudflare.com\n"
                        "https://uplands-site-induction.omegaleague.win\n"
                    ),
                    encoding="utf-8",
                )
                app_config.TUNNEL_LOG_PATH = tunnel_log_path
                self.assertEqual(
                    detect_public_tunnel_url_from_log(),
                    "https://uplands-site-induction.omegaleague.win",
                )
            finally:
                app_config.TUNNEL_LOG_PATH = original_tunnel_log_path

    def test_generate_site_induction_poster_returns_png_assets(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            logo_path = Path(temp_dir) / "uplands-logo.png"
            fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 50, 20), 0).save(str(logo_path))

            with patch.object(
                workspace_module,
                "get_local_ip_address",
                return_value="192.168.1.88",
            ):
                poster = generate_site_induction_poster(
                    site_name="NG Lovedean Substation",
                    logo_path=logo_path,
                )

        self.assertEqual(
            poster.induction_url,
            "https://uplands-site-induction.omegaleague.win/?station=induction&mode=kiosk",
        )
        self.assertTrue(poster.qr_code_png.startswith(b"\x89PNG\r\n\x1a\n"))
        self.assertTrue(poster.poster_png.startswith(b"\x89PNG\r\n\x1a\n"))


if __name__ == "__main__":
    unittest.main()
